"""Ten verification passes on the final Desktop thesis before delivery."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
METRICS = ROOT / "evaluation" / "metrics.json"
DEFAULT_PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)


def full_text(doc: Document) -> str:
    parts = [(p.text or "") for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def run_checks(paper: Path) -> tuple[int, list[str]]:
    if not paper.exists():
        return 0, [f"FAIL: missing file {paper}"]

    doc = Document(str(paper))
    text = full_text(doc)
    lower = text.lower()
    m = json.loads(METRICS.read_text(encoding="utf-8"))
    rf = f"{m['RandomForest']['accuracy'] * 100:.1f}"
    xgb = f"{m['XGBoost']['accuracy'] * 100:.1f}"
    iso = f"{m['IsolationForest']['accuracy'] * 100:.1f}"

    checks: list[tuple[str, bool, str]] = [
        (
            "1_structure",
            len(doc.paragraphs) >= 480 and len(doc.tables) >= 30,
            f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}",
        ),
        (
            "2_harvard_refs",
            "references" in lower and "(breiman, 2001)" in lower and "available at:" in lower,
            "Harvard-style references block",
        ),
        (
            "3_eval_12k",
            "12,000" in text and "39" in text and "six classes" in lower,
            "12k CICIDS-style evaluation dataset",
        ),
        (
            "4_metrics_match_app",
            rf in text and xgb in text and iso in text,
            f"RF {rf}% XGB {xgb}% IF {iso}%",
        ),
        (
            "5_sixteen_pages",
            "sixteen pages" in lower or "sixteen dashboard pages" in lower,
            "16 dashboard pages documented",
        ),
        (
            "6_app_modules",
            all(k in lower for k in ("copilot", "live evidence", "company")),
            "Copilot + Live Evidence + Company",
        ),
        (
            "7_no_email_channel",
            "smtp email" not in lower
            and not re.search(r"email alerts(?!\s+are not)", lower)
            and "notifications rely on telegram" in lower or "no smtp email" in lower,
            "Email not used as notification channel",
        ),
        (
            "8_objectives_o5_o6",
            "stored per-prediction" in lower or "per-prediction explainable" in lower,
            "O5 XAI wording",
        ),
        (
            "9_tables_11_12",
            any("Table 11" in (r.cells[0].text or "") for r in doc.tables[1].rows)
            and any("Table 12" in (r.cells[0].text or "") for r in doc.tables[1].rows),
            "LoT lists Table 11 and Table 12",
        ),
        (
            "10_figures_eval",
            any("Figure 14" in (r.cells[0].text or "") for r in doc.tables[0].rows)
            and any("Figure 17" in (r.cells[0].text or "") for r in doc.tables[0].rows)
            and any("Figure 21" in (r.cells[0].text or "") for r in doc.tables[0].rows),
            "LoF includes eval figures 14–21",
        ),
    ]

    o6_row = ""
    for row in doc.tables[3].rows:
        if (row.cells[0].text or "").strip() == "O6":
            o6_row = (row.cells[1].text or "").lower()
            break
    checks.append(
        (
            "11_o6_no_email",
            "email" not in o6_row or "no email" in o6_row,
            f"O6={o6_row[:80]}",
        )
    )

    tech = doc.tables[9]
    tech_blob = "\n".join(c.text for row in tech.rows for c in row.cells).lower()
    checks.append(
        (
            "12_tech_stack",
            "sqlite" in tech_blob
            and "streamlit" in tech_blob
            and "scapy" in tech_blob
            and ("no smtp email" in tech_blob or "telegram bot api" in tech_blob),
            "Technology table matches running stack",
        )
    )

    stale = "720 labeled samples" in text and "12,000" not in text
    checks.append(("13_no_stale_720_only", not stale, "720-only eval text removed"))

    checks.append(
        (
            "14_appendix_h",
            "appendix h" in lower,
            "Appendix H code section retained",
        )
    )

    passed = sum(1 for _, ok, _ in checks if ok)
    lines = []
    for name, ok, detail in checks:
        lines.append(f"{'PASS' if ok else 'FAIL'} {name}: {detail}")
    return passed, lines


def main() -> None:
    paper = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_PAPER
    passed, lines = run_checks(paper)
    total = len(lines)
    print(f"Verification {passed}/{total} on {paper}")
    for line in lines:
        print(line)
    if passed < total:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
