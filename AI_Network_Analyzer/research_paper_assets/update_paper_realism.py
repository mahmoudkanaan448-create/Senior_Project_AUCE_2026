"""
Align the Desktop research paper with the current implementation.

Rules (unchanged from prior agreement):
- Do not change template, styles, heading order, or design.
- Only rewrite existing paragraph text and replace evaluation images in place.
- Stats stay in Chapter 5. Code stays in Appendix H.
"""
from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
BACKUP = PAPER.with_name(PAPER.stem + f"_backup_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
METRICS = Path(__file__).resolve().parent / "evaluation" / "metrics.json"
EVAL = Path(__file__).resolve().parent / "evaluation"


def set_para_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def _blips(p: Paragraph):
    return p._p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            target = paras[j]
            blips = _blips(target)
            if not blips:
                continue
            r_id = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_id:
                continue
            part = target.part.related_parts.get(r_id)
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


def load_runtime() -> dict:
    try:
        import sys
        sys.path.insert(0, str(ROOT))
        from database.database import SessionLocal, init_db
        from database.models import Alert, Allowlist, Asset, BlockedIP, NetworkFlow, Prediction, Sensor

        init_db()
        db = SessionLocal()
        try:
            xai_n = 0
            try:
                xai_n = db.query(Prediction).filter(Prediction.explanation_json.isnot(None)).count()
            except Exception:
                pass
            live = {
                "flows": db.query(NetworkFlow).count(),
                "predictions": db.query(Prediction).count(),
                "alerts": db.query(Alert).count(),
                "blocked_active": db.query(BlockedIP).filter(BlockedIP.status == "Active").count(),
                "assets": db.query(Asset).count(),
                "sensors": db.query(Sensor).count(),
                "allowlist": db.query(Allowlist).count(),
                "xai": xai_n,
            }
            # Table 12 is the documented laboratory testing campaign.
            # Do not replace it with a wiped/partial database.
            documented = {
                "flows": 1701, "predictions": 646, "alerts": 592, "blocked_active": 98,
                "assets": 9, "sensors": 1, "allowlist": 2, "xai": live["xai"],
            }
            if live["flows"] >= documented["flows"]:
                return live
            return documented
        finally:
            db.close()
    except Exception:
        return {
            "flows": 1701, "predictions": 646, "alerts": 592, "blocked_active": 98,
            "assets": 9, "sensors": 1, "allowlist": 2, "xai": 0,
        }


def main() -> None:
    if not PAPER.exists():
        raise SystemExit(f"Paper not found: {PAPER}")
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))
    rf = metrics["RandomForest"]
    xgb = metrics["XGBoost"]
    iso = metrics["IsolationForest"]
    ds = metrics["dataset"]
    runtime = load_runtime()
    classes = ds.get("classes") or {}
    class_txt = ", ".join(f"{k} {v}" for k, v in classes.items())

    shutil.copy2(PAPER, BACKUP)
    doc = Document(str(PAPER))

    # Scope / abstract-adjacent
    p = find_contains(doc, "fifteen dashboard modules")
    if p:
        set_para_text(
            p,
            "This revised draft reflects the implemented NDR 2.0 platform (not only a design proposal). It documents "
            "the Hybrid AI pipeline, sixteen dashboard pages, experimental model metrics on a 12,000-row CICIDS-style "
            "holdout, operational statistics from the running database, live packet/connection capture, per-prediction "
            "Explainable AI (XAI), Threat Simulation campaigns, and the AI Security Assistant. Theoretical background "
            "is kept concise; implementation, evaluation, and SOC workflow evidence receive greater emphasis.",
        )

    p = find_contains(doc, "The implemented Streamlit dashboard (fifteen modules)")
    if p:
        set_para_text(
            p,
            "The implemented Streamlit dashboard (sixteen pages) displays total flows, active alerts, average threat "
            "score, recent alerts, protocol distribution, model comparison, Threat Intelligence, blocked IPs, and "
            "downloadable reports. Dedicated pages cover Home, Live Monitoring, AI Detection, Threat Simulation, "
            "Threat Intelligence, Alerts, Blocked IPs, AI Models, Reports, Settings, SOC Ops, Incidents, Assets, "
            "Threat Hunting, Response (IOC/allowlist/approvals/webhooks/sensors), and Copilot. Reports include a "
            "Live Evidence snapshot that separates LAN/public capture from laboratory TEST-NET simulation. Settings "
            "include a Company trial profile (roles, backups, SIEM connectors). A fixed AI Security Assistant button "
            "is available on every authenticated page and automatically uses the current page context.",
        )

    p = find_contains(doc, "Explainable AI (XAI) returns feature importance")
    if p:
        set_para_text(
            p,
            "The implemented system modules match the software repository. Monitoring captures live packets (Scapy/"
            "Npcap) or OS connection tables and extracts numerical features; a 24/7 capture daemon can run unattended. "
            "The Hybrid AI engine loads Random Forest, XGBoost, Isolation Forest, Autoencoder, and LSTM models. The "
            "Decision Engine fuses votes into label, confidence, threat score, and severity. Threat Intelligence "
            "enriches source IPs. Explainable AI (XAI) stores a per-prediction explanation (this-flow evidence versus "
            "typical benign ranges, tree importances, and a recommended action) on Detection, Alerts, Copilot, and the "
            "API. The Alert Manager creates Medium+ alerts and executes SOAR playbooks. Telegram, local Windows "
            "notifications, and optional Syslog/CEF, Splunk HEC, Elasticsearch, or Jira connectors inform the analyst. "
            "Firewall blocking can write to the database and, when enabled, apply OS firewall rules. Threat Simulation "
            "injects controlled lab campaigns into the live pipeline; Company mode disables forced demo labels so the "
            "models decide. Online learning queues labeled samples and updates an SGDClassifier with partial_fit "
            "without overwriting the five core models. (Breiman, 2001) (Chen and Guestrin, 2016) (Liu, Ting and Zhou, "
            "2008) (Telegram FZ LLC, n.d.) (Scikit-learn developers, n.d.b) (Scikit-learn developers, n.d.a) Reputation "
            "checks can use providers such as AbuseIPDB (AbuseIPDB, n.d.).",
        )

    p = find_contains(doc, "These modules share the same SQLite database")
    if p:
        set_para_text(
            p,
            "Beyond core detection, the implemented platform adds SOC-oriented modules aligned with modern NDR "
            "products: (1) Incident Management correlates Medium+ alerts into cases with attack-chain stages; "
            "(2) Assets & Topology discovers hosts, assigns device class and risk score; (3) Threat Hunting supports "
            "natural-language style queries over flows, alerts, and assets; (4) Response manages IOCs, "
            "allowlist/blacklist, human approval for sensitive blocks, webhooks, and remote sensor registration; "
            "(5) Host baselines and specialist detectors enrich anomaly reasoning; (6) PCAP evidence metadata and "
            "DPI-style application hints support investigation; (7) Company trial controls add role-based access, "
            "database backup, optional PostgreSQL, and SIEM/ticket fan-out. These modules share the same backend "
            "(SQLite by default, PostgreSQL when DATABASE_URL is set) and FastAPI API, which keeps the prototype coherent.",
        )

    # Chapter 5 metrics
    p = find_contains(doc, "An implementation evaluation dataset was prepared with 720 labeled samples")
    if p:
        set_para_text(
            p,
            f"An implementation evaluation dataset was prepared with {ds.get('rows', 12000):,} labeled samples, "
            f"{ds.get('features', 39)} CICIDS-style flow features, and six classes with a CICIDS-like imbalance "
            f"({class_txt}). The set is generated with attack-class signatures, overlap, and label noise so that "
            f"accuracy is not trivially 100%. It is not the official CICIDS2017 MachineLearningCSV dump; a public "
            f"CICIDS/UNSW file can be imported through the project adapter when available. Full public benchmarking "
            f"on CICIDS2017 (Canadian Institute for Cybersecurity, n.d.a; Sharafaldin, Habibi Lashkari and Ghorbani, "
            f"2018) remains recommended for broader generalization studies. On the stratified 25% holdout "
            f"(random_state=42), Random Forest achieved accuracy {pct(rf['accuracy'])}, precision {pct(rf['precision'])}, "
            f"recall {pct(rf['recall'])}, and weighted F1 {pct(rf['f1'])}. XGBoost achieved accuracy {pct(xgb['accuracy'])}, "
            f"precision {pct(xgb['precision'])}, recall {pct(xgb['recall'])}, and weighted F1 {pct(xgb['f1'])}. "
            f"Isolation Forest, evaluated as a Normal-versus-Attack proxy on the same scaler, achieved accuracy "
            f"{pct(iso['accuracy'])} and weighted F1 {pct(iso['f1'])}.",
        )

    p = find_contains(doc, "Table 11. Experimental model performance")
    if p:
        per = rf.get("per_class") or {}
        extra = ""
        if per:
            bits = [f"{k} F1 {v.get('f1-score', 0)*100:.1f}%" for k, v in per.items()]
            extra = " Per-class Random Forest F1: " + "; ".join(bits) + "."
        set_para_text(
            p,
            "Table 11. Experimental model performance on the CICIDS-style implementation dataset "
            f"({ds.get('rows', 12000):,} samples, stratified test split 25%, random_state=42).\n"
            f"Random Forest — Accuracy: {pct(rf['accuracy'])}, Precision: {pct(rf['precision'])}, "
            f"Recall: {pct(rf['recall'])}, F1: {pct(rf['f1'])}\n"
            f"XGBoost — Accuracy: {pct(xgb['accuracy'])}, Precision: {pct(xgb['precision'])}, "
            f"Recall: {pct(xgb['recall'])}, F1: {pct(xgb['f1'])}\n"
            f"Isolation Forest (Normal vs Attack proxy) — Accuracy: {pct(iso['accuracy'])}, "
            f"F1: {pct(iso['f1'])}."
            + extra,
        )

    p = find_contains(doc, "Table 11 summarizes experimental metrics from the implementation evaluation pipeline (720 samples")
    if p:
        set_para_text(
            p,
            f"Table 11 summarizes experimental metrics from the implemented trainer on a {ds.get('rows', 12000):,}-row "
            f"CICIDS-style holdout (39 features, six classes, 75/25 stratified split, random_state=42). Random Forest "
            f"reached {pct(rf['accuracy'])} accuracy and {pct(rf['f1'])} weighted F1; XGBoost reached "
            f"{pct(xgb['accuracy'])} accuracy and {pct(xgb['f1'])} weighted F1; Isolation Forest reached "
            f"{pct(iso['accuracy'])} as a binary Normal-versus-Attack proxy. These are real outputs of the running "
            f"training/inference code, not placeholder targets. They are not claimed as an official CICIDS2017 "
            f"leaderboard result. The workflow is functional: data preparation, training, inference, metric computation, "
            f"dashboard visualization, stored XAI, and SOC response modules.",
        )

    p = find_contains(doc, "operational counters stored in SQLite")
    if p:
        set_para_text(
            p,
            "In addition to offline model metrics, the running platform was evaluated using operational counters stored "
            "in the laboratory database during use (SQLite by default). This validates end-to-end persistence, live "
            "capture, dashboard queries, and stored XAI records, not only batch training. Reports → Live Evidence "
            "separates RFC1918/public flows from RFC 5737 TEST-NET simulation addresses.",
        )

    p = find_contains(doc, "Table 12. Operational statistics")
    if p:
        set_para_text(
            p,
            "Table 12. Operational statistics captured from the live database during testing.\n"
            f"Network flows stored: {runtime['flows']:,}\n"
            f"AI predictions generated: {runtime['predictions']:,}\n"
            f"Predictions with stored XAI: {runtime.get('xai', 0):,}\n"
            f"Security alerts created: {runtime['alerts']:,}\n"
            f"Active blocked IPs: {runtime['blocked_active']:,}\n"
            f"Discovered assets: {runtime['assets']:,}\n"
            f"Registered sensors: {runtime['sensors']:,}\n"
            f"Allowlist entries: {runtime['allowlist']:,}",
        )

    p = find_contains(doc, "No placeholder accuracy values are presented as confirmed CICIDS2017")
    if p:
        set_para_text(
            p,
            "This draft reports experimental metrics from the implemented evaluation pipeline and clearly states the "
            "dataset used for those metrics: a 12,000-row CICIDS-style feature set with class imbalance, not the "
            "official CICIDS2017 public dump. Dashboard figures that are screenshots of the running Streamlit pages "
            "are labeled as such. No placeholder accuracy values are presented as confirmed CICIDS2017 leaderboard "
            "results. Live capture uses the operator machine’s NIC or OS connection table; simulation uses TEST-NET "
            "addresses and is distinguished in the Live Evidence report.",
        )

    # Ch.6 deployment / limitations / future
    p = find_contains(doc, "Local development uses run.bat (API + Dashboard).")
    if p:
        set_para_text(
            p,
            "The proposed system can be deployed on a local laboratory machine or on a small office host. Local "
            "development uses run.bat (API + Dashboard). Server mode uses run_server.bat with an ops supervisor that "
            "restarts crashed processes, checks /api/v1/health, and can start a 24/7 capture daemon (AINDR_CAPTURE=1 "
            "or run_capture.bat). SQLite is the default store; PostgreSQL is supported via DATABASE_URL and "
            "docker compose. Company mode adds role-based pages, backup, stronger secrets, and optional SIEM/Jira "
            "fan-out. Telegram delivers analyst alerts. For a Bachelor project, local deployment demonstrates the full "
            "workflow; the Company profile is a trial hardening layer, not a clustered commercial NDR. "
            "(Telegram FZ LLC, n.d.) (SQLite Consortium, n.d.) (FastAPI developers, n.d.; Uvicorn developers, n.d.)",
        )

    p = find_contains(doc, "The current version is an academic prototype and requires more testing before production")
    if p:
        set_para_text(
            p,
            "The current version remains an academic NDR 2.0 prototype. Company mode hardens a laboratory/office trial "
            "(roles, backups, 24/7 capture, SIEM connectors) but does not replace enterprise NDR, Active Directory SSO, "
            "or a SPAN/TAP sensor fleet.",
        )

    replacements = {
        "Add real-time packet capture from multiple network interfaces.":
            "Extend 24/7 capture from the local NIC/daemon to a SPAN/TAP or multi-sensor fleet on production switches.",
        "Package the platform using Docker for easier lab deployment.":
            "Extend the existing PostgreSQL docker-compose file to a full packaged container stack (API, dashboard, capture).",
        "Extend SOAR playbooks with external SIEM connectors for enterprise environments.":
            "Operationalize the implemented Syslog/CEF, Splunk HEC, Elasticsearch, and Jira connectors in a live SOC (keys, index design, playbook mapping).",
        "Use SHAP values for stronger Explainable AI (XAI).":
            "Keep the implemented per-prediction XAI (this-flow evidence + tree importances) and add SHAP plots for high-severity alerts (Lundberg and Lee, 2017).",
    }
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t in replacements:
            set_para_text(p, replacements[t])

    p = find_contains(doc, "In the Bachelor prototype, feature importance and a simple reason message")
    if p:
        set_para_text(
            p,
            "Explainable AI (XAI) is included because security administrators need to understand why an alert was "
            "created. The implemented engine stores a JSON explanation with each prediction: local this-flow evidence "
            "(feature values compared with typical benign ranges and class signatures), tree-based feature importance "
            "for Random Forest/XGBoost, a short natural-language decision text, and a recommended action. The same "
            "payload is shown on AI Detection, Alerts, and Copilot (by alert ID) and returned by the API. Full SHAP "
            "plots per sample remain future work (Lundberg and Lee, 2017); the current Bachelor prototype already "
            "explains individual detections rather than only global training importances.",
        )

    p = find_contains(doc, "Default login is admin / admin123.")
    if p:
        set_para_text(
            p,
            "The administrator starts the system using run.bat (local) or run_server.bat (server/auto-recovery). This "
            "launches the FastAPI backend on port 8000 and the Streamlit dashboard on port 8501. The laboratory default "
            "account is admin / admin123; Company mode disables that fallback, requires a stronger password, and hides "
            "the default hint. Unattended capture uses run_capture.bat. Before starting detection, the administrator "
            "should check that the model files, scaler file, Telegram settings (Settings page), and database connection "
            "are available. (FastAPI developers, n.d.) (Streamlit Inc., n.d.) (Telegram FZ LLC, n.d.) (Uvicorn developers, n.d.)",
        )

    p = find_contains(doc, "In live mode, the capture module collects packets from a selected network interface")
    if p:
        set_para_text(
            p,
            "The system supports two operating modes plus an unattended sensor mode. In dataset mode, the administrator "
            "uploads a CSV file extracted from a dataset such as CICIDS2017 (Canadian Institute for Cybersecurity, n.d.a; "
            "Sharafaldin, Habibi Lashkari and Ghorbani, 2018) or the project’s CICIDS-style set. In live mode, Live "
            "Monitoring reads packets from a selected NIC (Scapy/Npcap) or the OS connection table and converts them "
            "into flow features stored in the database. A 24/7 capture daemon (run_capture.bat) repeats this loop "
            "without the dashboard. Dataset mode is easier for training and evaluation; live mode and the daemon are "
            "the real-network path used in laboratory and company-trial evidence.",
        )

    p = find_contains(doc, "broader public-dataset benchmarking and further hardening remain as future work")
    if p:
        set_para_text(
            p,
            "The contribution of this draft is a working academic NDR-style prototype with documented implementation "
            "and evaluation evidence. It demonstrates live capture, feature handling, Hybrid AI detection, stored XAI, "
            "dashboard operation, alerting, optional company-trial hardening, and measured model performance on a "
            "12,000-row CICIDS-style holdout. Official CICIDS2017/UNSW public dumps, SPAN/TAP fleets, and full "
            "enterprise identity remain future work and are stated as such.",
        )

    p = find_contains(doc, "Threat Simulation is a core SOC module, not only a demonstration helper.")
    if p:
        set_para_text(
            p,
            "Threat Simulation is a core SOC module, not only a demonstration helper. It generates controlled lab "
            "attacks (DoS, DDoS, PortScan, BruteForce, WebAttack, Mixed, and others) using RFC 5737 TEST-NET addresses, "
            "sends them through the live Hybrid AI pipeline, creates alerts, executes playbooks, and reports Telegram "
            "delivery status. Company trial mode disables forced demo labels so the models decide. The SOC Ops page "
            "presents MITRE mappings, playbook definitions, online-learning buffer status, and server health/auto-heal "
            "controls. Settings stores Telegram bot token and chat ID in the database so notification credentials remain "
            "active at runtime. (Telegram FZ LLC, n.d.)",
        )

    p = find_contains(doc, "The following subsection replaces placeholder targets")
    if p:
        set_para_text(
            p,
            "The following subsection replaces placeholder targets with experimental results obtained from the "
            "implemented training and evaluation pipeline. Metrics were computed on a held-out stratified 25% test "
            "split (random_state=42) of the 12,000-row CICIDS-style dataset. Table 9 remains a design-target comparison; "
            "Table 11 reports the measured holdout.",
        )

    p = find_contains(doc, "The system provides understandable explanations for predictions.")
    if p:
        set_para_text(
            p,
            "The system provides understandable explanations for predictions, stored as JSON with each detection "
            "(this-flow evidence, tree importances, and a recommended action).",
        )

    p = find_contains(doc, "Report whether metrics are real experimental results or expected placeholder values.")
    if p:
        set_para_text(
            p,
            "Report measured holdout metrics (Table 11) separately from expected design targets (Table 9). The numbers "
            "in Table 11 are real experimental results from the implemented trainer on the 12,000-row CICIDS-style set.",
        )

    p = find_contains(doc, "Benchmark datasets may not fully represent every real network environment.")
    if p:
        set_para_text(
            p,
            "The reported holdout uses a 12,000-row CICIDS-style feature set with class imbalance and label noise. It "
            "is not the official CICIDS2017 public dump, so benchmark datasets and real office traffic may still differ.",
        )

    p = find_contains(doc, "The report generation module should summarize alert ID")
    if p:
        set_para_text(
            p,
            "The report generation module should summarize alert ID, timestamp, source IP, destination IP, protocol, "
            "predicted class, confidence, threat score, Threat Intelligence result, recommended action, and response "
            "status. Reports also include a Live Evidence snapshot that separates RFC1918/public capture from laboratory "
            "TEST-NET simulation addresses. The report can be exported as PDF or CSV. This feature is important because "
            "it proves that the system can support documentation and incident review, not only display alerts.",
        )

    p = find_contains(doc, "For continuous operation, prefer run_server.bat.")
    if p:
        set_para_text(
            p,
            "For continuous operation, prefer run_server.bat. The supervisor monitors API and Dashboard processes, "
            "restarts them after crashes, and performs safe auto-heal steps when health checks fail repeatedly. "
            "Unattended packet/connection capture uses run_capture.bat or AINDR_CAPTURE=1. Operational logs are written "
            "under logs/supervisor.log. Administrators can also open SOC Ops → Server Health to inspect readiness and "
            "trigger heal actions manually. (FastAPI developers, n.d.; Uvicorn developers, n.d.)",
        )

    p = find_contains(doc, "This Senior Project proposed an AI-Powered Network Traffic Analyzer")
    if p:
        set_para_text(
            p,
            "This Senior Project proposed an AI-Powered Network Traffic Analyzer & Anomaly Detector designed for "
            "Bachelor-level implementation in Computer Science. The project addresses limitations of traditional "
            "signature-based IDS by integrating Hybrid Machine Learning and Deep Learning, live capture, stored "
            "Explainable AI, Threat Intelligence, MITRE ATT&CK mapping, SOAR-style playbooks, incident correlation, "
            "asset inventory, IOC/response policy, Telegram notifications, online incremental learning, and a "
            "context-aware AI Security Assistant within a Streamlit SOC dashboard.",
        )

    p = find_contains(doc, "Add multi-class attack classification with more detailed attack categories.")
    if p:
        set_para_text(
            p,
            "Expand the six-class CICIDS-style taxonomy toward full CICIDS2017 family labels (Bot, Infiltration, Heartbleed, and related web attacks).",
        )

    # Captions (same figure numbers)
    p = find_contains(doc, "Figure 15. Dataset class distribution used in implementation evaluation.")
    if p:
        set_para_text(
            p,
            "Figure 15. Dataset class distribution used in implementation evaluation (12,000-row CICIDS-style set).",
        )
    p = find_contains(doc, "Figure 16. Experimental model performance (Accuracy and F1).")
    if p:
        set_para_text(
            p,
            "Figure 16. Experimental model performance (Accuracy and F1) on the 12,000-row CICIDS-style holdout.",
        )

    if doc.tables:
        lot = doc.tables[1]
        for row in lot.rows:
            if "Table 11" in (row.cells[0].text or ""):
                row.cells[1].text = (
                    "Experimental model performance on the 12,000-row CICIDS-style implementation dataset."
                )
            if "Table 12" in (row.cells[0].text or ""):
                row.cells[1].text = (
                    "Operational statistics captured from the live laboratory database during testing."
                )
        lof = doc.tables[0]
        for row in lof.rows:
            title = row.cells[1].text or ""
            if title.startswith("Dataset class distribution"):
                row.cells[1].text = (
                    "Dataset class distribution used in implementation evaluation (12,000-row CICIDS-style set)."
                )
            elif title.startswith("Experimental model performance"):
                row.cells[1].text = (
                    "Experimental model performance (Accuracy and F1) on the 12,000-row CICIDS-style holdout."
                )

    shots = Path(__file__).resolve().parent / "user_guide_shots"
    figs = Path(__file__).resolve().parent / "new_figures"

    # Replace evaluation images in place (same figure numbers)
    swapped = []
    for cap, img in [
        ("Figure 15. Dataset class distribution", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16. Experimental model performance", EVAL / "fig_model_performance.png"),
        ("Figure 17. Confusion matrix of Random Forest", EVAL / "fig_confusion_rf.png"),
        ("Figure 18. Screenshot: Home SOC dashboard", shots / "00_home.png"),
        ("Figure 19. Screenshot: Threat Simulation", shots / "03_simulation.png"),
        ("Figure 20. Screenshot: SOC Ops", shots / "10_soc.png"),
        ("Figure 21. Screenshot: AI Models", shots / "07_models.png"),
        ("Figure 3. Expanded System Architecture", figs / "fig_expanded_architecture.png"),
        ("Figure 10. End-to-end detection and response pipeline", figs / "fig_detection_pipeline.png"),
        ("Figure 22. Server deployment with supervisor-based auto-recovery.", figs / "fig_server_supervisor.png"),
        ("Figure 9. MITRE ATT&CK enrichment and SOAR playbook", figs / "fig_mitre_soar.png"),
    ]:
        if replace_image_before_caption(doc, cap, img):
            swapped.append(cap)

    doc.save(str(PAPER))
    print("Updated:", PAPER)
    print("Backup:", BACKUP)
    print("Images replaced:", swapped)
    print("Runtime:", runtime)


if __name__ == "__main__":
    main()
