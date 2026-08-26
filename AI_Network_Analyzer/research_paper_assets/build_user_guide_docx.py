"""Capture every dashboard page and build an Arabic user-guide Word file on the Desktop."""

from __future__ import annotations

import base64
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.options import Options
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

BASE = "http://localhost:8501"
SHOTS = Path(__file__).resolve().parent / "user_guide_shots"
SHOTS.mkdir(parents=True, exist_ok=True)
OUT = Path(r"C:\Users\mohamad\OneDrive\Desktop\AI_Network_Analyzer_User_Guide_Pages.docx")

PAGES = [
    {
        "slug": "",
        "nav": "Home",
        "file": "00_home.png",
        "title": "1) الرئيسية — Home",
        "role": "لوحة SOC العامة: نظرة واحدة على حالة الشبكة والأمن.",
        "items": [
            "مؤشرات أعلى الصفحة: عدد التدفقات، التنبؤات، الهجمات، التنبيهات الجديدة، العناوين المحظورة، متوسط التهديد، الحوادث المفتوحة، الأصول/الحساسات.",
            "رسوم الحركة: توزيع طبيعي مقابل هجوم، أنواع الهجمات، وتوزيع الخطورة.",
            "مقارنة أداء نماذج الذكاء الاصطناعي (مخطط رادار + جدول الدقة).",
            "اختصارات سريعة لباقي الصفحات (محاكاة، مراقبة، حوادث، أصول، صيد، مساعد…).",
            "جداول أحدث التنبيهات، أحدث التنبؤات، والعناوين المحظورة.",
        ],
        "benefit": "تعرف فوراً إذا في هجوم، قديش الخطورة، ووين تروح بعدين.",
        "tabs": [],
    },
    {
        "slug": "/Live_Monitoring",
        "nav": "Live Monitoring",
        "file": "01_live.png",
        "title": "2) المراقبة المباشرة — Live Monitoring",
        "role": "مصدر بيانات الحركة الحية من جهازك.",
        "items": [
            "اختيار واجهة الشبكة (Wi-Fi / Ethernet).",
            "أزرار بدء وإيقاف الالتقاط المباشر.",
            "حالة الاتصال (متصل/خامل) وعداد الحزم والتدفقات.",
            "مؤشرات الحساسات والحزم في الثانية والحزم المفقودة.",
            "جدول الاتصالات النشطة + أحدث التدفقات المخزّنة في قاعدة البيانات.",
        ],
        "benefit": "بدون هالصفحة ما في حركة حقيقية للكشف. هي تغذي باقي النظام.",
        "tabs": [],
    },
    {
        "slug": "/AI_Detection",
        "nav": "AI Detection",
        "file": "02_detection.png",
        "title": "3) كشف الذكاء الاصطناعي — AI Detection",
        "role": "تشغيل محرك Hybrid AI على التدفقات غير المعالجة.",
        "items": [
            "مؤشرات: إجمالي التنبؤات، الهجمات، الحركة الطبيعية، متوسط الثقة.",
            "زر تحليل كل التدفقات غير المعالجة عبر النماذج الهجينة.",
            "رابط لمحاكاة التهديدات.",
            "جدول نتائج الكشف مع الثقة ودرجة التهديد والخطورة.",
            "رسوم توزيع التصنيفات وتوزيع الخطورة.",
        ],
        "benefit": "تحوّل الحركة الخام إلى قرار أمني: طبيعي أو هجوم، ونوعه، وهل يستاهل تنبيه.",
        "tabs": [],
    },
    {
        "slug": "/Threat_Simulation",
        "nav": "Threat Simulation",
        "file": "03_simulation.png",
        "title": "4) محاكاة التهديدات — Threat Simulation",
        "role": "حقن حملات هجوم مختبر في نفس مسار الكشف الحقيقي.",
        "items": [
            "اختيار نوع الحملة (DoS، PortScan، BruteForce، Mixed…).",
            "عدد الحوادث المراد توليدها.",
            "خيارات الاستجابة: إنشاء تنبيهات، حظر IP حرج، إشعار تيليغرام.",
            "زر إطلاق الحملة ونتائجها (تدفقات، هجمات، تنبيهات، تيليغرام).",
            "قائمة الحملات المتاحة مع وصف كل سيناريو.",
        ],
        "benefit": "تجرّب النظام أمام اللجنة من غير انتظار هجوم فعلي على الشبكة.",
        "tabs": [],
    },
    {
        "slug": "/Threat_Intelligence",
        "nav": "Threat Intelligence",
        "file": "04_ti.png",
        "title": "5) الاستخبارات التهديدية — Threat Intelligence",
        "role": "فحص سمعة العناوين وموقعها الجغرافي.",
        "items": [
            "حقل إدخال IP وزر البحث.",
            "معلومات السمعة ودرجة التهديد.",
            "الموقع الجغرافي على الخريطة.",
            "تقييم الخطورة (منخفض / متوسط / عالٍ / حرج).",
            "سجلات الاستخبارات المخزّنة ورسم أعلى الدول المصدر.",
        ],
        "benefit": "قبل الحظر تعرف إذا العنوان مصدر هجوم معروف أو عنوان عادي.",
        "tabs": [],
    },
    {
        "slug": "/Alerts",
        "nav": "Alerts",
        "file": "05_alerts.png",
        "title": "6) التنبيهات — Alerts",
        "role": "صندوق وارد المحلل لكل اكتشاف مهم.",
        "items": [
            "أعداد: الكل، جديد، قيد التحقيق، مغلق.",
            "تصفية حسب الحالة وجدول التنبيهات مع تعيين MITRE.",
            "تحديث حالة تنبيه معيّن.",
            "ملاحظات المحلل: هجوم حقيقي أو إيجابي كاذب (للتعلم المستمر).",
        ],
        "benefit": "تنظّم الشغل اليومي وتعلّم النظام يقلل الإنذارات الغلط.",
        "tabs": [],
    },
    {
        "slug": "/Blocked_IPs",
        "nav": "Blocked IPs",
        "file": "06_blocked.png",
        "title": "7) العناوين المحظورة — Blocked IPs",
        "role": "سجل الحظر اليدوي والتلقائي.",
        "items": [
            "أعداد الحظر الكلي والنشط والمُزال.",
            "نموذج حظر عنوان مع السبب.",
            "قائمة العناوين النشطة.",
            "إلغاء حظر عنوان بالخطأ.",
        ],
        "benefit": "تمنع مصدر الهجوم وتقدر ترجع عن الحظر الغلط.",
        "tabs": [],
    },
    {
        "slug": "/AI_Models",
        "nav": "AI Models",
        "file": "07_models.png",
        "title": "8) نماذج الذكاء الاصطناعي — AI Models",
        "role": "إدارة جودة النماذج وتدريبها.",
        "items": [
            "جدول الدقة وPrecision وRecall وF1 لكل نموذج.",
            "مخطط رادار ومقارنة الدقة.",
            "رفع ملف CSV وتدريب كل النماذج.",
            "تعلم مستمر (SGD) من عينات الهجوم المؤكدة دون استبدال النماذج الأساسية.",
        ],
        "benefit": "تعرف أي نموذج أقوى وتعيد التدريب عند تغيّر البيانات.",
        "tabs": [],
    },
    {
        "slug": "/Reports",
        "nav": "Reports",
        "file": "08_reports.png",
        "title": "9) التقارير — Reports",
        "role": "توثيق النتائج وتصديرها للتسليم.",
        "items": [
            "تبويب تقرير الكشف مع تنزيل CSV.",
            "تبويب تقارير الحوادث وإنشاء تقرير جديد.",
            "تبويب التصدير: تنبؤات / تنبيهات / محظور / حوادث (CSV و JSON).",
            "تبويب الامتثال: لقطة تدقيق أمني وملخص يومي.",
        ],
        "benefit": "عندك دليل مكتوب وأرقام جاهزة للدكتور أو الأرشيف.",
        "tabs": ["Detection Report", "Incident Reports", "Export", "Compliance",
                 "تقرير الكشف", "تقارير الحوادث", "تصدير", "الامتثال"],
    },
    {
        "slug": "/Settings",
        "nav": "Settings",
        "file": "09_settings.png",
        "title": "10) الإعدادات — Settings",
        "role": "التحكم بالمنصة من غير تعديل كود.",
        "items": [
            "عام: سرعة التحديث، صوت الهجوم، تغيير اسم المستخدم وكلمة المرور.",
            "تيليغرام: التوكن، معرّف المحادثة، اختبار، اكتشاف المحادثة.",
            "الذكاء الاصطناعي: عتبة الثقة وعتبة الحظر التلقائي.",
            "الاستجابة: تلقائي أو موافقة بشرية، مدة الاحتفاظ، التقارير المجدولة.",
            "مسح البيانات: اختيار النوع بعد تأكيد اسم المستخدم وكلمة المرور.",
            "معلومات النظام (الإصدار v1.0.0).",
        ],
        "benefit": "تظبط التنبيهات والأمان والحذف الآمن قبل العرض.",
        "tabs": ["General", "Telegram", "AI Config", "NDR / Response", "Clear Data",
                 "عام", "تيليغرام", "إعداد الذكاء الاصطناعي", "الاستجابة", "مسح البيانات"],
    },
    {
        "slug": "/SOC_Ops",
        "nav": "SOC Ops",
        "file": "10_soc.png",
        "title": "11) عمليات SOC — SOC Ops",
        "role": "غرفة عمليات: MITRE، أتمتة، صحة السيرفر.",
        "items": [
            "MITRE ATT&CK: ربط نوع الهجوم بالتكتيك والتقنية.",
            "أدلة SOAR: ماذا يحدث تلقائياً بعد التنبيه.",
            "التعلم المستمر: المخزن والتدريب التدريجي.",
            "صحة الخادم: جاهزية، إصلاح تلقائي، اختبار حمل.",
            "سجل النماذج ومجموعات البيانات وانحراف المضيف.",
        ],
        "benefit": "تبيّن أن المشروع SOC كامل وليس فقط تصنيف ML.",
        "tabs": ["MITRE ATT&CK", "SOAR Playbooks", "Online Learning", "Server Health", "ML Registry",
                 "أدلة SOAR", "التعلم المستمر", "صحة الخادم", "سجل النماذج"],
    },
    {
        "slug": "/Incidents",
        "nav": "Incidents",
        "file": "11_incidents.png",
        "title": "12) الحوادث — Incidents",
        "role": "تجميع التنبيهات المرتبطة في قضية واحدة.",
        "items": [
            "تصفية: مفتوح / قيد المعالجة / محلول.",
            "جدول الحوادث مع سلسلة الهجوم.",
            "تحديث الحالة والمالك والملاحظات.",
            "ملخص ذكي وتوصية استجابة.",
        ],
        "benefit": "تعالج الحادث كقضية SOC بدل التعامل مع كل تنبيه لوحده.",
        "tabs": [],
    },
    {
        "slug": "/Assets",
        "nav": "Assets",
        "file": "12_assets.png",
        "title": "13) الأصول — Assets",
        "role": "جرد الأجهزة المكتشفة ودرجة مخاطرها.",
        "items": [
            "جدول الجرد مع المخاطر والتصنيف الحرج.",
            "تعيين جهاز كحرج أو طبيعي.",
            "طوبولوجيا التدفقات بين المضيفين.",
            "ملف مضيف معيّن (تنبيهات + تدفقات).",
            "شذوذات الهوية والانحراف السلوكي.",
        ],
        "benefit": "تركّز الحماية على الأصول المهمة لا على كل IP بنفس الدرجة.",
        "tabs": ["Inventory", "Topology", "Host profile", "Identity / UEBA",
                 "الجرد", "الطوبولوجيا", "ملف المضيف", "الهوية / UEBA"],
    },
    {
        "slug": "/Hunting",
        "nav": "Hunting",
        "file": "13_hunting.png",
        "title": "14) الصيد التهديدي — Hunting",
        "role": "تحقيق يدوي أعمق من التنبيه.",
        "items": [
            "بحث حسب IP أو نوع الهجوم أو البروتوكول.",
            "نتائج التدفقات والتنبيهات والتنبؤات.",
            "أدلة PCAP للهجمات العالية/الحرجة.",
            "خدمات معروفة، تقييم نطاق، TLS، وهاش الملف.",
        ],
        "benefit": "لما التنبيه ما يكفي، تلاقي أثر الهجوم والدليل.",
        "tabs": ["Hunt", "Forensics PCAP", "Protocol / DNS / TLS",
                 "صيد", "أدلة PCAP", "بروتوكول / DNS / TLS"],
    },
    {
        "slug": "/Response",
        "nav": "Response",
        "file": "14_response.png",
        "title": "15) الاستجابة — Response",
        "role": "تنفيذ القرار الأمني: اسمح، احظر، وافق، اربط أنظمة ثانية.",
        "items": [
            "IOC: إضافة وتعطيل مؤشرات (IP / نطاق / هاش).",
            "قائمة سماح لا تُحظر تلقائياً.",
            "حظر يدوي أو مؤقت مع إمكانية التراجع.",
            "طابور موافقة بشرية إذا وضع الاستجابة Ask.",
            "Webhooks و تسجيل حساسات موزّعة.",
        ],
        "benefit": "تتحكم بالاستجابة وتمنع الحظر الغلط، وتوسّع المراقبة لأكثر من جهاز.",
        "tabs": ["IOC", "Allow / Block", "Approvals", "Webhooks", "Sensors",
                 "مؤشرات الاختراق", "سماح / حظر", "الموافقات", "الحساسات"],
    },
    {
        "slug": "/Copilot",
        "nav": "Copilot",
        "file": "15_copilot.png",
        "title": "16) المساعد — Copilot",
        "role": "مساحة تحقيق كاملة باللغة الطبيعية.",
        "items": [
            "سؤال حر لمركز الأمن (مضيفين مشبوهين، تنبيهات حرجة، أصول، حظر).",
            "ملخص حادث موجود.",
            "توصية استجابة: حظر / تحقيق / تجاهل.",
            "تفسير لماذا اكتُشف تصنيف معيّن مع دليل.",
        ],
        "benefit": "تشرح القرار بسرعة: ليش هالهجوم وشو الإجراء المناسب.",
        "tabs": [],
    },
]


def wait_server(timeout=60):
    import urllib.request

    deadline = time.time() + timeout
    last_err = None
    while time.time() < deadline:
        try:
            urllib.request.urlopen(BASE, timeout=3)
            return
        except Exception as exc:
            last_err = exc
            time.sleep(1.5)
    raise RuntimeError(f"Dashboard not reachable at {BASE}: {last_err}")


def wait_app(driver, seconds=20):
    WebDriverWait(driver, seconds).until(
        lambda d: d.find_elements(By.TAG_NAME, "input") or d.find_elements(By.TAG_NAME, "h1")
    )
    time.sleep(1.2)


def _fill(el, value):
    el.click()
    time.sleep(0.2)
    el.send_keys("\ue009", "a")  # Ctrl+A
    el.send_keys("\ue017")  # Delete
    el.send_keys(value)
    time.sleep(0.2)


def is_logged_in(driver) -> bool:
    try:
        src = driver.page_source or ""
    except Exception:
        return False
    if "Default account: admin / admin123" in src and "Sign In" in src:
        return False
    try:
        for btn in driver.find_elements(By.TAG_NAME, "button"):
            try:
                txt = (btn.text or "").strip()
            except Exception:
                continue
            if txt in ("Logout", "خروج"):
                return True
    except Exception:
        pass
    return "ONLINE" in src and "Secure Login" not in src


def expand_sidebar(driver):
    for sel in (
        '[data-testid="stSidebarCollapsedControl"]',
        'button[kind="headerNoPadding"]',
    ):
        for el in driver.find_elements(By.CSS_SELECTOR, sel):
            try:
                if el.is_displayed():
                    el.click()
                    time.sleep(0.8)
                    return
            except Exception:
                continue


def goto_page(driver, nav_label: str):
    expand_sidebar(driver)
    time.sleep(0.4)
    target = nav_label.lower()
    for a in driver.find_elements(By.TAG_NAME, "a"):
        txt = " ".join((a.text or "").split()).lower()
        if target == txt or txt.endswith(target) or target in txt:
            driver.execute_script("arguments[0].click();", a)
            time.sleep(3.2)
            if is_logged_in(driver):
                return
    raise RuntimeError(f"Could not open page via sidebar: {nav_label}")


def login(driver):
    driver.get(BASE)
    wait_app(driver, 30)
    time.sleep(2)
    if is_logged_in(driver):
        return
    pwds = driver.find_elements(By.CSS_SELECTOR, "input[type='password']")
    if not pwds:
        raise RuntimeError("Login form not found")
    user = None
    for sel in (
        "input[aria-label='Username']",
        "input[aria-label='اسم المستخدم']",
    ):
        found = driver.find_elements(By.CSS_SELECTOR, sel)
        if found:
            user = found[0]
            break
    if user is None:
        raise RuntimeError("Login username field not found")
    _fill(user, "admin")
    _fill(pwds[0], "admin123")
    clicked = False
    for btn in driver.find_elements(By.TAG_NAME, "button"):
        txt = (btn.text or "").strip()
        if txt in ("Sign In", "دخول", "Sign in"):
            driver.execute_script("arguments[0].click();", btn)
            clicked = True
            break
    if not clicked:
        pwds[0].send_keys("\n")
    WebDriverWait(driver, 25).until(lambda d: is_logged_in(d))
    time.sleep(2)
    expand_sidebar(driver)


def full_screenshot(driver, path: Path):
    driver.set_window_size(1600, 900)
    height = driver.execute_script(
        """
        if (!document.getElementById('ug-hide-chrome')) {
          const hide = document.createElement('style');
          hide.id = 'ug-hide-chrome';
          hide.innerHTML = `
            [data-testid="stToolbar"], [data-testid="stDecoration"],
            [data-testid="stStatusWidget"], .stDeployButton { display:none !important; }
            [data-testid="stAppViewContainer"], [data-testid="stMain"],
            section.main, .stApp, .stAppViewContainer, .block-container {
              height: auto !important; max-height: none !important;
              overflow: visible !important;
            }
          `;
          document.head.appendChild(hide);
        }
        const keys = ['stAppViewContainer','stMain','stApp','stVerticalBlock'];
        let maxH = Math.max(document.body.scrollHeight, document.documentElement.scrollHeight, 900);
        document.querySelectorAll('section, div, article, main').forEach(el => {
          const h = el.scrollHeight || 0;
          if (h > maxH && h < 20000) maxH = h;
        });
        return maxH;
        """
    )
    width = 1600
    height = min(max(int(height) + 120, 1000), 14000)
    driver.execute_cdp_cmd(
        "Emulation.setDeviceMetricsOverride",
        {"mobile": False, "width": width, "height": height, "deviceScaleFactor": 1},
    )
    time.sleep(0.9)
    png = driver.execute_cdp_cmd(
        "Page.captureScreenshot",
        {
            "fromSurface": True,
            "captureBeyondViewport": True,
            "format": "png",
            "clip": {"x": 0, "y": 0, "width": width, "height": height, "scale": 1},
        },
    )
    path.write_bytes(base64.b64decode(png["data"]))
    driver.execute_cdp_cmd("Emulation.clearDeviceMetricsOverride", {})


def click_all_tabs(driver, page_file: str, names: list[str] | None = None) -> list:
    if not names:
        return []
    taken = []
    seen = set()
    tabs = driver.find_elements(
        By.CSS_SELECTOR,
        "button[role='tab'], [data-baseweb='tab'], [data-testid='stTab'], div[data-testid='stTabs'] button",
    )
    labels = []
    allowed = set(names or [])
    for tab in tabs:
        label = " ".join((tab.text or "").split())
        if not label or label in seen:
            continue
        if allowed and label not in allowed:
            continue
        seen.add(label)
        labels.append((tab, label))
    for i, (tab, label) in enumerate(labels):
        try:
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", tab)
            driver.execute_script("arguments[0].click();", tab)
            time.sleep(1.6)
            safe = "".join(c if c.isalnum() else "_" for c in label)[:40]
            p = SHOTS / f"{Path(page_file).stem}_tab{i}_{safe}.png"
            full_screenshot(driver, p)
            taken.append((label, p))
        except Exception:
            continue
    return taken


def set_run(p, size=12, bold=False, color=None):
    p.font.name = "Calibri"
    p._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    p.font.size = Pt(size)
    p.bold = bold
    if color:
        p.font.color.rgb = color


def add_heading_ar(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run(r, size=size, bold=True, color=RGBColor(0x0C, 0x1A, 0x2E))
    return p


def add_body_ar(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def build_doc(captures: list[dict]):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    add_heading_ar(doc, "دليل صفحات نظام AI Network Analyzer", 26)
    add_body_ar(doc, "مشروع التخرج — محمود طلال كنعان — AUCE 2026 — الإصدار v1.0.0", 13, True)
    add_body_ar(
        doc,
        "هذا الملف يشرح كل صفحة في لوحة التحكم: لقطة شاشة حقيقية من النظام، ثم ماذا تحتوي الصفحة، "
        "ما وظيفتها، وما الفائدة منها. إذا كانت الصفحة مقسّمة إلى تبويبات، أُخذت لقطة لكل تبويب ظاهر.",
        12,
    )
    add_body_ar(doc, "تسجيل الدخول الافتراضي: admin / admin123", 12, True)

    for cap in captures:
        doc.add_page_break()
        add_heading_ar(doc, cap["title"], 18)
        add_body_ar(doc, "الوظيفة: " + cap["role"], 12, True)
        add_body_ar(doc, "ماذا تستفيد: " + cap["benefit"], 12)
        add_body_ar(doc, "محتويات الصفحة:", 12, True)
        for item in cap["items"]:
            add_body_ar(doc, "• " + item, 12)
        for img in cap["images"]:
            label, path = img
            if label:
                add_body_ar(doc, "لقطة: " + label, 11, True)
            if path.exists():
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic.add_run()
                run.add_picture(str(path), width=Inches(7.0))

    doc.save(str(OUT))
    return OUT


def main():
    wait_server()
    opts = Options()
    opts.add_argument("--window-size=1600,900")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--lang=en")
    driver = webdriver.Edge(options=opts)
    captures = []
    try:
        driver.get(BASE)
        wait_app(driver, 30)
        time.sleep(1.5)
        login_shot = SHOTS / "00_login.png"
        full_screenshot(driver, login_shot)
        captures.append({
            "title": "0) تسجيل الدخول — Login",
            "role": "بوابة الدخول إلى لوحة التحكم. بدونها ما بينفتح النظام.",
            "benefit": "تحمي المنصة: بس المستخدم المصرّح يقدر يشوف الحركة والتنبيهات.",
            "items": [
                "اختيار اللغة (إنجليزي / عربي) قبل الدخول.",
                "حقل اسم المستخدم وحقل كلمة المرور.",
                "زر Sign in / دخول.",
                "الحساب الافتراضي للعرض: admin / admin123.",
            ],
            "images": [("شاشة الدخول", login_shot)],
        })
        print("captured login")
        for attempt in range(3):
            try:
                login(driver)
                if is_logged_in(driver):
                    break
            except Exception as exc:
                print("login retry", attempt, type(exc).__name__, flush=True)
                time.sleep(2)
        else:
            raise RuntimeError("Login did not reach the dashboard")
        for page in PAGES:
            try:
                goto_page(driver, page["nav"])
                time.sleep(1.8)
                if not is_logged_in(driver):
                    print("session lost before", page["file"], "re-login", flush=True)
                    login(driver)
                    goto_page(driver, page["nav"])
                    time.sleep(1.8)
                main_path = SHOTS / page["file"]
                full_screenshot(driver, main_path)
                images = [("الصفحة كاملة", main_path)]
                extra = click_all_tabs(driver, page["file"], page.get("tabs") or [])
                if extra:
                    images.extend(extra)
                captures.append({**page, "images": images})
                print("captured", page["file"], "images=", len(images), flush=True)
            except Exception as exc:
                print("FAILED", page["file"], type(exc).__name__, str(exc)[:200], flush=True)
                captures.append({**page, "images": []})
    finally:
        try:
            driver.quit()
        except Exception:
            pass

    out = build_doc(captures)
    print("WROTE", out)


if __name__ == "__main__":
    main()
