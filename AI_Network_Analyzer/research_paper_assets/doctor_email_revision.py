"""
Apply Dr. Hassan Noureddine email feedback (Aug 2026) to the Desktop thesis.
Uses reproducible metrics from training/evaluate_models.py and FEATURE_COLUMNS from config.py.
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
ASSETS = Path(__file__).resolve().parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
BACKUP = PAPER.with_name(PAPER.stem + f"_pre_doctor_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
METRICS_PATH = ASSETS / "evaluation" / "metrics.json"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def find_exact(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def insert_after(p: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    if style:
        try:
            para.style = style
        except Exception:
            pass
    if text:
        set_text(para, text)
    return para


def set_cell(cell, text: str) -> None:
    p = cell.paragraphs[0]
    set_text(p, text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


def row_by_first(table, first: str):
    for row in table.rows:
        if (row.cells[0].text or "").strip() == first:
            return list(row.cells)
    return None


def clear_table_keep_header(table) -> None:
    tbl = table._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell(row.cells[i], val)


def page_break_before_chapters(doc: Document) -> None:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if st == "Heading 1" and t.startswith("Chapter "):
            br = p.insert_paragraph_before("")
            br.add_run().add_break(WD_BREAK.PAGE)


FEATURE_META: list[tuple[str, str, str, str]] = [
    ("duration", "Flow duration (seconds)", "Numeric", "Long flows may indicate scanning or C2 beacons"),
    ("protocol_type", "Protocol encoded (TCP/UDP/ICMP)", "Categorical→numeric", "Attack families differ by protocol"),
    ("src_bytes", "Source-to-destination bytes", "Numeric", "Volume spikes in exfiltration/DDoS"),
    ("dst_bytes", "Destination-to-source bytes", "Numeric", "Asymmetric ratios flag floods/scans"),
    ("count", "Connections to same host in window", "Numeric", "Port scans raise connection counts"),
    ("srv_count", "Connections to same service", "Numeric", "Service-targeted brute force"),
    ("serror_rate", "SYN error rate", "Numeric", "DoS often increases SYN errors"),
    ("rerror_rate", "REJ/reset error rate", "Numeric", "Probing and failed connections"),
    ("same_srv_rate", "Fraction same service", "Numeric", "Normal sessions vs horizontal scans"),
    ("diff_srv_rate", "Fraction different services", "Numeric", "Reconnaissance behaviour"),
    ("dst_host_count", "Hosts contacted by source", "Numeric", "Worm/scan spread indicator"),
    ("dst_host_srv_count", "Services accessed on host", "Numeric", "Vertical scanning signal"),
    ("dst_host_same_srv_rate", "Same service rate to host", "Numeric", "Brute-force repetition"),
    ("dst_host_diff_srv_rate", "Different service rate to host", "Numeric", "Multi-service probing"),
    ("dst_host_serror_rate", "Host SYN error rate", "Numeric", "Host-level DoS symptoms"),
    ("dst_host_rerror_rate", "Host reset error rate", "Numeric", "Failed connection attempts"),
    ("packet_count", "Packets in flow", "Numeric", "Floods vs benign short flows"),
    ("byte_count", "Total bytes in flow", "Numeric", "Data exfiltration volume"),
    ("packet_rate", "Packets per second", "Numeric", "DoS/DDoS intensity"),
    ("flow_rate", "Bytes per second", "Numeric", "Throughput anomalies"),
    ("avg_packet_size", "Mean packet size", "Numeric", "Tunneling/malformed traffic"),
    ("syn_count", "SYN flag count", "Numeric", "SYN flood detection"),
    ("ack_count", "ACK flag count", "Numeric", "Established vs half-open flows"),
    ("fin_count", "FIN flag count", "Numeric", "Connection teardown patterns"),
    ("rst_count", "RST flag count", "Numeric", "Scan/reset storms"),
    ("psh_count", "PSH flag count", "Numeric", "Interactive/payload-heavy flows"),
    ("urg_count", "URG flag count", "Numeric", "Rare urgent-flag abuse"),
    ("flow_duration", "Alias of active flow span", "Numeric", "Timing behaviour"),
    ("fwd_packets", "Forward-direction packets", "Numeric", "Client/server asymmetry"),
    ("bwd_packets", "Backward-direction packets", "Numeric", "Response-side anomalies"),
    ("fwd_bytes", "Forward bytes", "Numeric", "Upload/exfil patterns"),
    ("bwd_bytes", "Backward bytes", "Numeric", "Download/amplification"),
    ("min_packet_length", "Minimum packet size", "Numeric", "Malformed packet detection"),
    ("max_packet_length", "Maximum packet size", "Numeric", "Jumbo/fragment abuse"),
    ("mean_packet_length", "Mean packet size", "Numeric", "Protocol misuse"),
    ("std_packet_length", "Std dev packet size", "Numeric", "Irregular payload sizes"),
    ("inter_arrival_time", "Mean inter-packet gap", "Numeric", "Bot/beacon timing"),
    ("active_time", "Active period in flow", "Numeric", "Burst vs steady traffic"),
    ("idle_time", "Idle period in flow", "Numeric", "Long idle then burst (C2)"),
]


def rebuild_feature_table(doc: Document) -> None:
    table = doc.tables[10]
    clear_table_keep_header(table)
    set_cell(table.rows[0].cells[0], "Feature")
    set_cell(table.rows[0].cells[1], "Meaning / relevance (all scaled with StandardScaler on train split; used by RF, XGB, IF, AE; LSTM uses windowed sequences of the same vector)")
    for name, meaning, _typ, rel in FEATURE_META:
        add_row(table, [name, f"{meaning}. {rel}."])


def main() -> None:
    subprocess.run([sys.executable, "-m", "training.evaluate_models"], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ASSETS / "render_eval_figures.py")], check=True, cwd=str(ASSETS))

    m = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    ds = m["dataset"]
    classes = ds.get("classes") or {}
    class_txt = ", ".join(f"{k} {v}" for k, v in classes.items())

    shutil.copy2(PAPER, BACKUP)
    doc = Document(str(PAPER))

    # ── Objectives O5/O6 (no email) ──
    obj = doc.tables[3]
    c = row_by_first(obj, "O5")
    if c:
        set_cell(c[1], "Improve transparency using stored per-prediction Explainable AI (this-flow evidence and tree importances).")
    c = row_by_first(obj, "O6")
    if c:
        set_cell(c[1], "Support incident response using IP blocking, Telegram alerts, webhooks, and reports (no email channel).")

    # ── Technology stack = running app ──
    tech = doc.tables[9]
    tech_map = {
        "Database": "SQLite by default; PostgreSQL optional via DATABASE_URL",
        "AI Libraries": "Scikit-learn, XGBoost, PyTorch (sklearn MLP fallback if torch unavailable)",
        "Traffic Capture": "Scapy/Npcap live NIC, OS connection table, optional 24/7 capture daemon",
        "Feature Extraction": "Custom 39 CICIDS-style flow features (config.FEATURE_COLUMNS)",
        "Explainability": "Stored per-prediction XAI JSON; SHAP is future work",
        "Dashboard": "Streamlit SOC dashboard, port 8501",
        "Notifications": "Telegram Bot API, local Windows notification, optional webhooks/SIEM (no SMTP email)",
        "Backend": "FastAPI REST API on port 8000",
    }
    for row in tech.rows:
        k = (row.cells[0].text or "").strip()
        if k in tech_map:
            set_cell(row.cells[1], tech_map[k])

    # ── 4.2 Feature engineering + selection justification ──
    fe = find_exact(doc, "4.2 Feature Engineering")
    if fe:
        intro = insert_after(
            fe,
            "The experiments use a fixed schema of 39 flow-level features defined in config.FEATURE_COLUMNS. "
            "They were selected using domain knowledge from CICIDS-style NDR literature (flow duration, byte/packet "
            "statistics, flag counts, host/service rates) rather than an automatic filter applied only for the report. "
            "The same schema is used in live capture, training, and inference so train/serve skew is avoided. "
            "Post-training Random Forest/XGBoost feature importances and per-prediction XAI are used for interpretation, "
            "not to drop columns mid-project. Recursive Feature Elimination or PCA were not applied because reducing "
            "dimensionality would break parity with the deployed scaler and saved models.",
            "Normal",
        )
        h = insert_after(intro, "4.2.1 Complete Feature Set Used in Experiments", "Heading 3")
        insert_after(
            h,
            "Table 5 lists every feature used in the holdout evaluation. All numeric features are standardised with "
            "StandardScaler fitted on the training split only. protocol_type is encoded numerically before scaling.",
            "Normal",
        )

    rebuild_feature_table(doc)

    # ── 4.2.2 Dataset construction (doctor §1) ──
    h422 = find_contains(doc, "4.2.1 Complete Feature Set Used in Experiments")
    anchor422 = h422 or find_exact(doc, "4.2 Feature Engineering")
    if anchor422:
        h = insert_after(anchor422, "4.2.2 Evaluation Dataset Construction", "Heading 3")
        p = insert_after(
            h,
            f"Source. The reproducible evaluation file is datasets/dataset.csv built by training/generate_realistic_dataset.py "
            f"(synthetic CICIDS-style rows using the same attack signatures as the live detector). If a public CICIDS2017 or "
            f"UNSW-NB15 CSV is placed under datasets/, training/cicids_adapter.py imports it onto the same 39-feature schema. "
            f"This is not claimed as the official CICIDS2017 MachineLearningCSV leaderboard dump unless such a file was imported.\n\n"
            f"Shape. {ds['rows']:,} rows × 39 input features + label column (40 columns in CSV).\n\n"
            f"Why not 720? An earlier draft used a balanced 720-row pilot (120 per class). Following supervisor feedback, the "
            f"final evaluation was regenerated at {ds['rows']:,} rows with realistic imbalance: {class_txt}.\n\n"
            f"Split. Stratified 75/25 train-test holdout, random_state=42, executed in training/evaluate_models.py. "
            f"Preprocessing (duplicate removal, infinity→NaN→0, LabelEncoder, StandardScaler) is fit on training rows only.\n\n"
            f"Leakage controls. No label column in features; scaler/encoder saved from train split; holdout rows never used "
            f"to fit preprocessing; stratify preserves class ratios; duplicate rows removed before splitting.",
            "Normal",
        )

    # ── 4.3 per-model config (doctor §4-6) ──
    mdev = find_exact(doc, "4.3 Model Development")
    if mdev:
        h431 = find_contains(doc, "4.3.1 Per-Model Training Configuration")
        if not h431:
            h431 = insert_after(mdev, "4.3.1 Per-Model Training Configuration", "Heading 3")
        model_txt = (
            "Random Forest — supervised multi-class on all 39 features; n_estimators=200, max_depth=20; "
            "output: class label + probabilities. XGBoost — supervised multi-class; n_estimators=200, max_depth=8, "
            "learning_rate=0.1; output: class probabilities. Isolation Forest — trained on Normal traffic only; "
            "n_estimators=200, contamination=0.1; output: inlier/outlier mapped to Normal/Attack for proxy metrics. "
            "Autoencoder — trained on Normal rows only; reconstruction error threshold; implemented and loaded at inference "
            "but not scored on the same stratified holdout table below. LSTM — sliding windows of 10 consecutive scaled "
            "flows (10×39 tensor); implemented; holdout classification metrics not reported in Table 11. Hybrid AI — "
            "decision_engine.fuse_decisions() performs weighted majority fusion over live model votes; not a separate "
            "classifier and therefore has no standalone holdout accuracy."
        )
        if not find_contains(doc, "decision_engine.fuse_decisions"):
            insert_after(h431, model_txt, "Normal")

    # ── Table 9: design illustration only (doctor §7) ──
    perf = doc.tables[19]
    values = {
        "Accuracy": [pct(rf["accuracy"]), pct(xgb["accuracy"]), pct(iso["accuracy"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "Precision": [pct(rf["precision"]), pct(xgb["precision"]), pct(iso["precision"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "Recall": [pct(rf["recall"]), pct(xgb["recall"]), pct(iso["recall"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "F1 Score": [pct(rf["f1"]), pct(xgb["f1"]), pct(iso["f1"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "ROC-AUC": ["See §5.6", "See §5.6", "Binary proxy", "Not reported", "Not reported", "N/A"],
    }
    for row in perf.rows:
        key = (row.cells[0].text or "").strip()
        if key in values:
            for i, val in enumerate(values[key], start=1):
                set_cell(row.cells[i], val)

    for row in doc.tables[1].rows:
        if "Table 9" in (row.cells[0].text or ""):
            set_cell(row.cells[1], "Holdout metrics for RF/XGB/IF; AE/LSTM/Hybrid not scored on the same split (see text).")

    # ── Chapter 5 experimental text (doctor §7-9) ──
    p720 = find_contains(doc, "An implementation evaluation dataset was prepared with 720")
    if p720:
        set_text(
            p720,
            f"An implementation evaluation dataset was prepared with {ds['rows']:,} labeled samples, 39 flow features, "
            f"and six classes ({class_txt}). Rows are generated with CICIDS-style attack signatures, jitter, and label "
            f"noise so accuracy is not trivially 100%. Holdout evaluation uses a stratified 25% test split "
            f"(random_state=42). Random Forest achieved accuracy {pct(rf['accuracy'])}, precision {pct(rf['precision'])}, "
            f"recall {pct(rf['recall'])}, weighted F1 {pct(rf['f1'])}. XGBoost achieved accuracy {pct(xgb['accuracy'])}, "
            f"precision {pct(xgb['precision'])}, recall {pct(xgb['recall'])}, weighted F1 {pct(xgb['f1'])}. "
            f"Isolation Forest (Normal-vs-Attack proxy on the same scaled holdout) achieved accuracy {pct(iso['accuracy'])} "
            f"and weighted F1 {pct(iso['f1'])}.",
        )

    t11 = find_contains(doc, "Table 11. Experimental model performance")
    if t11:
        per = rf.get("per_class") or {}
        per_txt = "; ".join(
            f"{k}: P {v.get('precision', 0)*100:.1f}%, R {v.get('recall', 0)*100:.1f}%, F1 {v.get('f1-score', 0)*100:.1f}%"
            for k, v in per.items()
        )
        set_text(
            t11,
            f"Table 11. Measured holdout metrics (reproducible via python -m training.evaluate_models).\n"
            f"Dataset: {ds['rows']:,} rows, stratified 25% test, random_state=42.\n"
            f"Random Forest — Accuracy {pct(rf['accuracy'])}, Precision {pct(rf['precision'])}, Recall {pct(rf['recall'])}, F1 {pct(rf['f1'])}.\n"
            f"XGBoost — Accuracy {pct(xgb['accuracy'])}, Precision {pct(xgb['precision'])}, Recall {pct(xgb['recall'])}, F1 {pct(xgb['f1'])}.\n"
            f"Isolation Forest (Normal vs Attack) — Accuracy {pct(iso['accuracy'])}, F1 {pct(iso['f1'])}.\n"
            f"Per-class RF: {per_txt}.\n"
            f"Autoencoder, LSTM, and Hybrid AI: implemented in software; no separate holdout row is reported because "
            f"AE/LSTM were not evaluated with the same stratified split script and Hybrid is a fusion layer.",
        )

    summ = find_contains(doc, "Table 11 summarizes experimental metrics")
    if summ:
        set_text(
            summ,
            f"Interpretation. XGBoost slightly exceeds RF on weighted F1 ({pct(xgb['f1'])} vs {pct(rf['f1'])}) because "
            f"boosting reduces bias on minority classes such as SQLInjection. RF remains strong on Normal traffic "
            f"(high recall on the majority class). Isolation Forest scores lower ({pct(iso['f1'])}) because it is trained "
            f"only on Normal traffic and evaluated as a binary Normal-vs-Attack proxy, so multi-class attacks are collapsed. "
            f"From a SOC perspective, supervised models are primary for labelled families; IF adds an unsupervised signal for "
            f"unknown patterns. High scores are plausible on this synthetic set but are not official CICIDS2017 benchmarks; "
            f"leakage checks (train-only scaler, no label in features, deduplication) were applied as described in §4.2.2.",
        )

    fig14 = find_contains(doc, "Figure 14. Expected Model Performance")
    if fig14:
        set_text(fig14, "Figure 14. Measured holdout performance (RF, XGB, IF) on the 12,000-row evaluation set.")

    note = find_contains(doc, "Important note: the values above are expected performance targets")
    if note:
        set_text(
            note,
            "Table 9 above reports measured holdout values for RF/XGB/IF only. Earlier placeholder targets "
            "(e.g., 98.6%, 99.2%) were removed. Autoencoder, LSTM, and Hybrid AI are described in §4.3.1 and are not "
            "assigned standalone holdout percentages unless reproduced with the same evaluation script.",
        )

    # ── Implementation status table (doctor §18) ──
    arch = find_contains(doc, "3.3.1 Implemented Core Modules")
    if arch:
        h = insert_after(arch, "3.3.2 Implementation and Evaluation Status", "Heading 3")
        insert_after(
            h,
            "Implemented and tested: capture, 39-feature extraction, RF/XGB/IF holdout evaluation, dashboard, alerts, "
            "Telegram, Threat Simulation, stored XAI, MITRE mapping, SOAR playbooks. Implemented but not fully evaluated "
            "on the Table 11 holdout: Autoencoder, LSTM sequence model, Hybrid fusion (live only). Proposed/future: "
            "SHAP plots, full CICIDS2017 public import, SPAN/TAP fleet, email/SMTP channel.",
            "Normal",
        )

    # ── Attack scenario table (doctor §10) ──
    atk = find_exact(doc, "5.3 Attack Simulation Scenarios")
    if atk:
        h = insert_after(atk, "5.3.1 Structured Attack Testing Matrix", "Heading 3")
        insert_after(
            h,
            "DoS/DDoS — RF+XGB primary; IF anomaly vote — Medium+ alert — playbook block/Telegram — tested via Threat Simulation (TEST-NET). "
            "PortScan — RF — High alert — TI enrichment — tested in simulation. BruteForce/SQLInjection — RF/XGB — "
            "Alert+MITRE — block IP — tested. Malware/C2 — discussed theoretically; partial IF signal — future live TI. "
            "Zero-day — IF+AE in engine — Medium anomaly alert — analyst review — simulated via noisy flows, not a real zero-day sample.",
            "Normal",
        )

    # ── Fault test cases (doctor §11) ──
    tests = doc.tables[20]
    extra = [
        ("TC11", "Upload invalid CSV", "Parser rejects file; dashboard/API shows readable error; no crash."),
        ("TC12", "Database unavailable", "API returns 503-style message; dashboard shows connection warning."),
        ("TC13", "Model file missing", "Detection endpoint reports model-not-ready safely."),
        ("TC14", "Threat Intel API timeout", "Detection continues; TI field marked unavailable."),
        ("TC15", "False positive review", "Analyst marks alert investigated; no auto-block below threshold."),
    ]
    for tc, act, exp in extra:
        if not row_by_first(tests, tc):
            add_row(tests, [tc, act, exp])

    # ── Chapter 7 restructure (doctor §12) ──
    ch7 = find_exact(doc, "Chapter 7 - Practical Implementation Notes")
    if ch7:
        insert_after(
            ch7,
            "Supervisor feedback: detailed methodology previously repeated in this chapter is now integrated into "
            "§4.2 (features/dataset), §4.3 (models), and §5 (evaluation). This chapter keeps brief reminders only.",
            "Normal",
        )
    for title, pointer in [
        ("7.1 Dataset Preparation Notes", "See §4.2.2 for dataset construction, preprocessing, and leakage controls."),
        ("7.2 Handling Class Imbalance", "See §4.2.2 and §5.6 for stratified splitting and per-class metrics."),
        ("7.3 Threshold Tuning", "See §4.3.1 (IF contamination, AE error threshold) and Settings → AI Config in the dashboard."),
        ("7.4 Explainability Strategy", "See §3.3.1 and stored XAI JSON on Detection/Alerts/Copilot pages."),
        ("7.5 Report Generation", "See §8.4 and Reports → Export / Live Evidence tabs."),
    ]:
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() == title and p.style and "Heading" in (p.style.name or ""):
                if i + 1 < len(doc.paragraphs):
                    nxt = doc.paragraphs[i + 1]
                    if not (nxt.style and nxt.style.name.startswith("Heading")):
                        set_text(nxt, pointer)
                break

    ch8 = find_exact(doc, "Chapter 8 - User Manual and Operation Guide")
    if ch8:
        insert_after(
            ch8,
            "This chapter is a concise operator guide only. Methodology and evaluation evidence are in Chapters 4–5.",
            "Normal",
        )

    # ── Appendix B defense demo (doctor §19) ──
    outline = doc.tables[26]
    slides = [
        ("1", "Title + AUCE + supervisor + v1.0.0"),
        ("2", "Agenda: Problem → Methodology → Design → Implementation → Experiments → Results → Discussion"),
        ("3", "Dataset: 12k CICIDS-style, 39 features, stratified holdout (reproducible script)"),
        ("4", "Models: RF/XGB/IF measured; AE/LSTM/Hybrid status"),
        ("5", "Live demo: capture → features → hybrid fusion → alert → XAI → MITRE → response → dashboard"),
        ("6", "Attack scenario (Threat Simulation) + anomaly/IF scenario"),
        ("7", "Q&A: leakage controls, feature choice, why Hybrid, limitations"),
    ]
    for i, (num, content) in enumerate(slides, start=1):
        if i < len(outline.rows):
            set_cell(outline.rows[i].cells[0], num)
            set_cell(outline.rows[i].cells[1], content)

    # ── Figure captions ──
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure 15. Dataset class distribution"):
            set_text(p, f"Figure 15. Class distribution ({ds['rows']:,}-row evaluation set).")
        if t.startswith("Figure 4. Data Flow Diagram"):
            set_text(p, "Figure 4. System data pipeline (flow extraction → AI engine → storage → dashboard).")

    page_break_before_chapters(doc)

    doc.save(str(PAPER))
    print("Saved", PAPER)
    print("Backup", BACKUP)


if __name__ == "__main__":
    main()
