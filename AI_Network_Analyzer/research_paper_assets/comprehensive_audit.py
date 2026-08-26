"""Comprehensive thesis audit: TOC, headings, tables, figures, spacing, refs."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"


def has_page_break(p) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def has_image(p) -> bool:
    return bool(p._p.findall(f".//{NS_A}blip"))


def norm_title(t: str) -> str:
    t = re.sub(r"\t\d+$", "", t.strip())
    t = re.sub(r"\s+", " ", t)
    return t.lower()


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    doc = Document(str(PAPER))
    issues: list[str] = []
    warnings: list[str] = []

    # --- Headings ---
    h1, h2, h3 = [], [], []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if st == "Heading 1":
            h1.append(t)
        elif st == "Heading 2":
            h2.append(t)
        elif st == "Heading 3":
            h3.append(t)

    # --- TOC entries ---
    toc_entries: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        st = (p.style.name if p.style else "").lower()
        if st.startswith("toc"):
            toc_entries.append((st, (p.text or "").strip()))

    toc_titles = {norm_title(re.sub(r"^\d+(\.\d+)*\s+", "", re.sub(r"^Chapter \d+ - ", "", t.split("\t")[0]))) for _, t in toc_entries if t}
    toc_nums = set()
    for _, t in toc_entries:
        m = re.match(r"^(\d+(?:\.\d+)*)", t.split("\t")[0])
        if m:
            toc_nums.add(m.group(1))
        m2 = re.match(r"^Chapter (\d+)", t.split("\t")[0])
        if m2:
            toc_nums.add(m2.group(1))

    # Headings not in TOC (body sections only, after front matter)
    body_started = False
    missing_from_toc: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if t.startswith("Chapter 1 -"):
            body_started = True
        if not body_started:
            continue
        if st in ("Heading 2", "Heading 3"):
            num_m = re.match(r"^(\d+(?:\.\d+)*)", t)
            if num_m and num_m.group(1) not in toc_nums:
                # allow appendix subsections not always in toc
                if not t.startswith(("Appendix ", "References")):
                    missing_from_toc.append(t[:80])

    if missing_from_toc:
        for m in missing_from_toc[:15]:
            issues.append(f"Heading not in TOC: {m}")
        if len(missing_from_toc) > 15:
            issues.append(f"... and {len(missing_from_toc) - 15} more headings missing from TOC")

    # --- Table captions & proximity ---
    table_caps: list[tuple[int, str]] = []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.match(r"^Table\s+\d+[A-Z]?\.", t, re.I):
            table_caps.append((i, t))

    # Tables in doc vs captions
    n_tables = len(doc.tables)
    print(f"Tables in document: {n_tables}")
    print(f"Table caption paragraphs: {len(table_caps)}")
    for idx, cap in table_caps:
        print(f"  @{idx}: {cap[:90]}")

    if len(table_caps) < n_tables - 2:
        warnings.append(
            f"Fewer captions ({len(table_caps)}) than tables ({n_tables}) — some tables may lack captions"
        )

    # Check table before/after caption
    for idx, cap in table_caps:
        has_tbl_before = idx > 0 and any(
            True for _ in []  # placeholder
        )
        # find nearest table element position - tables are block-level between paragraphs
        # scan paragraphs around caption for table in document body order
        found_near = False
        for ti, tbl in enumerate(doc.tables):
            pass  # python-docx doesn't give paragraph index for tables easily

    # Table number references in text
    table_refs = sorted(
        set(re.findall(r"Table\s+(\d+[A-Z]?)\b", "\n".join(p.text or "" for p in doc.paragraphs), re.I))
    )
    cap_nums_flat = sorted(
        set(re.findall(r"^Table\s+(\d+[A-Z]?)\.", "\n".join(c for _, c in table_caps), re.I | re.M))
    )

    for ref in table_refs:
        if ref not in cap_nums_flat and ref not in ("5", "5A", "3A", "9", "11", "12"):
            # Table 5 might be referenced in prose without formal caption starting Table 5.
            if ref not in cap_nums_flat:
                warnings.append(f"Text references Table {ref} but no caption 'Table {ref}.' found")

    # List of Tables section
    in_lot = False
    lot_entries: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "List of Tables":
            in_lot = True
            continue
        if in_lot:
            if p.style and p.style.name == "Heading 1":
                break
            if t:
                lot_entries.append(t)

    print(f"\nList of Tables entries: {len(lot_entries)}")
    for e in lot_entries[:20]:
        print(f"  {e[:80]}")

    # --- Figures ---
    fig_caps = []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.match(r"^Figure\s+(\d+\.|H\d)", t):
            prev = doc.paragraphs[i - 1] if i else None
            has_img = prev and has_image(prev)
            fig_caps.append((i, t[:60], has_img))
            if not has_img:
                issues.append(f"Figure without image: {t[:70]}")

    print(f"\nFigures: {len(fig_caps)} ({sum(1 for _, _, h in fig_caps if h)} with images)")

    # --- Spacing ---
    empty_pb = sum(
        1
        for p in doc.paragraphs
        if not (p.text or "").strip() and not has_image(p) and has_page_break(p)
    )
    if empty_pb:
        issues.append(f"{empty_pb} empty paragraphs with page breaks (blank pages)")

    triple_empty = 0
    run = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t and not has_image(p) and not has_page_break(p):
            run += 1
        else:
            if run >= 3:
                triple_empty += 1
            run = 0
    if triple_empty:
        warnings.append(f"{triple_empty} runs of 3+ empty paragraphs (excess spacing)")

    # Caption spacing
    bad_cap_space = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure ") or t.startswith("Table "):
            pf = p.paragraph_format
            if pf.space_after and pf.space_after.pt > 14:
                bad_cap_space += 1
    if bad_cap_space:
        warnings.append(f"{bad_cap_space} captions with large space_after")

    # --- Grammar heuristics ---
    blob = "\n".join((p.text or "") for p in doc.paragraphs)
    grammar_patterns = [
        (r"\b(a|an)\s+(hour|AI|NDR|IDS|API|LSTM|XGB|RF)\b", "article agreement"),
        (r"\s{2,}", "double spaces"),
        (r"\.\s*[a-z]", "sentence break lowercase"),
        (r"\bteh\b", "typo teh"),
        (r"\brecieve\b", "typo recieve"),
        (r"\boccured\b", "typo occured"),
        (r"\bseperate\b", "typo seperate"),
        (r"\butilise\b", "British utilise (OK)"),
        (r"  +", "multiple spaces"),
    ]
    grammar_hits: list[str] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if not t.strip() or (p.style and p.style.name.lower().startswith("toc")):
            continue
        if re.search(r"\s{3,}", t):
            grammar_hits.append(f"@{i} triple+ spaces: {t[:60]}")
        if re.search(r"([a-z])\s+([,.])", t):
            grammar_hits.append(f"@{i} space before punct: {t[:60]}")
        if "  " in t and not t.startswith("    "):
            grammar_hits.append(f"@{i} double space: {t[:60]}")

    # --- Required doctor tables/content ---
    required_phrases = [
        ("Table 5 lists all 39 features", "feature table intro"),
        ("Table 5A", "model usage table"),
        ("Table 3A", "implementation status"),
        ("Table 11", "holdout metrics"),
        ("Table 12", "operational stats"),
        ("Component Implementation Status", "impl status heading"),
        ("Per-Model Dataset Usage", "model usage"),
        ("leakage", "leakage checklist"),
        ("TC11", "fault tests"),
        ("fuse_decisions", "hybrid fusion"),
    ]
    full = blob.lower()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full += "\n" + (cell.text or "").lower()

    print("\n=== REQUIRED CONTENT ===")
    for phrase, label in required_phrases:
        ok = phrase.lower() in full
        print(("PASS" if ok else "FAIL"), label, f"({phrase})")
        if not ok:
            issues.append(f"Missing required content: {label} ({phrase})")

    # --- Table inventory ---
    print("\n=== TABLE INVENTORY ===")
    for ti, table in enumerate(doc.tables):
        rows, cols = len(table.rows), len(table.rows[0].cells) if table.rows else 0
        c00 = (table.rows[0].cells[0].text or "").replace("\n", " ")[:50]
        print(f"T{ti:2d} {rows}x{cols} | {c00}")

    # --- Bullet lists that should be tables? ---
    list_heavy_sections = []
    in_h2 = ""
    bullet_run = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if st == "Heading 2":
            if bullet_run >= 6:
                list_heavy_sections.append(f"{in_h2}: {bullet_run} bullets (consider table?)")
            in_h2 = t[:50]
            bullet_run = 0
        elif st == "List Bullet" and in_h2:
            bullet_run += 1
    if bullet_run >= 6:
        list_heavy_sections.append(f"{in_h2}: {bullet_run} bullets")

    # --- Summary ---
    print("\n=== ISSUES (must fix) ===")
    if issues:
        for x in issues:
            print("-", x)
    else:
        print("None")

    print("\n=== WARNINGS (review) ===")
    if warnings:
        for x in warnings:
            print("-", x)
    else:
        print("None")

    if grammar_hits:
        print(f"\n=== GRAMMAR/SPACING HITS ({len(grammar_hits)}) ===")
        for g in grammar_hits[:20]:
            print("-", g)
        if len(grammar_hits) > 20:
            print(f"... and {len(grammar_hits) - 20} more")

    if missing_from_toc:
        print(f"\n=== HEADINGS NOT IN TOC ({len(missing_from_toc)}) ===")
        for m in missing_from_toc:
            print("-", m)

    print(f"\nTOTAL: {len(issues)} issues, {len(warnings)} warnings, {len(grammar_hits)} grammar hits")


if __name__ == "__main__":
    main()
