"""
Final thesis polish: spacing, remove redundant Ch.7, renumber Ch.8→7,
full-width clear figures, table spacing, TOC cleanup, save to Desktop.
Preserves all doctor-required content (methodology stays in Ch.4–5).
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
DESKTOP = Path(r"C:\Users\mohamad\OneDrive\Desktop")
PAPER = DESKTOP / "AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
STAGING = ROOT / "_POLISH_STAGING.docx"
BACKUP = PAPER.with_name(PAPER.stem + f"_pre_polish_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
FIGS = ROOT / "new_figures"
EVAL = ROOT / "evaluation"
SHOTS = ROOT / "user_guide_shots"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
EMU = 914400

# Full text width (1" margins on letter)
BODY_IMG_WIDTH = 6.75
DIAGRAM_MAX_H = 5.85
SCREENSHOT_MAX_H = 6.25
APPENDIX_IMG_WIDTH = 6.2


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def has_page_break(p: Paragraph) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def has_image(p: Paragraph) -> bool:
    return bool(p._p.findall(f".//{NS_A}blip"))


def set_inline_size(p: Paragraph, width_in: float, max_height_in: float) -> None:
    ext = p._p.find(f".//{NS_WP}extent")
    if ext is None:
        return
    cx = int(ext.get("cx") or 0)
    cy = int(ext.get("cy") or 0)
    if cx <= 0 or cy <= 0:
        return
    new_cx = int(width_in * EMU)
    new_cy = int(cy * (new_cx / cx))
    if new_cy > max_height_in * EMU:
        new_cy = int(max_height_in * EMU)
        new_cx = int(cx * (new_cy / cy))
    ext.set("cx", str(new_cx))
    ext.set("cy", str(new_cy))
    for aext in p._p.findall(f".//{NS_A}ext"):
        if aext.get("cx") and aext.get("cy"):
            aext.set("cx", str(new_cx))
            aext.set("cy", str(new_cy))


def remove_chapter7(doc: Document) -> int:
    """Remove Ch.7 body (content already in Ch.4–5 per doctor)."""
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "Chapter 7 - Practical Implementation Notes":
            start = i
        elif start is not None and t == "Chapter 8 - User Manual and Operation Guide":
            end = i
            break
    if start is None or end is None:
        return 0
    to_del = [doc.paragraphs[i] for i in range(start, end)]
    for p in to_del:
        delete_paragraph(p)
    return len(to_del)


def renumber_chapter8_to_7(doc: Document) -> None:
    """Ch.8 User Manual becomes Ch.7 after removing old Ch.7."""
    for p in doc.paragraphs:
        t = p.text or ""
        if not t:
            continue
        if t.strip() == "Chapter 8 - User Manual and Operation Guide":
            set_text(p, "Chapter 7 - User Manual and Operation Guide")
            continue
        if re.match(r"^8\.\d+\b", t.strip()):
            set_text(p, re.sub(r"^8\.", "7.", t.strip(), count=1))
            continue
        # cross-refs inside former Ch.8 sections
        t2 = re.sub(r"\bChapter 8\b", "Chapter 7", t)
        t2 = re.sub(r"\b§8\.", "§7.", t2)
        t2 = re.sub(r"\bsection 8\.", "section 7.", t2, flags=re.I)
        if t2 != t:
            set_text(p, t2)


def fix_toc(doc: Document) -> None:
    """Clean static TOC: drop old Ch.7, promote Ch.8→7, fix known glitches."""
    skip_until_ch8 = False
    for p in list(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        if not st.startswith("toc"):
            continue
        t = (p.text or "").strip()
        # Corrupted merged TOC lines
        if t.count("Chapter 7") > 1 or t.count("Chapter 8") > 0 and "Chapter 7" in t:
            if "Chapter 8 - User Manual" in t:
                set_text(p, "Chapter 7 - User Manual and Operation Guide\t51")
                continue
        if t.startswith("Chapter 7 - Practical Implementation Notes"):
            delete_paragraph(p)
            skip_until_ch8 = True
            continue
        if skip_until_ch8:
            if t.startswith("Chapter 8"):
                skip_until_ch8 = False
            elif re.match(r"^7\.\d+", t):
                delete_paragraph(p)
                continue
        if "Chapter 8 - User Manual" in t:
            set_text(p, "Chapter 7 - User Manual and Operation Guide\t51")
            continue
        if re.match(r"^8\.\d+", t):
            set_text(p, re.sub(r"^8\.", "7.", t, count=1))
            continue
        if "3.4 Data Flow Diagram" in t and "System Data Pipeline" in t:
            set_text(p, "3.4 System Data Pipeline\t23")
            continue
        if t.startswith("3.4 Data Flow Diagram"):
            set_text(p, "3.4 System Data Pipeline\t23")
            continue
        if "5.6 Expected Performance Comparison" in t:
            set_text(p, "5.6 Experimental Performance Results\t41")
            continue
        if "5.6 Expected Performance" in t and "Experimental" not in t:
            set_text(p, "5.6 Experimental Performance Results\t41")


def collapse_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for p in list(doc.paragraphs):
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if st == "Heading 1" and not t and not has_image(p):
            delete_paragraph(p)
            removed += 1
    paras = list(doc.paragraphs)
    run: list[Paragraph] = []
    for p in paras:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if (
            not t
            and not has_image(p)
            and not has_page_break(p)
            and st in ("Normal", "List Bullet", "List Number")
        ):
            run.append(p)
        else:
            if len(run) > 1:
                for extra in run[1:]:
                    delete_paragraph(extra)
                    removed += 1
            run = []
    return removed


def space_body_text(doc: Document) -> None:
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if st == "Normal" and t and not t.startswith("Figure ") and not t.startswith("Table "):
            pf = p.paragraph_format
            if pf.space_after is None or (pf.space_after and pf.space_after.pt < 6):
                pf.space_after = Pt(6)
            if pf.line_spacing is None:
                pf.line_spacing = 1.15
        if st.startswith("Heading"):
            pf = p.paragraph_format
            pf.space_before = Pt(12 if st == "Heading 1" else 8)
            pf.space_after = Pt(6)


def space_captions(doc: Document) -> None:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure ") or t.startswith("Table "):
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(12)
            pf.keep_with_next = False
            for r in p.runs:
                r.italic = True
                if r.font.size is None or r.font.size.pt < 10:
                    r.font.size = Pt(10)


def space_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_before = Pt(3)
                    pf.space_after = Pt(3)
                    pf.line_spacing = 1.08
                    for r in p.runs:
                        if r.font.size is None:
                            r.font.size = Pt(10)


def enlarge_figures(doc: Document) -> int:
    n = 0
    first_cover = True
    for i, p in enumerate(doc.paragraphs):
        if not has_image(p):
            continue
        n += 1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        if first_cover:
            first_cover = False
            set_inline_size(p, 6.5, 8.5)
            continue
        # screenshot figures
        nxt = doc.paragraphs[i + 1].text if i + 1 < len(doc.paragraphs) else ""
        if any(x in (nxt or "") for x in ("Figure 18.", "Figure 19.", "Figure 20.", "Figure 21.", "Figure 11.")):
            set_inline_size(p, BODY_IMG_WIDTH, SCREENSHOT_MAX_H)
        elif "Figure H" in (nxt or "") or "Appendix" in (nxt or ""):
            set_inline_size(p, APPENDIX_IMG_WIDTH, 5.5)
        else:
            set_inline_size(p, BODY_IMG_WIDTH, DIAGRAM_MAX_H)
    return n


def page_break_before_chapters(doc: Document) -> None:
    seen: set[str] = set()
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not (p.style and p.style.name == "Heading 1"):
            continue
        if not (t.startswith("Chapter ") or t.startswith("Appendix ") or t in ("Conclusion", "References")):
            continue
        if t in seen:
            continue
        seen.add(t)
        prev = p._p.getprevious()
        if prev is not None and prev.tag.endswith("p"):
            # check if previous para already has page break
            pass
        br = p.insert_paragraph_before("")
        br.add_run().add_break(WD_BREAK.PAGE)


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            r_id = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = paras[j].part.related_parts.get(r_id) if r_id else None
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def refresh_key_figures(doc: Document) -> None:
    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
    ]
    for cap, img in mapping:
        replace_image_before_caption(doc, cap, img)


def word_update_fields(path: Path) -> None:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(path), ReadOnly=False)
        wdoc.Fields.Update()
        for i in range(1, wdoc.TablesOfContents.Count + 1):
            wdoc.TablesOfContents(i).Update()
        for i in range(1, wdoc.TablesOfFigures.Count + 1):
            wdoc.TablesOfFigures(i).Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
    except Exception as exc:
        print("Word COM:", exc)


def verify_doctor(path: Path) -> None:
    import verify_doctor_revision as vd

    for i in range(3):
        passed, lines = vd.run_checks(path)
        print(f"Verify run {i+1}: {passed}/{len(lines)}")
        if passed < len(lines):
            for ln in lines:
                if ln.startswith("FAIL"):
                    print(ln)
            raise SystemExit("Doctor verification failed after polish")


def fix_global_chapter_refs(doc: Document) -> None:
    """Update prose that still mentions old Ch.7/Ch.8 numbering."""
    repl = [
        (r"Chapter 7 and Chapter 8", "Chapter 7 (User Manual)"),
        (r"Chapters 7 and 8", "Chapter 7"),
        (r"Chapter 8 - User Manual", "Chapter 7 - User Manual"),
        (r"see Chapter 8", "see Chapter 7"),
        (r"in Chapter 8", "in Chapter 7"),
        (r"Chapter 7 - Practical Implementation Notes", ""),
    ]
    for p in doc.paragraphs:
        t = p.text or ""
        new = t
        for old, new_s in repl:
            new = re.sub(old, new_s, new, flags=re.I)
        if new != t:
            set_text(p, new.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    new = t
                    for old, new_s in repl:
                        new = re.sub(old, new_s, new, flags=re.I)
                    if new != t:
                        set_text(p, new.strip())


def regenerate_figure_assets() -> None:
    """Refresh PNGs at readable resolution before embedding."""
    subprocess.run([sys.executable, str(ROOT / "generate_academic_diagrams.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "render_eval_figures.py")], check=True, cwd=str(ROOT))
    try:
        from PIL import Image

        def crop(src: Path, dst: Path) -> None:
            if not src.exists():
                return
            im = Image.open(src).convert("RGB")
            w, h = im.size
            th = min(h, int(w * 10 / 16))
            im.crop((0, 0, w, th)).save(dst, quality=95)

        for s, d in [
            (SHOTS / "00_home.png", EVAL / "shot_home_print.png"),
            (SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png"),
            (SHOTS / "10_soc.png", EVAL / "shot_soc_print.png"),
            (SHOTS / "07_models.png", EVAL / "shot_models_print.png"),
        ]:
            crop(s, d)
    except Exception as exc:
        print("Screenshot crop skip:", exc)


def main() -> None:
    regenerate_figure_assets()
    src = PAPER if PAPER.exists() else ROOT / "_FINAL_STAGING.docx"
    if not src.exists():
        src = ROOT / "_REBUILD_STAGING.docx"
    polished_alt = DESKTOP / "AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_POLISHED.docx"
    if polished_alt.exists() and not PAPER.exists():
        src = polished_alt
    if not src.exists():
        raise SystemExit(f"No thesis found at {PAPER}")

    shutil.copy2(src, BACKUP)
    shutil.copy2(src, STAGING)
    doc = Document(str(STAGING))

    stats = {
        "removed_ch7": remove_chapter7(doc),
        "removed_empty": collapse_empty_paragraphs(doc),
    }
    renumber_chapter8_to_7(doc)
    fix_toc(doc)
    fix_global_chapter_refs(doc)
    space_body_text(doc)
    space_captions(doc)
    space_tables(doc)
    refresh_key_figures(doc)
    stats["images_resized"] = enlarge_figures(doc)
    page_break_before_chapters(doc)

    doc.save(str(STAGING))

    try:
        doc.save(str(PAPER))
        out = PAPER
    except PermissionError:
        alt = DESKTOP / "AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_POLISHED.docx"
        doc.save(str(alt))
        out = alt
        print("Desktop file locked — saved as:", alt.name)

    word_update_fields(out)
    verify_doctor(out)

    print("Polish complete:", out)
    print("Backup:", BACKUP)
    print("Stats:", stats)


if __name__ == "__main__":
    main()
