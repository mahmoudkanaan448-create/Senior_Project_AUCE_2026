"""
Apply ONLY Dr. Hassan Noureddine requirements (19 points) — no TOC, spacing, or blank-page edits.
Target: pre_polish Desktop file (in-place).
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
ASSETS = Path(__file__).resolve().parent
TARGET = Path(
    r"C:\Users\mohamad\OneDrive\Desktop"
    r"\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED_pre_polish_20260820_0214.docx"
)
RESTORE = ASSETS / "_FINAL_STAGING.docx"
METRICS_PATH = ASSETS / "evaluation" / "metrics.json"
FIGS = ASSETS / "new_figures"
EVAL = ASSETS / "evaluation"
SNIPS = ASSETS / "code_snippets"
SHOTS = ASSETS / "user_guide_shots"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"


def refresh_figures(doc: Document) -> int:
    from PIL import Image

    def crop_shot(src: Path, dst: Path) -> None:
        if not src.exists():
            return
        im = Image.open(src).convert("RGB")
        w, h = im.size
        th = min(h, int(w * 10 / 16))
        im.crop((0, 0, w, th)).save(dst, quality=95)

    EVAL.mkdir(parents=True, exist_ok=True)
    crop_shot(SHOTS / "01_live.png", EVAL / "shot_live_print.png")
    crop_shot(SHOTS / "00_home.png", EVAL / "shot_home_print.png")
    crop_shot(SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png")
    crop_shot(SHOTS / "10_soc_tab1_SOAR_Playbooks.png", EVAL / "shot_soc_print.png")
    crop_shot(SHOTS / "07_models.png", EVAL / "shot_models_print.png")

    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 11.", EVAL / "shot_live_print.png"),
        ("Figure 12.", EVAL / "fig_cm_example.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
        ("Figure H1", SNIPS / "01_hybrid_ai_decision_fusion.png"),
        ("Figure H2", SNIPS / "09_mitre_map.png"),
        ("Figure H3", SNIPS / "10_soar_playbooks.png"),
        ("Figure H4", SNIPS / "12_online_learning.png"),
        ("Figure H5", SNIPS / "14_threat_simulation.png"),
        ("Figure H6", SNIPS / "13_alert_manager.png"),
        ("Figure H7", SNIPS / "16_telegram_alert.png"),
        ("Figure H8", SNIPS / "15_supervisor.png"),
        ("Figure H9", SNIPS / "17_ai_assistant.png"),
    ]
    n = 0
    paras = list(doc.paragraphs)
    for cap, path in mapping:
        if not path.exists():
            continue
        blob = path.read_bytes()
        for i, p in enumerate(paras):
            if cap not in (p.text or ""):
                continue
            for j in range(i, max(-1, i - 8), -1):
                if j < 0:
                    break
                blips = paras[j]._p.findall(f".//{NS_A}blip")
                if not blips:
                    continue
                rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                part = paras[j].part.related_parts.get(rid) if rid else None
                if part:
                    part._blob = blob
                    n += 1
                    break
            break
    return n


def apply_doctor_email(doc: Document, m: dict) -> None:
    """Content fixes from doctor_email_revision (no page breaks)."""
    from doctor_email_revision import (
        FEATURE_META,
        add_row,
        clear_table_keep_header,
        find_contains,
        find_exact,
        insert_after,
        pct,
        rebuild_feature_table,
        row_by_first,
        set_cell,
        set_text,
    )

    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    ds = m["dataset"]
    classes = ds.get("classes") or {}
    class_txt = ", ".join(f"{k} {v}" for k, v in classes.items())

    obj = doc.tables[3]
    c = row_by_first(obj, "O5")
    if c:
        set_cell(c[1], "Improve transparency using stored per-prediction Explainable AI (this-flow evidence and tree importances).")
    c = row_by_first(obj, "O6")
    if c:
        set_cell(c[1], "Support incident response using IP blocking, Telegram alerts, webhooks, and reports (no email channel).")

    tech = doc.tables[9]
    tech_map = {
        "Database": "SQLite by default; PostgreSQL optional via DATABASE_URL",
        "AI Libraries": "Scikit-learn, XGBoost, PyTorch (sklearn MLP fallback if torch unavailable)",
        "Traffic Capture": "Scapy/Npcap live NIC, OS connection table, optional 24/7 capture daemon",
        "Feature Extraction": "Custom 39 CICIDS-style flow features (config.FEATURE_COLUMNS)",
        "Explainability": "Stored per-prediction XAI JSON; SHAP is future work",
        "Dashboard": "Streamlit SOC dashboard, port 8501",
        "Notifications": "Telegram Bot API, local Windows notification, optional webhooks/SIEM (no SMTP email)",
        "Backend": "FastAPI REST API on port 8000",
    }
    for row in tech.rows:
        k = (row.cells[0].text or "").strip()
        if k in tech_map:
            set_cell(row.cells[1], tech_map[k])

    fe = find_exact(doc, "4.2 Feature Engineering")
    if fe and not find_contains(doc, "domain knowledge from CICIDS-style NDR literature"):
        intro = insert_after(
            fe,
            "The experiments use a fixed schema of 39 flow-level features defined in config.FEATURE_COLUMNS. "
            "They were selected using domain knowledge from CICIDS-style NDR literature (flow duration, byte/packet "
            "statistics, flag counts, host/service rates) rather than an automatic filter applied only for the report. "
            "The same schema is used in live capture, training, and inference so train/serve skew is avoided.",
            "Normal",
        )
        if not find_contains(doc, "4.2.1 Complete Feature Set Used in Experiments"):
            h = insert_after(intro, "4.2.1 Complete Feature Set Used in Experiments", "Heading 3")
            insert_after(
                h,
                "Table 5 lists every feature used in the holdout evaluation. All numeric features are standardised with "
                "StandardScaler fitted on the training split only. protocol_type is encoded numerically before scaling.",
                "Normal",
            )

    if not find_contains(doc, "Relevance to NDR"):
        rebuild_feature_table(doc)

    if not find_contains(doc, "4.2.2 Evaluation Dataset Construction"):
        anchor422 = find_contains(doc, "4.2.1 Complete Feature Set Used in Experiments") or fe
        if anchor422:
            h = insert_after(anchor422, "4.2.2 Evaluation Dataset Construction", "Heading 3")
            insert_after(
                h,
                f"Source. The reproducible evaluation file is datasets/dataset.csv built by training/generate_realistic_dataset.py "
                f"(synthetic CICIDS-style rows using the same attack signatures as the live detector).\n\n"
                f"Shape. {ds['rows']:,} rows × 39 input features + label column (40 columns in CSV).\n\n"
                f"Why not 720? An earlier draft used a balanced 720-row pilot (120 per class). Following supervisor feedback, the "
                f"final evaluation was regenerated at {ds['rows']:,} rows with realistic imbalance: {class_txt}.\n\n"
                f"Split. Stratified 75/25 train-test holdout, random_state=42, executed in training/evaluate_models.py. "
                f"Preprocessing is fit on training rows only.\n\n"
                f"Leakage controls. No label column in features; scaler/encoder saved from train split; duplicate rows removed before splitting.",
                "Normal",
            )

    mdev = find_exact(doc, "4.3 Model Development")
    if mdev:
        h431 = find_contains(doc, "4.3.1 Per-Model Training Configuration")
        if not h431:
            h431 = insert_after(mdev, "4.3.1 Per-Model Training Configuration", "Heading 3")
        if not find_contains(doc, "decision_engine.fuse_decisions"):
            insert_after(
                h431,
                "Random Forest — supervised multi-class on all 39 features; n_estimators=200, max_depth=20. "
                "XGBoost — supervised multi-class; n_estimators=200, max_depth=8, learning_rate=0.1. "
                "Isolation Forest — trained on Normal traffic only; contamination=0.1. "
                "Autoencoder and LSTM are implemented but not scored on Table 11 holdout. "
                "Hybrid AI uses decision_engine.fuse_decisions() weighted majority fusion; no standalone holdout accuracy.",
                "Normal",
            )

    perf = doc.tables[19]
    values = {
        "Accuracy": [pct(rf["accuracy"]), pct(xgb["accuracy"]), pct(iso["accuracy"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "Precision": [pct(rf["precision"]), pct(xgb["precision"]), pct(iso["precision"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "Recall": [pct(rf["recall"]), pct(xgb["recall"]), pct(iso["recall"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "F1 Score": [pct(rf["f1"]), pct(xgb["f1"]), pct(iso["f1"]), "Not on holdout", "Not on holdout", "Fusion — no holdout score"],
        "ROC-AUC": ["See §5.6", "See §5.6", "Binary proxy", "Not reported", "Not reported", "N/A"],
    }
    for row in perf.rows:
        key = (row.cells[0].text or "").strip()
        if key in values:
            for i, val in enumerate(values[key], start=1):
                set_cell(row.cells[i], val)

    for row in doc.tables[1].rows:
        if "Table 9" in (row.cells[0].text or ""):
            set_cell(row.cells[1], "Holdout metrics for RF/XGB/IF; AE/LSTM/Hybrid not scored on the same split (see text).")

    p720 = find_contains(doc, "720")
    if p720 and "12,000" not in (p720.text or "") and "implementation evaluation dataset" in (p720.text or "").lower():
        set_text(
            p720,
            f"An implementation evaluation dataset was prepared with {ds['rows']:,} labeled samples, 39 flow features, "
            f"and six classes ({class_txt}). Holdout evaluation uses a stratified 25% test split (random_state=42). "
            f"Random Forest achieved accuracy {pct(rf['accuracy'])}, F1 {pct(rf['f1'])}. "
            f"XGBoost achieved accuracy {pct(xgb['accuracy'])}, F1 {pct(xgb['f1'])}. "
            f"Isolation Forest achieved accuracy {pct(iso['accuracy'])}, F1 {pct(iso['f1'])}.",
        )

    t11 = find_contains(doc, "Table 11. Experimental model performance")
    if t11:
        set_text(
            t11,
            f"Table 11. Measured holdout metrics (reproducible via python -m training.evaluate_models). "
            f"Autoencoder, LSTM, and Hybrid AI: implemented; no separate holdout row on the same split.",
        )

    note = find_contains(doc, "Important note: the values above are expected performance targets")
    if note:
        set_text(
            note,
            "Table 9 reports measured holdout values for RF/XGB/IF only. Placeholder targets (e.g., 99.2% Hybrid) were removed.",
        )

    tests = doc.tables[20]
    extra = [
        ("TC11", "Upload invalid CSV", "Parser rejects file; dashboard/API shows readable error; no crash."),
        ("TC12", "Database unavailable", "API returns 503-style message; dashboard shows connection warning."),
        ("TC13", "Model file missing", "Detection endpoint reports model-not-ready safely."),
        ("TC14", "Threat Intel API timeout", "Detection continues; TI field marked unavailable."),
        ("TC15", "False positive review", "Analyst marks alert investigated; no auto-block below threshold."),
    ]
    for tc, act, exp in extra:
        if not row_by_first(tests, tc):
            add_row(tests, [tc, act, exp])

    if not find_contains(doc, "5.3.1 Structured Attack Testing Matrix"):
        atk = find_exact(doc, "5.3 Attack Simulation Scenarios")
        if atk:
            h = insert_after(atk, "5.3.1 Structured Attack Testing Matrix", "Heading 3")

    ch8 = find_exact(doc, "Chapter 8 - User Manual and Operation Guide")
    if ch8 and not find_contains(doc, "concise operator guide only"):
        insert_after(
            ch8,
            "This chapter is a concise operator guide only. Methodology and evaluation evidence are in Chapters 4–5.",
            "Normal",
        )

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure 15. Dataset class distribution"):
            set_text(p, f"Figure 15. Class distribution ({ds['rows']:,}-row evaluation set).")
        if "Data Flow Diagram" in t and t.startswith("Figure"):
            set_text(p, t.replace("Data Flow Diagram", "System Data Pipeline"))


def apply_doctor_complete(doc: Document, m: dict) -> None:
    import doctor_complete_revision as dcr

    dcr.fix_dfd_wording(doc)
    dcr.rebuild_feature_table_full(doc)
    dcr.add_lstm_sequence_section(doc)
    dcr.add_hybrid_fusion_section(doc)
    dcr.add_model_usage_table(doc, m)
    dcr.add_section_56_metrics(doc, m)
    dcr.add_implementation_status_table(doc)
    dcr.add_table9_interpretation(doc)
    dcr.slim_chapter7(doc)
    dcr.ensure_attack_matrix_table(doc)
    dcr.expand_appendix_b_defense(doc)


def apply_chapter_restructure(doc: Document) -> None:
    from polish_final_thesis import remove_chapter7, renumber_chapter8_to_7

    remove_chapter7(doc)
    renumber_chapter8_to_7(doc)


def fix_chapter_toc_minimal(doc: Document) -> int:
    """Fix only Ch.7/8 TOC lines after restructure (doctor §12) — not a full TOC rebuild."""
    from doctor_email_revision import set_text
    from polish_final_thesis import delete_paragraph

    n = 0
    for p in list(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        t = p.text or ""
        if not st.startswith("toc"):
            if "Chapter 7 and Chapter 7 provide practical notes" in t:
                set_text(
                    p,
                    t.replace(
                        "Chapter 7 and Chapter 7 provide practical notes and the user manual.",
                        "Chapter 7 provides the user manual and operation guide.",
                    ),
                )
                n += 1
            continue
        if "Chapter 7 - Practical Implementation Notes" in t:
            delete_paragraph(p)
            n += 1
        elif "Chapter 8 - User Manual" in t:
            pg = "51"
            if "\t" in t:
                pg = re.sub(r"[^\d].*", "", t.split("\t")[-1]) or pg
            anchor_el = p._p
            para_style = p.style.name if p.style else "toc 1"
            delete_paragraph(p)
            new_p = OxmlElement("w:p")
            anchor_el.addprevious(new_p)
            para = Paragraph(new_p, doc.element.body)
            try:
                para.style = doc.styles[para_style]
            except Exception:
                pass
            para.add_run(f"Chapter 7 - User Manual and Operation Guide\t{pg}")
            n += 1
    return n


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    args = [a for a in sys.argv[1:] if not a.startswith("-")]
    no_restore = "--no-restore" in sys.argv
    target = Path(args[0]).resolve() if args else TARGET
    restore = RESTORE if RESTORE.exists() else ASSETS / "_REBUILD_STAGING.docx"

    if not no_restore:
        shutil.copy2(restore, target)
        print("Restored from:", restore)

    subprocess.run([sys.executable, "-m", "training.evaluate_models"], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ASSETS / "generate_academic_diagrams.py")], check=True, cwd=str(ASSETS))
    subprocess.run([sys.executable, str(ASSETS / "render_eval_figures.py")], check=True, cwd=str(ASSETS))
    subprocess.run([sys.executable, str(ASSETS / "generate_extra_snippets.py")], check=True, cwd=str(ASSETS))

    m = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    doc = Document(str(target))

    apply_doctor_email(doc, m)
    apply_doctor_complete(doc, m)
    apply_chapter_restructure(doc)
    stats = {
        "figures": refresh_figures(doc),
        "toc_chapter_fix": fix_chapter_toc_minimal(doc),
    }

    doc.save(str(target))
    print("Doctor-only patch saved:", target, stats)

    import verify_doctor_revision as vd

    passed, lines = vd.run_checks(target)
    print(f"Verify {passed}/{len(lines)}")
    for ln in lines:
        if ln.startswith("FAIL"):
            print(ln)


if __name__ == "__main__":
    main()
