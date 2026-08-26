"""
Ultimate quality fix: TOC gaps, List of Tables completeness, captions, spacing, grammar.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def style_caption(p: Paragraph) -> None:
    p.alignment = 1  # center
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(10)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(10)


def clear_table_keep_header(table) -> None:
    tbl = table._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = val


# Canonical List of Tables (Tables 1–12 + 3A + 5A) then acronym rows preserved separately
CANONICAL_TABLES = [
    ("Table 1", "List of Acronyms"),
    ("Table 2", "Research Objectives and Research Questions"),
    ("Table 3", "Comparison of 20 Cybersecurity Platforms"),
    ("Table 3A", "Component Implementation Status"),
    ("Table 4", "Dataset Options"),
    ("Table 5", "Extracted Network Features (39 holdout features)"),
    ("Table 5A", "Per-Model Dataset Usage Summary"),
    ("Table 6", "AI Models and Roles"),
    ("Table 7", "Database Tables"),
    ("Table 8", "Performance Metrics Definitions"),
    ("Table 9", "Technology Stack"),
    ("Table 10", "Risk and Mitigation Plan"),
    ("Table 11", "Measured holdout metrics (reproducible via python -m training.evaluate_models)"),
    ("Table 12", "Operational statistics captured from the live SQLite database during testing"),
]

ACRONYMS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("ATT&CK", "Adversarial Tactics, Techniques, and Common Knowledge"),
    ("AUC", "Area Under Curve"),
    ("CSV", "Comma-Separated Values"),
    ("DDoS", "Distributed Denial of Service"),
    ("DFD", "Data Flow Diagram"),
    ("DL", "Deep Learning"),
    ("DoS", "Denial of Service"),
    ("IDS", "Intrusion Detection System"),
    ("IF", "Isolation Forest"),
    ("IPS", "Intrusion Prevention System"),
    ("JWT", "JSON Web Token"),
    ("LSTM", "Long Short-Term Memory"),
    ("MITRE", "MITRE Corporation (ATT&CK knowledge base)"),
    ("ML", "Machine Learning"),
    ("NDR", "Network Detection and Response"),
    ("NIDS", "Network Intrusion Detection System"),
    ("REST", "Representational State Transfer"),
    ("RF", "Random Forest"),
    ("ROC", "Receiver Operating Characteristic"),
    ("SGD", "Stochastic Gradient Descent"),
    ("SIEM", "Security Information and Event Management"),
    ("SOAR", "Security Orchestration, Automation and Response"),
    ("SOC", "Security Operations Center"),
    ("SQL", "Structured Query Language"),
    ("SQLite", "Embedded Relational Database Engine"),
    ("STRIDE", "Spoofing, Tampering, Repudiation, Information Disclosure, DoS, Elevation of Privilege"),
    ("TCP", "Transmission Control Protocol"),
    ("TI", "Threat Intelligence"),
    ("UDP", "User Datagram Protocol"),
    ("UML", "Unified Modeling Language"),
    ("XAI", "Explainable Artificial Intelligence"),
    ("XDR", "Extended Detection and Response"),
    ("XGB", "Extreme Gradient Boosting (XGBoost)"),
]


def replace_corrupted_563_toc(doc: Document) -> int:
    """Fix merged 5.6.1/5.6.3 TOC corruption via delete + insert."""
    anchor = None
    removed = 0
    for p in list(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        if not st.startswith("toc"):
            continue
        t = (p.text or "").strip()
        if "Operational Deployment Statistics" in t or (
            "5.6.3" in t and "5.6.1" in t
        ):
            if anchor is None:
                # remember paragraph after 5.6.2 to insert before next toc entry
                pass
            delete_paragraph(p)
            removed += 1
    # find 5.6.2 to insert 5.6.3 after it
    for i, p in enumerate(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        t = (p.text or "").strip()
        if st == "toc 3" and t.startswith("5.6.2"):
            from docx.oxml import OxmlElement

            new_p = OxmlElement("w:p")
            p._p.addnext(new_p)
            para = Paragraph(new_p, p._parent)
            try:
                para.style = doc.styles["toc 3"]
            except Exception:
                pass
            set_text(para, "5.6.3 Operational Deployment Statistics\t41")
            return removed + 1
    return removed


def fix_toc_563(doc: Document) -> int:
    return replace_corrupted_563_toc(doc)


def rebuild_lot(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    clear_table_keep_header(table)
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Table Title"
    for no, title in CANONICAL_TABLES:
        add_table_row(table, [no, title])
    for acr, meaning in ACRONYMS:
        add_table_row(table, [acr, meaning])


def ensure_table_captions(doc: Document) -> int:
    """Ensure formal Table 5 / 5A caption lines exist before their tables."""
    n = 0
    targets = {
        "Table 5 lists all 39 features actually passed to training/inference": (
            "Table 5. Extracted network features (39 holdout features)."
        ),
        "Table 5A summarises how each model consumes the same underlying CSV": (
            "Table 5A. Per-model dataset usage summary."
        ),
    }
    existing = {(p.text or "").strip() for p in doc.paragraphs}
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        for needle, caption in targets.items():
            if needle in t:
                if caption in existing:
                    break
                from docx.oxml import OxmlElement

                new_p = OxmlElement("w:p")
                p._p.addprevious(new_p)
                para = Paragraph(new_p, p._parent)
                set_text(para, caption)
                style_caption(para)
                existing.add(caption)
                n += 1
                break
    return n


def remove_duplicate_captions(doc: Document) -> int:
    """Remove consecutive duplicate Table 5 / 5A caption lines."""
    removed = 0
    prev = ""
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("Table 5.") or t.startswith("Table 5A."):
            if t == prev:
                el = p._p
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    removed += 1
                continue
            prev = t
        else:
            prev = ""
    return removed


def fix_grammar_spacing(doc: Document) -> int:
    fixed = 0
    for p in doc.paragraphs:
        t = p.text or ""
        if "Signature:" in t and "Date:" in t:
            new = re.sub(r"\s{2,}", " ", t)
            new = new.replace("Signature: ______________________ Date:", "Signature: ______________________    Date:")
            if new != t:
                set_text(p, new)
                fixed += 1
        # collapse triple+ spaces in normal body text (not code/indented)
        st = p.style.name if p.style else ""
        if st == "Normal" and re.search(r" {3,}", t) and not t.startswith("    "):
            new = re.sub(r" {2,}", " ", t)
            if new != t:
                set_text(p, new)
                fixed += 1
    return fixed


def dedupe_prose_table5(doc: Document) -> int:
    """Keep one Table 5 intro paragraph (remove duplicate older wording)."""
    removed = 0
    seen_better = False
    to_delete = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if "Table 5 lists all 39 features actually passed" in t:
            if seen_better:
                to_delete.append(p)
            else:
                seen_better = True
        elif "Table 5 lists every feature used in the holdout evaluation" in t:
            to_delete.append(p)
    for p in to_delete:
        el = p._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    return removed


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    doc = Document(str(PAPER))
    stats = {
        "toc_563": fix_toc_563(doc),
        "lot_rebuilt": 1,
        "captions": ensure_table_captions(doc),
        "dup_captions_removed": remove_duplicate_captions(doc),
        "grammar": fix_grammar_spacing(doc),
        "dedupe": dedupe_prose_table5(doc),
    }
    rebuild_lot(doc)
    doc.save(str(PAPER))
    print("Ultimate quality fix OK:", stats)


if __name__ == "__main__":
    main()
