"""
Rebuild List of Figures and List of Tables; renumber figures/tables consistently.
Keeps document template unchanged; fixes numbering gaps after Ch.4 code removal.
"""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
ASSETS = Path(__file__).resolve().parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop"
    r"\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
WORKING = ASSETS / "WORKING_LOF.docx"
SNIPS = ASSETS / "code_snippets"

APPENDIX_FIGURES = [
    ("H1", "01_hybrid_ai_decision_fusion.png", "Code excerpt: Hybrid AI decision fusion."),
    ("H2", "09_mitre_map.png", "Code excerpt: MITRE ATT&CK mapping module."),
    ("H3", "10_soar_playbooks.png", "Code excerpt: SOAR playbook definitions."),
    ("H4", "12_online_learning.png", "Code excerpt: online incremental learning."),
    ("H5", "14_threat_simulation.png", "Code excerpt: threat simulation engine."),
    ("H6", "13_alert_manager.png", "Code excerpt: alert manager with playbook hook."),
    ("H7", "16_telegram_alert.png", "Code excerpt: Telegram alert delivery."),
    ("H8", "15_supervisor.png", "Code excerpt: server auto-recovery supervisor."),
    ("H9", "17_ai_assistant.png", "Code excerpt: context-aware AI Security Assistant."),
]

LOT_EXTRA = [
    ("Table 11", "Experimental model performance on the implementation evaluation dataset."),
    ("Table 12", "Operational statistics captured from the live SQLite database during testing."),
]


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return new_para


def insert_picture_after(
    paragraph: Paragraph, image_path: Path, width_in: float, caption: str
) -> Paragraph:
    img_p = insert_after(paragraph, "")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = insert_after(img_p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)
    return cap


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def set_para_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def paragraph_has_image(p: Paragraph) -> bool:
    ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    return bool(p._p.findall(f".//{{{ns}}}blip"))


def clear_table_keep_header(table) -> None:
    tbl = table._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = val


def is_body_figure_caption(text: str) -> bool:
    return bool(re.match(r"^Figure\s+\d+\.\s*.+", text.strip()))


def is_appendix_figure_caption(text: str) -> bool:
    return bool(
        re.match(r"^(Appendix\s+)?Figure\s+H\d+\.\s*.+", text.strip(), re.I)
    )


def appendix_label(text: str) -> str | None:
    m = re.match(r"^(?:Appendix\s+)?Figure\s+(H\d+)\.\s*(.+)$", text.strip(), re.I)
    if m:
        return m.group(1).upper()
    return None


def renumber_body_figures(doc: Document) -> list[tuple[int, str]]:
    """Renumber Figure 1..N in document order (main chapters only)."""
    ch1 = find_para(doc, "Chapter 1 - Introduction")
    refs = find_para(doc, "References")
    entries: list[tuple[int, str]] = []
    n = 0
    started = ch1 is None

    for p in doc.paragraphs:
        if ch1 and p._p is ch1._p:
            started = True
        if not started:
            continue
        if refs and p._p is refs._p:
            break

        t = (p.text or "").strip()
        if not is_body_figure_caption(t):
            continue
        m = re.match(r"^Figure\s+\d+\.\s*(.+)$", t)
        if not m:
            continue
        n += 1
        title = m.group(1).strip()
        set_para_text(p, f"Figure {n}. {title}")
        entries.append((n, title))

    return entries


def fix_table_numbering(doc: Document) -> None:
  replacements = [
      (
          "Table 3. Experimental model performance",
          "Table 11. Experimental model performance",
      ),
      (
          "Table 4. Operational statistics captured",
          "Table 12. Operational statistics captured",
      ),
      ("Table 3 summarizes experimental metrics", "Table 11 summarizes experimental metrics"),
  ]
  for p in doc.paragraphs:
      t = p.text or ""
      new = t
      for old, new_val in replacements:
          if old in new:
              new = new.replace(old, new_val)
      if new != t:
          set_para_text(p, new)


def remove_appendix_h_figures(doc: Document) -> Paragraph | None:
    """Remove existing appendix H figure images/captions; return intro paragraph."""
    h = find_contains(doc, "Appendix H - Implementation Code Excerpts")
    if not h:
        return None

    intro = None
    to_delete: list[Paragraph] = []
    in_section = False
    for p in doc.paragraphs:
        if h and p._p is h._p:
            in_section = True
            continue
        if not in_section:
            continue
        t = (p.text or "").strip()
        if t.startswith("Appendix I") or t == "References":
            break
        if is_appendix_figure_caption(t) or paragraph_has_image(p):
            to_delete.append(p)
            continue
        if t and intro is None:
            intro = p

    for p in to_delete:
        delete_paragraph(p)
    return intro or h


def rebuild_appendix_h_figures(doc: Document) -> list[tuple[str, str]]:
    intro = remove_appendix_h_figures(doc)
    anchor = intro or find_contains(doc, "Appendix H - Implementation Code Excerpts")
    if not anchor:
        return []

    entries: list[tuple[str, str]] = []
    cur = anchor
    for label, fname, title in APPENDIX_FIGURES:
        img = SNIPS / fname
        if not img.exists():
            continue
        caption = f"Figure {label}. {title}"
        cur = insert_picture_after(cur, img, 6.0, caption)
        entries.append((label, title))
    return entries


def rebuild_lof_table(doc: Document, body_figs: list[tuple[int, str]], app_figs: list[tuple[str, str]]) -> None:
    table = doc.tables[0]
    clear_table_keep_header(table)
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Figure Title"
    for num, title in body_figs:
        add_table_row(table, [f"Figure {num}", title])
    for label, title in app_figs:
        add_table_row(table, [f"Figure {label}", title])


def rebuild_lot_table(doc: Document) -> None:
    table = doc.tables[1]
    existing: list[tuple[str, str]] = []
    for row in table.rows[1:]:
        no = row.cells[0].text.strip()
        title = row.cells[1].text.strip()
        if no and title and title != "List of Acronyms":
            existing.append((no, title))
        elif no == "Table 1" and title == "List of Acronyms":
            existing.append((no, title))

    # Re-read from scratch to preserve original 1-10
    existing = []
    for row in table.rows[1:]:
        existing.append((row.cells[0].text.strip(), row.cells[1].text.strip()))

    # Drop duplicate experimental tables if a prior run added wrong numbers
    filtered = []
    seen = set()
    for no, title in existing:
        key = (no, title)
        if key in seen:
            continue
        if "Experimental model performance" in title and no != "Table 11":
            continue
        if "Operational statistics captured" in title and no != "Table 12":
            continue
        seen.add(key)
        filtered.append((no, title))

    has_11 = any(no == "Table 11" for no, _ in filtered)
    has_12 = any(no == "Table 12" for no, _ in filtered)
    if not has_11 or not has_12:
        for no, title in LOT_EXTRA:
            if no == "Table 11" and not has_11:
                filtered.append((no, title))
            if no == "Table 12" and not has_12:
                filtered.append((no, title))

    clear_table_keep_header(table)
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Table Title"
    for no, title in filtered:
        add_table_row(table, [no, title])


def remove_loose_lot_line(doc: Document) -> None:
    lof = find_para(doc, "List of Figures")
    lot = find_para(doc, "List of Tables")
    ch1 = find_para(doc, "Chapter 1 - Introduction")
    if not lot:
        return
    to_delete = []
    started = False
    for p in doc.paragraphs:
        if lot and p._p is lot._p:
            started = True
            continue
        if not started:
            continue
        if ch1 and p._p is ch1._p:
            break
        t = (p.text or "").strip()
        if t == "List of Acronyms":
            to_delete.append(p)
    for p in to_delete:
        delete_paragraph(p)


def main() -> None:
    src = PAPER if PAPER.exists() else WORKING
    if not src.exists():
        raise SystemExit(f"Paper not found: {PAPER}")

    shutil.copy2(src, WORKING)
    backup = PAPER.with_name(
        PAPER.stem + f"_backup_lof_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    )
    shutil.copy2(src, backup)

    doc = Document(str(WORKING))
    remove_loose_lot_line(doc)
    fix_table_numbering(doc)
    body_figs = renumber_body_figures(doc)
    app_figs = rebuild_appendix_h_figures(doc)
    rebuild_lof_table(doc, body_figs, app_figs)
    rebuild_lot_table(doc)

    doc.save(str(WORKING))
    shutil.copy2(WORKING, PAPER)

    print("Saved:", PAPER)
    print("Backup:", backup)
    print("Body figures:", len(body_figs))
    for n, t in body_figs:
        print(f"  Figure {n}: {t[:70]}")
    print("Appendix figures:", len(app_figs))
    for label, t in app_figs:
        print(f"  Figure {label}: {t[:70]}")
    print("LoT rows:", len(doc.tables[1].rows) - 1)


if __name__ == "__main__":
    main()
