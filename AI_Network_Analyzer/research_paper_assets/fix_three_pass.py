"""
Pass 1: blank pages + TOC
Pass 2: heading intros + table spacing
Pass 3: refresh all figures/diagrams
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
WORK = ROOT / "_fix_work.docx"
STAGING = ROOT / "_fix_staging.docx"
DESKTOP = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
TARGET_PRE_POLISH = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED_pre_polish_20260820_0214.docx"
)
FIGS = ROOT / "new_figures"
EVAL = ROOT / "evaluation"
SNIPS = ROOT / "code_snippets"
SHOTS = ROOT / "user_guide_shots"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

# Intro paragraph before table-only headings
HEADING_INTROS: dict[str, str] = {
    "1.4 Research Objectives": (
        "Table 2 summarises the six research objectives that guided the design and evaluation of the AI-NDR platform."
    ),
    "1.5 Research Questions": (
        "Table 2 (continued) lists the research questions that each objective was designed to answer."
    ),
    "2.3 Machine Learning and Deep Learning Techniques": (
        "Table 6 compares the supervised and unsupervised models selected for this project and their roles in the hybrid engine."
    ),
    "2.4 Datasets Used in Network Intrusion Detection": (
        "Table 4 documents the public and synthetic datasets considered; the holdout evaluation uses the 12,000-row CICIDS-style file described in §4.2.2."
    ),
    "3.3 Main Modules": (
        "The implemented platform is organised into the modules listed below; §3.3.1 and §3.3.2 map each module to its evaluation status."
    ),
    "4.1 Technology Stack": (
        "Table 9 lists the programming languages, libraries, and services used in the FastAPI + Streamlit implementation."
    ),
    "4.4 Decision Fusion": (
        "Hybrid AI is implemented as weighted majority fusion (decision_engine.fuse_decisions), not a sixth trained classifier. Figure 2 and §4.4.1 show the live pipeline."
    ),
    "4.5 Database Design": (
        "Table 7 describes the SQLite tables that persist users, flows, alerts, incidents, and model metadata."
    ),
    "4.7 Implementation Workflow": (
        "Table 13 summarises the end-to-end workflow from dataset preparation through model training to dashboard deployment."
    ),
    "4.9 Algorithmic Workflow": (
        "Table 15 presents the step-by-step algorithm executed at inference time for each captured or uploaded flow."
    ),
    "4.11 Project Folder Structure": (
        "Table 16 maps the repository folders to their responsibilities in the senior project codebase."
    ),
    "5.2 Testing Environment": (
        "Table 17 lists the hardware, software, and network configuration used during functional and experimental testing."
    ),
    "5.3 Attack Simulation Scenarios": (
        "Table 18 defines the attack scenarios exercised through the Threat Simulation module; §5.3.1 provides the structured testing matrix required by the supervisor."
    ),
    "5.4 Evaluation Metrics": (
        "Table 8 defines accuracy, precision, recall, F1, ROC-AUC, FPR, and FNR used to report holdout results in §5.6."
    ),
    "5.6 Experimental Performance Results": (
        "This section reports measured holdout metrics only for Random Forest, XGBoost, and Isolation Forest (Table 11). Autoencoder, LSTM, and Hybrid AI are described honestly as implemented but not scored on the same split."
    ),
    "6.2 Security Considerations": (
        "Table 10 identifies deployment risks (credential storage, live blocking, API exposure) and the mitigations applied in the prototype."
    ),
    "7.5 Maintenance Procedures": (
        "Table 24 recommends maintenance tasks and frequencies for keeping the SOC dashboard and models operational."
    ),
    "7.6 Minimal Functional Requirements": (
        "Table 25 lists the minimum functional requirements verified during acceptance testing."
    ),
    "7.7 Non-Functional Requirements": (
        "Table 26 summarises non-functional requirements for performance, reliability, and usability."
    ),
    "3.3.1 Implemented Core Modules (Aligned with Software)": (
        "The implemented system modules match the software repository. Monitoring captures live connections or "
        "CSV flows and extracts 39 numerical features. The Hybrid AI engine loads Random Forest, XGBoost, Isolation "
        "Forest, Autoencoder, and LSTM models. The Decision Engine fuses votes into label, confidence, threat score, "
        "and severity. Threat Intelligence enriches source IPs. Explainable AI returns feature importance and recommended "
        "actions. The Alert Manager creates Medium+ alerts and executes SOAR playbooks."
    ),
    "3.3.2 Implementation and Evaluation Status": (
        "Implemented and tested: capture, 39-feature extraction, RF/XGB/IF holdout evaluation, dashboard, alerts, "
        "Telegram, Threat Simulation, stored XAI, MITRE mapping, SOAR playbooks. Implemented but not fully evaluated "
        "on the Table 11 holdout: Autoencoder, LSTM sequence model, Hybrid fusion (live only)."
    ),
    "4.2.1 Complete Feature Set Used in Experiments": (
        "This subsection documents the complete input schema used in training and inference. Every column in Table 5 "
        "is defined in config.FEATURE_COLUMNS and extracted identically during live capture, CSV import, and holdout "
        "evaluation using CICIDS-style NDR domain knowledge."
    ),
    "4.2.2 Evaluation Dataset Construction": (
        "The holdout file datasets/dataset.csv contains 12,000 labelled flows (75/25 stratified split, random_state=42). "
        "This subsection documents how the file was built, class distribution, and leakage controls applied before reporting §5.6 metrics."
    ),
    "4.3 Model Development": (
        "Model development: load dataset.csv, apply train-only preprocessing (LabelEncoder, StandardScaler), train each "
        "model, persist artifacts under models/ for FastAPI/Streamlit inference."
    ),
    "4.3.1 Per-Model Training Configuration": (
        "Random Forest — supervised multi-class on 39 scaled features (n_estimators=200, max_depth=20). XGBoost — "
        "supervised multi-class (n_estimators=200, max_depth=8, learning_rate=0.1). Isolation Forest — trained on Normal "
        "traffic only (contamination=0.1). Autoencoder — Normal-only reconstruction threshold; not on Table 11 holdout. "
        "Hybrid AI — fuse_decisions() weighted majority fusion; no standalone holdout accuracy."
    ),
    "4.3.2 LSTM Sequence Input Construction": (
        "LSTM uses training/data_preprocessing.prepare_sequences() with window_size=10. Each sample is a (10, 39) tensor; "
        "the label is the class of the flow after the window. Row-based models use one flow per sample; LSTM uses a sliding "
        "window. LSTM is implemented but not scored in Table 11 because evaluate_models.py is row-based."
    ),
    "4.3.3 Per-Model Dataset Usage Summary": (
        "Table 5A below shows how each model consumes the same CSV: RF/XGB/IF/AE use one 39-feature row; LSTM uses "
        "ten consecutive rows; Hybrid combines live votes without a separate training file."
    ),
    "4.3.4 Implemented Preprocessing and Training Configuration": (
        "The preprocessing pipeline loads CSV flows, removes duplicates, replaces infinite values, fits LabelEncoder and "
        "StandardScaler on the training split only, and persists encoders/scalers for identical inference-time transforms."
    ),
    "4.4.1 Hybrid Decision Fusion Pipeline": (
        "Hybrid AI is not a sixth trained classifier. Individual models produce label/confidence votes; "
        "decision_engine.fuse_decisions() merges them into final_label, confidence, threat_score, and severity before "
        "alert_manager creates SOC alerts (see Figure 2 and detection/decision_engine.py)."
    ),
    "5.3.1 Structured Attack Testing Matrix": (
        "Table 19 maps each simulated attack type to expected detector behaviour, alert severity, and SOAR response "
        "actions exercised during Threat Simulation campaigns."
    ),
    "5.6.1 Complete Holdout Metrics (Reproducible)": (
        "Table 11 reports measured holdout metrics for Random Forest, XGBoost, and Isolation Forest only. "
        "Values are reproduced by training/evaluate_models.py on datasets/dataset.csv with the split documented in §4.2.2."
    ),
    "5.6.2 High-Score Justification and Leakage Checklist": (
        "High accuracy is explained by controlled 12,000-row evaluation, stratified splitting, train-only scaling, "
        "and class-weight settings — not by inflating results with test-set leakage."
    ),
    "5.6.3 Operational Deployment Statistics": (
        "Table 12 summarises live operational counters captured from the SQLite database during SOC dashboard use "
        "(alerts raised, flows processed, simulation runs)."
    ),
    "Table 3A. Component Implementation Status": (
        "Table 3A maps each NDR 2.0 module to its implementation and evaluation status (green = tested on holdout or live demo)."
    ),
}


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def has_page_break(p) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def has_image(p) -> bool:
    return bool(p._p.findall(f".//{NS_A}blip"))


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = anchor.part.document.styles[style]
    except Exception:
        pass
    set_text(para, text)
    pf = para.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    return para


def insert_before(doc: Document, anchor: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass
    set_text(para, text)
    return para


def remove_blank_pages(doc: Document) -> int:
    n = 0
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if not t and not has_image(p) and has_page_break(p):
            delete_paragraph(p)
            n += 1
    # Empty Heading 1/2/3 with no text (often stray page breaks)
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if not t and not has_image(p) and st.startswith("Heading"):
            delete_paragraph(p)
            n += 1
    run: list[Paragraph] = []
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        empty = not t and not has_image(p) and not has_page_break(p) and st in ("Normal", "List Bullet")
        if empty:
            run.append(p)
        else:
            if len(run) >= 2:
                for e in run:
                    delete_paragraph(e)
                    n += 1
            run = []
    return n


def heading_has_intro(p: Paragraph, min_len: int = 25) -> bool:
    el = p._p
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag.endswith("}tbl"):
            return False
        if nxt.tag.endswith("}p"):
            para = Paragraph(nxt, p._parent)
            t = (para.text or "").strip()
            st = para.style.name if para.style else ""
            if st.startswith("Heading"):
                return False
            if t and len(t) >= min_len and not re.match(r"^Table \d+\.", t):
                return True
        nxt = nxt.getnext()
    return False


def add_heading_intros(doc: Document) -> int:
    added = 0
    for exact, intro in HEADING_INTROS.items():
        for p in doc.paragraphs:
            if (p.text or "").strip() != exact:
                continue
            if heading_has_intro(p):
                break
            insert_after(p, intro)
            added += 1
            break
    return added


def is_toc_style(style: str) -> bool:
    return (style or "").lower().startswith("toc")


def rebuild_toc(doc: Document) -> int:
    """Delete every toc 1/2/3 paragraph before List of Figures, then insert one clean block."""
    page = 1
    entries: list[tuple[int, str, int]] = []
    started = False
    for p in doc.paragraphs:
        if has_page_break(p):
            page += 1
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if t.startswith("Chapter 1 - Introduction") and st == "Heading 1":
            started = True
        if not started or not st.startswith("Heading"):
            continue
        if t in ("List of Figures", "List of Tables"):
            continue
        lvl = 1 if t.startswith("Chapter ") or t in ("Conclusion", "References") or t.startswith("Appendix") else (
            2 if st == "Heading 2" else 3
        )
        if st == "Heading 1" and not (t.startswith("Chapter") or t in ("Conclusion", "References") or t.startswith("Appendix")):
            continue
        entries.append((lvl, t, page))

    anchor = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == "List of Figures" and p.style and p.style.name == "Heading 1":
            anchor = p
            break
    if anchor is None:
        return 0

    # Remove ALL toc-styled paragraphs immediately before List of Figures
    removed = 0
    prev = anchor._p.getprevious()
    while prev is not None and prev.tag.endswith("}p"):
        para = Paragraph(prev, anchor._parent)
        st = (para.style.name if para.style else "").lower()
        if not is_toc_style(st):
            break
        to_del = prev
        prev = prev.getprevious()
        to_del.getparent().remove(to_del)
        removed += 1

    front = [
        (1, "Declaration", 2),
        (1, "Acknowledgment", 2),
        (1, "Abstract", 2),
        (1, "List of Figures", 3),
        (1, "List of Tables", 3),
    ]
    style_map = {1: "toc 1", 2: "toc 2", 3: "toc 3"}
    for lvl, title, pg in reversed(front + entries):
        insert_before(doc, anchor, f"{title}\t{pg}", style_map[lvl])
    return len(front) + len(entries)


def replace_image_before_caption(doc: Document, cap: str, path: Path) -> bool:
    if not path.exists():
        return False
    blob = path.read_bytes()
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if cap not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = paras[j].part.related_parts.get(rid) if rid else None
            if part:
                part._blob = blob
                return True
    return False


def refresh_figures(doc: Document) -> int:
    # Ensure print-ready screenshots from user_guide_shots
    from PIL import Image

    def crop_shot(src: Path, dst: Path) -> None:
        if not src.exists():
            return
        im = Image.open(src).convert("RGB")
        w, h = im.size
        th = min(h, int(w * 10 / 16))
        im.crop((0, 0, w, th)).save(dst, quality=95)

    EVAL.mkdir(parents=True, exist_ok=True)
    # Fresh app screenshots — different pages from previous pass (live, home, sim, SOAR tab, models)
    crop_shot(SHOTS / "01_live.png", EVAL / "shot_live_print.png")
    crop_shot(SHOTS / "00_home.png", EVAL / "shot_home_print.png")
    crop_shot(SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png")
    crop_shot(SHOTS / "10_soc_tab1_SOAR_Playbooks.png", EVAL / "shot_soc_print.png")
    crop_shot(SHOTS / "07_models.png", EVAL / "shot_models_print.png")

    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 11.", EVAL / "shot_live_print.png"),
        ("Figure 12.", EVAL / "fig_cm_example.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
        ("Figure H1", SNIPS / "01_hybrid_ai_decision_fusion.png"),
        ("Figure H2", SNIPS / "09_mitre_map.png"),
        ("Figure H3", SNIPS / "10_soar_playbooks.png"),
        ("Figure H4", SNIPS / "12_online_learning.png"),
        ("Figure H5", SNIPS / "14_threat_simulation.png"),
        ("Figure H6", SNIPS / "13_alert_manager.png"),
        ("Figure H7", SNIPS / "16_telegram_alert.png"),
        ("Figure H8", SNIPS / "15_supervisor.png"),
        ("Figure H9", SNIPS / "17_ai_assistant.png"),
    ]
    n = 0
    for cap, img in mapping:
        if replace_image_before_caption(doc, cap, img):
            n += 1
    return n


def table_spacing(doc: Document) -> int:
    body = doc.element.body
    count = 0
    tbl_i = 0
    for child in list(body):
        if not child.tag.endswith("}tbl"):
            continue
        if tbl_i >= 2:
            nxt = child.getnext()
            if nxt is not None and nxt.tag.endswith("}p"):
                pf = Paragraph(nxt, body).paragraph_format
                if pf.space_before is None or pf.space_before.pt < 10:
                    pf.space_before = Pt(14)
                    count += 1
        tbl_i += 1
    return count


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    target = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else (
        TARGET_PRE_POLISH if TARGET_PRE_POLISH.exists() else DESKTOP
    )
    if not target.exists():
        raise SystemExit(f"Target file not found: {target}")

    shutil.copy2(target, STAGING)
    print("Target:", target)

    subprocess.run([sys.executable, str(ROOT / "generate_academic_diagrams.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "render_eval_figures.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "generate_extra_snippets.py")], check=True, cwd=str(ROOT))

    doc = Document(str(STAGING))

    from master_reorder import fix_duplicate_431, reorder_section_42, reorder_section_43

    stats = {
        "blank_removed": remove_blank_pages(doc),
        "heading_intros": add_heading_intros(doc),
        "reorder_42": reorder_section_42(doc),
        "reorder_43": reorder_section_43(doc),
        "dup431": fix_duplicate_431(doc),
        "table_space": table_spacing(doc),
        "toc_entries": rebuild_toc(doc),
        "figures_swapped": refresh_figures(doc),
    }
    doc.save(str(STAGING))
    shutil.copy2(STAGING, target)

    import verify_doctor_revision as vd

    passed, lines = vd.run_checks(target)
    stats["verify"] = f"{passed}/{len(lines)}"
    print("FIX THREE-PASS OK:", stats)
    for ln in lines:
        if ln.startswith("FAIL"):
            print(ln)
    print("Saved:", target)


if __name__ == "__main__":
    main()
