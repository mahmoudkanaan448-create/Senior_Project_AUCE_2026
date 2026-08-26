# -*- coding: utf-8 -*-
"""Build a comprehensive Arabic Word user guide for AI Network Analyzer pages."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SHOTS = Path(r"c:\Users\mohamad\OneDrive\Desktop\AI_NDR_Page_Guide_Screenshots")
OUT = Path(r"c:\Users\mohamad\OneDrive\Desktop\دليل_شرح_كل_صفحات_AI_Network_Analyzer.docx")
IMG_WIDTH = Inches(6.2)


def _set_run_rtl(run) -> None:
    rpr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:bidi"), "ar-SA")
    rpr.append(lang)


def _set_para_rtl(paragraph, align_right: bool = True) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    ppr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _find_shot(*names: str) -> Path | None:
    for name in names:
        for ext in (".jpg", ".png"):
            p = SHOTS / f"{name}{ext}"
            if p.exists():
                return p
    return None


def add_ar(paragraph, text: str, *, bold: bool = False, size: int = 12, color=None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    if color is not None:
        run.font.color.rgb = color
    _set_run_rtl(run)


def para(doc: Document, text: str, *, bold: bool = False, size: int = 12, space_after: int = 8) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_ar(p, text, bold=bold, size=size)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(level=level)
    _set_para_rtl(h)
    h.clear()
    add_ar(h, text, bold=True, size=16 if level == 1 else 14, color=RGBColor(0x0C, 0x1A, 0x2E))


def section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    add_ar(p, text, bold=True, size=13, color=RGBColor(0x1E, 0x3A, 0x5F))


def bullets(doc: Document, items: list[str], *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        _set_para_rtl(p)
        p.clear()
        add_ar(p, item, size=11)
        p.paragraph_format.space_after = Pt(3)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    add_ar(p, text, bold=False, size=10, color=RGBColor(0x47, 0x55, 0x69))


def add_image(doc: Document, path: Path | None, cap: str) -> None:
    if path and path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=IMG_WIDTH)
        caption(doc, cap)
    else:
        para(doc, f"[لم تُعثر على لقطة الشاشة: {cap}]", size=10)


def page_block(
    doc: Document,
    *,
    title: str,
    shot_names: list[str],
    shot_caption: str,
    function_paras: list[str],
    elements: list[str],
    how_to: list[str],
    tabs: list[tuple[str, list[str]]] | None = None,
    extra_shots: list[tuple[list[str], str]] | None = None,
) -> None:
    heading(doc, title, level=1)
    shot = _find_shot(*shot_names)
    add_image(doc, shot, shot_caption)
    if extra_shots:
        for names, cap in extra_shots:
            add_image(doc, _find_shot(*names), cap)

    section_title(doc, "وظيفة الصفحة")
    for t in function_paras:
        para(doc, t)

    section_title(doc, "شرح كل عنصر بالحرف")
    bullets(doc, elements)

    if tabs:
        section_title(doc, "التبويبات الداخلية")
        for tab_name, tab_items in tabs:
            para(doc, tab_name, bold=True, size=12)
            bullets(doc, tab_items)

    section_title(doc, "كيف تستخدمها عملياً")
    bullets(doc, how_to, numbered=True)
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    # RTL section
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)

    # ── Cover ──────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_ar(p, "دليل شرح واجهات نظام", bold=True, size=22, color=RGBColor(0x0C, 0x1A, 0x2E))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_ar(p2, "AI Network Analyzer", bold=True, size=20, color=RGBColor(0x1E, 0x3A, 0x5F))
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_ar(p3, "(NDR Platform v1.0.0)", bold=True, size=14, color=RGBColor(0x47, 0x55, 0x69))
    doc.add_paragraph()
    for line in (
        "مشروع تخرج — الجامعة الأمريكية للثقافة والتعليم AUCE",
        "العام الأكاديمي 2026",
        "",
        "إعداد الطالب: محمود طلال كنعان",
        "Mahmoud Talal Kanaan",
        "",
        "إشراف: د. حسن نور الدين",
        "Dr. Hassan Noureddine",
        "",
        "هذا الدليل يشرح كل صفحة في لوحة التحكم بالتفصيل:",
        "وظيفة الصفحة، معنى كل عنصر ظاهر، والتبويبات الداخلية، وخطوات الاستخدام العملي.",
    ):
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_ar(cp, line, bold=False, size=12)
    doc.add_page_break()

    # ── TOC-like intro ─────────────────────────────────────────────────
    heading(doc, "مقدمة الدليل وترتيب الصفحات", level=1)
    para(
        doc,
        "يهدف هذا الدليل إلى شرح واجهات نظام AI Network Analyzer (منصة كشف الشبكة والتعامل مع الحوادث NDR Platform v1.0.0) "
        "صفحةً بصفحة، بالعربية الواضحة، مع لقطات شاشة حقيقية. الترتيب أدناه هو ترتيب التنقل المقصود في المنصة.",
    )
    bullets(
        doc,
        [
            "Login — تسجيل الدخول",
            "Home — الرئيسية",
            "Live Monitoring — المراقبة المباشرة",
            "AI Detection — كشف الذكاء الاصطناعي",
            "Threat Simulation — محاكاة التهديدات",
            "Threat Intelligence — الاستخبارات التهديدية",
            "Alerts — التنبيهات",
            "Blocked IPs — العناوين المحظورة",
            "AI Models — نماذج الذكاء الاصطناعي",
            "Reports — التقارير",
            "System Settings — إعدادات النظام",
            "SOC Ops — عمليات مركز العمليات الأمنية",
            "Incidents — الحوادث",
            "Assets — الأصول",
            "Hunting — الصيد التهديدي",
            "Response — الاستجابة",
            "Copilot — المساعد الأمني",
        ],
        numbered=True,
    )
    para(
        doc,
        "ملاحظة أسلوبية: أسماء الأزرار والحقول الإنجليزية الظاهرة في الواجهة مكتوبة كما هي "
        "داخل النص العربي لتسهيل المطابقة مع الشاشة.",
    )
    doc.add_page_break()

    # ── Shared chrome ──────────────────────────────────────────────────
    heading(doc, "العناصر المشتركة في كل الصفحات — Shared Chrome", level=1)
    para(
        doc,
        "بعد تسجيل الدخول تظهر في كل صفحة تقريباً نفس «الإطار الخارجي» (Chrome): الشريط الجانبي، "
        "ومؤشرات الحالة، ومساعد الذكاء الاصطناعي العائم. فهم هذه العناصر مرة واحدة يغني عن تكرارها في كل صفحة.",
    )

    section_title(doc, "الشريط الجانبي — Sidebar")
    bullets(
        doc,
        [
            "Brand / اسم النظام: يظهر أعلى الشريط (مثل AI Network Analyzer أو اسم الشركة إن ضُبط). يدل على أنك داخل منصة NDR وليس تطبيقاً آخر.",
            "ONLINE: شارة حالة خضراء تعني أن الواجهة متصلة بالخادم المحلي وأن الجلسة نشطة.",
            "Language / اللغة: قائمة اختيار English أو العربية. تبدّل نصوص الواجهة فوراً وتُحفظ تفضيل اللغة.",
            "Pages / قائمة الصفحات: روابط التنقل الرسمية لكل الوحدات (Home، Live Monitoring، … Copilot). بعض الروابط قد تُخفى حسب دور المستخدم (Admin / Analyst / Viewer).",
            "Time / Date: ساعة وتاريخ الجلسة الحالية لمساعدة المحلل على توثيق وقت الأحداث.",
            "User / Role: يعرض اسم المستخدم المسجّل ودوره (مثل admin / Admin). الدور يحدد صلاحيات الإعدادات والحذف والاستجابة.",
            "Logout: ينهي الجلسة ويعيدك إلى شاشة Login. استخدمه بعد انتهاء العرض أو عند تبديل المستخدم.",
            "Attack sound / صوت الهجوم: خانة أو زر في الشريط لتفعيل تنبيهات صوتية محلية عند وصول تنبيه Medium/High/Critical. زر Enable Attack Alerts يطلب تفاعلاً من المستخدم لفتح الصوت في المتصفح/ويندوز.",
        ],
    )

    section_title(doc, "مساعد الذكاء الاصطناعي العائم — AI Assistant Popover")
    bullets(
        doc,
        [
            "زر 🤖 AI Assistant أعلى يمين منطقة المحتوى يفتح نافذة منبثقة (Popover) متاحة من كل صفحة مصادَق عليها.",
            "Roles caption: يوضح أن المساعد يعمل كمحلل سياق من قاعدة البيانات وليس نموذجاً منفصلاً مستقلاً.",
            "Page: يذكر اسم الصفحة الحالية لربط الإجابة بالسياق.",
            "Ask / حقل السؤال: اكتب سؤالاً بالعربية أو الإنجليزية عن تنبيه أو تدفق أو IP.",
            "Ask button: يرسل السؤال للمساعد ويعرض ملخصاً، نوع التهديد، الخطورة، سبب الكشف، أدلة، MITRE، وتوصية.",
            "Analyze current focus: يحلل العنصر الذي ركّزت عليه الصفحة (مثل Alert #ID أو IP).",
            "Clear chat: يمسح سجل المحادثة داخل الجلسة.",
            "الإجابات قد تعرض Threat و Severity و Why و Evidence و Recommended action لمساعدة التحقيق السريع دون مغادرة الصفحة.",
        ],
    )
    para(
        doc,
        "الفرق بين AI Assistant العائم وصفحة Copilot: المساعد العائم سريع وسياقي داخل أي صفحة؛ "
        "أما Copilot فصفحة مخصصة لاستعلامات أوسع وتلخيص حوادث وتوصيات استجابة وشرح الكشف.",
    )
    doc.add_page_break()

    # ── 1 Login ────────────────────────────────────────────────────────
    page_block(
        doc,
        title="1) تسجيل الدخول — Login",
        shot_names=["00_login"],
        shot_caption="الشكل 1 — شاشة تسجيل الدخول Login",
        function_paras=[
            "صفحة Login هي بوابة الأمان الأولى للمنصة. لا يمكن الوصول إلى أي صفحة SOC قبل مصادقة ناجحة.",
            "تعرض عنوان النظام ووسم NDR Platform مع رقم الإصدار، ثم نموذج إدخال بيانات الدخول.",
            "تدعم اختيار اللغة قبل الدخول حتى تظهر الواجهة بالعربية أو الإنجليزية منذ البداية.",
            "الحساب الافتراضي للتجربة الأكاديمية هو admin / admin123 ما لم يُغيَّر لاحقاً من Settings.",
        ],
        elements=[
            "عنوان AI Network Analyzer: اسم المنتج الظاهر بوضوح على شاشة الدخول.",
            "NDR Platform v1.0.0 · Secure Login: يوضح أن المنصة من نوع Network Detection & Response والإصدار الحالي.",
            "Language: قائمة تبديل اللغة (English / العربية) قبل تسجيل الدخول.",
            "Username: حقل اسم المستخدم. يجب إدخاله كما هو مخزَّن في قاعدة المستخدمين.",
            "Password: حقل كلمة المرور المخفي. لا تظهر الأحرف أثناء الكتابة.",
            "Sign In / دخول: زر إرسال النموذج. عند النجاح تُفتح الجلسة وتنتقل إلى Home.",
            "Default account: admin / admin123: تذكير تعليمي بالحساب الافتراضي للتجربة المحلية.",
            "رسائل الخطأ المحتملة: Invalid username or password، Account is disabled، Login failed، أو طلب إدخال الحقول.",
        ],
        how_to=[
            "افتح المنصة على المنفذ المحلي (عادةً Streamlit).",
            "اختر اللغة إن رغبت.",
            "أدخل Username = admin وكلمة المرور admin123 (أو حسابك المحدَّث).",
            "اضغط Sign In وانتظر ظهور الشريط الجانبي وصفحة Home.",
        ],
    )

    # ── 2 Home ─────────────────────────────────────────────────────────
    page_block(
        doc,
        title="2) الرئيسية — Home",
        shot_names=["01_home"],
        shot_caption="الشكل 2 — لوحة الرئيسية Home",
        function_paras=[
            "Home هي غرفة القيادة العامة لمركز العمليات: ملخص فوري لصحة الشبكة والكشف والاستجابة.",
            "تجمع أرقاماً حيّة من قاعدة البيانات (تدفقات، تنبؤات، تنبيهات، حظر، حوادث، أصول) في صف مقاييس واحد.",
            "تعرض رسوماً لتوزيع الحركة وأنواع الهجمات والخطورة، إضافة إلى رادار أداء نماذج AI.",
            "توفّر Quick Actions للانتقال السريع إلى أهم صفحات العمل، وجداول «الأحدث» للمتابعة اليومية.",
        ],
        elements=[
            "Overview / نظرة عامة: عنوان قسم المقاييس الثمانية.",
            "Total Flows: إجمالي سجلات حركة الشبكة (flows) المخزّنة.",
            "Predictions: عدد قرارات التصنيف التي أصدرتها النماذج.",
            "Attacks: عدد التنبؤات غير الطبيعية (ليست Normal)، مع إشارة لعدد الحركة الطبيعية.",
            "Active Alerts: التنبيهات النشطة/الجديدة التي تحتاج متابعة.",
            "Blocked IPs: عدد عناوين IP المحظورة حالياً أو المسجّلة كمحظورة.",
            "Avg Threat: متوسط درجة التهديد من 0 إلى 10 عبر التنبؤات.",
            "Open Incidents: الحوادث المفتوحة أو قيد المعالجة في وحدة Incidents.",
            "Assets/Sensors: عدد الأصول المسجّلة مقابل عدد الحساسات (مثل 12/1).",
            "Traffic Distribution: رسم دائري Normal مقابل Attack يوضح نسبة الحركة السليمة إلى المشبوهة.",
            "Attack Types: أعمدة لأنواع الهجمات المكتشفة (DoS، PortScan، …).",
            "Severity Distribution: توزيع الخطورة Low / Medium / High / Critical.",
            "AI Model Performance Comparison: مخطط رادار Accuracy / Precision / Recall / F1 لكل نموذج، مع جدول القيم.",
            "Quick Actions: روابط سريعة مثل Threat Simulation، Live Monitoring، Incidents، Assets، Hunting، Copilot، Response، SOC Ops، AI Detection، Alerts.",
            "Recent Alerts table: أعمدة ID، Type، Priority، Status، Message، Time لأحدث التنبيهات.",
            "Recent Predictions table: أحدث قرارات الكشف مع التصنيف والثقة.",
            "Recent Blocked table: أحدث عناوين IP المحظورة وسبب/وقت الحظر.",
        ],
        how_to=[
            "بعد الدخول افتح Home وامسح صف المقاييس بسرعة: هل ارتفعت Attacks أو Active Alerts؟",
            "راجع الرسوم الثلاثة لفهم طبيعة التهديد (نوع الهجوم والخطورة).",
            "إن احتجت عملاً فورياً استخدم Quick Actions بدل البحث في الشريط الجانبي.",
            "افتح جدولاً حديثاً (Alerts أو Predictions) وانتقِ عنصراً للمتابعة في صفحته المختصة.",
        ],
    )

    # ── 3 Live Monitoring ──────────────────────────────────────────────
    page_block(
        doc,
        title="3) المراقبة المباشرة — Live Monitoring",
        shot_names=["02_live_monitoring"],
        shot_caption="الشكل 3 — صفحة Live Monitoring",
        function_paras=[
            "Live Monitoring هي مصدر البيانات الحي للمنصة: تلتقط حزم الشبكة من واجهة الجهاز وتبني منها flows تُخزَّن في قاعدة البيانات.",
            "بدون تشغيل الالتقاط لا تتغذى صفحات الكشف والتقارير على حركة حقيقية جديدة.",
            "تعرض حالة الاتصال CONNECTED·LIVE أو IDLE ومؤشرات الحساسات ومعدل الحزم والتدفقات المخزّنة.",
            "تتيح إنشاء Snapshot لأدلة الجلسة (live-session evidence) لاستخدامها لاحقاً في Reports.",
        ],
        elements=[
            "Network interface: قائمة اختيار واجهة الشبكة (Wi-Fi / Ethernet / غيرها). اختر الواجهة التي تحمل الحركة المراد مراقبتها.",
            "Start Live Capture: يبدأ الالتقاط المستمر ويحوّل الحالة إلى وضع حي.",
            "Stop Live Capture: يوقف الالتقاط ويعيد الحالة إلى IDLE.",
            "Status: CONNECTED·LIVE يعني الالتقاط يعمل؛ IDLE يعني متوقف.",
            "Mode / Last batch / Packets seen: تفاصيل تشغيلية عن وضع الالتقاط وآخر دفعة وعدد الحزم المرصودة.",
            "Sensors: عدد حساسات الالتقاط المسجّلة/النشطة.",
            "pkt/s (packets per second): معدل الحزم في الثانية تقريباً.",
            "Dropped: حزم فُقدت ولم تُعالَج (مؤشر ضغط أو أداء).",
            "Stored Flows: عدد التدفقات المحفوظة المرتبطة بصحة الالتقاط.",
            "Live Active Connections table: اتصالات نشطة لحظياً (مصدر/وجهة/منافذ/بروتوكول).",
            "Recent Stored Flows: أحدث التدفقات التي كُتبت في قاعدة البيانات بعد البناء من الحزم.",
            "Snapshot / live-session evidence: زر يكتب تقرير أدلة الجلسة لاستخدامه في التوثيق والتقارير.",
        ],
        how_to=[
            "اختر Network interface المناسبة (غالباً Wi-Fi أو Ethernet).",
            "اضغط Start Live Capture وانتظر تحول Status إلى CONNECTED·LIVE.",
            "راقب pkt/s وStored Flows وجداول الاتصالات/التدفقات.",
            "عند الحاجة للتوثيق اضغط Snapshot ثم أوقف الالتقاط بـ Stop عند الانتهاء.",
        ],
    )

    # ── 4 AI Detection ─────────────────────────────────────────────────
    page_block(
        doc,
        title="4) كشف الذكاء الاصطناعي — AI Detection",
        shot_names=["03_ai_detection"],
        shot_caption="الشكل 4 — صفحة AI Detection",
        function_paras=[
            "AI Detection تشغّل محرك الكشف الهجين على التدفقات غير المعالجة وتحوّلها إلى تنبؤات وتنبيهات.",
            "تجمع قرارات النماذج عبر طبقة fusion ثم تخزّن النتيجة مع درجة الثقة والتهديد والخطورة وشرح XAI.",
            "تنشئ تنبيهات للمستويات Medium فأعلى وتربط النتائج برسوم توزيع التصنيف والخطورة.",
            "توفر رابطاً إلى Threat Simulation لحقن سيناريوهات مختبر عندما لا تتوفر حركة هجوم حقيقية.",
        ],
        elements=[
            "Total Predictions: إجمالي قرارات الكشف المخزّنة.",
            "Attacks: عدد التصنيفات غير Normal.",
            "Normal: عدد الحركة المصنّفة طبيعية.",
            "Avg Confidence: متوسط ثقة النموذج كنسبة مئوية.",
            "Analyze All Unprocessed Flows: الزر الرئيسي؛ يمر على حتى 500 تدفق بلا تنبؤ ويشغّل التحميل + predict + fuse + XAI + تنبيه عند الحاجة.",
            "Detection Results table: جدول النتائج (معرّف، نموذج، تصنيف، ثقة، درجة تهديد، خطورة، وقت).",
            "XAI — Why this detection: لوحة تفسير توضح لماذا اتُخذ القرار، السمات المؤثرة، والإجراء المقترح.",
            "Prediction chart: توزيع تسميات التنبؤ (أنواع الهجوم مقابل Normal).",
            "Severity chart: توزيع Low/Medium/High/Critical لنتائج الكشف.",
            "رابط Threat Simulation: للانتقال السريع لتوليد حملات مختبر تغذي نفس مسار الكشف.",
        ],
        how_to=[
            "تأكد أن لديك تدفقات من Live Monitoring أو من محاكاة.",
            "اضغط Analyze All Unprocessed Flows وانتظر شريط التقدم.",
            "راجع Detection Results وافتح شرح XAI لأي صف مشبوه.",
            "إن ظهرت تنبيهات Medium+ انتقل إلى Alerts لمتابعة الحالة.",
        ],
    )

    # ── 5 Threat Simulation ────────────────────────────────────────────
    page_block(
        doc,
        title="5) محاكاة التهديدات — Threat Simulation",
        shot_names=["04_threat_simulation"],
        shot_caption="الشكل 5 — صفحة Threat Simulation",
        function_paras=[
            "Threat Simulation وحدة SOC أساسية لحقن حملات هجوم مخبرية عبر نفس مسار الإنتاج: تدفق → كشف → تنبيه → استجابة/تيليغرام.",
            "تسمح بعرض النظام أمام اللجنة دون انتظار هجوم حقيقي على الشبكة.",
            "تستخدم عناوين مختبر آمنة من نطاق 203.0.113.x (TEST-NET) حتى لا تُستهدف أنظمة إنتاج خارجية.",
            "بعد الإطلاق تظهر ملخصات: تدفقات مولَّدة، هجمات مكتشفة، تنبيهات، وإرسال تيليغرام.",
        ],
        elements=[
            "Campaign selector: اختيار نوع الحملة: DoS، DDoS، PortScan، BruteForce، SQLInjection، WebAttack، Botnet، Malware، Mixed.",
            "وصف السيناريو: نص قصير تحت القائمة يشرح ماذا تولّد الحملة المختارة.",
            "Incidents slider: شريط من 1 إلى 20 يحدد عدد الحوادث/العينات المراد توليدها في الجولة.",
            "Create alerts: إن مفعّل يُنشئ سجلات تنبيه في قاعدة البيانات للنتائج المؤهلة.",
            "Block Critical: إن مفعّل يحاول حظر عناوين مصنّفة Critical ضمن سياسة المختبر.",
            "Ensure classification: في وضع المختبر يفرض تسميات عرض واضحة للتجربة؛ يُعطَّل خارج وضع العرض القسري.",
            "Telegram: يرسل إشعارات تيليغرام عند اكتشاف هجمات (يتطلب إعداد التوكن والمحادثة في Settings).",
            "Launch Threat Campaign: يطلق الحملة عبر مسار الكشف الحقيقي ويعرض الملخص والنتائج.",
            "Lab IPs 203.0.113.x: عناوين مصدر/وجهة مخبرية آمنة وفق RFC 5737.",
            "نتائج الحملة: أعداد flows / attacks_detected / alerts_created / telegram_sent وغيرها.",
        ],
        how_to=[
            "اختر Campaign مناسباً للعرض (مثلاً Mixed أو DoS).",
            "اضبط Incidents slider (مثلاً 5).",
            "فعّل Create alerts وTelegram حسب الحاجة، وBlock Critical بحذر في المختبر.",
            "اضغط Launch Threat Campaign ثم راجع Home/Alerts/Blocked IPs للتأكد من ظهور الأثر.",
        ],
    )

    # ── 6 Threat Intelligence ──────────────────────────────────────────
    page_block(
        doc,
        title="6) الاستخبارات التهديدية — Threat Intelligence",
        shot_names=["05_threat_intelligence"],
        shot_caption="الشكل 6 — صفحة Threat Intelligence",
        function_paras=[
            "Threat Intelligence تفحص سمعة عنوان IP وموقعه الجغرافي لدعم قرار الحظر أو التحقيق.",
            "تخزّن نتائج البحث في جدول TI داخلي لمراجعتها لاحقاً دون إعادة الاستعلام في كل مرة.",
            "تعرض رسوم الدول الأكثر ظهوراً كمصدر تهديد محتمل ضمن السجلات المخزّنة.",
            "تكمل صفحة Alerts وBlocked IPs بمعلومة خارجية عن العنوان قبل/بعد الحظر.",
        ],
        elements=[
            "Enter IP Lookup: حقل إدخال عنوان IP (مثال توضيحي مثل 185.220.101.1).",
            "Lookup button: ينفّذ فحص السمعة والجغرافيا ويعرض النتيجة.",
            "Reputation: تقييم السمعة/الخطورة المرتبط بالعنوان.",
            "Geo: الدولة/المدينة/الإحداثيات إن توفرت، وقد تظهر على خريطة.",
            "Stored TI table: سجل عمليات البحث السابقة والنتائج المحفوظة.",
            "Country charts: رسوم توضح أعلى الدول تكراراً في بيانات الاستخبارات المخزّنة.",
        ],
        how_to=[
            "انسخ IP مشبوهاً من Alert أو Detection.",
            "الصقه في Enter IP Lookup واضغط زر البحث.",
            "راجع Reputation وGeo ثم قرر الانتقال إلى Blocked IPs أو Response.",
            "عد لاحقاً لجدول Stored TI لمقارنة العناوين المتكررة.",
        ],
    )

    # ── 7 Alerts ───────────────────────────────────────────────────────
    page_block(
        doc,
        title="7) التنبيهات — Alerts",
        shot_names=["06_alerts"],
        shot_caption="الشكل 7 — صفحة Alerts",
        function_paras=[
            "Alerts هي صندوق وارد المحلل الأمني: كل اكتشاف متوسط الخطورة فأعلى يصل هنا كحالة قابلة للتحقيق.",
            "تدعم تصفية الحالات وتحديث الحالة وربط MITRE ATT&CK وشرح XAI.",
            "تسمح بتغذية راجعة True Attack أو False Positive لتحسين التعلم المستمر لاحقاً.",
            "ترتبط مباشرة بصفحات Incidents وResponse وCopilot لإكمال دورة التعامل مع الحدث.",
        ],
        elements=[
            "Total / New / Investigating / Closed: مقاييس أعداد التنبيهات حسب دورة الحياة.",
            "Filter status: قائمة All / New / Investigating / Closed لتصفية الجدول.",
            "جدول التنبيهات: أعمدة ID، Type، Priority، Status، MITRE، Message، Time.",
            "MITRE: تعيين مختصر للتكتيك/التقنية حسب نوع الهجوم.",
            "XAI panel: أدخل Alert ID لعرض Why this detection المرتبط بالتنبؤ الأصلي.",
            "Update status: اختر التنبيه والحالة الجديدة ثم Update لتغيير New→Investigating→Closed.",
            "True Attack feedback: يعلّم النظام أن التنبيه صحيح (مفيد للتعلم عبر الإنترنت).",
            "False Positive feedback: يعلّم النظام أن التنبيه إنذار كاذب لتقليل التكرار.",
        ],
        how_to=[
            "افتح Alerts وابدأ بفلتر New.",
            "اختر تنبيهاً، راجع MITRE وXAI.",
            "حدّث الحالة إلى Investigating أثناء العمل ثم Closed عند الانتهاء.",
            "سجّل True Attack أو False Positive حسب الحقيقة الميدانية.",
        ],
    )

    # ── 8 Blocked IPs ──────────────────────────────────────────────────
    page_block(
        doc,
        title="8) العناوين المحظورة — Blocked IPs",
        shot_names=["07_blocked_ips"],
        shot_caption="الشكل 8 — صفحة Blocked IPs",
        function_paras=[
            "Blocked IPs سجل مركزي لعناوين IP المحظورة يدوياً أو تلقائياً بعد حملات/سياسات الاستجابة.",
            "تسمح بحظر سريع لعنوان مصدر هجوم مع سبب مكتوب، وبعكس الحظر عند الخطأ.",
            "تعرض أعداد الحظر الكلي والنشط والمُزال لمراقبة حجم سياسة المنع.",
            "وهي تتكامل مع صفحة Response التي تدير مؤشرات IOC والموافقات ومدة الحظر.",
        ],
        elements=[
            "مقاييس Total / Active / Unblocked: إحصاء سجل الحظر.",
            "Block IP: حقل عنوان IP المراد منعه.",
            "Reason: سبب الحظر (مثال DDoS attack) لأغراض التدقيق.",
            "زر الحظر: يضيف العنوان إلى القائمة النشطة ويحدّث الحالة.",
            "Active blocks table: جدول العناوين النشطة مع نوع الهجوم/السبب والوقت والحالة.",
            "Unblock: حقل + زر لإزالة حظر خاطئ أو منتهٍ.",
        ],
        how_to=[
            "انسخ IP من تنبيه Critical/High.",
            "أدخله مع Reason واضح ثم نفّذ Block IP.",
            "تحقق من ظهوره في Active blocks.",
            "إن كان الحظر خطأً استخدم Unblock فوراً.",
        ],
    )

    # ── 9 AI Models ────────────────────────────────────────────────────
    page_block(
        doc,
        title="9) نماذج الذكاء الاصطناعي — AI Models",
        shot_names=["08_ai_models"],
        shot_caption="الشكل 9 — صفحة AI Models",
        function_paras=[
            "AI Models تعرض جودة النماذج الأساسية وتتيح إعادة التدريب من ملف CSV وإدارة مخزن التعلم عبر الإنترنت.",
            "الجدول والرادار يوضحان Accuracy وPrecision وRecall وF1 لكل نموذج مسجّل.",
            "التدريب من CSV يحدّث النماذج الأساسية عند توفر بيانات جديدة.",
            "Online SGD buffer يتعلم تدريجياً من عينات مؤكدة دون استبدال النماذج الأساسية (core models).",
            "عند التدريب/التقييم المقاس على holdout سُجّلت تقريباً: RandomForest 97.8٪، XGBoost 98.0٪، IsolationForest 93.9٪.",
        ],
        elements=[
            "Metrics table: صفوف النماذج وأعمدة Accuracy / Precision / Recall / F1 Score.",
            "Radar chart: مقارنة متعددة المحاور لأداء النماذج معاً.",
            "Train from CSV: رفع ملف بيانات وتدريب النماذج دفعة واحدة.",
            "Start training button: يطلق عملية التدريب ويعرض التقدم/النتيجة.",
            "Online SGD buffer: يعرض حجم العينات المتراكمة للتعلم التدريجي.",
            "Train online: يدرّب مصنّف SGD على المخزن دون overwrite للنماذج الأساسية.",
            "ملاحظة الأداء المقاس: RF ≈ 97.8٪، XGB ≈ 98.0٪، IF ≈ 93.9٪ على مجموعة holdout عند التدريب المقاس.",
        ],
        how_to=[
            "افتح الصفحة وراجع جدول المقاييس والرادار أولاً.",
            "لتحديث النماذج ارفع CSV مناسباً ثم Start training.",
            "للتعلم المستمر راقب حجم Online SGD واضغط Train online عند امتلاء المخزن.",
            "تحقق من نتائج التقييم أيضاً في Reports → Live Evidence / model evaluation.",
        ],
    )

    # ── 10 Reports ─────────────────────────────────────────────────────
    page_block(
        doc,
        title="10) التقارير — Reports",
        shot_names=["09_reports"],
        shot_caption="الشكل 10 — صفحة Reports (التبويب الرئيسي الظاهر)",
        function_paras=[
            "Reports وحدة التوثيق والتصدير: تحوّل نتائج الكشف والحوادث والامتثال وأدلة الجلسة إلى ملفات قابلة للتسليم.",
            "تضم خمسة تبويبات داخلية تغطي الكشف، تقارير الحوادث، التصدير، الامتثال، وأدلة البث الحي مع تقييم النماذج.",
            "مفيدة للدكتور/اللجنة والأرشيف الأمني اليومي دون الدخول يدوياً إلى قاعدة البيانات.",
        ],
        elements=[
            "عنوان الصفحة Reports: مدخل وحدة التقارير.",
            "شريط التبويبات الخمسة: Detection Report، Incident Reports، Export، Compliance، Live Evidence.",
            "أزرار Download CSV / Download JSON داخل التبويبات حسب نوع التقرير.",
            "جداول المعاينة قبل التنزيل لتأكيد المحتوى.",
        ],
        tabs=[
            (
                "Detection Report — تقرير الكشف",
                [
                    "يعرض أحدث التنبؤات (حتى مئات السجلات) بأعمدة ID، Flow، Model، Label، Confidence، Threat Score، Severity، Time.",
                    "زر تنزيل CSV باسم detection_report_YYYYMMDD.csv.",
                    "يفيد في تسليم نتائج الكشف الخام للمحكّمين أو الأرشفة.",
                ],
            ),
            (
                "Incident Reports — تقارير الحوادث",
                [
                    "جدول تقارير الحوادث المكتوبة: ID، Summary، Status، Analyst، Created.",
                    "نموذج Create: حقل Summary ونص Analyst ثم زر الإنشاء.",
                    "يُستخدم لتوثيق تحقيق مكتوب مرتبط بحادثة تشغيلية.",
                ],
            ),
            (
                "Export — التصدير",
                [
                    "اختيار نوع البيانات: Predictions / Alerts / Blocked IPs / Incidents.",
                    "Generate ثم معاينة الجدول.",
                    "تنزيل CSV و/أو JSON.",
                    "زر Daily digest لتوليد ملخص يومي آلي إن وُجد المسار.",
                ],
            ),
            (
                "Compliance — الامتثال",
                [
                    "يعرض لقطة JSON لملخص الامتثال/التدقيق الأمني.",
                    "تنزيل ملف compliance_YYYYMMDD.json.",
                    "مفيد لإظهار جاهزية الضوابط والسياسات أمام المراجعة.",
                ],
            ),
            (
                "Live Evidence — أدلة الجلسة الحية (+ تقييم النماذج)",
                [
                    "زر Snapshot لكتابة live_session_evidence.json.",
                    "مقاييس: flows_total، live_or_lan_flows، simulation_flows، evidence_level.",
                    "عرض JSON كامل وتنزيله.",
                    "قسم Model evaluation holdout: يقرأ eval_metrics.json ويعرض Accuracy/Precision/Recall/F1 لنماذج RandomForest وXGBoost وIsolationForest، مع جدول per-class إن وُجد.",
                ],
            ),
        ],
        how_to=[
            "للتقديم الأكاديمي: افتح Detection Report ونزّل CSV، ثم Live Evidence لأخذ Snapshot.",
            "لتوثيق تحقيق: أنشئ Incident Report بملخص واضح واسم المحلل.",
            "للأرشفة الشاملة: استخدم Export لكل من Predictions وAlerts وBlocked IPs.",
            "راجع Compliance قبل جلسة المراجعة الرسمية.",
        ],
    )

    # ── 11 Settings ────────────────────────────────────────────────────
    page_block(
        doc,
        title="11) إعدادات النظام — System Settings",
        shot_names=["10_settings"],
        shot_caption="الشكل 11 — صفحة Settings",
        function_paras=[
            "System Settings تتيح ضبط المنصة تشغيلياً دون تعديل الكود: التحديث، الصوت، تيليغرام، عتبات AI، سياسات NDR، المسح، وبيانات الشركة.",
            "التغييرات هنا تؤثر على التنبيهات والاستجابة والأمان والصلاحيات الظاهرة للمستخدم.",
            "تتضمن ست تبويبات داخلية يجب فهم كل منها قبل العرض الحي أو التسليم.",
        ],
        elements=[
            "عنوان Settings ومدخل التبويبات الستة.",
            "حقول الحفظ داخل كل تبويب (Save) لتثبيت القيم في قاعدة الإعدادات.",
            "معلومات الإصدار NDR Platform v1.0.0 ضمن معلومات النظام/الشركة.",
        ],
        tabs=[
            (
                "General — عام",
                [
                    "إعدادات عامة مثل سرعة/سلوك التحديث في الواجهة.",
                    "Attack sound: تفعيل أو إيقاف صوت الهجوم المحلي.",
                    "تغيير Username/Password الحالي بعد إدخال كلمة المرور الحالية والتأكيد.",
                    "Save general لحفظ التعديلات.",
                ],
            ),
            (
                "Telegram — تيليغرام",
                [
                    "Bot Token: توكن بوت تيليغرام.",
                    "Chat ID: معرّف المحادثة المستهدفة.",
                    "Save / Test / Detect chat: حفظ الإعداد، إرسال رسالة اختبار، ومحاولة اكتشاف المحادثة تلقائياً.",
                    "ضروري لعمل إشعارات Threat Simulation والتنبيهات الحرجة.",
                ],
            ),
            (
                "AI Config — إعداد الذكاء الاصطناعي",
                [
                    "عتبة الثقة Confidence threshold لقبول/إبراز القرارات.",
                    "عتبة الحظر التلقائي Auto-block threshold.",
                    "Save لحفظ سياسة الكشف/الحظر.",
                ],
            ),
            (
                "NDR / Response — الاستجابة",
                [
                    "وضع الاستجابة: تلقائي أو موافقة بشرية (human approval).",
                    "Retention: مدة الاحتفاظ بالبيانات.",
                    "Scheduled reports cadence: جدولة التقارير.",
                    "Syslog host اختياري لترحيل السجلات.",
                    "Audit log وعناصر purge وفق السياسة.",
                ],
            ),
            (
                "Clear Data — مسح البيانات",
                [
                    "اختيار أنواع البيانات المراد مسحها (تدفقات/تنبؤات/تنبيهات… حسب الواجهة).",
                    "يتطلب تأكيد اسم المستخدم وكلمة المرور قبل الحذف لحماية العرض من المسح العرضي.",
                    "استخدمه فقط لإعادة تهيئة المختبر.",
                ],
            ),
            (
                "Company — الشركة",
                [
                    "بيانات المؤسسة الظاهرة في الواجهة/التقارير.",
                    "إعدادات متعلقة بوضع المختبر/العرض إن وُجدت.",
                    "معلومات النظام والإصدار.",
                ],
            ),
        ],
        how_to=[
            "اضبط Telegram أولاً إن كنت ستعرض الإشعارات.",
            "راجع AI Config وNDR/Response قبل تشغيل حملات Simulation.",
            "فعّل Attack sound من General إن أردت تنبيهاً سمعياً.",
            "لا تستخدم Clear Data إلا بعد أخذ نسخ تقارير إن لزم.",
        ],
    )

    # ── 12 SOC Ops ─────────────────────────────────────────────────────
    page_block(
        doc,
        title="12) عمليات SOC — SOC Ops",
        shot_names=["11_soc_ops"],
        shot_caption="الشكل 12 — صفحة SOC Ops",
        function_paras=[
            "SOC Ops هي غرفة عمليات المنصة: تربط الكشف بإطار MITRE، وأدلة SOAR، والتعلم المستمر، وصحة الخادم، وسجل النماذج.",
            "تُظهر أن المشروع ليس مصنّفاً معزولاً بل منظومة SOC متكاملة.",
            "لكل تبويب وظيفة تشغيلية مستقلة يمكن شرحها حتى بدون لقطة منفصلة لكل تبويب.",
        ],
        elements=[
            "عنوان SOC Ops وشريط التبويبات: MITRE ATT&CK، SOAR Playbooks، Online Learning، Server Health، ML Registry.",
            "محددات Inspect/Select داخل التبويبات لاختيار نوع هجوم أو ملف نموذج.",
            "أزرار Force train / Heal / Benchmark حسب التبويب.",
        ],
        extra_shots=[
            (["11_soc_ops_b"], "الشكل 12ب — لقطة إضافية لـ SOC Ops إن توفرت"),
        ],
        tabs=[
            (
                "MITRE ATT&CK",
                [
                    "يعرض خريطة ربط تسمية الهجوم بتكتيكات وتقنيات ATT&CK.",
                    "Inspect attack label: اختيار نوع هجوم لعرض التفاصيل.",
                    "يفيد المحلل في صياغة التقرير بلغة معيارية صناعية.",
                ],
            ),
            (
                "SOAR Playbooks",
                [
                    "قائمة أدلة الأتمتة: ماذا يحدث بعد تنبيه معيّن (تنبيه، تيليغرام، حظر، تصعيد…).",
                    "توضح سلسلة الاستجابة الآلية/شبه الآلية للمنصة.",
                ],
            ),
            (
                "Online Learning",
                [
                    "يعرض حالة مخزن التعلم التدريجي ومقاييسه.",
                    "Force train: تشغيل تدريب SGD يدوياً على العينات المتراكمة.",
                    "لا يستبدل النماذج الأساسية؛ يعمل كطبقة تكيفية مساعدة.",
                ],
            ),
            (
                "Server Health",
                [
                    "Overall status: جاهزية الخدمات.",
                    "Heal: محاولة إصلاح/استعادة تلقائية.",
                    "Capture health ومؤشرات مرتبطة بالالتقاط.",
                    "Benchmark: اختبار حمل بسيط لقياس الاستجابة.",
                ],
            ),
            (
                "ML Registry",
                [
                    "سجل ملفات النماذج/المجموعات المحلية.",
                    "Register file: تسجيل نموذج في السجل.",
                    "Drift: مؤشرات انحراف المضيف/البيانات إن توفرت.",
                ],
            ),
        ],
        how_to=[
            "ابدأ بـ MITRE لربط نوع هجوم ظاهر في Alerts بإطار ATT&CK.",
            "راجع SOAR Playbooks لشرح ماذا يحدث تلقائياً بعد التنبيه.",
            "من Online Learning راقب المخزن ونفّذ Force train عند الحاجة.",
            "قبل العرض افتح Server Health للتأكد من Overall status السليم.",
        ],
    )

    # ── 13 Incidents ───────────────────────────────────────────────────
    page_block(
        doc,
        title="13) الحوادث — Incidents",
        shot_names=["12_incidents"],
        shot_caption="الشكل 13 — صفحة Incidents",
        function_paras=[
            "Incidents تجمع أحداثاً مترابطة في «حالة» واحدة (correlated cases) بدل التعامل مع كل تنبيه بمعزل.",
            "تتيح تعيين الحالة والمالك والملاحظات وتوليد ملخص AI للحادثة.",
            "تربط بين الكشف والاستجابة الإدارية داخل SOC.",
            "وهي تتكامل مع Alerts: التنبيه إشارة، والحادثة ملف تحقيق.",
        ],
        elements=[
            "Filter: All / Open / In Progress / Resolved لتصفية الحالات.",
            "مقاييس أعداد الحالات حسب الحالة.",
            "جدول الحوادث المترابطة: معرّف، عنوان، خطورة، حالة، سلسلة هجوم، وقت…",
            "Update section: اختيار حالة جديدة (Open / In Progress / Resolved).",
            "Owner: اسم المالك/المحلل المسؤول.",
            "Notes/Save: حفظ التحديثات على الحادثة.",
            "AI summary: اختيار Incident ID وعرض ملخص ذكي للتحقيق.",
        ],
        how_to=[
            "صفِّ Open أولاً لتحديد ما يحتاج عملاً.",
            "افتح حادثة، عيّن Owner، واضبط الحالة إلى In Progress.",
            "أضف ملاحظات التحقيق ثم Save.",
            "استخدم AI summary لصياغة ملخص سريع قبل الإغلاق Resolved.",
        ],
    )

    # ── 14 Assets ──────────────────────────────────────────────────────
    page_block(
        doc,
        title="14) الأصول — Assets",
        shot_names=["13_assets"],
        shot_caption="الشكل 14 — صفحة Assets",
        function_paras=[
            "Assets تدير جرد الأجهزة/المضيفين وعلاقاتهم وملفاتهم السلوكية ضمن شبكة المراقبة.",
            "تساعد على معرفة «من الجهاز المتأثر؟» وليس فقط «ما الهجوم؟».",
            "تشمل تبويبات Inventory وTopology وHost profile وIdentity/UEBA.",
        ],
        elements=[
            "مقاييس عامة لعدد الأصول/التصنيفات الظاهرة أعلى الصفحة.",
            "تبويبات داخلية أربعة تغطي الجرد والطوبولوجيا وملف المضيف والهوية السلوكية.",
            "حقول IP لتعليم أصل كحرج أو عادي، ولاستعلام ملف مضيف.",
        ],
        tabs=[
            (
                "Inventory — الجرد",
                [
                    "قائمة الأصول المعروفة (IP/اسم/أهمية).",
                    "Tag IP ثم Mark critical أو Mark normal لتصنيف الأهمية.",
                    "يساعد على إعطاء أولوية أعلى للأصول الحرجة عند التنبيه.",
                ],
            ),
            (
                "Topology — الطوبولوجيا",
                [
                    "عرض علاقات/هيكل الشبكة بين الأصول والحساسات قدر الإمكان.",
                    "يفيد في شرح مسار الحركة الجانبية أو نقاط التجميع.",
                ],
            ),
            (
                "Host profile — ملف المضيف",
                [
                    "إدخال Host IP لعرض نشاط المضيف.",
                    "Recent alerts وRecent flows المرتبطة بهذا المضيف.",
                    "ملف سريع للتحقيق المركّز على جهاز واحد.",
                ],
            ),
            (
                "Identity / UEBA",
                [
                    "قسم المصادقة/الهوية Auth إن توفرت أحداث دخول.",
                    "Drift: انحراف سلوكي للمضيف عن نمطه المعتاد.",
                    "يدعم كشف سلوك غير طبيعي على مستوى المستخدم/المضيف.",
                ],
            ),
        ],
        how_to=[
            "من تنبيه انسخ IP الداخلية وافتح Assets → Host profile.",
            "صنّف الأصول الحرجة من Inventory قبل العرض.",
            "راجع Topology عند شرح البنية للجنة.",
            "راقب Drift في UEBA إذا ظهرت سلوكيات شاذة متكررة.",
        ],
    )

    # ── 15 Hunting ─────────────────────────────────────────────────────
    page_block(
        doc,
        title="15) الصيد التهديدي — Hunting",
        shot_names=["14_hunting"],
        shot_caption="الشكل 15 — صفحة Hunting",
        function_paras=[
            "Hunting تمكّن المحلل من البحث الاستباقي في البيانات بدل انتظار التنبيه فقط.",
            "تدعم صيداً بالعنوان/النوع، وتحليلاً جنائياً لملفات PCAP، وفحص بروتوكولات DNS/TLS.",
            "وهي تتكامل مع Detection: الكشف آلي؛ والصيد بشري موجّه بفرضية.",
        ],
        elements=[
            "تبويبات Hunt / Forensics PCAP / Protocol·DNS·TLS.",
            "حقول إدخال IP أو نوع هجوم أو نطاق أو SNI أو hash حسب التبويب.",
            "زر Run لإطلاق الاستعلام وعرض مقاييس flows/alerts المطابقة.",
        ],
        tabs=[
            (
                "Hunt — الصيد",
                [
                    "أدخل IP و/أو نوع هجوم ثم Run.",
                    "يعرض النتائج المرتبطة من التدفقات والتنبيهات.",
                    "مفيد عندما تشك في نمط معيّن غير ظاهر بعد في صندوق Alerts.",
                ],
            ),
            (
                "Forensics PCAP",
                [
                    "أدوات مرتبطة بتحليل ملفات الالتقاط PCAP/الخدمات الجنائية المتاحة في المنصة.",
                    "تُستخدم لتعميق الدليل على مستوى الحزم وليس فقط الـ flow الملخص.",
                ],
            ),
            (
                "Protocol / DNS / TLS",
                [
                    "Domain: بحث استعلامات DNS عن نطاق مشبوه.",
                    "SNI / Issuer: فحص إشارات TLS (اسم الخادم الظاهر وجهة الإصدار).",
                    "Hash: البحث ببصمة إن توفرت.",
                    "يساعد على كشف C2 أو تصفح مشبوه أو شهادات غير معتادة.",
                ],
            ),
        ],
        how_to=[
            "شكّل فرضية (مثلاً اتصالات DNS غريبة).",
            "استخدم تبويب Protocol/DNS/TLS بالبحث عن النطاق أو SNI.",
            "أو ابحث في Hunt بعنوان المصدر من التنبيه.",
            "إن توفر PCAP انتقل لتبويب Forensics لتعميق الدليل.",
        ],
    )

    # ── 16 Response ────────────────────────────────────────────────────
    page_block(
        doc,
        title="16) الاستجابة — Response",
        shot_names=["15_response"],
        shot_caption="الشكل 16 — صفحة Response",
        function_paras=[
            "Response هي غرفة تنفيذ الإجراءات الدفاعية: مؤشرات IOC، السماح/الحظر، الموافقات، Webhooks، وتسجيل الحساسات.",
            "تربط قرار المحلل بفعل قابل للتدقيق مع دعم وضع الموافقة البشرية.",
            "وهي أوسع من Blocked IPs: تشمل مدة الحظر وقائمة السماح والتكامل الخارجي.",
        ],
        elements=[
            "تبويبات: IOC، Allow/Block، Approvals، Webhooks، Sensors.",
            "نماذج إدخال قيمة المؤشر، السبب، المدة، اسم Webhook، عنوان URL، اسم الحساس…",
            "أزرار Add / Queue / Approve / Reject / Register حسب السياق.",
        ],
        tabs=[
            (
                "IOC — مؤشرات الاقتحام",
                [
                    "إضافة مؤشر (قيمة + وصف) إلى قائمة IOC النشطة.",
                    "إمكانية تعطيل مؤشر لم يعد صالحاً.",
                    "تغذي سياسات المطابقة والحظر اللاحق.",
                ],
            ),
            (
                "Allow / Block",
                [
                    "Allowlist: إضافة قيمة موثوقة حتى لا تُحظر خطأً.",
                    "Manual block: إدخال IP ومدة (5m / 1h / 24h / permanent) وسبب ثم Queue.",
                    "Unblock/Rollback لعنوان سبق حظره.",
                ],
            ),
            (
                "Approvals — الموافقات",
                [
                    "قائمة إجراءات بانتظار موافقة بشرية عندما تكون سياسة NDR على وضع approval.",
                    "Approve لتنفيذ الإجراء أو Reject لرفضه.",
                    "ضروري في البيئات التي تمنع الحظر التلقائي الكامل.",
                ],
            ),
            (
                "Webhooks",
                [
                    "إضافة Webhook باسم وURL ونوع حدث (alert / incident / all).",
                    "ترسل إشعارات خارجية لأنظمة أخرى عند وقوع الأحداث.",
                ],
            ),
            (
                "Sensors — الحساسات",
                [
                    "تسجيل حساس جديد: الاسم، الموقع/الفرع، الواجهات (مثل eth0).",
                    "يربط نقاط الالتقاط الموزعة بلوحة المراقبة المركزية.",
                ],
            ),
        ],
        how_to=[
            "أضف IOC للعناوين/القيم المؤكدة خبيثة.",
            "للسماح باستثناء موثوق استخدم Allowlist قبل الحظر الواسع.",
            "إن ظهرت Approvals راجعها وApprove/Reject بوعي.",
            "اضبط Webhook إن رغبت بربط المنصة بأداة خارجية، وسجّل Sensors للفروع.",
        ],
    )

    # ── 17 Copilot ─────────────────────────────────────────────────────
    page_block(
        doc,
        title="17) المساعد الأمني — Copilot",
        shot_names=["16_copilot"],
        shot_caption="الشكل 17 — صفحة Copilot",
        function_paras=[
            "Copilot صفحة المساعد الأمني المخصصة لاستعلامات بلغة طبيعية وتلخيص الحوادث وتوصية الاستجابة وشرح سبب الكشف.",
            "تعتمد على سياق قاعدة البيانات ووظائف SOC المساعدة وليست بديلاً عن النماذج الأساسية في الكشف.",
            "وهي تتكامل مع المساعد العائم، مع أدوات أوسع وأكثر تنظيماً للتحقيق.",
            "مفيدة جداً أثناء العرض لشرح «لماذا اعتُبر هذا هجوماً؟» و«ماذا نفعل الآن؟».",
        ],
        elements=[
            "Ask the SOC: حقل سؤال حر (placeholder يوضح أمثلة الصيد/الاستعلام).",
            "Go / تشغيل السؤال: يعيد جواباً نصياً وقد يعرض جدول صفوف مطابقة.",
            "Incident summary: اختيار Incident ID وعرض ملخص نصي للحادثة.",
            "Recommend response: حقول Attack type وSeverity وSource IP ثم زر Recommend لعرض توصية JSON/منظمة.",
            "Why was this detected: إما Alert ID لشرح تنبيه مخزّن، أو Label + Evidence لشرح سيناريو عام.",
            "Explain button: يطبع تفسيراً مقروءاً يربط السلوك بالقرار الأمني.",
        ],
        how_to=[
            "اسأل سؤالاً واضحاً في Ask the SOC مثل البحث عن تنبيهات Critical الأخيرة.",
            "اختر حادثة من Incident summary لقراءة الملخص.",
            "أدخل نوع هجوم وخطورة ومصدر ثم Recommend response.",
            "لشرح كشف معيّن أدخل Alert ID أو Label+Evidence واضغط Explain.",
        ],
    )

    # Closing
    heading(doc, "خاتمة الدليل", level=1)
    para(
        doc,
        "بهذا يكتمل شرح صفحات AI Network Analyzer (NDR Platform v1.0.0) من Login حتى Copilot، "
        "مع العناصر المشتركة (Sidebar وAI Assistant) ومعنى الضوابط الظاهرة في كل واجهة. "
        "يُنصح عند العرض الحي باتباع مسار: Login → Live Monitoring أو Threat Simulation → AI Detection → Alerts → "
        "Blocked IPs/Response → Reports، مع الاستعانة بـ SOC Ops وCopilot للتوضيح المعياري واللغوي للجنة.",
    )
    para(
        doc,
        "الطالب: محمود طلال كنعان — AUCE 2026 — إشراف د. حسن نور الدين.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    size = path.stat().st_size
    import sys
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(f"OUTPUT={path}")
    print(f"SIZE_BYTES={size}")
    print(f"SIZE_KB={size/1024:.1f}")
