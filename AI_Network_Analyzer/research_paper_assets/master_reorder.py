"""
Master reorder pass: TOC, chapter spacing, table spacing, empty headings, section 4 order.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
WORK = ROOT / "_master_work.docx"
if not WORK.exists():
    WORK = ROOT / "_REBUILD_STAGING.docx"
if not WORK.exists():
    WORK = ROOT / "_POLISH_STAGING.docx"
OUT = ROOT / "FINAL_MASTER.docx"
DESKTOP = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

TEXT_331 = (
    "The implemented system modules match the software repository. Monitoring captures live connections or "
    "CSV flows and extracts 39 numerical features. The Hybrid AI engine loads Random Forest, XGBoost, Isolation "
    "Forest, Autoencoder, and LSTM models. The Decision Engine fuses votes into label, confidence, threat score, "
    "and severity. Threat Intelligence enriches source IPs. Explainable AI returns feature importance and recommended "
    "actions. The Alert Manager creates Medium+ alerts and executes SOAR playbooks. Telegram notifications inform "
    "the analyst. Threat Simulation injects controlled lab campaigns. Online learning updates an SGDClassifier with "
    "partial_fit without overwriting the five core models."
)

TEXT_332 = (
    "Implemented and tested: capture, 39-feature extraction, RF/XGB/IF holdout evaluation, dashboard, alerts, "
    "Telegram, Threat Simulation, stored XAI, MITRE mapping, SOAR playbooks. Implemented but not fully evaluated "
    "on the Table 11 holdout: Autoencoder, LSTM sequence model, Hybrid fusion (live only). Proposed/future: "
    "SHAP plots, full CICIDS2017 public import, SPAN/TAP fleet, email/SMTP channel."
)

TEXT_421 = (
    "This subsection documents the complete input schema used in training and inference. Every column in Table 5 "
    "is defined in config.FEATURE_COLUMNS and extracted identically during live capture, CSV import, and holdout "
    "evaluation using CICIDS-style NDR domain knowledge."
)

TEXT_431 = (
    "Random Forest — supervised multi-class on 39 scaled features (n_estimators=200, max_depth=20). XGBoost — "
    "supervised multi-class (n_estimators=200, max_depth=8, learning_rate=0.1). Isolation Forest — trained on Normal "
    "traffic only (contamination=0.1). Autoencoder — Normal-only reconstruction threshold; not on Table 11 holdout. "
    "Hybrid AI — fuse_decisions() weighted majority fusion; no standalone holdout accuracy."
)

TEXT_432 = (
    "LSTM uses training/data_preprocessing.prepare_sequences() with window_size=10. Each sample is a (10, 39) tensor; "
    "the label is the class of the flow after the window. Row-based models use one flow per sample; LSTM uses a sliding "
    "window. LSTM is implemented but not scored in Table 11 because evaluate_models.py is row-based."
)

TEXT_433 = (
    "Table 5A below shows how each model consumes the same CSV: RF/XGB/IF/AE use one 39-feature row; LSTM uses "
    "ten consecutive rows; Hybrid combines live votes without a separate training file."
)

TEXT_43 = (
    "Model development: load dataset.csv, apply train-only preprocessing (LabelEncoder, StandardScaler), train each "
    "model, persist artifacts under models/ for FastAPI/Streamlit inference."
)


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = anchor.part.document.styles[style]
    except Exception:
        pass
    set_text(para, text)
    pf = para.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    return para


def insert_before(doc: Document, anchor: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass
    set_text(para, text)
    return para


def has_page_break(p) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def heading_has_body(p: Paragraph, min_len: int = 30) -> bool:
    el = p._p
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag.endswith("}tbl"):
            return True
        if nxt.tag.endswith("}p"):
            para = Paragraph(nxt, p._parent)
            t = (para.text or "").strip()
            st = para.style.name if para.style else ""
            if st.startswith("Heading"):
                return False
            if t and len(t) >= min_len and not re.match(r"^Table \d+\.", t):
                return True
        nxt = nxt.getnext()
    return False


def move_after(element, anchor_el) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    anchor_el.addnext(element)


def ensure_under(heading: str, text: str, doc: Document) -> bool:
    p = find_para(doc, heading)
    if not p or heading_has_body(p):
        return False
    insert_after(p, text)
    return True


def fix_duplicate_431(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "4.3.1 Implemented Preprocessing and Training Configuration":
            set_text(p, "4.3.4 Implemented Preprocessing and Training Configuration")
            n += 1
    return n


def remove_orphan_dupes(doc: Document) -> int:
    """Remove duplicate LSTM/RF/model-dev paragraphs left after reorder."""
    seen: set[str] = set()
    removed = 0
    keys = (
        "LSTM requires sequential input, not a single flow row",
        "Random Forest — supervised multi-class on all 39 features",
        "Random Forest - supervised multi-class on all 39 features",
        "Model development begins with data loading",
    )
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        for k in keys:
            if t.startswith(k):
                if k in seen:
                    delete_paragraph(p)
                    removed += 1
                else:
                    seen.add(k)
                break
    return removed


def reorder_section_42(doc: Document) -> bool:
    h421 = find_para(doc, "4.2.1 Complete Feature Set Used in Experiments")
    if not h421:
        return False
    cap_el = desc_el = tbl_el = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "Table 5. Extracted network features (39 holdout features).":
            cap_el = p._p
        if t.startswith("Table 5 lists all 39 features actually passed"):
            desc_el = p._p
    if cap_el is None:
        return False
    el = (desc_el or cap_el).getnext()
    while el is not None:
        if el.tag.endswith("}tbl"):
            tbl_el = el
            break
        if el.tag.endswith("}p"):
            pt = Paragraph(el, doc.element.body).text or ""
            if pt.strip().startswith("4.3") or pt.strip().startswith("Chapter"):
                break
        el = el.getnext()
    if tbl_el is None:
        return False
    anchor = h421._p
    nxt = anchor.getnext()
    if nxt is not None and nxt.tag.endswith("}p"):
        t = Paragraph(nxt, doc.element.body).text or ""
        if "This subsection documents" in t:
            anchor = nxt
    for block in (cap_el, desc_el, tbl_el):
        if block is not None:
            move_after(block, anchor)
            anchor = block
    return True


def reorder_section_43(doc: Document) -> int:
    moved = 0
    h431 = find_para(doc, "4.3.1 Per-Model Training Configuration")
    h432 = find_para(doc, "4.3.2 LSTM Sequence Input Construction")
    h433 = find_para(doc, "4.3.3 Per-Model Dataset Usage Summary")
    if not all([h431, h432, h433]):
        return 0

    rf_p = lstm_p = cap5a = desc5a = interp = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Random Forest") and "n_estimators=200" in t:
            rf_p = p
        elif t.startswith("LSTM requires sequential input"):
            lstm_p = p
        elif t == "Table 5A. Per-model dataset usage summary.":
            cap5a = p
        elif t.startswith("Table 5A summarises how each model"):
            desc5a = p
        elif t.startswith("Interpretation. Supervised RF and XGB share"):
            interp = p

    body = doc.element.body
    if rf_p and not heading_has_body(h431):
        move_after(rf_p._p, h431._p)
        moved += 1
    if lstm_p and not heading_has_body(h432):
        move_after(lstm_p._p, h432._p)
        moved += 1
    if cap5a:
        anchor = h433._p
        for block in (cap5a._p, desc5a._p if desc5a else None):
            if block is not None:
                move_after(block, anchor)
                anchor = block
                moved += 1
        el = anchor.getnext()
        while el is not None:
            if el.tag.endswith("}tbl"):
                move_after(el, anchor)
                anchor = el
                moved += 1
                break
            if el.tag.endswith("}p"):
                pt = (Paragraph(el, body).text or "").strip()
                if pt.startswith("4.3.4") or pt.startswith("4.4"):
                    break
            el = el.getnext()
        if interp:
            move_after(interp._p, anchor)
            moved += 1
    return moved


def add_table_spacing(doc: Document) -> int:
    body = doc.element.body
    count = 0
    tbl_i = 0
    for child in list(body):
        if not child.tag.endswith("}tbl"):
            continue
        if tbl_i >= 2:
            nxt = child.getnext()
            if nxt is not None and nxt.tag.endswith("}p"):
                para = Paragraph(nxt, body)
                pf = para.paragraph_format
                if pf.space_before is None or pf.space_before.pt < 10:
                    pf.space_before = Pt(14)
                    count += 1
            else:
                spacer = OxmlElement("w:p")
                child.addnext(spacer)
                sp = Paragraph(spacer, body)
                sp.paragraph_format.space_after = Pt(14)
                count += 1
            # space after table via tblPr margin if possible - use preceding para space_after
            prev = child.getprevious()
            if prev is not None and prev.tag.endswith("}p"):
                pf = Paragraph(prev, body).paragraph_format
                pf.space_after = Pt(8)
        tbl_i += 1
    return count


def format_captions(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if re.match(r"^Table \d+[A-Z]?\.", t):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(8)
            for r in p.runs:
                r.bold = True
                r.italic = False
                r.font.size = Pt(11)
            n += 1
    return n


def chapter_spacing(doc: Document) -> int:
    changed = 0
    titles = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if p.style and p.style.name == "Heading 1" and (
            t.startswith("Chapter ") or t in ("Conclusion", "References") or t.startswith("Appendix A")
        ):
            titles.append(t)
    for title in titles:
        if title.startswith("Chapter 1"):
            continue
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() != title or not (p.style and p.style.name == "Heading 1"):
                continue
            pf = p.paragraph_format
            pf.space_before = Pt(18)
            pf.space_after = Pt(10)
            if i > 0:
                prev = doc.paragraphs[i - 1]
                if not has_page_break(prev):
                    run = prev.runs[0] if prev.runs else prev.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    changed += 1
            break
    return changed


def remove_empty_page_breaks(doc: Document) -> int:
    removed = 0
    NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        has_img = bool(p._p.findall(f".//{NS_A}blip"))
        if not t and not has_img and has_page_break(p) and st == "Normal":
            delete_paragraph(p)
            removed += 1
    return removed


def rebuild_toc(doc: Document) -> int:
    page = 1
    entries: list[tuple[int, str, int]] = []
    started = False
    for p in doc.paragraphs:
        if has_page_break(p):
            page += 1
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if t.startswith("Chapter 1 - Introduction") and st == "Heading 1":
            started = True
        if not started or not st.startswith("Heading"):
            continue
        if t in ("List of Figures", "List of Tables"):
            continue
        lvl = 1 if t.startswith("Chapter ") or t in ("Conclusion", "References") or t.startswith("Appendix") else (
            2 if st == "Heading 2" else 3
        )
        if st == "Heading 1" and not (t.startswith("Chapter") or t in ("Conclusion", "References") or t.startswith("Appendix")):
            continue
        entries.append((lvl, t, page))

    toc_start = toc_end = None
    for i, p in enumerate(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        t = (p.text or "").strip()
        if toc_start is None and st == "toc 1" and "Declaration" in t:
            toc_start = i
        if toc_start is not None and st == "heading 1" and t == "List of Figures":
            toc_end = i
            break
    if toc_start is None or toc_end is None:
        return 0

    anchor = doc.paragraphs[toc_end]
    n_del = toc_end - toc_start
    for _ in range(n_del):
        delete_paragraph(doc.paragraphs[toc_start])
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "List of Figures" and p.style and p.style.name == "Heading 1":
            anchor = p
            break

    front = [
        (1, "Declaration", 2),
        (1, "Acknowledgment", 2),
        (1, "Abstract", 2),
        (1, "List of Figures", 3),
        (1, "List of Tables", 3),
    ]
    for lvl, title, pg in reversed(front + entries):
        insert_before(doc, anchor, f"{title}\t{pg}", "toc 1" if lvl == 1 else ("toc 2" if lvl == 2 else "toc 3"))
    return len(front) + len(entries)


def rebuild_lot(doc: Document) -> None:
    from ultimate_quality_fix import rebuild_lot as _rebuild

    _rebuild(doc)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    src = WORK if WORK.exists() else DESKTOP
    doc = Document(str(src))

    stats = {
        "331": ensure_under("3.3.1 Implemented Core Modules (Aligned with Software)", TEXT_331, doc),
        "332": ensure_under("3.3.2 Implementation and Evaluation Status", TEXT_332, doc),
        "421": ensure_under("4.2.1 Complete Feature Set Used in Experiments", TEXT_421, doc),
        "43": ensure_under("4.3 Model Development", TEXT_43, doc),
        "431": ensure_under("4.3.1 Per-Model Training Configuration", TEXT_431, doc),
        "432": ensure_under("4.3.2 LSTM Sequence Input Construction", TEXT_432, doc),
        "433": ensure_under("4.3.3 Per-Model Dataset Usage Summary", TEXT_433, doc),
        "dup431": fix_duplicate_431(doc),
        "re42": reorder_section_42(doc),
        "re43": reorder_section_43(doc),
        "dupes": remove_orphan_dupes(doc),
        "tbl_space": add_table_spacing(doc),
        "captions": format_captions(doc),
        "chapters": chapter_spacing(doc),
        "empty_pb": remove_empty_page_breaks(doc),
    }
    rebuild_lot(doc)
    stats["toc"] = rebuild_toc(doc)
    doc.save(str(OUT))

    # verify
    import verify_doctor_revision as vd

    passed, lines = vd.run_checks(OUT)
    stats["verify"] = f"{passed}/{len(lines)}"

    for dest in (DESKTOP, ROOT / "FINAL_FIXED.docx"):
        try:
            shutil.copy2(OUT, dest)
            print("Copied ->", dest)
            break
        except Exception as exc:
            print("Copy failed", dest, exc)

    print("MASTER REORDER OK:", OUT)
    print(stats)
    for ln in lines:
        if ln.startswith("FAIL"):
            print(ln)


if __name__ == "__main__":
    main()
