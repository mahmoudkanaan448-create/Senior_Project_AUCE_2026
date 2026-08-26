"""Tidy the Desktop thesis in place: spacing, readable figures, TOC. No new docx."""
from __future__ import annotations

import json
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
ROOT = Path(__file__).resolve().parent
EVAL = ROOT / "evaluation"
SHOTS = ROOT / "user_guide_shots"
METRICS = EVAL / "metrics.json"
EMU = 914400
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def font(size: int, bold: bool = False):
    names = ("calibrib.ttf", "arialbd.ttf") if bold else ("calibri.ttf", "arial.ttf", "segoeui.ttf")
    for n in names:
        try:
            return ImageFont.truetype(n, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _text(d, xy, text, fill, size, bold=False):
    d.text(xy, text, fill=fill, font=font(size, bold))


def save_dataset_light(classes: dict, path: Path) -> None:
    w, h = 1400, 720
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 70), fill=(12, 26, 46))
    _text(d, (36, 18), "Figure 15  ·  Class distribution  ·  12,000 CICIDS-style samples, 39 features", (240, 193, 75), 26, True)
    order = ["Normal", "DDoS", "DoS", "PortScan", "BruteForce", "SQLInjection"]
    items = [(k, int(classes[k])) for k in order if k in classes]
    mx = max(v for _, v in items)
    ox, oy, maxh, bw, gap = 90, 620, 430, 150, 50
    d.line((ox, oy, w - 50, oy), fill=(30, 41, 59), width=2)
    for i, (lab, v) in enumerate(items):
        x = ox + 30 + i * (bw + gap)
        hh = int(v / mx * maxh)
        d.rectangle((x, oy - hh, x + bw, oy), fill=(29, 78, 216), outline=(12, 26, 46), width=2)
        _text(d, (x + 28, oy - hh - 36), f"{v:,}", (15, 23, 42), 22, True)
        _text(d, (x + 8, oy + 12), lab, (15, 23, 42), 20, True)
    _text(d, (90, 90), "Imbalanced on purpose (Normal ~49%). Not the official CICIDS2017 dump.", (71, 85, 105), 20)
    img.save(path, quality=95)


def save_bars_light(rf, xgb, iso, path: Path) -> None:
    w, h = 1400, 720
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 70), fill=(12, 26, 46))
    _text(d, (36, 18), "Figure 14 / 16  ·  Holdout metrics  ·  75/25 split, random_state=42", (240, 193, 75), 26, True)
    models = [
        ("Random Forest", rf["accuracy"] * 100, rf["f1"] * 100),
        ("XGBoost", xgb["accuracy"] * 100, xgb["f1"] * 100),
        ("Isolation Forest\n(Normal vs Attack)", iso["accuracy"] * 100, iso["f1"] * 100),
    ]
    ox, oy, maxh = 110, 620, 430
    d.line((ox, oy, w - 40, oy), fill=(30, 41, 59), width=2)
    d.line((ox, oy, ox, oy - maxh - 10), fill=(30, 41, 59), width=2)
    for pct in (90, 95, 100):
        y = oy - int(pct / 100 * maxh)
        d.line((ox, y, w - 40, y), fill=(226, 232, 240), width=1)
        _text(d, (40, y - 10), f"{pct}%", (100, 116, 139), 16)
    bw = 90
    for i, (name, acc, f1) in enumerate(models):
        x = 200 + i * 380
        ha = int(acc / 100 * maxh)
        hf = int(f1 / 100 * maxh)
        d.rectangle((x, oy - ha, x + bw, oy), fill=(37, 99, 235))
        d.rectangle((x + bw + 16, oy - hf, x + 2 * bw + 16, oy), fill=(202, 138, 4))
        _text(d, (x + 8, oy - ha - 34), f"{acc:.1f}%", (30, 64, 175), 20, True)
        _text(d, (x + bw + 20, oy - hf - 34), f"{f1:.1f}%", (146, 64, 14), 20, True)
        lines = name.split("\n")
        for li, line in enumerate(lines):
            _text(d, (x, oy + 12 + li * 24), line, (15, 23, 42), 18, True)
    d.rectangle((200, 88, 230, 112), fill=(37, 99, 235))
    _text(d, (240, 88), "Accuracy", (15, 23, 42), 18)
    d.rectangle((380, 88, 410, 112), fill=(202, 138, 4))
    _text(d, (420, 88), "Weighted F1", (15, 23, 42), 18)
    img.save(path, quality=95)


def save_confusion_light(cm: np.ndarray, labels: list[str], path: Path, title: str) -> None:
    n = len(labels)
    cell = 92
    left, top = 160, 120
    w = left + n * cell + 40
    h = top + n * cell + 70
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 56), fill=(12, 26, 46))
    _text(d, (20, 14), title, (240, 193, 75), 22, True)
    mx = int(cm.max()) or 1
    _text(d, (left, 70), "Predicted →", (71, 85, 105), 16)
    _text(d, (18, top - 8), "True ↓", (71, 85, 105), 16)
    for i, lab in enumerate(labels):
        short = lab[:11]
        _text(d, (18, top + i * cell + 34), short, (15, 23, 42), 16, True)
        _text(d, (left + i * cell + 8, top - 28), short, (15, 23, 42), 15, True)
    for i in range(n):
        for j in range(n):
            v = int(cm[i, j])
            x1 = left + j * cell
            y1 = top + i * cell
            t = v / mx
            if i == j:
                fill = (219, 234, 254) if t < 0.5 else (37, 99, 235)
                tc = (15, 23, 42) if t < 0.5 else (255, 255, 255)
            else:
                fill = (254, 226, 226) if v else (248, 250, 252)
                tc = (127, 29, 29) if v else (148, 163, 184)
            d.rectangle((x1, y1, x1 + cell - 6, y1 + cell - 6), fill=fill, outline=(15, 23, 42))
            _text(d, (x1 + 28, y1 + 32), str(v), tc, 20, True)
    img.save(path, quality=95)


def save_example_cm(path: Path) -> None:
    cm = np.array([[85, 8], [6, 41]])
    save_confusion_light(cm, ["Normal", "Attack"], path, "Figure 12  ·  How to read a confusion matrix (example)")


def save_example_roc(path: Path) -> None:
    w, h = 1100, 720
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 56), fill=(12, 26, 46))
    _text(d, (24, 14), "Figure 13  ·  How to read an ROC curve (example)", (240, 193, 75), 22, True)
    ox, oy, s = 140, 640, 480
    d.rectangle((ox, oy - s, ox + s, oy), outline=(15, 23, 42), width=2)
    d.line((ox, oy, ox + s, oy - s), fill=(203, 213, 225), width=2)
    pts = []
    for i in range(21):
        fpr = i / 20
        tpr = 1 - (1 - fpr) ** 3
        pts.append((ox + int(fpr * s), oy - int(tpr * s)))
    d.line(pts, fill=(37, 99, 235), width=5)
    _text(d, (ox + 140, oy + 16), "False positive rate →", (15, 23, 42), 18)
    _text(d, (24, oy - s // 2), "TPR", (15, 23, 42), 18)
    _text(d, (ox + 40, 90), "Closer to the top-left corner = better separation of attack vs normal.", (71, 85, 105), 18)
    img.save(path, quality=95)


def crop_readable(src: Path, dst: Path) -> None:
    im = Image.open(src).convert("RGB")
    w, h = im.size
    # Keep the top of the dashboard (nav + main panel), 16:10 so UI text stays readable.
    th = min(h, int(w * 10 / 16))
    im.crop((0, 0, w, th)).save(dst, quality=92)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def set_inline_size(p: Paragraph, width_in: float, max_height_in: float) -> None:
    ext = p._p.find(f".//{NS_WP}extent")
    if ext is None:
        return
    cx = int(ext.get("cx") or 0)
    cy = int(ext.get("cy") or 0)
    if cx <= 0 or cy <= 0:
        return
    new_cx = int(width_in * EMU)
    new_cy = int(cy * (new_cx / cx))
    if new_cy > max_height_in * EMU:
        new_cy = int(max_height_in * EMU)
        new_cx = int(cx * (new_cy / cy))
    ext.set("cx", str(new_cx))
    ext.set("cy", str(new_cy))
    for aext in p._p.findall(f".//{NS_A}ext"):
        if aext.get("cx") and aext.get("cy"):
            aext.set("cx", str(new_cx))
            aext.set("cy", str(new_cy))


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 6), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            r_id = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            part = paras[j].part.related_parts.get(r_id) if r_id else None
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def insert_picture_before(caption: Paragraph, image_path: Path, width_in: float) -> None:
    new_p = OxmlElement("w:p")
    caption._p.addprevious(new_p)
    para = Paragraph(new_p, caption._parent)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))


def tidy_spacing(doc: Document) -> dict:
    removed = 0
    paras = list(doc.paragraphs)
    # Remove empty Heading 1 (they pollute the TOC)
    for p in paras:
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        has_img = bool(p._p.findall(f".//{NS_A}blip"))
        if st == "Heading 1" and not t and not has_img:
            delete_paragraph(p)
            removed += 1
    # Collapse consecutive empty Normal paragraphs (keep page-break paras)
    paras = list(doc.paragraphs)
    empty_run = []
    for p in paras:
        t = (p.text or "").strip()
        has_img = bool(p._p.findall(f".//{NS_A}blip"))
        has_page = any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))
        if not t and not has_img and not has_page and (p.style.name if p.style else "") in ("Normal", "List Bullet"):
            empty_run.append(p)
        else:
            # keep none of a run of empty normals (page-break already separates)
            for extra in empty_run:
                delete_paragraph(extra)
                removed += 1
            empty_run = []
    for extra in empty_run:
        delete_paragraph(extra)
        removed += 1
    return {"removed_empty": removed}


def main() -> None:
    m = json.loads(METRICS.read_text(encoding="utf-8"))
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    labels = m["labels"]
    cm = np.array(rf["confusion_matrix"], dtype=int)

    save_dataset_light(m["dataset"]["classes"], EVAL / "fig_dataset_distribution.png")
    save_bars_light(rf, xgb, iso, EVAL / "fig_model_performance.png")
    save_confusion_light(cm, labels, EVAL / "fig_confusion_rf.png", "Figure 17  ·  Random Forest confusion matrix (12,000-row holdout)")
    save_example_cm(EVAL / "fig_cm_example.png")
    save_example_roc(EVAL / "fig_roc_example.png")
    for src, dst in [
        (SHOTS / "00_home.png", EVAL / "shot_home_print.png"),
        (SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png"),
        (SHOTS / "10_soc.png", EVAL / "shot_soc_print.png"),
        (SHOTS / "07_models.png", EVAL / "shot_models_print.png"),
        (SHOTS / "02_detection.png", EVAL / "shot_detect_print.png"),
    ]:
        if src.exists():
            crop_readable(src, dst)

    doc = Document(str(PAPER))
    stats = tidy_spacing(doc)

    # Figure 12 had a caption but no image
    has12 = False
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if (p.text or "").startswith("Figure 12."):
            prev = paras[i - 1] if i else None
            if prev is None or not prev._p.findall(f".//{NS_A}blip"):
                insert_picture_before(p, EVAL / "fig_cm_example.png", 5.6)
                stats["fig12_inserted"] = True
            has12 = True
            break
    stats["fig12"] = has12

    swapped = []
    for cap, img in [
        ("Figure 13. ROC Curve Example.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 11.", EVAL / "shot_home_print.png"),
    ]:
        if replace_image_before_caption(doc, cap, img):
            swapped.append(cap.split(".")[0])
    stats["swapped"] = swapped

    # Readable on the page: cap body figures, skip cover (first image ~full page)
    first_img = True
    for p in doc.paragraphs:
        if not p._p.findall(f".//{NS_A}blip"):
            continue
        if first_img:
            first_img = False
            continue
        set_inline_size(p, width_in=6.3, max_height_in=4.55)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure ") or t.startswith("Table "):
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = False

    # TOC glitch leftover in static/field result
    for p in doc.paragraphs:
        t = p.text or ""
        if "5.6 Expected Performance Comparison" in t and "Experimental Performance Results" in t:
            if p.runs:
                p.runs[0].text = "5.6 Experimental Performance Results\t34"
                for r in p.runs[1:]:
                    r.text = ""

    doc.save(str(PAPER))
    print("Saved", PAPER)
    print(stats)

    # Refresh TOC / LoF page numbers in Word
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(PAPER), ReadOnly=False)
        wdoc.Fields.Update()
        for i in range(1, wdoc.TablesOfContents.Count + 1):
            wdoc.TablesOfContents(i).Update()
        for i in range(1, wdoc.TablesOfFigures.Count + 1):
            wdoc.TablesOfFigures(i).Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
        print("Word TOC/LoF updated")
    except Exception as exc:
        print("Word COM skipped:", type(exc).__name__, exc)


if __name__ == "__main__":
    main()
