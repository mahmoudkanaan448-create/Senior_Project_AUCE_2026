"""Deep audit of final thesis docx."""
import sys
import re
from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')
PAPER = Path(r'C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx')
d = Document(str(PAPER))
NS_A = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

issues = []

# Chapters
chapters = []
for i, p in enumerate(d.paragraphs):
    t = (p.text or '').strip()
    if p.style and p.style.name == 'Heading 1' and (t.startswith('Chapter ') or t in ('Conclusion', 'References') or t.startswith('Appendix')):
        chapters.append((i, t))

print('=== CHAPTERS ===')
for idx, t in chapters:
    print(idx, t)

# Empty normals between chapters
for ci in range(len(chapters) - 1):
    start, end = chapters[ci][0], chapters[ci + 1][0]
    empties = []
    for j in range(start + 1, end):
        p = d.paragraphs[j]
        t = (p.text or '').strip()
        has_img = bool(p._p.findall(f'.//{NS_A}blip'))
        has_pb = any(br.get(f'{NS_W}type') == 'page' for br in p._p.findall(f'.//{NS_W}br'))
        if not t and not has_img and not has_pb:
            empties.append(j)
    if len(empties) > 2:
        issues.append(f'Many empty paras between {chapters[ci][1][:30]} and {chapters[ci+1][1][:30]}: {len(empties)}')

# Double page breaks before chapters
for idx, t in chapters:
    if idx > 0:
        prev = d.paragraphs[idx - 1]
        prev2 = d.paragraphs[idx - 2] if idx >= 2 else None
        pb1 = any(br.get(f'{NS_W}type') == 'page' for br in prev._p.findall(f'.//{NS_W}br'))
        pb2 = prev2 and any(br.get(f'{NS_W}type') == 'page' for br in prev2._p.findall(f'.//{NS_W}br'))
        if not t and pb1:
            issues.append(f'Empty para with page break before: {t or chapters[ci][1]}')
        if pb1 and pb2:
            issues.append(f'Double page break before chapter near para {idx}')

# TOC issues
print('\n=== TOC ISSUES ===')
for p in d.paragraphs:
    st = p.style.name if p.style else ''
    t = (p.text or '').strip()
    if not st.lower().startswith('toc'):
        continue
    if 'Data Flow Diagram' in t and 'System Data Pipeline' not in t:
        issues.append(f'TOC DFD: {t[:80]}')
        print('BAD TOC:', t[:100])
    if 'Chapter 8' in t:
        issues.append(f'TOC still Ch8: {t}')
        print('BAD TOC:', t)
    if 'Practical Implementation Notes' in t:
        issues.append(f'TOC old Ch7: {t}')
        print('BAD TOC:', t)
    if t.count('Chapter 7') > 1:
        issues.append(f'TOC merged: {t[:80]}')
        print('BAD TOC:', t[:120])
    if 'Expected Performance Comparison' in t:
        issues.append(f'TOC old 5.6: {t}')
        print('BAD TOC:', t)

# Images
print('\n=== IMAGES ===')
widths = []
small = []
for i, p in enumerate(d.paragraphs):
    if not p._p.findall(f'.//{NS_A}blip'):
        continue
    ext = p._p.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
    if ext is not None:
        w_in = int(ext.get('cx', 0)) / 914400
        widths.append(w_in)
        if w_in < 5.5:
            cap = d.paragraphs[i+1].text[:50] if i+1 < len(d.paragraphs) else ''
            small.append((w_in, cap))

print(f'images: {len(widths)}, avg width {sum(widths)/len(widths):.2f} in' if widths else 'no images')
if small:
    print('SMALL images (<5.5in):')
    for w, c in small[:8]:
        print(f'  {w:.2f}in before {c}')

# Figures without image
print('\n=== FIGURES WITHOUT IMAGE ===')
missing_fig = []
paras = list(d.paragraphs)
for i, p in enumerate(paras):
    t = (p.text or '').strip()
    if re.match(r'^Figure (\d+\.|H\d)', t):
        prev = paras[i-1] if i else None
        has = prev and prev._p.findall(f'.//{NS_A}blip')
        if not has:
            missing_fig.append(t[:60])
for m in missing_fig:
    print('MISSING:', m)
    issues.append(f'Figure missing image: {m}')

# Tables
print(f'\n=== TABLES: {len(d.tables)} ===')
wide_tables = []
for ti, t in enumerate(d.tables):
    cols = len(t.rows[0].cells) if t.rows else 0
    rows = len(t.rows)
    if cols >= 7:
        wide_tables.append((ti, cols, rows, (t.rows[0].cells[0].text or '')[:30]))
print('Wide tables (7+ cols):', len(wide_tables))
for x in wide_tables[:5]:
    print(' ', x)

# Doctor key phrases
blob = '\n'.join((p.text or '') for p in d.paragraphs).lower()
for tbl in d.tables:
    for row in tbl.rows:
        for cell in row.cells:
            blob += '\n' + (cell.text or '').lower()

checks = [
    ('12,000', '12,000' in blob),
    ('4.2.2', '4.2.2' in blob),
    ('5.6.1', '5.6.1' in blob),
    ('fpr', 'fpr' in blob),
    ('window_size=10', 'window_size=10' in blob),
    ('fuse_decisions', 'fuse_decisions' in blob),
    ('attack simulation', 'tested (threat simulation)' in blob),
    ('no hybrid 992', 'hybrid ai: 99.2' not in blob),
    ('ch7 user manual', 'chapter 7 - user manual' in blob),
    ('no old ch7', 'practical implementation notes' not in blob),
]
print('\n=== DOCTOR CHECKS ===')
for name, ok in checks:
    print(('PASS' if ok else 'FAIL'), name)

print('\n=== ISSUES ===')
if issues:
    for x in issues:
        print('-', x)
else:
    print('No major structural issues detected')

print('\nempty paragraphs total:', sum(1 for p in d.paragraphs if not (p.text or '').strip()))
