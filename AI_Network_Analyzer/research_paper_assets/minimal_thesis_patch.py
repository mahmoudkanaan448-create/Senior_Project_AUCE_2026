"""
Minimal patch ONLY — no TOC rebuild, no spacing, no blank-page removal, no reorder.
1. Restore from clean backup (_FINAL_STAGING)
2. Add explanation paragraph under headings that have none
3. Swap in new app screenshot images
4. Insert doctor-required tables only if missing from body
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
RESTORE_FROM = ROOT / "_FINAL_STAGING.docx"
TARGET = Path(
    r"C:\Users\mohamad\OneDrive\Desktop"
    r"\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED_pre_polish_20260820_0214.docx"
)
FIGS = ROOT / "new_figures"
EVAL = ROOT / "evaluation"
SNIPS = ROOT / "code_snippets"
SHOTS = ROOT / "user_guide_shots"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

HEADING_INTROS: dict[str, str] = {
    "1.4 Research Objectives": (
        "Table 2 summarises the six research objectives that guided the design and evaluation of the AI-NDR platform."
    ),
    "1.5 Research Questions": (
        "Table 2 (continued) lists the research questions that each objective was designed to answer."
    ),
    "2.3 Machine Learning and Deep Learning Techniques": (
        "Table 6 compares the supervised and unsupervised models selected for this project and their roles in the hybrid engine."
    ),
    "2.4 Datasets Used in Network Intrusion Detection": (
        "Table 4 documents the public and synthetic datasets considered; the holdout evaluation uses the 12,000-row CICIDS-style file described in §4.2.2."
    ),
    "3.3 Main Modules": (
        "The implemented platform is organised into the modules listed below; §3.3.1 and §3.3.2 map each module to its evaluation status."
    ),
    "3.3.1 Implemented Core Modules (Aligned with Software)": (
        "The implemented system modules match the software repository: live/CSV capture, 39-feature extraction, "
        "Hybrid AI engine (RF, XGB, IF, AE, LSTM), decision fusion, threat intelligence, XAI, alerts, Telegram, "
        "Threat Simulation, MITRE mapping, and SOAR playbooks."
    ),
    "3.3.2 Implementation and Evaluation Status": (
        "Implemented and tested on holdout or live demo: capture, RF/XGB/IF evaluation, dashboard, alerts, Telegram, "
        "Threat Simulation, XAI, MITRE, SOAR. Autoencoder, LSTM, and Hybrid fusion are implemented but not scored in Table 11."
    ),
    "4.1 Technology Stack": (
        "Table 9 lists the programming languages, libraries, and services used in the FastAPI + Streamlit implementation."
    ),
    "4.2.1 Complete Feature Set Used in Experiments": (
        "This subsection documents the complete input schema used in training and inference. Every column in Table 5 "
        "is defined in config.FEATURE_COLUMNS and extracted identically during live capture, CSV import, and holdout evaluation."
    ),
    "4.3 Model Development": (
        "Model development loads dataset.csv, applies train-only preprocessing (LabelEncoder, StandardScaler), "
        "trains each model, and persists artifacts under models/ for FastAPI/Streamlit inference."
    ),
    "4.3.1 Per-Model Training Configuration": (
        "Random Forest — supervised multi-class on 39 scaled features (n_estimators=200, max_depth=20). XGBoost — "
        "supervised multi-class (n_estimators=200, max_depth=8, learning_rate=0.1). Isolation Forest — Normal-only training "
        "(contamination=0.1). Autoencoder and LSTM are implemented; Hybrid uses fuse_decisions() at inference."
    ),
    "4.3.2 LSTM Sequence Input Construction": (
        "LSTM uses training/data_preprocessing.prepare_sequences() with window_size=10. Each sample is a (10, 39) tensor; "
        "the label is the class of the flow after the window. LSTM is not scored in Table 11 because evaluate_models.py is row-based."
    ),
    "4.4 Decision Fusion": (
        "Hybrid AI is weighted majority fusion (decision_engine.fuse_decisions), not a sixth trained classifier. See §4.4.1 and Figure 2."
    ),
    "4.5 Database Design": (
        "Table 7 describes the SQLite tables that persist users, flows, alerts, incidents, and model metadata."
    ),
    "4.7 Implementation Workflow": (
        "Table 13 summarises the end-to-end workflow from dataset preparation through model training to dashboard deployment."
    ),
    "4.9 Algorithmic Workflow": (
        "Table 15 presents the step-by-step algorithm executed at inference time for each captured or uploaded flow."
    ),
    "4.11 Project Folder Structure": (
        "Table 16 maps the repository folders to their responsibilities in the senior project codebase."
    ),
    "5.2 Testing Environment": (
        "Table 17 lists the hardware, software, and network configuration used during functional and experimental testing."
    ),
    "5.3 Attack Simulation Scenarios": (
        "Table 18 defines the attack scenarios exercised through Threat Simulation; §5.3.1 provides the structured testing matrix."
    ),
    "5.4 Evaluation Metrics": (
        "Table 8 defines accuracy, precision, recall, F1, ROC-AUC, FPR, and FNR used to report holdout results in §5.6."
    ),
    "5.6 Experimental Performance Results": (
        "Measured holdout metrics are reported for Random Forest, XGBoost, and Isolation Forest (Table 11). "
        "Autoencoder, LSTM, and Hybrid AI are described as implemented but not scored on the same split."
    ),
    "6.2 Security Considerations": (
        "Table 10 identifies deployment risks (credential storage, live blocking, API exposure) and the mitigations applied."
    ),
    "7.5 Maintenance Procedures": (
        "Table 24 recommends maintenance tasks and frequencies for keeping the SOC dashboard and models operational."
    ),
    "7.6 Minimal Functional Requirements": (
        "Table 25 lists the minimum functional requirements verified during acceptance testing."
    ),
    "7.7 Non-Functional Requirements": (
        "Table 26 summarises non-functional requirements for performance, reliability, and usability."
    ),
    "8.5 Maintenance Procedures": (
        "Table 24 recommends maintenance tasks and frequencies for keeping the SOC dashboard and models operational."
    ),
    "8.6 Minimal Functional Requirements": (
        "Table 25 lists the minimum functional requirements verified during acceptance testing."
    ),
    "8.7 Non-Functional Requirements": (
        "Table 26 summarises non-functional requirements for performance, reliability, and usability."
    ),
}


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = anchor.part.document.styles[style]
    except Exception:
        pass
    set_text(para, text)
    para.paragraph_format.space_after = Pt(6)
    return para


def heading_has_intro(p: Paragraph, min_len: int = 25) -> bool:
    nxt = p._p.getnext()
    while nxt is not None:
        if nxt.tag.endswith("}tbl"):
            return False
        if nxt.tag.endswith("}p"):
            para = Paragraph(nxt, p._parent)
            t = (para.text or "").strip()
            st = para.style.name if para.style else ""
            if st.startswith("Heading"):
                return False
            if t and len(t) >= min_len and not re.match(r"^Table \d+\.", t):
                return True
        nxt = nxt.getnext()
    return False


def add_heading_intros(doc: Document) -> int:
    added = 0
    for exact, intro in HEADING_INTROS.items():
        for p in doc.paragraphs:
            st = p.style.name if p.style else ""
            if st not in ("Heading 2", "Heading 3"):
                continue
            if (p.text or "").strip() != exact:
                continue
            if heading_has_intro(p):
                break
            insert_after(p, intro)
            added += 1
            break
    return added


def replace_image_before_caption(doc: Document, cap: str, path: Path) -> bool:
    if not path.exists():
        return False
    blob = path.read_bytes()
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if cap not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = paras[j].part.related_parts.get(rid) if rid else None
            if part:
                part._blob = blob
                return True
    return False


def refresh_figures(doc: Document) -> int:
    from PIL import Image

    def crop_shot(src: Path, dst: Path) -> None:
        if not src.exists():
            return
        im = Image.open(src).convert("RGB")
        w, h = im.size
        th = min(h, int(w * 10 / 16))
        im.crop((0, 0, w, th)).save(dst, quality=95)

    EVAL.mkdir(parents=True, exist_ok=True)
    crop_shot(SHOTS / "01_live.png", EVAL / "shot_live_print.png")
    crop_shot(SHOTS / "00_home.png", EVAL / "shot_home_print.png")
    crop_shot(SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png")
    crop_shot(SHOTS / "10_soc_tab1_SOAR_Playbooks.png", EVAL / "shot_soc_print.png")
    crop_shot(SHOTS / "07_models.png", EVAL / "shot_models_print.png")

    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 11.", EVAL / "shot_live_print.png"),
        ("Figure 12.", EVAL / "fig_cm_example.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
        ("Figure H1", SNIPS / "01_hybrid_ai_decision_fusion.png"),
        ("Figure H2", SNIPS / "09_mitre_map.png"),
        ("Figure H3", SNIPS / "10_soar_playbooks.png"),
        ("Figure H4", SNIPS / "12_online_learning.png"),
        ("Figure H5", SNIPS / "14_threat_simulation.png"),
        ("Figure H6", SNIPS / "13_alert_manager.png"),
        ("Figure H7", SNIPS / "16_telegram_alert.png"),
        ("Figure H8", SNIPS / "15_supervisor.png"),
        ("Figure H9", SNIPS / "17_ai_assistant.png"),
    ]
    return sum(1 for cap, img in mapping if replace_image_before_caption(doc, cap, img))


def ensure_required_tables(doc: Document) -> list[str]:
    """Call doctor_complete helpers only when a required table/section is absent."""
    import json

    import doctor_complete_revision as dcr

    blob = " ".join((p.text or "") for p in doc.paragraphs).lower()
    actions: list[str] = []
    m = json.loads(dcr.METRICS_PATH.read_text(encoding="utf-8"))

    if "table 5a summarises" not in blob and "per-model dataset usage summary" not in blob:
        dcr.add_model_usage_table(doc, m)
        actions.append("added_table_5a")
    if "table 3a. component implementation status" not in blob and "component implementation status" not in blob:
        dcr.add_implementation_status_table(doc)
        actions.append("added_table_3a")
    if "5.3.1" in blob and "attack type" not in blob and "expected detector" not in blob and "dos/ddos" not in blob:
        dcr.ensure_attack_matrix_table(doc)
        actions.append("added_attack_matrix")
    if "4.3.2 lstm sequence input construction" not in blob:
        dcr.add_lstm_sequence_section(doc)
        actions.append("added_lstm_section")
    if "4.4.1 hybrid decision fusion pipeline" not in blob:
        dcr.add_hybrid_fusion_section(doc)
        actions.append("added_hybrid_section")

    return actions


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    target = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else TARGET
    restore = RESTORE_FROM if RESTORE_FROM.exists() else ROOT / "_REBUILD_STAGING.docx"
    if not restore.exists():
        raise SystemExit(f"No restore backup: {restore}")

    shutil.copy2(restore, target)
    print("Restored from:", restore)
    print("Target:", target)

    subprocess.run([sys.executable, str(ROOT / "generate_academic_diagrams.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "render_eval_figures.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "generate_extra_snippets.py")], check=True, cwd=str(ROOT))

    doc = Document(str(target))
    stats = {
        "heading_intros": add_heading_intros(doc),
        "figures_swapped": refresh_figures(doc),
        "tables_fixed": ensure_required_tables(doc),
    }
    doc.save(str(target))
    print("MINIMAL PATCH OK:", stats)
    print("Saved:", target)


if __name__ == "__main__":
    main()
