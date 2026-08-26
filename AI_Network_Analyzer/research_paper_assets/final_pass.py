"""
Final quality pass: fix TOC corruption, blank pages, spacing, images, tables.
Saves in-place to Desktop FINAL_UPDATED.docx (no backup files left on Desktop).
"""
from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
FIGS = ROOT / "new_figures"
EVAL = ROOT / "evaluation"
SNIPS = ROOT / "code_snippets"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
EMU = 914400
BODY_W = 6.75
DIAG_H = 5.85
SHOT_H = 6.25


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def has_page_break(p: Paragraph) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def remove_page_breaks(p: Paragraph) -> None:
    for br in list(p._p.findall(f".//{NS_W}br")):
        if br.get(f"{NS_W}type") == "page":
            br.getparent().remove(br)


def has_image(p: Paragraph) -> bool:
    return bool(p._p.findall(f".//{NS_A}blip"))


def insert_before(doc: Document, anchor: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass
    set_text(para, text)
    return para


def replace_corrupted_ch7_toc(doc: Document) -> int:
    """Delete field-corrupted TOC lines and insert clean static entries."""
    clean = [
        ("toc 1", "Chapter 7 - User Manual and Operation Guide\t51"),
        ("toc 2", "7.1 Starting the System\t51"),
        ("toc 2", "7.2 Uploading Data or Capturing Live Traffic\t52"),
        ("toc 2", "7.3 Reading the Dashboard\t52"),
        ("toc 2", "7.4 Responding to Alerts\t53"),
        ("toc 2", "7.5 Maintenance Procedures\t53"),
        ("toc 3", "7.5.1 Reliability and Auto-Recovery\t53"),
        ("toc 2", "7.6 Minimal Functional Requirements\t54"),
        ("toc 2", "7.7 Non-Functional Requirements\t54"),
    ]
    concl_idx = None
    for i, p in enumerate(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        if st == "toc 1" and (p.text or "").strip().startswith("Conclusion"):
            concl_idx = i
            break
    if concl_idx is None:
        return 0
    anchor = doc.paragraphs[concl_idx]
    # Walk backwards from Conclusion: remove toc lines until chapter 6 entry
    k = concl_idx - 1
    removed = 0
    while k >= 0:
        p = doc.paragraphs[k]
        st = (p.style.name if p.style else "").lower()
        t = (p.text or "").strip()
        if not st.startswith("toc"):
            break
        if t.startswith("6.") or t.startswith("Chapter 6"):
            break
        delete_paragraph(p)
        removed += 1
        k -= 1
        concl_idx -= 1
        anchor = doc.paragraphs[concl_idx]
    for style, text in clean:
        insert_before(doc, anchor, text, style)
    return len(clean)


def rebuild_ch7_toc_block(doc: Document) -> int:
    """Replace corrupted Chapter 7 static TOC entries (field-safe)."""
    return replace_corrupted_ch7_toc(doc)


def fix_toc_line(t: str) -> str:
    """Repair duplicated static TOC lines from repeated polish passes."""
    t = t.strip()
    if not t:
        return t

    # Chapter 7 block corruption
    if "User Manual" in t and ("Chapter 8" in t or t.count("Chapter 7") > 1):
        return "Chapter 7 - User Manual and Operation Guide\t51"

    # Prefer last valid "7.x Title\tpage" pair embedded in duplicated string
    pairs = re.findall(
        r"(7\.\d+(?:\.\d+)?\s+[^\t\d]+?)\t(\d{1,3})(?:\D|$)",
        t.replace("78", "\t78 ").replace("37.", "\t7."),
    )
    if pairs:
        title, page = pairs[-1]
        title = re.sub(r"\s+", " ", title).strip()
        return f"{title}\t{page}"

    # 8.x → 7.x with trailing page number
    m = re.match(r"^8\.(\d+(?:\.\d+)?)\s+(.+)$", t.split("\t")[0])
    if m:
        title = m.group(2).split("\t")[0].split("7.")[0].strip()
        page = "51"
        for part in reversed(t.split("\t")):
            if part.isdigit() and len(part) <= 3:
                page = part
                break
        return f"7.{m.group(1)} {title}\t{page}"

    if t.startswith("Chapter 8 - User Manual"):
        page = "51"
        for part in reversed(t.split("\t")):
            if part.isdigit():
                page = part
                break
        return f"Chapter 7 - User Manual and Operation Guide\t{page}"

    if "Data Flow Diagram" in t and "System Data Pipeline" not in t:
        return "3.4 System Data Pipeline\t23"
    if "Expected Performance Comparison" in t:
        return "5.6 Experimental Performance Results\t41"
    if t.startswith("5.6.1 Operational Deployment Statistics"):
        return "5.6.3 Operational Deployment Statistics\t41"
    if "5.6.3" in t and "5.6.1 Operational" in t:
        return "5.6.3 Operational Deployment Statistics\t41"

    return t


def rebuild_static_toc(doc: Document) -> int:
    fixed = 0
    for p in doc.paragraphs:
        st = (p.style.name if p.style else "").lower()
        if not st.startswith("toc"):
            continue
        old = (p.text or "").strip()
        if not old:
            continue
        new = fix_toc_line(old)
        if new != old:
            set_text(p, new)
            fixed += 1
    return fixed


def remove_empty_page_break_paragraphs(doc: Document) -> int:
    """Remove blank Normal paragraphs that contain only a page break (cause blank pages)."""
    removed = 0
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if not t and not has_image(p) and has_page_break(p) and st == "Normal":
            delete_paragraph(p)
            removed += 1
    return removed


def normalize_chapter_breaks(doc: Document) -> int:
    """One page break before each major chapter — no double breaks / blank pages."""
    changed = 0
    titles: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if p.style and p.style.name == "Heading 1" and (
            t.startswith("Chapter ") or t.startswith("Appendix ") or t in ("Conclusion", "References")
        ):
            titles.append(t)

    for title in titles:
        if title.startswith("Chapter 1"):
            continue
        # Appendices B–H stay on same page flow as Appendix A (no extra breaks)
        if title.startswith("Appendix ") and title != "Appendix A - Defense Preparation Notes":
            continue
        for k, p in enumerate(list(doc.paragraphs)):
            if (p.text or "").strip() != title or not (p.style and p.style.name == "Heading 1"):
                continue
            while k > 0:
                prev = doc.paragraphs[k - 1]
                pt = (prev.text or "").strip()
                pst = prev.style.name if prev.style else ""
                if pt or has_image(prev) or pst.startswith("Heading"):
                    remove_page_breaks(prev)
                    break
                delete_paragraph(prev)
                changed += 1
                k -= 1
            for k2, hp in enumerate(doc.paragraphs):
                if (hp.text or "").strip() != title or not (hp.style and hp.style.name == "Heading 1"):
                    continue
                if k2 > 0:
                    prev = doc.paragraphs[k2 - 1]
                    if not has_page_break(prev):
                        run = prev.runs[0] if prev.runs else prev.add_run()
                        run.add_break(WD_BREAK.PAGE)
                        changed += 1
                break
            break
    return changed


def remove_extra_appendix_breaks(doc: Document) -> int:
    """Remove page breaks inserted before Appendix B–H (keep section continuous)."""
    n = 0
    for k, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not (
            p.style
            and p.style.name == "Heading 1"
            and t.startswith("Appendix ")
            and not t.startswith("Appendix A")
        ):
            continue
        if k > 0:
            pr = doc.paragraphs[k - 1]
            if has_page_break(pr):
                remove_page_breaks(pr)
                n += 1
    return n


def trim_empty_before_headings(doc: Document) -> int:
    """Remove empty Normal paragraphs directly before Heading 1 (not content spacing mid-chapter)."""
    removed = 0
    for i, p in enumerate(list(doc.paragraphs)):
        if not (p.style and p.style.name == "Heading 1"):
            continue
        k = i - 1
        while k >= 0:
            prev = doc.paragraphs[k]
            pt = (prev.text or "").strip()
            pst = prev.style.name if prev.style else ""
            if pt or has_image(prev) or pst.startswith("Heading"):
                break
            if has_page_break(prev):
                break
            delete_paragraph(prev)
            removed += 1
            k -= 1
            i -= 1
    return removed


def trim_empty_between_sections(doc: Document) -> int:
    """Remove runs of 3+ empty Normal paragraphs anywhere in body."""
    removed = 0
    run: list[Paragraph] = []
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        is_empty = (
            not t
            and not has_image(p)
            and not has_page_break(p)
            and st in ("Normal", "List Bullet", "List Number")
        )
        if is_empty:
            run.append(p)
        else:
            if len(run) >= 3:
                for extra in run:
                    delete_paragraph(extra)
                    removed += 1
            run = []
    if len(run) >= 3:
        for extra in run:
            delete_paragraph(extra)
            removed += 1
    return removed


def set_inline_size(p: Paragraph, width_in: float, max_h: float) -> None:
    ext = p._p.find(f".//{NS_WP}extent")
    if ext is None:
        return
    cx, cy = int(ext.get("cx") or 0), int(ext.get("cy") or 0)
    if cx <= 0 or cy <= 0:
        return
    new_cx = int(width_in * EMU)
    new_cy = int(cy * (new_cx / cx))
    if new_cy > max_h * EMU:
        new_cy = int(max_h * EMU)
        new_cx = int(cx * (new_cy / cy))
    ext.set("cx", str(new_cx))
    ext.set("cy", str(new_cy))
    for aext in p._p.findall(f".//{NS_A}ext"):
        if aext.get("cx"):
            aext.set("cx", str(new_cx))
            aext.set("cy", str(new_cy))


def replace_image_before_caption(doc: Document, cap: str, path: Path) -> bool:
    if not path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = path.read_bytes()
    for i, p in enumerate(paras):
        if cap not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = paras[j].part.related_parts.get(rid) if rid else None
            if part:
                part._blob = blob
                return True
    return False


def refresh_all_figures(doc: Document) -> None:
    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 11.", EVAL / "shot_home_print.png"),
        ("Figure 12.", EVAL / "fig_cm_example.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
        ("Figure H1", SNIPS / "01_hybrid_ai_decision_fusion.png"),
        ("Figure H2", SNIPS / "09_mitre_map.png"),
        ("Figure H3", SNIPS / "10_soar_playbooks.png"),
        ("Figure H4", SNIPS / "12_online_learning.png"),
        ("Figure H5", SNIPS / "14_threat_simulation.png"),
        ("Figure H6", SNIPS / "13_alert_manager.png"),
        ("Figure H7", SNIPS / "16_telegram_alert.png"),
        ("Figure H8", SNIPS / "15_supervisor.png"),
        ("Figure H9", SNIPS / "17_ai_assistant.png"),
    ]
    for cap, img in mapping:
        replace_image_before_caption(doc, cap, img)

    first = True
    for i, p in enumerate(doc.paragraphs):
        if not has_image(p):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        if first:
            first = False
            set_inline_size(p, 6.5, 8.5)
            continue
        nxt = doc.paragraphs[i + 1].text if i + 1 < len(doc.paragraphs) else ""
        if any(x in (nxt or "") for x in ("Figure 18.", "Figure 19.", "Figure 20.", "Figure 21.", "Figure 11.")):
            set_inline_size(p, BODY_W, SHOT_H)
        elif "Figure H" in (nxt or ""):
            set_inline_size(p, 6.4, 5.2)
        else:
            set_inline_size(p, BODY_W, DIAG_H)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        # allow row to break across pages if needed but keep readable font
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
                    pf.line_spacing = 1.08
                    pf.widow_control = True
                    for r in p.runs:
                        if r.font.size is None or r.font.size.pt > 11:
                            r.font.size = Pt(9 if len(table.rows[0].cells) >= 7 else 10)


def format_captions_and_body(doc: Document) -> None:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if t.startswith("Figure ") or t.startswith("Table "):
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(10)
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(10)
        elif st == "Normal" and t:
            pf = p.paragraph_format
            pf.space_after = Pt(6)
            pf.line_spacing = 1.15


def word_update(path: Path, update_toc: bool = False) -> None:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(path), ReadOnly=False)
        wdoc.Fields.Update()
        if update_toc:
            for i in range(1, wdoc.TablesOfContents.Count + 1):
                wdoc.TablesOfContents(i).Update()
            for i in range(1, wdoc.TablesOfFigures.Count + 1):
                wdoc.TablesOfFigures(i).Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
    except Exception as exc:
        print("Word COM:", exc)


def verify(path: Path) -> None:
    import verify_doctor_revision as vd

    for i in range(5):
        passed, lines = vd.run_checks(path)
        if passed < len(lines):
            for ln in lines:
                if ln.startswith("FAIL"):
                    print(ln)
            raise SystemExit(f"Verify failed {passed}/{len(lines)}")
    print("Doctor verify 5/5 passes:", passed, "/", len(lines))


def main() -> None:
    if not PAPER.exists():
        raise SystemExit(f"Missing {PAPER}")

    subprocess.run([sys.executable, str(ROOT / "generate_academic_diagrams.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "render_eval_figures.py")], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ROOT / "ultimate_quality_fix.py")], check=True, cwd=str(ROOT))

    doc = Document(str(PAPER))
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "5.6.1 Operational Deployment Statistics":
            set_text(p, "5.6.3 Operational Deployment Statistics")

    stats = {
        "toc_fixed": rebuild_static_toc(doc) + rebuild_ch7_toc_block(doc),
        "empties_removed": trim_empty_between_sections(doc)
        + trim_empty_before_headings(doc)
        + remove_empty_page_break_paragraphs(doc),
        "chapter_breaks": normalize_chapter_breaks(doc),
        "appendix_breaks_removed": remove_extra_appendix_breaks(doc),
    }
    format_captions_and_body(doc)
    format_tables(doc)
    refresh_all_figures(doc)

    doc.save(str(PAPER))
    word_update(PAPER, update_toc=False)

    # Re-apply static TOC fixes (Word TOC auto-update merges duplicate Ch7/Ch8 entries)
    doc = Document(str(PAPER))
    stats["toc_fixed"] += rebuild_static_toc(doc) + rebuild_ch7_toc_block(doc)
    from ultimate_quality_fix import replace_corrupted_563_toc, remove_duplicate_captions

    stats["toc_563"] = replace_corrupted_563_toc(doc)
    stats["dup_caps"] = remove_duplicate_captions(doc)
    doc.save(str(PAPER))

    verify(PAPER)

    print("Final pass OK:", PAPER)
    print(stats)


if __name__ == "__main__":
    main()
