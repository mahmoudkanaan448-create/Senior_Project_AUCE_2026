import sys, re
from pathlib import Path
from zipfile import ZipFile
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')
p = Path(r'C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx')
d = Document(str(p))
print('=== TOC ===')
for i in range(118, 132):
    print(i, repr(d.paragraphs[i].text))
with ZipFile(p) as z:
    emus = [int(x) for x in re.findall(r'cx="([0-9]+)"', z.read('word/document.xml').decode('utf-8'))]
    print('avg img width in', round(sum(emus) / len(emus) / 914400, 2))
    print('media files', len([n for n in z.namelist() if n.startswith('word/media/')]))
NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
figs = 0
for i, para in enumerate(d.paragraphs):
    t = (para.text or '').strip()
    if t.startswith('Figure') and i > 0:
        if d.paragraphs[i - 1]._p.findall(f'.//{NS}blip'):
            figs += 1
print('figures with image before caption:', figs)
bad_toc = sum(1 for para in d.paragraphs if 'Chapter 8' in (para.text or ''))
print('Chapter 8 mentions:', bad_toc)
