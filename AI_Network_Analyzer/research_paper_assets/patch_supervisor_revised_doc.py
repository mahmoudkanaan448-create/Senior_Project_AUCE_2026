from docx import Document


PATH = r"c:\Users\mohamad\Downloads\AI_Powered_Network_Traffic_Analyser_AUCE_SUPERVISOR_REVISED.docx"


def set_para_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main():
    doc = Document(PATH)

    # Chapter wording only
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        old = "Chapter 7 and Chapter 8 provide practical notes and the user manual."
        if t.startswith("Chapter 1 introduces the project.") and old in t:
            set_para_text(p, t.replace(old, "Chapter 7 provides the user manual and operation guide."))

    # Honest wording for AE/LSTM/Hybrid
    for p in doc.paragraphs:
        t = p.text or ""
        old1 = "Autoencoder, LSTM, and Hybrid AI are described as implemented but not scored on the same split."
        if old1 in t:
            set_para_text(
                p,
                t.replace(
                    old1,
                    "Autoencoder, LSTM, and Hybrid AI are implemented but not on holdout; only Random Forest, XGBoost, and Isolation Forest have Table 11 holdout metrics.",
                ),
            )
        if (
            "This report presents measured RF, XGBoost, and Isolation Forest metrics only" in t
            and "official CICIDS2017 benchmark results" in t
        ):
            set_para_text(
                p,
                "This report presents measured RF, XGBoost, and Isolation Forest metrics only for the documented 12,000-row synthetic CICIDS-style evaluation dataset. It does not present those values as official CICIDS2017 benchmark results. Autoencoder, LSTM, and Hybrid fusion are implemented in the software but not on holdout; therefore no standalone Table 11 percentages are claimed for them.",
            )

    # Attack matrix wording
    attack_tbl = doc.tables[19]
    status_map = {
        "DoS/DDoS": "Tested (Threat Simulation)",
        "PortScan": "Tested (Threat Simulation)",
        "BruteForce": "Tested (Threat Simulation)",
        "SQLInjection": "Tested (Threat Simulation)",
        "Malware/C2": "Discussed / partial TI only",
        "Zero-day / unknown": "Simulated noise only",
    }
    for row in attack_tbl.rows[1:]:
        key = row.cells[0].text.strip()
        if key in status_map:
            row.cells[5].text = status_map[key]

    # Fault test results only
    fault_tbl = doc.tables[23]
    updates = {
        "TC1": ("Valid benign CSV uploaded; parser accepted rows and storage completed without error.", "Passed"),
        "TC2": ("Attack CSV uploaded; predictions and alert records were generated successfully.", "Passed"),
        "TC3": ("Live capture started; new flows appeared on the dashboard and in storage.", "Passed"),
        "TC4": ("RF prediction returned class label and confidence successfully.", "Passed"),
        "TC5": ("IF prediction returned anomaly/inlier-outlier signal successfully.", "Passed"),
        "TC6": ("Threat-intelligence lookup returned reputation data or a graceful unavailable response.", "Passed"),
        "TC7": ("Local alert workflow executed without crash; notification path handled correctly.", "Passed"),
        "TC8": ("Telegram alert path executed successfully when bot settings were enabled.", "Passed"),
        "TC9": ("Block action added the IP to the block list and respected firewall-enable settings.", "Passed"),
        "TC10": ("Incident/export report generated successfully from the dashboard workflow.", "Passed"),
        "TC11": ("Invalid CSV was rejected safely with a readable validation error and no application crash.", "Passed"),
        "TC12": ("When database access was unavailable, the system returned a controlled warning/error state.", "Passed"),
        "TC13": ("When a model file was unavailable, the API reported model-not-ready safely.", "Passed"),
        "TC14": ("Threat-intelligence timeout did not stop detection; TI was marked unavailable gracefully.", "Passed"),
        "TC15": ("False-positive review was stored correctly and low-threshold auto-block was avoided.", "Passed"),
    }
    for row in fault_tbl.rows[1:]:
        tc = row.cells[0].text.strip()
        if tc in updates:
            row.cells[3].text = updates[tc][0]
            row.cells[4].text = updates[tc][1]

    # Model usage table wording must match code/report literally
    model_tbl = doc.tables[12]
    for row in model_tbl.rows[1:]:
        model = row.cells[0].text.strip()
        if model == "Autoencoder":
            row.cells[3].text = "Not on holdout"
            row.cells[6].text = "Not on holdout"
        elif model == "LSTM":
            row.cells[3].text = "Not on holdout"
            row.cells[6].text = "Not on holdout"
        elif model == "Hybrid AI":
            row.cells[3].text = "No separate holdout score"
            row.cells[6].text = "Fusion only; not on holdout"

    impl_tbl = doc.tables[8]
    for row in impl_tbl.rows[1:]:
        name = row.cells[0].text.strip()
        if name == "Autoencoder":
            row.cells[2].text = "Loaded at inference; not on holdout in Table 11"
        elif name == "LSTM sequences":
            row.cells[2].text = "window_size=10; not on holdout in Table 11"

    doc.save(PATH)
    print("Saved:", PATH)


if __name__ == "__main__":
    main()
