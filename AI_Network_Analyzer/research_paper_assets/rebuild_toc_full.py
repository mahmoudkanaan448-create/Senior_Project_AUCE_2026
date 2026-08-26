"""Rebuild entire static TOC from body headings with page numbers."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PAPER = Path(__file__).resolve().parent / "FINAL_FIXED.docx"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def has_page_break(p) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def insert_before(doc: Document, anchor: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass
    set_text(para, text)
    return para


def toc_style(level: int) -> str:
    if level == 1:
        return "toc 1"
    if level == 2:
        return "toc 2"
    return "toc 3"


def heading_level(style: str) -> int:
    if style == "Heading 1":
        return 1
    if style == "Heading 2":
        return 2
    return 3


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    doc = Document(str(PAPER))

    # Build entries from body headings (skip front-matter duplicates before Ch1 body)
    page = 1
    entries: list[tuple[int, str, int]] = []
    started = False
    for p in doc.paragraphs:
        if has_page_break(p):
            page += 1
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if t.startswith("Chapter 1 - Introduction") and st == "Heading 1":
            started = True
        if not started:
            continue
        if not st.startswith("Heading"):
            continue
        if t in ("List of Figures", "List of Tables"):
            continue
        if t.startswith("Appendix") or t in ("Conclusion", "References"):
            lvl = heading_level(st)
        elif t.startswith("Chapter "):
            lvl = 1
        else:
            lvl = heading_level(st)
        entries.append((lvl, t, page))

    # Find TOC block: from first toc 1 through last toc before List of Figures body heading
    toc_start = None
    toc_end = None
    for i, p in enumerate(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        t = (p.text or "").strip()
        if toc_start is None and st == "toc 1" and "Declaration" in t:
            toc_start = i
        if toc_start is not None and st == "heading 1" and t == "List of Figures":
            toc_end = i
            break

    if toc_start is None or toc_end is None:
        raise SystemExit("Could not locate TOC block")

    anchor = doc.paragraphs[toc_end]
    for _ in range(toc_end - toc_start):
        delete_paragraph(doc.paragraphs[toc_start])

    # Re-find anchor after deletions
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "List of Figures" and p.style and p.style.name == "Heading 1":
            anchor = p
            break

    for lvl, title, pg in reversed(entries):
        # also add front matter entries manually at top later
        pass

    # Front matter (fixed pages approximate from first section)
    front = [
        (1, "Declaration", 2),
        (1, "Acknowledgment", 2),
        (1, "Abstract", 2),
        (1, "List of Figures", 3),
        (1, "List of Tables", 3),
    ]
    all_entries = front + entries

    for lvl, title, pg in reversed(all_entries):
        insert_before(doc, anchor, f"{title}\t{pg}", toc_style(lvl))

    doc.save(str(PAPER))
    print(f"Rebuilt TOC: {len(all_entries)} entries -> {PAPER}")


if __name__ == "__main__":
    main()
