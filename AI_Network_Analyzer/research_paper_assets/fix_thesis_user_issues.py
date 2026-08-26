"""
Fix TOC page numbers, empty subheadings, table spacing, duplicate numbering, section 4 order.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

WORK = Path(__file__).resolve().parent / "_working_thesis.docx"
DESKTOP = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
PAPER = WORK if WORK.exists() else DESKTOP

TEXT_421 = (
    "This subsection documents the complete input schema used in training and inference. "
    "Every column listed in Table 5 is defined in config.FEATURE_COLUMNS and is extracted "
    "identically during live capture, offline CSV import, and holdout evaluation. "
    "Features were chosen using CICIDS-style NDR domain knowledge (flow duration, byte/packet "
    "statistics, TCP flag counts, host/service rates) rather than post-hoc automatic selection."
)

TEXT_431 = (
    "Each model consumes the same underlying flow CSV but with different tensor shapes and "
    "training objectives. Random Forest — supervised multi-class on 39 scaled features; "
    "n_estimators=200, max_depth=20. XGBoost — supervised multi-class; n_estimators=200, "
    "max_depth=8, learning_rate=0.1. Isolation Forest — trained on Normal traffic only; "
    "contamination=0.1; outputs inlier/outlier mapped to Normal/Attack. Autoencoder — trained "
    "on Normal rows; reconstruction-error threshold at inference; not scored on Table 11 holdout. "
    "Hybrid AI — decision_engine.fuse_decisions() weighted majority fusion over live votes; "
    "no standalone holdout accuracy."
)

TEXT_432 = (
    "LSTM requires sequential input, not a single flow row. training/data_preprocessing.prepare_sequences() "
    "groups consecutive scaled flows with window_size=10 (configurable in train_lstm.py). Each sample "
    "is a (10, 39) tensor and the label is the class of the flow immediately after the window. "
    "Random Forest, XGBoost, and Isolation Forest use one row per flow; only LSTM uses the sliding "
    "window. Holdout metrics for LSTM are not in Table 11 because evaluate_models.py scores row-based "
    "models; LSTM needs a window-aware holdout and is marked implemented but not fully evaluated."
)

TEXT_433 = (
    "Table 5A summarises how each model consumes the same CSV: row-level models (RF, XGB, IF, AE) "
    "use one 39-feature vector per flow; LSTM uses ten consecutive vectors; Hybrid fusion combines "
    "live votes without retraining on a separate dataset."
)

TEXT_43_INTRO = (
    "Model development follows a fixed pipeline: load dataset.csv, apply train-only preprocessing "
    "(duplicate removal, LabelEncoder, StandardScaler), train each model with the configuration below, "
    "then persist artifacts under models/ for FastAPI/Streamlit inference."
)


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    try:
        para.style = anchor.part.document.styles[style]
    except Exception:
        pass
    set_text(para, text)
    pf = para.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    return para


def next_sibling_element(el):
    return el.getnext()


def move_after(element, anchor_el) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    anchor_el.addnext(element)


def heading_has_body(p: Paragraph) -> bool:
    el = p._p
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag.endswith("}tbl"):
            return True
        if nxt.tag.endswith("}p"):
            para = Paragraph(nxt, p._parent)
            t = (para.text or "").strip()
            st = para.style.name if para.style else ""
            if st.startswith("Heading"):
                return False
            if t and len(t) > 25 and not re.match(r"^Table \d+\.", t):
                return True
        nxt = nxt.getnext()
    return False


def ensure_text_after_heading(doc: Document, heading: str, text: str) -> bool:
    p = find_para(doc, heading)
    if not p or heading_has_body(p):
        return False
    insert_after(p, text)
    return True


def fix_duplicate_431(doc: Document) -> int:
    n = 0
    seen = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "4.3.1 Per-Model Training Configuration":
            if seen:
                set_text(p, "4.3.4 Implemented Preprocessing and Training Configuration")
                n += 1
            else:
                seen = True
        elif t == "4.3.1 Implemented Preprocessing and Training Configuration":
            set_text(p, "4.3.4 Implemented Preprocessing and Training Configuration")
            n += 1
    for p in doc.paragraphs:
        st = (p.style.name if p.style else "").lower()
        if st.startswith("toc") and "Implemented Preprocessing" in (p.text or ""):
            set_text(p, "4.3.4 Implemented Preprocessing and Training Configuration\t41")
            n += 1
    return n


def reorder_section_42(doc: Document) -> bool:
    """Move Table 5 block to immediately follow §4.2.1 (before §4.2.2)."""
    h421 = find_para(doc, "4.2.1 Complete Feature Set Used in Experiments")
    h422 = find_para(doc, "4.2.2 Evaluation Dataset Construction")
    if not h421 or not h422:
        return False
    body = doc.element.body
    # collect elements from h422 until 4.3 - find table 5 caption block
    cap_el = None
    tbl_el = None
    desc_el = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "Table 5. Extracted network features (39 holdout features).":
            cap_el = p._p
        if t.startswith("Table 5 lists all 39 features actually passed"):
            desc_el = p._p
    if cap_el is None:
        return False
    # find table after desc or cap
    walk = desc_el or cap_el
    el = walk.getnext()
    while el is not None:
        if el.tag.endswith("}tbl"):
            tbl_el = el
            break
        if el.tag.endswith("}p"):
            pt = Paragraph(el, body).text or ""
            if pt.strip().startswith("Chapter") or pt.strip().startswith("4.3"):
                break
        el = el.getnext()
    if tbl_el is None:
        return False
    # move cap, desc, tbl after h421 (after intro if present)
    anchor = h421._p
    nxt = anchor.getnext()
    if nxt is not None and nxt.tag.endswith("}p"):
        t = Paragraph(nxt, body).text or ""
        if "This subsection documents" in t:
            anchor = nxt
    for block in (cap_el, desc_el, tbl_el):
        if block is not None:
            move_after(block, anchor)
            anchor = block
    return True


def reorder_section_43_content(doc: Document) -> int:
    """Move misplaced paragraphs under correct 4.3.x headings."""
    moved = 0
    h431 = find_para(doc, "4.3.1 Per-Model Training Configuration")
    h432 = find_para(doc, "4.3.2 LSTM Sequence Input Construction")
    h433 = find_para(doc, "4.3.3 Per-Model Dataset Usage Summary")
    if not all([h431, h432, h433]):
        return 0

    # Find orphan content paragraphs by signature
    lstm_p = None
    rf_p = None
    cap5a = None
    desc5a = None
    interp5a = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("LSTM requires sequential input, not a single flow row"):
            lstm_p = p
        elif t.startswith("Random Forest") and "n_estimators=200" in t:
            rf_p = p
        elif t == "Table 5A. Per-model dataset usage summary.":
            cap5a = p
        elif t.startswith("Table 5A summarises how each model"):
            desc5a = p
        elif t.startswith("Interpretation. Supervised RF and XGB share identical"):
            interp5a = p

    body = doc.element.body
    if rf_p and not heading_has_body(h431):
        move_after(rf_p._p, h431._p)
        moved += 1
    if lstm_p and not heading_has_body(h432):
        move_after(lstm_p._p, h432._p)
        moved += 1
    if cap5a:
        anchor = h433._p
        for block in (cap5a._p, desc5a._p if desc5a else None):
            if block is not None:
                move_after(block, anchor)
                anchor = block
                moved += 1
        # move table 5A after desc
        el = anchor.getnext()
        while el is not None:
            if el.tag.endswith("}tbl"):
                move_after(el, anchor)
                anchor = el
                moved += 1
                break
            if el.tag.endswith("}p"):
                pt = (Paragraph(el, body).text or "").strip()
                if pt.startswith("4.3.4") or pt.startswith("4.4"):
                    break
            el = el.getnext()
        if interp5a:
            move_after(interp5a._p, anchor)
            moved += 1
    return moved


def add_table_spacing(doc: Document) -> int:
    """Insert spacer paragraph after each body table (except LoF/LoT)."""
    body = doc.element.body
    count = 0
    tbl_i = 0
    for child in list(body):
        if not child.tag.endswith("}tbl"):
            continue
        if tbl_i < 2:  # skip List of Figures / List of Tables
            tbl_i += 1
            continue
        nxt = child.getnext()
        if nxt is not None and nxt.tag.endswith("}p"):
            para = Paragraph(nxt, body)
            t = (para.text or "").strip()
            pf = para.paragraph_format
            if pf.space_before is None or pf.space_before.pt < 8:
                pf.space_before = Pt(12)
                count += 1
        else:
            spacer = OxmlElement("w:p")
            child.addnext(spacer)
            sp = Paragraph(spacer, body)
            sp.paragraph_format.space_after = Pt(12)
            count += 1
        tbl_i += 1
    return count


def style_table_captions(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if re.match(r"^Table \d+[A-Z]?\.", t):
            pf = p.paragraph_format
            pf.space_before = Pt(10)
            pf.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.italic = False
                r.font.size = Pt(11)
            n += 1
    return n


def fix_toc_static(doc: Document) -> int:
    from ultimate_quality_fix import replace_corrupted_563_toc
    from final_pass import rebuild_ch7_toc_block, rebuild_static_toc

    n = rebuild_static_toc(doc) + rebuild_ch7_toc_block(doc) + replace_corrupted_563_toc(doc)
    # remove duplicate toc 4.3.1 preprocessing if 4.3.4 exists
    seen_prep = False
    for p in list(doc.paragraphs):
        st = (p.style.name if p.style else "").lower()
        t = (p.text or "").strip()
        if st.startswith("toc") and "Implemented Preprocessing" in t:
            if seen_prep or t.startswith("4.3.1 Implemented"):
                delete_paragraph(p)
                n += 1
            else:
                set_text(p, "4.3.4 Implemented Preprocessing and Training Configuration\t41")
                seen_prep = True
    return n


def word_rebuild_toc(path: Path) -> None:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(path.resolve()), ReadOnly=False)
        # Update all fields + TOC page numbers
        wdoc.Fields.Update()
        for i in range(1, wdoc.TablesOfContents.Count + 1):
            wdoc.TablesOfContents(i).Update()
        wdoc.Repaginate()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
        print("Word TOC updated OK")
    except Exception as exc:
        print("Word COM:", exc)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    if not PAPER.exists():
        raise SystemExit(f"Missing {PAPER}")

    doc = Document(str(PAPER))
    stats = {
        "421": ensure_text_after_heading(doc, "4.2.1 Complete Feature Set Used in Experiments", TEXT_421),
        "43_intro": ensure_text_after_heading(doc, "4.3 Model Development", TEXT_43_INTRO),
        "431": ensure_text_after_heading(doc, "4.3.1 Per-Model Training Configuration", TEXT_431),
        "432": ensure_text_after_heading(doc, "4.3.2 LSTM Sequence Input Construction", TEXT_432),
        "433": ensure_text_after_heading(doc, "4.3.3 Per-Model Dataset Usage Summary", TEXT_433),
        "dup431": fix_duplicate_431(doc),
        "reorder42": reorder_section_42(doc),
        "reorder43": reorder_section_43_content(doc),
        "table_space": add_table_spacing(doc),
        "captions": style_table_captions(doc),
    }
    doc.save(str(PAPER))

    fix_toc_static(Document(str(PAPER)))  # pre-fix corruption
    doc = Document(str(PAPER))
    doc.save(str(PAPER))

    word_rebuild_toc(PAPER)

    doc = Document(str(PAPER))
    stats["toc_static"] = fix_toc_static(doc)
    doc.save(str(PAPER))

    # Save to desktop
    try:
        import shutil

        shutil.copy2(PAPER, DESKTOP)
        print("Saved to Desktop:", DESKTOP)
    except Exception as exc:
        print("Desktop copy:", exc)

    print("Fix complete:", stats)


if __name__ == "__main__":
    main()
