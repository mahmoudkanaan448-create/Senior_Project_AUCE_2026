"""
Update the AUCE research paper DOCX to match the implemented AI-NDR application.

Rules followed:
- Keep original writing style and all important existing content
- Add missing implemented features (MITRE, SOAR, Online Learning, Threat Simulation,
  Streamlit, SQLite, Telegram, server supervisor)
- Number references [1]..[n] and place matching inline citations
- Insert expanded diagrams + key code figures
- Stay reasonably compact (target well under 80 pages)
"""

from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

SRC = Path(r"c:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL.docx")
DST = Path(r"c:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx")
ASSETS = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
FIGS = ASSETS / "new_figures"
SNIPS = ASSETS / "code_snippets"


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_in: float = 6.2, caption: str = "") -> Paragraph:
    p = insert_after(paragraph, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    if caption:
        cap = insert_after(p, caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(10)
        return cap
    return p


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def find_para_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def set_para_text(p: Paragraph, text: str) -> None:
    # Keep first run formatting if possible
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def add_citation_markers(doc: Document) -> None:
    """Add [n] markers next to key claims without rewriting whole paragraphs."""
    replacements = [
        ("Random forests", "Random forests [2]"),
        ("Random Forest and XGBoost", "Random Forest [2] and XGBoost [5]"),
        ("Isolation Forest and Autoencoder", "Isolation Forest [15] and Autoencoder"),
        ("Isolation forest", "Isolation forest [15]"),
        ("CICIDS2017", "CICIDS2017 [3][24]"),
        ("NSL-KDD", "NSL-KDD [4]"),
        ("UNSW-NB15", "UNSW-NB15 [17]"),
        ("NIST Cybersecurity Framework", "NIST Cybersecurity Framework [18]"),
        ("scikit-learn", "scikit-learn [20]"),
        ("FastAPI is selected", "FastAPI [33] is selected"),
        ("Explainable AI", "Explainable AI (XAI)"),
        ("signature-based detection", "signature-based detection [14][26]"),
        ("Outside the closed world", "Outside the closed world [26]"),
    ]
    # Apply carefully once per paragraph
    for p in doc.paragraphs:
        t = p.text or ""
        if not t or t.startswith("["):
            continue
        new = t
        for old, rep in replacements:
            if old in new and rep not in new:
                # avoid double-numbering references section itself
                if p.style and p.style.name == "Heading 1" and "References" in t:
                    continue
                # skip reference bibliography lines (author year pattern at start of refs)
                if new.startswith("Ahmad,") or new.startswith("Breiman,") or new.startswith("Canadian") \
                        or new.startswith("Chen,") or new.startswith("Cisco.") or new.startswith("Corelight") \
                        or new.startswith("Darktrace") or new.startswith("Elastic") or new.startswith("ExtraHop") \
                        or new.startswith("Google") or new.startswith("IBM.") or new.startswith("Khraisat") \
                        or new.startswith("Liu,") or new.startswith("Microsoft") or new.startswith("Moustafa") \
                        or new.startswith("National") or new.startswith("Palo Alto") or new.startswith("Pedregosa") \
                        or new.startswith("Rapid7") or new.startswith("Scikit-learn") or new.startswith("Security Onion") \
                        or new.startswith("Sharafaldin") or new.startswith("Snort") or new.startswith("Sommer") \
                        or new.startswith("Splunk") or new.startswith("Suricata") or new.startswith("Trend Micro") \
                        or new.startswith("Vectra") or new.startswith("Wazuh") or new.startswith("Zeek"):
                    continue
                new = new.replace(old, rep, 1)
        if new != t:
            set_para_text(p, new)


def number_references(doc: Document) -> None:
    """Prefix existing references with [n] and append new implementation-related refs."""
    ref_heading = find_para(doc, "References")
    if not ref_heading:
        return

    refs = []
    started = False
    for p in doc.paragraphs:
        # Compare underlying XML nodes (Paragraph wrappers are not stable identities)
        if p._p is ref_heading._p:
            started = True
            continue
        if not started:
            continue
        txt = (p.text or "").strip()
        if not txt:
            continue
        # Skip extras if script re-run
        if txt.startswith("[33]") or txt.startswith("[34]"):
            continue
        refs.append(p)

    for i, p in enumerate(refs, start=1):
        t = (p.text or "").strip()
        if t.startswith("[") and t[1:2].isdigit():
            # already numbered
            continue
        set_para_text(p, f"[{i}] {t}")

    last = refs[-1] if refs else ref_heading
    extra = [
        "[33] FastAPI developers. (n.d.). FastAPI documentation. https://fastapi.tiangolo.com/",
        "[34] Streamlit Inc. (n.d.). Streamlit documentation. https://docs.streamlit.io/",
        "[35] Telegram FZ LLC. (n.d.). Telegram Bot API. https://core.telegram.org/bots/api",
        "[36] MITRE Corporation. (n.d.). MITRE ATT&CK. https://attack.mitre.org/",
        "[37] The SQLite Consortium. (n.d.). SQLite documentation. https://www.sqlite.org/docs.html",
        "[38] Uvicorn developers. (n.d.). Uvicorn - ASGI server. https://www.uvicorn.org/",
        "[39] AbuseIPDB. (n.d.). AbuseIPDB API documentation. https://docs.abuseipdb.com/",
        "[40] Scikit-learn developers. (n.d.). SGDClassifier. https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.SGDClassifier.html",
    ]
    # Avoid duplicating extras on re-run
    existing_tail = "\n".join((p.text or "") for p in doc.paragraphs[-12:])
    if "[40]" in existing_tail and "SGDClassifier" in existing_tail:
        return
    cur = last
    for line in extra:
        cur = insert_after(cur, line)


def update_existing_text(doc: Document) -> None:
    # Scope / stack accuracy
    p = find_para_contains(doc, "The project is designed as an academic Bachelor-level prototype.")
    if p:
        set_para_text(
            p,
            "The project is designed as an academic Bachelor-level prototype. It focuses on packet/flow capture, "
            "feature extraction, Hybrid AI-based detection, Streamlit SOC dashboard visualization [34], Telegram "
            "notifications [35], MITRE ATT&CK enrichment [36], SOAR-style playbooks, online incremental learning, "
            "and response simulation with optional firewall blocking. It does not claim to replace enterprise "
            "commercial systems, but it demonstrates a complete AI-assisted NDR workflow using open-source tools "
            "(FastAPI [33], SQLite [37], scikit-learn [20]).",
        )

    p = find_para_contains(doc, "within the required 40-50 pages")
    if p:
        set_para_text(
            p,
            "The full senior project reviewed twenty cybersecurity systems and platforms. The condensed report keeps "
            "the comparison in a compact form so that the project remains within a practical academic length "
            "(target below 80 pages) while preserving the main research gap and how the proposed system addresses it.",
        )

    p = find_para_contains(doc, "The backend API connects the dashboard")
    if p:
        set_para_text(
            p,
            "The backend API connects the Streamlit dashboard [34], the AI engine, the SQLite database [37], and the "
            "response engine. FastAPI [33] is selected because it is lightweight, suitable for Python AI projects, and "
            "easy to document. The API receives traffic features, calls the model service, stores predictions, and "
            "returns results to the dashboard. Health and SOC endpoints also expose MITRE mappings, playbooks, online "
            "learning status, and auto-heal actions for server readiness.",
        )

    p = find_para_contains(doc, "The dashboard should display total flows")
    if p:
        set_para_text(
            p,
            "The implemented Streamlit dashboard displays total flows, active alerts, average threat score, recent "
            "alerts, protocol distribution, model comparison, Threat Intelligence, blocked IPs, and downloadable "
            "reports. Dedicated pages cover Live Monitoring, AI Detection, Threat Simulation, Threat Intelligence, "
            "Alerts, Blocked IPs, AI Models, Reports, Settings, and SOC Ops (MITRE / SOAR / Online Learning / Health). "
            "This makes the system useful for both technical evaluation and project defense.",
        )

    p = find_para_contains(doc, "The proposed system can be deployed on a local laboratory machine")
    if p:
        set_para_text(
            p,
            "The proposed system can be deployed on a local laboratory machine or on a small server. Local development "
            "uses run.bat (API + Dashboard). Server mode uses run_server.bat with an ops supervisor that restarts "
            "crashed processes, checks /api/v1/health, and runs lightweight auto-heal actions (directories + database "
            "init). The network capture service collects traffic, the backend processes requests, the AI engine "
            "performs predictions, SQLite stores events, Telegram delivers alerts, and the dashboard provides visual "
            "monitoring. For a Bachelor project, local deployment is sufficient, while server mode demonstrates "
            "production readiness.",
        )

    p = find_para_contains(doc, "The administrator starts the database service, then starts the FastAPI backend")
    if p:
        set_para_text(
            p,
            "The administrator starts the system using run.bat (local) or run_server.bat (server/auto-recovery). This "
            "launches the FastAPI backend on port 8000 and the Streamlit dashboard on port 8501. Default login is "
            "admin / admin123. Before starting detection, the administrator should check that the model files, scaler "
            "file, Telegram settings (Settings page), and database connection are available.",
        )

    # Future work: remove items already implemented, keep remaining
    replacements_fw = {
        "Integrate SIEM/SOAR connectors for enterprise environments.":
            "Extend SOAR playbooks with external SIEM connectors for enterprise environments.",
        "Add continuous learning and model drift monitoring.":
            "Extend online learning with formal model-drift monitoring and scheduled revalidation.",
    }
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t in replacements_fw:
            set_para_text(p, replacements_fw[t])

    # Conclusion refresh (keep tone, add implemented modules)
    p = find_para_contains(doc, "This Senior Project proposed an AI-Powered Network Traffic Analyzer")
    if p:
        set_para_text(
            p,
            "This Senior Project proposed an AI-Powered Network Traffic Analyzer & Anomaly Detector designed for "
            "Bachelor-level implementation in Computer Science. The project addresses the limitations of traditional "
            "signature-based IDS by integrating Machine Learning, Deep Learning, anomaly detection, Threat "
            "Intelligence, Explainable AI, MITRE ATT&CK mapping [36], SOAR-style playbooks, Telegram alerts [35], "
            "online incremental learning, automated response, and Streamlit dashboard monitoring [34].",
        )

    p = find_para_contains(doc, "The expected contribution of the project is not only a single ML classifier")
    if p:
        set_para_text(
            p,
            "The expected contribution of the project is not only a single ML classifier, but a complete academic "
            "NDR-style prototype aligned with the implemented software. It demonstrates the full process from traffic "
            "capture and Threat Simulation to feature extraction, Hybrid AI detection, risk scoring, MITRE enrichment, "
            "SOAR response, explanation, storage, visualization, notifications, and optional IP blocking. With real "
            "experimental metrics captured from the running system, this project provides a strong and realistic "
            "graduation project suitable for AUCE Computer Science requirements.",
        )


def insert_new_sections(doc: Document) -> None:
    """Insert implemented-feature sections without deleting existing chapters."""

    # After 3.3 Main Modules
    anchor = find_para(doc, "3.3 Main Modules")
    if anchor:
        # find last content before 3.4
        cur = anchor
        # insert heading + body after a short scan: place right before 3.4
        before = find_para(doc, "3.4 Data Flow Diagram")
        if before:
            # insert just before 3.4 by inserting after previous paragraph of before
            # We insert after 3.3 heading a new subsection block using before's previous
            pass
        h = insert_after(anchor, "3.3.1 Implemented Core Modules (Aligned with Software)", "Heading 3")
        p = insert_after(
            h,
            "The implemented system modules match the software repository. Monitoring captures live connections or "
            "CSV flows and extracts numerical features. The Hybrid AI engine loads Random Forest, XGBoost, Isolation "
            "Forest, Autoencoder, and LSTM models. The Decision Engine fuses votes into label, confidence, threat "
            "score, and severity. Threat Intelligence enriches source IPs. Explainable AI returns feature importance "
            "and recommended actions. The Alert Manager creates Medium+ alerts and executes SOAR playbooks. Telegram "
            "and local Windows notifications inform the analyst. Firewall blocking can write to the database and, when "
            "enabled, apply OS firewall rules. Threat Simulation injects controlled lab campaigns into the live "
            "pipeline. Online learning queues labeled samples and updates an SGDClassifier with partial_fit without "
            "overwriting the five core models. SOC Ops and Settings pages manage MITRE views, playbooks, online status, "
            "and Telegram credentials.",
        )
        # Architecture figure
        img = FIGS / "fig_expanded_architecture.png"
        if img.exists():
            insert_picture_after(p, img, 6.3, "Figure A. Expanded System Architecture aligned with the implemented AI-NDR software.")

    # After 3.6 Threat Model - add MITRE/SOAR section before Chapter 4
    ch4 = find_para(doc, "Chapter 4 - Implementation Methodology")
    if ch4:
        # Insert before chapter 4: get paragraph before by inserting after previous content of 3.6
        # Easiest: insert after STRIDE figure caption if present
        stride = find_para_contains(doc, "Figure 7. Simplified STRIDE Threat Model")
        base = stride or find_para(doc, "3.6 Threat Model") or ch4
        # If base is ch4, we still insert_after something before it - use stride/3.6
        if base is ch4:
            base = find_para(doc, "3.6 Threat Model")
        h = insert_after(base, "3.7 MITRE ATT&CK Mapping and SOAR Playbooks", "Heading 2")
        p1 = insert_after(
            h,
            "To improve SOC interpretability, each attack label is mapped to MITRE ATT&CK tactics and technique IDs "
            "[36]. For example, PortScan maps to Network Service Discovery (T1046), BruteForce maps to Brute Force "
            "(T1110), and Ransomware maps to Data Encrypted for Impact (T1486). The mapping is additive enrichment: it "
            "does not change the detection decision, but it attaches a standard vocabulary used by security analysts.",
        )
        p2 = insert_after(
            p1,
            "SOAR-style playbooks then convert detections into ordered response steps. A typical playbook creates an "
            "alert, enriches MITRE/Threat Intelligence, sends Telegram and local notifications, optionally blocks the "
            "source IP according to severity rules, and queues the sample for online learning. This provides an "
            "academic mini-SOAR workflow without depending on commercial SIEM/SOAR platforms.",
        )
        img = FIGS / "fig_mitre_soar.png"
        cur = p2
        if img.exists():
            cur = insert_picture_after(p2, img, 6.0, "Figure B. MITRE ATT&CK enrichment and SOAR playbook example (Ransomware).")

        h2 = insert_after(cur, "3.8 Online Incremental Learning", "Heading 2")
        insert_after(
            h2,
            "The system stores confirmed/simulated labeled feature vectors in an online buffer. When enough samples "
            "are available, an SGDClassifier is updated using partial_fit [40] and saved as models/online_sgd.pkl. "
            "This online model may contribute an optional sixth vote during detection. Importantly, online learning "
            "does not overwrite Random Forest, XGBoost, Isolation Forest, Autoencoder, or LSTM models, which preserves "
            "stability of the core Hybrid AI pipeline.",
        )

    # Pipeline figure near 4.4 Decision Fusion
    fuse = find_para(doc, "4.4 Decision Fusion")
    if fuse:
        # after the fusion paragraph
        nxt = None
        found = False
        for p in doc.paragraphs:
            if p is fuse:
                found = True
                continue
            if found and (p.text or "").strip():
                nxt = p
                break
        if nxt:
            img = FIGS / "fig_detection_pipeline.png"
            if img.exists():
                insert_picture_after(nxt, img, 6.3, "Figure C. End-to-end detection and response pipeline implemented in software.")

    # After 4.6 dashboard - code snippets + threat simulation
    dash = find_para_contains(doc, "The implemented Streamlit dashboard displays total flows")
    if not dash:
        dash = find_para(doc, "4.6 Dashboard Design")
    if dash:
        h = insert_after(dash, "4.6.1 Threat Simulation and SOC Operations Pages", "Heading 3")
        p = insert_after(
            h,
            "Threat Simulation is a core SOC module, not only a demonstration helper. It generates controlled lab "
            "attacks (DoS, DDoS, PortScan, BruteForce, WebAttack, Mixed, and others), sends them through the live "
            "Hybrid AI pipeline, creates alerts, executes playbooks, and reports Telegram delivery status. The SOC Ops "
            "page presents MITRE mappings, playbook definitions, online-learning buffer status, and server health/"
            "auto-heal controls. Settings stores Telegram bot token and chat ID in the database so notification "
            "credentials remain active at runtime.",
        )
        # Insert selected code figures
        order = [
            (SNIPS / "01_hybrid_ai_decision_fusion.png", "Figure D. Code excerpt: Hybrid AI decision fusion."),
            (SNIPS / "09_mitre_map.png", "Figure E. Code excerpt: MITRE ATT&CK mapping module."),
            (SNIPS / "10_soar_playbooks.png", "Figure F. Code excerpt: SOAR playbook definitions."),
            (SNIPS / "12_online_learning.png", "Figure G. Code excerpt: online incremental learning."),
            (SNIPS / "14_threat_simulation.png", "Figure H. Code excerpt: threat simulation engine."),
            (SNIPS / "13_alert_manager.png", "Figure I. Code excerpt: alert manager with playbook hook."),
            (SNIPS / "16_telegram_alert.png", "Figure J. Code excerpt: Telegram alert delivery."),
            (SNIPS / "15_supervisor.png", "Figure K. Code excerpt: server auto-recovery supervisor."),
        ]
        cur = p
        for path, cap in order:
            if path.exists():
                cur = insert_picture_after(cur, path, 6.0, cap)

    # Chapter 5 attack simulation section update
    sim = find_para(doc, "5.3 Attack Simulation Scenarios")
    if sim:
        insert_after(
            sim,
            "In the implemented software, attack simulation is executed from the Threat Simulation page and from the "
            "API endpoint /api/v1/simulate-attack. Each campaign creates flows, runs Hybrid AI prediction, creates "
            "Medium+ alerts, maps MITRE techniques, runs the matching SOAR playbook, sends Telegram notifications when "
            "configured, and may block Critical/High sources according to playbook policy. Lab source IPs use "
            "documentation ranges (for example 203.0.113.x) to avoid harming real networks.",
        )

    # Chapter 6 deployment figure
    dep = find_para_contains(doc, "Server mode uses run_server.bat")
    if not dep:
        dep = find_para(doc, "6.1 Deployment Architecture")
    if dep:
        img = FIGS / "fig_server_supervisor.png"
        if img.exists():
            insert_picture_after(dep, img, 5.8, "Figure L. Server deployment with supervisor-based auto-recovery.")

    # Chapter 8 additions
    maint = find_para(doc, "8.5 Maintenance Procedures")
    if maint:
        h = insert_after(maint, "8.5.1 Reliability and Auto-Recovery", "Heading 3")
        insert_after(
            h,
            "For continuous operation, prefer run_server.bat. The supervisor monitors API and Dashboard processes, "
            "restarts them after crashes, and performs safe auto-heal steps when health checks fail repeatedly. "
            "Operational logs are written under logs/supervisor.log. Administrators can also open SOC Ops → Server "
            "Health to inspect readiness and trigger heal actions manually.",
        )

    # Demo scenario update
    demo = find_para_contains(doc, "Trigger email or Telegram notification.")
    if demo:
        set_para_text(demo, "Trigger Telegram notification and verify local sound/banner alerts.")
    # Insert extra demo steps near final scenario list
    end_demo = find_para_contains(doc, "Explain the model comparison results.")
    if end_demo:
        a = insert_after(end_demo, "- Open Threat Simulation and launch a Mixed campaign.")
        b = insert_after(a, "- Verify MITRE technique IDs on the created alerts and Telegram message.")
        c = insert_after(b, "- Open SOC Ops and review playbook steps and online-learning buffer growth.")
        insert_after(c, "- On server mode, stop one process and confirm supervisor auto-restart.")


def update_lists(doc: Document) -> None:
    lof = find_para(doc, "List of Figures")
    if lof:
        extras = [
            "Figure A. Expanded System Architecture aligned with the implemented AI-NDR software.",
            "Figure B. MITRE ATT&CK enrichment and SOAR playbook example (Ransomware).",
            "Figure C. End-to-end detection and response pipeline implemented in software.",
            "Figure D. Code excerpt: Hybrid AI decision fusion.",
            "Figure E. Code excerpt: MITRE ATT&CK mapping module.",
            "Figure F. Code excerpt: SOAR playbook definitions.",
            "Figure G. Code excerpt: online incremental learning.",
            "Figure H. Code excerpt: threat simulation engine.",
            "Figure I. Code excerpt: alert manager with playbook hook.",
            "Figure J. Code excerpt: Telegram alert delivery.",
            "Figure K. Code excerpt: server auto-recovery supervisor.",
            "Figure L. Server deployment with supervisor-based auto-recovery.",
        ]
        # Insert after LoF heading (before LoT)
        lot = find_para(doc, "List of Tables")
        cur = lof
        for line in extras:
            cur = insert_after(cur, line)

    acr = find_para(doc, "List of Acronyms")
    if acr:
        extras = [
            "ATT&CK - Adversarial Tactics, Techniques, and Common Knowledge",
            "MITRE - MITRE Corporation knowledge base for adversary behavior",
            "SOAR - Security Orchestration, Automation and Response",
            "SGD - Stochastic Gradient Descent",
            "SOC - Security Operations Center",
            "SQLite - Embedded relational database engine",
        ]
        cur = acr
        for line in extras:
            cur = insert_after(cur, line)

    # Thesis organization update
    org = find_para_contains(doc, "Chapter 1 introduces the project.")
    if org:
        set_para_text(
            org,
            "Chapter 1 introduces the project. Chapter 2 reviews related work and compares twenty applications. "
            "Chapter 3 presents the proposed architecture, including MITRE mapping, SOAR playbooks, and online "
            "learning. Chapter 4 explains implementation methodology and key code modules. Chapter 5 presents testing "
            "and evaluation, including Threat Simulation. Chapter 6 discusses deployment with auto-recovery, "
            "limitations, and future work. Chapter 7 and Chapter 8 provide practical notes and the user manual. The "
            "report ends with conclusion, references, and appendices.",
        )


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Source not found: {SRC}")
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    update_existing_text(doc)
    insert_new_sections(doc)
    update_lists(doc)
    add_citation_markers(doc)
    number_references(doc)

    doc.save(str(DST))
    # Stats
    words = sum(len((p.text or "").split()) for p in doc.paragraphs)
    imgs = sum(1 for r in doc.part.rels.values() if "image" in r.reltype)
    print("Saved:", DST)
    print("paragraphs", len(doc.paragraphs), "words~", words, "images", imgs)


if __name__ == "__main__":
    main()
