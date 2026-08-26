"""Align paper tables/captions with the running app. Template and heading order unchanged."""
from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
BACKUP = PAPER.with_name(PAPER.stem + f"_align_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
EVAL = Path(__file__).resolve().parent / "evaluation"
METRICS = EVAL / "metrics.json"
FIGS = Path(__file__).resolve().parent / "new_figures"
SHOTS = Path(__file__).resolve().parent / "user_guide_shots"


def set_para_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def set_cell(cell: _Cell, text: str) -> None:
    p = cell.paragraphs[0]
    set_para_text(p, text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if not blips:
                continue
            r_id = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_id:
                continue
            part = paras[j].part.related_parts.get(r_id)
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


def row_by_first(table: Table, first: str) -> list | None:
    for row in table.rows:
        if (row.cells[0].text or "").strip() == first:
            return list(row.cells)
    return None


def main() -> None:
    if not PAPER.exists():
        raise SystemExit(f"Missing {PAPER}")
    shutil.copy2(PAPER, BACKUP)
    m = json.loads(METRICS.read_text(encoding="utf-8"))
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    doc = Document(str(PAPER))

    # --- Objectives O5/O6 match the software (no email) ---
    obj = doc.tables[3]
    cells = row_by_first(obj, "O5")
    if cells:
        set_cell(
            cells[1],
            "Improve transparency using stored per-prediction Explainable AI (this-flow evidence and tree importances).",
        )
    cells = row_by_first(obj, "O6")
    if cells:
        set_cell(
            cells[1],
            "Support incident response using IP blocking, Telegram alerts, webhooks, and reports (no email channel).",
        )

    # --- 4.1 technology table = running stack ---
    tech = doc.tables[9]
    tech_map = {
        "Database": "SQLite by default; PostgreSQL optional via DATABASE_URL",
        "AI Libraries": "Scikit-learn, XGBoost, PyTorch (sklearn neural fallback if torch is unavailable)",
        "Traffic Capture": "Scapy/Npcap live NIC, OS connection table, and optional 24/7 capture daemon",
        "Feature Extraction": "Custom 39 CICIDS-style flow features used in training and inference",
        "Explainability": "Stored per-prediction XAI JSON (local evidence + tree importances); SHAP is future work",
        "Dashboard": "Streamlit SOC dashboard, 16 pages, port 8501",
        "Notifications": "Telegram Bot API, local Windows notification, optional webhooks/SIEM (no SMTP email)",
        "Backend": "FastAPI REST API on port 8000",
    }
    for row in tech.rows:
        key = (row.cells[0].text or "").strip()
        if key in tech_map:
            set_cell(row.cells[1], tech_map[key])

    # --- Table 9 metrics: same holdout as Table 11 ---
    perf = doc.tables[19]
    # rows: header, acc, prec, rec, f1, roc
    values = {
        "Accuracy": [pct(rf["accuracy"]), pct(xgb["accuracy"]), pct(iso["accuracy"]), "In engine", "In engine", "Fusion in engine"],
        "Precision": [pct(rf["precision"]), pct(xgb["precision"]), pct(iso["precision"]), "In engine", "In engine", "Fusion in engine"],
        "Recall": [pct(rf["recall"]), pct(xgb["recall"]), pct(iso["recall"]), "In engine", "In engine", "Fusion in engine"],
        "F1 Score": [pct(rf["f1"]), pct(xgb["f1"]), pct(iso["f1"]), "In engine", "In engine", "Fusion in engine"],
        "ROC-AUC": ["—", "—", "—", "—", "—", "—"],
    }
    for row in perf.rows:
        key = (row.cells[0].text or "").strip()
        if key in values:
            for i, val in enumerate(values[key], start=1):
                set_cell(row.cells[i], val)

    # --- Test cases: Telegram/local notify, not email ---
    tests = doc.tables[20]
    cells = row_by_first(tests, "TC7")
    if cells:
        set_cell(cells[1], "Generate local Windows / dashboard notification")
        set_cell(cells[2], "Administrator sees an on-host alert for a Medium+ detection.")
    cells = row_by_first(tests, "TC3")
    if cells:
        set_cell(cells[1], "Run live capture (NIC or 24/7 daemon)")
        set_cell(cells[2], "System stores new flows; Live Evidence separates LAN/public from TEST-NET simulation.")
    cells = row_by_first(tests, "TC10")
    if cells:
        set_cell(cells[1], "Export report / Live Evidence")
        set_cell(cells[2], "System generates a detection/incident export and a live-versus-simulation snapshot.")

    # --- Appendix B outline = paper chapters = defense PPT story ---
    outline = doc.tables[26]
    slides = [
        ("1", "Title: name, AUCE, supervisor, NDR v1.0.0 (same cover facts as the paper)"),
        ("2", "Agenda = paper chapters 1–8"),
        ("3", "Ch.1 Introduction — problem statement"),
        ("4", "Ch.1 Research objectives O1–O6 (same wording as Table 2)"),
        ("5", "Ch.2 Literature gap (20 platforms)"),
        ("6", "Ch.3 Architecture and hybrid AI pipeline"),
        ("7", "Ch.4 Implementation — stack and 16 dashboard pages"),
        ("8", "Ch.4 Screenshots of the running Streamlit SOC"),
        ("9", "Ch.5 Evaluation — 12,000-row CICIDS-style holdout; RF 97.8%, XGB 98.0%, IF 93.9%"),
        ("10", "Ch.6 Deployment, limitations, and future work"),
        ("11", "Conclusion — same claims as the paper conclusion"),
        ("12", "Questions"),
    ]
    for i, (num, content) in enumerate(slides, start=1):
        if i < len(outline.rows):
            set_cell(outline.rows[i].cells[0], num)
            set_cell(outline.rows[i].cells[1], content)

    # --- Final checklist ---
    chk = doc.tables[30]
    cells = row_by_first(chk, "Abstract")
    if cells:
        set_cell(cells[1], "Summarizes problem, solution, models, and measured 12,000-row holdout.")
    cells = row_by_first(chk, "Results")
    if cells:
        set_cell(
            cells[1],
            "Table 11 reports measured holdout metrics. Table 9 uses the same RF/XGB/IF numbers; AE/LSTM/Hybrid are marked in-engine.",
        )
    cells = row_by_first(chk, "References")
    if cells:
        set_cell(cells[1], "Harvard-style numbered references; sources are relevant.")

    # LoT / LoF titles (same numbers)
    for row in doc.tables[1].rows:
        if "Table 9" in (row.cells[0].text or ""):
            set_cell(
                row.cells[1],
                "Model performance comparison on the implemented holdout (RF/XGB/IF measured; AE/LSTM/Hybrid in engine).",
            )
    for row in doc.tables[0].rows:
        if "Figure 14" in (row.cells[0].text or ""):
            set_cell(
                row.cells[1],
                "Model performance comparison on the 12,000-row CICIDS-style holdout.",
            )

    # Body captions / glue text
    p = find_contains(doc, "The testing stage also measures the expected performance")
    if p:
        set_para_text(
            p,
            "Testing verifies that the proposed system can capture traffic, extract features, classify attacks, detect "
            "anomalies, trigger alerts, store results, update the dashboard, and generate reports. Chapter 5 then reports "
            "measured holdout metrics from the implemented trainer (Table 9 compact view; Table 11 detailed view).",
        )

    p = find_contains(doc, "Figure 14. Expected Model Performance Comparison.")
    if p:
        set_para_text(
            p,
            "Figure 14. Model performance comparison on the 12,000-row CICIDS-style holdout (same numbers as Table 9 / Table 11).",
        )

    p = find_contains(doc, "Table 9 remains a design-target comparison")
    if p:
        set_para_text(
            p,
            "Table 9 and Table 11 report the same measured holdout: Random Forest 97.8% accuracy/F1, XGBoost 98.0% "
            "accuracy/F1, Isolation Forest 93.9% as a Normal-versus-Attack proxy (12,000 rows, 39 features, six classes, "
            "75/25 stratified split, random_state=42). Autoencoder, LSTM, and hybrid fusion run in the live engine; they "
            "are marked “in engine” on Table 9 because they are not scored as a separate leaderboard on this holdout. "
            "Figures 12–13 remain schematic examples; Figures 15–17 are the measured evaluation charts.",
        )

    # empty para after 5.5 heading
    for i, para in enumerate(doc.paragraphs):
        if (para.text or "").strip() == "5.5 Confusion Matrix and ROC Curve" and para.style and "Heading" in (para.style.name or ""):
            if i + 1 < len(doc.paragraphs) and not (doc.paragraphs[i + 1].text or "").strip():
                set_para_text(
                    doc.paragraphs[i + 1],
                    "Figures 12 and 13 are schematic examples of how a confusion matrix and ROC curve are read. The measured "
                    "Random Forest confusion matrix from the 12,000-row holdout is Figure 17.",
                )
            break

    p = find_contains(doc, "The implemented Streamlit dashboard (sixteen pages)")
    if p:
        set_para_text(
            p,
            "The implemented Streamlit dashboard (sixteen pages, same order as the sidebar) displays total flows, active "
            "alerts, average threat score, recent alerts, protocol distribution, model comparison, Threat Intelligence, "
            "blocked IPs, and downloadable reports. Pages: (1) Home, (2) Live Monitoring, (3) AI Detection, (4) Threat "
            "Simulation, (5) Threat Intelligence, (6) Alerts, (7) Blocked IPs, (8) AI Models, (9) Reports — Detection, "
            "Incidents, Export, Compliance, Live Evidence, (10) Settings — General, Telegram, AI Config, NDR/Response, "
            "Clear Data, Company, (11) SOC Ops, (12) Incidents, (13) Assets, (14) Threat Hunting, (15) Response, "
            "(16) Copilot. A fixed AI Security Assistant button is available on every authenticated page and uses the "
            "current page context.",
        )

    # TOC leftover title (static text until Word refreshes fields)
    p = find_contains(doc, "5.6 Expected Performance Comparison")
    if p and "Heading" not in ((p.style.name or "") if p.style else ""):
        set_para_text(p, "5.6 Experimental Performance Results\t34")

    swapped = []
    for cap, img in [
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15. Dataset class distribution", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16. Experimental model performance", EVAL / "fig_model_performance.png"),
        ("Figure 17. Confusion matrix of Random Forest", EVAL / "fig_confusion_rf.png"),
        ("Figure 18. Screenshot: Home SOC dashboard", SHOTS / "00_home.png"),
        ("Figure 19. Screenshot: Threat Simulation", SHOTS / "03_simulation.png"),
        ("Figure 20. Screenshot: SOC Ops", SHOTS / "10_soc.png"),
        ("Figure 21. Screenshot: AI Models", SHOTS / "07_models.png"),
        ("Figure 3. Expanded System Architecture", FIGS / "fig_expanded_architecture.png"),
        ("Figure 10. End-to-end detection and response pipeline", FIGS / "fig_detection_pipeline.png"),
        ("Figure 22. Server deployment", FIGS / "fig_server_supervisor.png"),
    ]:
        if replace_image_before_caption(doc, cap, img):
            swapped.append(cap)

    doc.save(str(PAPER))
    print("Updated", PAPER)
    print("Backup", BACKUP)
    print("Images", swapped)


if __name__ == "__main__":
    main()
