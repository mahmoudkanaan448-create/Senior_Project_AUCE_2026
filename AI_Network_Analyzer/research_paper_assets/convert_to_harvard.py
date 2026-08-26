"""Convert paper citations and reference list from numbered APA-like to Harvard style."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
SRC = ROOT / "WORKING_DOCTOR_FEEDBACK.docx"
if not SRC.exists():
    SRC = ROOT / "WORKING_UPDATED.docx"
OUT = ROOT / "WORKING_HARVARD.docx"
DESK = Path(r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx")

# number -> (in_text, harvard_reference_line)
# in_text WITHOUT outer parentheses so we can combine multiples
REFS: dict[int, tuple[str, str]] = {
    1: (
        "Ahmad et al., 2021",
        "Ahmad, Z., Khan, A.S., Wai Shiang, C., Abdullah, J. and Ahmad, F. (2021) 'Network intrusion detection system: A systematic study of machine learning and deep learning approaches', Transactions on Emerging Telecommunications Technologies, 32(1), e4150. Available at: https://doi.org/10.1002/ett.4150",
    ),
    2: (
        "Breiman, 2001",
        "Breiman, L. (2001) 'Random forests', Machine Learning, 45, pp. 5-32. Available at: https://doi.org/10.1023/A:1010933404324",
    ),
    3: (
        "Canadian Institute for Cybersecurity, n.d.a",
        "Canadian Institute for Cybersecurity (n.d.a) Intrusion detection evaluation dataset (CICIDS2017). University of New Brunswick. Available at: https://www.unb.ca/cic/datasets/ids-2017.html (Accessed: 6 August 2026)",
    ),
    4: (
        "Canadian Institute for Cybersecurity, n.d.b",
        "Canadian Institute for Cybersecurity (n.d.b) NSL-KDD dataset. University of New Brunswick. Available at: https://www.unb.ca/cic/datasets/nsl.html (Accessed: 6 August 2026)",
    ),
    5: (
        "Chen and Guestrin, 2016",
        "Chen, T. and Guestrin, C. (2016) 'XGBoost: A scalable tree boosting system', in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. pp. 785-794. Available at: https://doi.org/10.1145/2939672.2939785",
    ),
    6: (
        "Cisco, n.d.a",
        "Cisco (n.d.a) Cisco Secure Network Analytics. Available at: https://www.cisco.com/site/us/en/products/security/secure-network-analytics/index.html (Accessed: 6 August 2026)",
    ),
    7: (
        "Cisco, n.d.b",
        "Cisco (n.d.b) Cisco XDR. Available at: https://www.cisco.com/site/us/en/products/security/xdr/index.html (Accessed: 6 August 2026)",
    ),
    8: (
        "Corelight, n.d.",
        "Corelight (n.d.) Open NDR platform. Available at: https://corelight.com/ (Accessed: 6 August 2026)",
    ),
    9: (
        "Darktrace, n.d.",
        "Darktrace (n.d.) Darktrace ActiveAI Security Platform. Available at: https://darktrace.com/ (Accessed: 6 August 2026)",
    ),
    10: (
        "Elastic, n.d.",
        "Elastic (n.d.) Elastic Security. Available at: https://www.elastic.co/security (Accessed: 6 August 2026)",
    ),
    11: (
        "ExtraHop, n.d.",
        "ExtraHop (n.d.) RevealX network detection and response. Available at: https://www.extrahop.com/products/revealx/ (Accessed: 6 August 2026)",
    ),
    12: (
        "Google Cloud, n.d.",
        "Google Cloud (n.d.) Google Security Operations. Available at: https://cloud.google.com/security/products/security-operations (Accessed: 6 August 2026)",
    ),
    13: (
        "IBM, n.d.",
        "IBM (n.d.) IBM QRadar SIEM. Available at: https://www.ibm.com/products/qradar-siem (Accessed: 6 August 2026)",
    ),
    14: (
        "Khraisat et al., 2019",
        "Khraisat, A., Gondal, I., Vamplew, P. and Kamruzzaman, J. (2019) 'Survey of intrusion detection systems: Techniques, datasets and challenges', Cybersecurity, 2, Article 20. Available at: https://doi.org/10.1186/s42400-019-0038-7",
    ),
    15: (
        "Liu, Ting and Zhou, 2008",
        "Liu, F.T., Ting, K.M. and Zhou, Z.H. (2008) 'Isolation forest', in 2008 Eighth IEEE International Conference on Data Mining. pp. 413-422. Available at: https://doi.org/10.1109/ICDM.2008.17",
    ),
    16: (
        "Microsoft, n.d.",
        "Microsoft (n.d.) Microsoft Sentinel documentation. Available at: https://learn.microsoft.com/azure/sentinel/ (Accessed: 6 August 2026)",
    ),
    17: (
        "Moustafa and Slay, 2015",
        "Moustafa, N. and Slay, J. (2015) 'UNSW-NB15: A comprehensive data set for network intrusion detection systems', in 2015 Military Communications and Information Systems Conference (MilCIS). IEEE, pp. 1-6. Available at: https://doi.org/10.1109/MilCIS.2015.7348942",
    ),
    18: (
        "NIST, 2024",
        "National Institute of Standards and Technology (NIST) (2024) The NIST Cybersecurity Framework (CSF) 2.0 (NIST CSWP 29). Available at: https://doi.org/10.6028/NIST.CSWP.29",
    ),
    19: (
        "Palo Alto Networks, n.d.",
        "Palo Alto Networks (n.d.) Cortex XSIAM. Available at: https://www.paloaltonetworks.com/cortex/cortex-xsiam (Accessed: 6 August 2026)",
    ),
    20: (
        "Pedregosa et al., 2011",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M. and Duchesnay, E. (2011) 'Scikit-learn: Machine learning in Python', Journal of Machine Learning Research, 12, pp. 2825-2830",
    ),
    21: (
        "Rapid7, n.d.",
        "Rapid7 (n.d.) InsightIDR. Available at: https://www.rapid7.com/products/insightidr/ (Accessed: 6 August 2026)",
    ),
    22: (
        "Scikit-learn developers, n.d.a",
        "Scikit-learn developers (n.d.a) IsolationForest. Scikit-learn documentation. Available at: https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html (Accessed: 6 August 2026)",
    ),
    23: (
        "Security Onion Solutions, n.d.",
        "Security Onion Solutions (n.d.) Security Onion documentation. Available at: https://docs.securityonion.net/ (Accessed: 6 August 2026)",
    ),
    24: (
        "Sharafaldin, Habibi Lashkari and Ghorbani, 2018",
        "Sharafaldin, I., Habibi Lashkari, A. and Ghorbani, A.A. (2018) 'Toward generating a new intrusion detection dataset and intrusion traffic characterization', in Proceedings of the 4th International Conference on Information Systems Security and Privacy. pp. 108-116. Available at: https://doi.org/10.5220/0006639801080116",
    ),
    25: (
        "Snort, n.d.",
        "Snort (n.d.) Snort 3. Available at: https://www.snort.org/snort3 (Accessed: 6 August 2026)",
    ),
    26: (
        "Sommer and Paxson, 2010",
        "Sommer, R. and Paxson, V. (2010) 'Outside the closed world: On using machine learning for network intrusion detection', in 2010 IEEE Symposium on Security and Privacy. pp. 305-316. Available at: https://doi.org/10.1109/SP.2010.25",
    ),
    27: (
        "Splunk, n.d.",
        "Splunk (n.d.) Splunk Enterprise Security. Available at: https://www.splunk.com/en_us/products/enterprise-security.html (Accessed: 6 August 2026)",
    ),
    28: (
        "Suricata, n.d.",
        "Suricata (n.d.) Suricata user guide. Available at: https://docs.suricata.io/ (Accessed: 6 August 2026)",
    ),
    29: (
        "Trend Micro, n.d.",
        "Trend Micro (n.d.) Trend Vision One. Available at: https://www.trendmicro.com/en_us/business/products/one-platform.html (Accessed: 6 August 2026)",
    ),
    30: (
        "Vectra AI, n.d.",
        "Vectra AI (n.d.) Vectra AI platform. Available at: https://www.vectra.ai/ (Accessed: 6 August 2026)",
    ),
    31: (
        "Wazuh, n.d.",
        "Wazuh (n.d.) Wazuh documentation. Available at: https://documentation.wazuh.com/ (Accessed: 6 August 2026)",
    ),
    32: (
        "Zeek, n.d.",
        "Zeek (n.d.) Zeek documentation. Available at: https://docs.zeek.org/ (Accessed: 6 August 2026)",
    ),
    33: (
        "FastAPI developers, n.d.",
        "FastAPI developers (n.d.) FastAPI documentation. Available at: https://fastapi.tiangolo.com/ (Accessed: 6 August 2026)",
    ),
    34: (
        "Streamlit Inc., n.d.",
        "Streamlit Inc. (n.d.) Streamlit documentation. Available at: https://docs.streamlit.io/ (Accessed: 6 August 2026)",
    ),
    35: (
        "Telegram FZ LLC, n.d.",
        "Telegram FZ LLC (n.d.) Telegram Bot API. Available at: https://core.telegram.org/bots/api (Accessed: 6 August 2026)",
    ),
    36: (
        "MITRE Corporation, n.d.",
        "MITRE Corporation (n.d.) MITRE ATT&CK. Available at: https://attack.mitre.org/ (Accessed: 6 August 2026)",
    ),
    37: (
        "SQLite Consortium, n.d.",
        "The SQLite Consortium (n.d.) SQLite documentation. Available at: https://www.sqlite.org/docs.html (Accessed: 6 August 2026)",
    ),
    38: (
        "Uvicorn developers, n.d.",
        "Uvicorn developers (n.d.) Uvicorn - ASGI server. Available at: https://www.uvicorn.org/ (Accessed: 6 August 2026)",
    ),
    39: (
        "AbuseIPDB, n.d.",
        "AbuseIPDB (n.d.) AbuseIPDB API documentation. Available at: https://docs.abuseipdb.com/ (Accessed: 6 August 2026)",
    ),
    40: (
        "Scikit-learn developers, n.d.b",
        "Scikit-learn developers (n.d.b) SGDClassifier. Available at: https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.SGDClassifier.html (Accessed: 6 August 2026)",
    ),
    41: (
        "CrowdStrike, n.d.",
        "CrowdStrike (n.d.) CrowdStrike Falcon platform. Available at: https://www.crowdstrike.com/platform/ (Accessed: 6 August 2026)",
    ),
}


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_citations(text: str) -> str:
    """Replace [n] or [n][m]... clusters with Harvard (Author, Year; ...)."""

    def repl_cluster(m: re.Match) -> str:
        nums = [int(x) for x in re.findall(r"\[(\d+)\]", m.group(0))]
        parts = []
        for n in nums:
            if n in REFS:
                parts.append(REFS[n][0])
        if not parts:
            return m.group(0)
        return "(" + "; ".join(parts) + ")"

    # Replace contiguous citation clusters first
    text = re.sub(r"(?:\[\d+\])+", repl_cluster, text)
    return text


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    # 1) Body + all paragraphs except bibliography lines: convert [n]
    refs_started = False
    bib_paras = []
    for p in doc.paragraphs:
        t = p.text or ""
        if t.strip() == "References":
            refs_started = True
            continue
        if refs_started:
            if t.strip():
                bib_paras.append(p)
            continue
        # Declaration APA mention
        if "APA 7" in t:
            set_text(p, t.replace("APA 7", "Harvard"))
            t = p.text or ""
        if "according to APA" in t:
            set_text(p, t.replace("according to APA", "according to Harvard style"))
            t = p.text or ""
        if "[" in t and re.search(r"\[\d+\]", t):
            set_text(p, replace_citations(t))

    # Also tables may contain citations rarely - skip for now

    # 2) Replace bibliography with Harvard alphabetical list (no numbers)
    harvard_lines = sorted((REFS[n][1] for n in REFS), key=lambda s: s.lower())

    # Clear existing bib paragraphs and rewrite
    for i, p in enumerate(bib_paras):
        if i < len(harvard_lines):
            set_text(p, harvard_lines[i])
        else:
            set_text(p, "")

    # If fewer existing paras than needed, append after last
    if bib_paras:
        last = bib_paras[-1]
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        from docx.shared import Pt

        cur = last
        for line in harvard_lines[len(bib_paras):]:
            new_p = OxmlElement("w:p")
            cur._p.addnext(new_p)
            np = Paragraph(new_p, cur._parent)
            run = np.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            cur = np
            bib_paras.append(np)

    # 3) Ensure References heading note
    for p in doc.paragraphs:
        if (p.text or "").strip() == "References":
            # keep heading; optional style note in next empty? skip
            break

    doc.save(str(OUT))
    shutil.copy2(OUT, DESK)
    shutil.copy2(OUT, ROOT / "WORKING_UPDATED.docx")

    # verify
    doc2 = Document(str(OUT))
    body = []
    started = False
    bib = []
    for p in doc2.paragraphs:
        t = (p.text or "").strip()
        if t == "References":
            started = True
            continue
        if not started:
            body.append(t)
        elif t:
            bib.append(t)
    body_txt = "\n".join(body)
    leftover = re.findall(r"\[\d+\]", body_txt)
    print("leftover_numbered_in_body", leftover[:10], "count", len(leftover))
    print("harvard_refs", len(bib))
    print("sample_cite", re.findall(r"\([^)]*\d{4}[^)]*\)", body_txt)[:3])
    print("sample_nd", [b[:70] for b in bib[:3]])
    print("Saved", OUT)
    print("Desktop", DESK)


if __name__ == "__main__":
    main()
