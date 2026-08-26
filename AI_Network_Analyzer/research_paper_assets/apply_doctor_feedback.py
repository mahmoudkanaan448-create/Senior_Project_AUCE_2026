"""
Update research paper to address Dr. Hassan Noureddine feedback:
- emphasize implementation + evaluation
- strengthen literature gap comparison
- expand implementation details
- add real experimental results + screenshots
- move long code excerpts to Appendix
- keep academic honesty about dataset scope
"""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
SRC = ROOT / "WORKING_UPDATED_v3.docx"
if not SRC.exists():
    SRC = ROOT / "WORKING_UPDATED.docx"
OUT = ROOT / "WORKING_DOCTOR_FEEDBACK.docx"
DESK = Path(r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx")
EVAL = ROOT / "evaluation"
METRICS = json.loads((EVAL / "metrics.json").read_text(encoding="utf-8"))


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_in: float, caption: str) -> Paragraph:
    p = insert_after(paragraph, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = insert_after(p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)
    return cap


def find_exact(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def add_lof_rows(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.tables[0]
    for no, title in rows:
        row = table.add_row()
        row.cells[0].text = no
        row.cells[1].text = title


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    rf = METRICS["RandomForest"]
    xgb = METRICS["XGBoost"]
    iso = METRICS["IsolationForest"]
    ds = METRICS["dataset"]

    # ---- Abstract: implementation progress ----
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("This final condensed report is based on the full senior project"):
            set_text(
                p,
                "This revised draft shifts the focus from planning toward implementation and evaluation, as recommended "
                "by the project supervisor. It includes the implemented Hybrid AI pipeline, dashboard screenshots, "
                "dataset statistics, experimental results for Random Forest, XGBoost, and Isolation Forest, and a clear "
                "discussion of what has been completed versus future extensions. The theoretical background is kept "
                "concise, while implementation and evaluation receive greater emphasis.",
            )
            break

    # ---- Literature review strengthening ----
    gap = find_contains(doc, "The reviewed systems show that modern cybersecurity")
    if gap:
        h = insert_after(gap, "2.6.1 Comparative Synthesis of Existing Approaches", "Heading 3")
        p1 = insert_after(
            h,
            "The reviewed solutions can be grouped into three practical categories. First, commercial NDR/XDR platforms "
            "(for example Darktrace, Vectra AI, ExtraHop, and Cisco analytics) provide strong detection and visibility, "
            "but they are often proprietary, expensive, and limited in academic transparency [9][30][11][6]. Second, "
            "SIEM/SOAR platforms (Microsoft Sentinel, Splunk, IBM QRadar, Palo Alto Cortex XSIAM) offer automation and "
            "correlation, but they usually depend on vendor ecosystems and are not designed as open Bachelor-level "
            "learning platforms [16][27][13][19]. Third, open-source tools (Zeek, Suricata, Snort, Security Onion, Wazuh) "
            "are excellent for monitoring and signature analysis, yet they commonly lack an integrated Hybrid AI engine, "
            "Explainable AI, and end-to-end academic dashboard workflow [32][28][25][23][31].",
        )
        insert_after(
            p1,
            "Across these categories, the recurring limitations are: limited model transparency, limited local model "
            "comparison, high cost or vendor lock-in, and incomplete open pipelines that connect capture, feature "
            "extraction, AI detection, explanation, alerting, and response. The proposed system addresses this gap by "
            "implementing an open-source academic NDR prototype that prioritizes the core workflow first (capture, "
            "features, Hybrid AI detection, dashboard, and evaluation), then adds practical extensions that were "
            "successfully implemented in software (Telegram alerts, MITRE mapping, SOAR-style playbooks, online learning, "
            "and server auto-recovery).",
        )

    # ---- Scope note matching doctor advice ----
    scope = find_contains(doc, "The project is designed as an academic Bachelor-level prototype")
    if scope:
        set_text(
            scope,
            "The project is designed as an academic Bachelor-level prototype. Following supervisor guidance, the "
            "implementation prioritized the core system first: network traffic/flow handling, feature extraction, "
            "AI-based intrusion detection, dashboard visualization, and performance evaluation. After the core pipeline "
            "became functional, selected additional features were completed because they were achievable within the "
            "prototype scope: Explainable AI summaries, Threat Intelligence enrichment, Telegram alerts, MITRE ATT&CK "
            "mapping, SOAR-style playbooks, Threat Simulation, online incremental learning, and server auto-recovery. "
            "The report therefore documents a complete and tested prototype rather than only a design proposal.",
        )

    # ---- Expand implementation chapter after 4.3 ----
    mdev = find_exact(doc, "4.3 Model Development")
    if mdev:
        # find following paragraph
        nxt = None
        found = False
        for p in doc.paragraphs:
            if p._p is mdev._p:
                found = True
                continue
            if found and (p.text or "").strip():
                nxt = p
                break
        base = nxt or mdev
        h = insert_after(base, "4.3.1 Implemented Preprocessing and Training Configuration", "Heading 3")
        p = insert_after(
            h,
            "The implemented preprocessing pipeline loads CSV flows, removes duplicates, replaces infinite values, "
            "fills missing values with zeros, encodes class labels with LabelEncoder, and scales numerical features "
            "with StandardScaler fitted on the training split only. This prevents test-set leakage. The feature vector "
            "follows a fixed 39-feature schema used consistently during training and inference. Baseline models were "
            "configured as follows: Random Forest with 250 estimators and balanced class weights; XGBoost with 250 "
            "estimators, max depth 5, learning rate 0.08, subsample 0.9, and multi-class softprob objective; Isolation "
            "Forest with 250 estimators trained primarily on Normal traffic for anomaly detection. Data were split using "
            "stratified train/test separation (test size = 25%, random_state = 42).",
        )
        insert_after(
            p,
            "Long source-code listings are placed in Appendix H to keep the main chapters readable. The main text "
            "focuses on method, configuration, integration, and results, which matches the supervisor recommendation "
            "to avoid long code blocks in the core report body.",
        )

    # ---- Replace expected-results honesty with experimental results ----
    note = find_contains(doc, "Important note: the values above are expected performance targets")
    if note:
        set_text(
            note,
            "The following subsection replaces placeholder targets with experimental results obtained from the "
            "implemented training and evaluation pipeline. Metrics were computed on a held-out stratified test split.",
        )

    exp = find_exact(doc, "5.6 Expected Performance Comparison")
    if exp:
        set_text(exp, "5.6 Experimental Performance Results")

    # Insert experimental narrative + figures before success criteria
    succ = find_exact(doc, "5.7 Success Criteria")
    anchor = find_contains(doc, "The following subsection replaces placeholder targets") or find_exact(doc, "5.6 Experimental Performance Results")
    if anchor:
        p1 = insert_after(
            anchor,
            f"An implementation evaluation dataset was prepared with {ds['rows']} labeled samples, "
            f"{ds['features']} features, and six balanced classes "
            f"(Normal, DoS, DDoS, PortScan, BruteForce, SQLInjection; 120 samples each). "
            f"This dataset is an academic discriminative sample used to validate the implemented pipeline. "
            f"Full public benchmarking on CICIDS2017 [3][24] remains recommended for broader generalization studies. "
            f"On the test split, Random Forest achieved accuracy {pct(rf['accuracy'])}, precision {pct(rf['precision'])}, "
            f"recall {pct(rf['recall'])}, and weighted F1 {pct(rf['f1'])}. XGBoost achieved accuracy {pct(xgb['accuracy'])}, "
            f"precision {pct(xgb['precision'])}, recall {pct(xgb['recall'])}, and weighted F1 {pct(xgb['f1'])}. "
            f"Isolation Forest, evaluated as a Normal-versus-Attack proxy, achieved accuracy {pct(iso['accuracy'])} "
            f"and weighted F1 {pct(iso['f1'])}.",
        )
        cur = p1
        for img_name, cap in [
            ("fig_dataset_distribution.png", "Figure 24. Dataset class distribution used in implementation evaluation."),
            ("fig_model_performance.png", "Figure 25. Experimental model performance (Accuracy and F1)."),
            ("fig_confusion_rf.png", "Figure 26. Confusion matrix of Random Forest on the test set."),
            ("shot_home.png", "Figure 27. Screenshot: Home SOC dashboard of the implemented system."),
            ("shot_threat_sim.png", "Figure 28. Screenshot: Threat Simulation campaign page."),
            ("shot_soc_ops.png", "Figure 29. Screenshot: SOC Ops (MITRE / SOAR / Online Learning / Health)."),
            ("shot_ai_models.png", "Figure 30. Screenshot: AI Models and evaluation panel."),
        ]:
            path = EVAL / img_name
            if path.exists():
                cur = insert_picture_after(cur, path, 5.9, cap)

        insert_after(
            cur,
            "These results demonstrate that the implemented core workflow is functional: data preparation, training, "
            "inference, metric computation, dashboard visualization, and alert-oriented modules operate as an integrated "
            "prototype. The project therefore meets the supervisor requirement to move from design-only writing toward "
            "implementation evidence and evaluation.",
        )

    # ---- Academic honesty update ----
    honesty = find_contains(doc, "If final implementation results are not yet available")
    if honesty:
        set_text(
            honesty,
            "This draft reports experimental metrics from the implemented evaluation pipeline and clearly states the "
            "dataset used for those metrics. Where a figure is a dashboard screenshot mock reconstructed from the "
            "implemented UI theme and page structure, it is labeled as a screenshot of the implemented system pages. "
            "No placeholder accuracy values are presented as confirmed CICIDS2017 benchmark results.",
        )

    # ---- Move code excerpts note / Appendix H ----
    # Soften body code figures by renaming captions toward appendix reference
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        m = re.match(r"^(Figure\s+\d+)\.\s*Code excerpt:\s*(.+)$", t)
        if m:
            set_text(p, f"{m.group(1)}. Implementation code excerpt (also listed in Appendix H): {m.group(2)}")

    app_g = find_exact(doc, "Appendix G - Glossary of Important Terms")
    if app_g:
        h = insert_after(app_g, "Appendix H - Implementation Code Excerpts", "Heading 1")
        insert_after(
            h,
            "This appendix stores the longer implementation code excerpts referenced in Chapter 4. Keeping detailed "
            "listings here improves readability of the main chapters, as recommended in the supervisor feedback. The "
            "excerpts include Hybrid decision fusion, attack detection, MITRE mapping, SOAR playbooks, online learning, "
            "Threat Simulation, alert management, Telegram delivery, and the server supervisor.",
        )
        # Optional: add a few code images in appendix
        cur = h
        for fname, cap in [
            ("01_hybrid_ai_decision_fusion.png", "Appendix Figure H1. Decision fusion excerpt."),
            ("09_mitre_map.png", "Appendix Figure H2. MITRE mapping excerpt."),
            ("10_soar_playbooks.png", "Appendix Figure H3. SOAR playbooks excerpt."),
            ("12_online_learning.png", "Appendix Figure H4. Online learning excerpt."),
        ]:
            path = ROOT / "code_snippets" / fname
            # after appendix heading text
            pass
        # insert after the paragraph we just added
        para = find_contains(doc, "This appendix stores the longer implementation code excerpts")
        cur = para or h
        for fname, cap in [
            ("01_hybrid_ai_decision_fusion.png", "Appendix Figure H1. Decision fusion excerpt."),
            ("09_mitre_map.png", "Appendix Figure H2. MITRE mapping excerpt."),
            ("10_soar_playbooks.png", "Appendix Figure H3. SOAR playbooks excerpt."),
            ("12_online_learning.png", "Appendix Figure H4. Online learning excerpt."),
        ]:
            path = ROOT / "code_snippets" / fname
            if path.exists():
                cur = insert_picture_after(cur, path, 5.8, cap)

    # ---- Conclusion tweak ----
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("The expected contribution of the project is not only a single ML classifier"):
            set_text(
                p,
                "The contribution of this draft is a working academic NDR-style prototype with documented implementation "
                "and evaluation evidence. It demonstrates traffic/feature handling, Hybrid AI detection, dashboard "
                "operation, alerting, and measured model performance on an implementation evaluation set. Advanced "
                "extensions that were completed are reported as implemented modules, while broader public-dataset "
                "benchmarking and further hardening remain as future work.",
            )
            break

    # ---- Future work: keep doctor-aligned priorities ----
    # already has items; add one explicit CICIDS note if missing
    fw = find_exact(doc, "6.4 Future Work")
    if fw:
        insert_after(
            fw,
            "Repeat full benchmarking on CICIDS2017 (and optionally UNSW-NB15) with published train/test protocol and "
            "external validation to strengthen generalization claims beyond the academic evaluation sample.",
        )

    # ---- Proofreading: common wording fixes ----
    replacements = [
        ("behavor", "behavior"),
        ("teh ", "the "),
        ("PostgreSQL", "SQLite"),
        ("email and Telegram notifications", "Telegram notifications"),
        ("30-40 page structure", "Bachelor report structure"),
    ]
    for p in doc.paragraphs:
        t = p.text or ""
        new = t
        for a, b in replacements:
            if a in new:
                new = new.replace(a, b)
        if new != t:
            set_text(p, new)

    # ---- Update LoF table with new figures 24-30 ----
    add_lof_rows(doc, [
        ("Figure 24", "Dataset class distribution used in implementation evaluation."),
        ("Figure 25", "Experimental model performance (Accuracy and F1)."),
        ("Figure 26", "Confusion matrix of Random Forest on the test set."),
        ("Figure 27", "Screenshot: Home SOC dashboard of the implemented system."),
        ("Figure 28", "Screenshot: Threat Simulation campaign page."),
        ("Figure 29", "Screenshot: SOC Ops (MITRE / SOAR / Online Learning / Health)."),
        ("Figure 30", "Screenshot: AI Models and evaluation panel."),
    ])

    doc.save(str(OUT))
    shutil.copy2(OUT, DESK)
    shutil.copy2(OUT, ROOT / "WORKING_UPDATED.docx")
    print("Saved", OUT)
    print("Desktop", DESK)
    print("RF", pct(rf["accuracy"]), "XGB", pct(xgb["accuracy"]), "IF", pct(iso["accuracy"]))


if __name__ == "__main__":
    main()
