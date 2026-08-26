"""Update static TOC page numbers by counting page breaks before each heading."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

WORK = Path(__file__).resolve().parent / "_working_thesis.docx"
DESKTOP = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def has_page_break(p) -> bool:
    return any(br.get(f"{NS_W}type") == "page" for br in p._p.findall(f".//{NS_W}br"))


def norm_key(t: str) -> str:
    return re.sub(r"\s+", " ", t.strip().lower())


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    src = WORK if WORK.exists() else DESKTOP
    doc = Document(str(src))

    page = 1
    pages: dict[str, int] = {}

    for p in doc.paragraphs:
        if has_page_break(p):
            page += 1
        st = p.style.name if p.style else ""
        if not st.startswith("Heading"):
            continue
        t = (p.text or "").strip()
        if not t:
            continue
        key = norm_key(t)
        pages.setdefault(key, page)
        m = re.match(r"^(\d+(?:\.\d+)+)\s", t)
        if m:
            pages.setdefault(m.group(1), page)

    updated = 0
    for p in doc.paragraphs:
        st = (p.style.name if p.style else "").lower()
        if not st.startswith("toc"):
            continue
        raw = (p.text or "").strip()
        if not raw:
            continue
        title = raw.split("\t")[0].strip()
        key = norm_key(title)
        pg = pages.get(key)
        if pg is None:
            m = re.match(r"^(\d+(?:\.\d+)*)", title)
            if m:
                pg = pages.get(m.group(1))
        if pg is None:
            continue
        new = f"{title}\t{pg}"
        if new != raw:
            set_text(p, new)
            updated += 1

    out = DESKTOP
    try:
        doc.save(str(out))
    except PermissionError:
        doc.save(str(src))
        out = src

    check = Document(str(out))
    stale = sum(
        1
        for p in check.paragraphs
        if (p.style.name if p.style else "").lower().startswith("toc")
        and (p.text or "").strip().endswith("\t3")
    )
    print(f"Updated {updated} TOC lines. Stale page-3 entries: {stale}. Saved: {out}")


if __name__ == "__main__":
    main()
