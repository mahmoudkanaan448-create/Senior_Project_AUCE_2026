"""Audit current UPDATED paper: figures, acronyms, references, citations."""
from docx import Document
from pathlib import Path
import re

DOC = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets\WORKING_UPDATED.docx")
OUT = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets\_audit2.txt")

d = Document(str(DOC))
lines = []

def add(s=""):
    lines.append(s)
    print(s)

add(f"paragraphs={len(d.paragraphs)} tables={len(d.tables)}")

# Find LoF / LoA / refs sections
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    st = p.style.name if p.style else ""
    if any(k in t for k in ("List of Figures", "List of Tables", "List of Acronyms", "References", "Figure ", "Fig.")):
        if st.startswith("Heading") or t.startswith("Figure") or t.startswith("List of") or t == "References":
            add(f"[{i}|{st}] {t[:160]}")

add("\n=== FIGURE CAPTIONS IN BODY ===")
figs = []
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    m = re.match(r"^(Figure\s+[A-Z0-9]+[.:)]?\s*.+)$", t, re.I)
    if m:
        figs.append((i, t))
        add(f"{i}: {t[:140]}")

add(f"\nfigure_captions_count={len(figs)}")

add("\n=== ACRONYM-LIKE LINES NEAR LoA ===")
in_loa = False
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    if t == "List of Acronyms":
        in_loa = True
        add(f"START LoA @{i}")
        continue
    if in_loa:
        if p.style and p.style.name.startswith("Heading") and t and t != "List of Acronyms":
            add(f"END LoA @{i} -> {t}")
            break
        if t:
            add(f"  {t[:120]}")

add("\n=== TABLES PREVIEW (first cell) ===")
for ti, table in enumerate(d.tables):
    try:
        c00 = table.rows[0].cells[0].text.strip().replace("\n", " ")[:60]
        cols = len(table.columns)
        rows = len(table.rows)
        add(f"T{ti}: {rows}x{cols} | {c00}")
    except Exception as e:
        add(f"T{ti}: err {e}")

add("\n=== REFERENCES ===")
refs = False
ref_map = {}
for p in d.paragraphs:
    t = (p.text or "").strip()
    if t == "References":
        refs = True
        continue
    if refs and t:
        m = re.match(r"^\[(\d+)\]\s*(.+)$", t)
        if m:
            ref_map[int(m.group(1))] = m.group(2)[:100]
            add(f"[{m.group(1)}] {m.group(2)[:110]}")
        else:
            add(f"UNNUMBERED: {t[:110]}")

add("\n=== INLINE CITATIONS USED ===")
body = "\n".join(p.text or "" for p in d.paragraphs)
# exclude references section for citation scan
body_only = body.split("References")[0] if "References" in body else body
cites = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", body_only)))
add(f"inline_cites={cites}")
missing_defs = [c for c in cites if c not in ref_map]
orphan_defs = [n for n in ref_map if n not in cites]
add(f"cited_but_no_ref={missing_defs}")
add(f"ref_but_never_cited={orphan_defs}")

# Claims that often need citations but may lack them
add("\n=== POSSIBLE UNCITED CLAIMS ===")
needles = [
    "Random Forest", "XGBoost", "Isolation Forest", "LSTM", "Autoencoder",
    "CICIDS", "NSL-KDD", "UNSW", "Streamlit", "FastAPI", "SQLite", "Telegram",
    "MITRE", "ATT&CK", "scikit-learn", "Snort", "Suricata", "Zeek", "Darktrace",
    "NIST", "SOAR", "SGD",
]
for needle in needles:
    hits = []
    for i, p in enumerate(d.paragraphs):
        t = p.text or ""
        if needle in t and "References" not in t[:20]:
            # skip bibliography lines
            if re.match(r"^\[\d+\]", t.strip()):
                continue
            if t.strip().startswith(needle.split()[0]) and "(n.d.)" in t:
                continue
            has_cite = bool(re.search(r"\[\d+\]", t))
            if not has_cite and len(t) > 40:
                hits.append((i, t[:100]))
    if hits:
        add(f"{needle}: {len(hits)} uncited-ish paras")
        for h in hits[:2]:
            add(f"   @{h[0]} {h[1]}")

OUT.write_text("\n".join(lines), encoding="utf-8")
print("WROTE", OUT)
