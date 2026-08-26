"""
Complete Dr. Hassan Noureddine Aug 2026 feedback (19 points) — literal compliance pass.
Runs after doctor_email_revision.py on the staging thesis.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
ASSETS = Path(__file__).resolve().parent
PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
METRICS_PATH = ASSETS / "evaluation" / "metrics.json"

# Re-use helpers from doctor_email_revision
from doctor_email_revision import (  # noqa: E402
    FEATURE_META,
    add_row,
    clear_table_keep_header,
    find_contains,
    find_exact,
    insert_after,
    pct,
    row_by_first,
    set_cell,
    set_text,
)

STATUS_GREEN = "Implemented and tested"
STATUS_YELLOW = "Implemented but not fully evaluated"
STATUS_BLUE = "Proposed / future work"


def slim_chapter7(doc: Document) -> None:
    """Doctor §12: Ch7 brief pointers only — methodology lives in Ch4–5."""
    pointers = {
        "7.1 Dataset Preparation Notes": (
            "See §4.2.2 for dataset source, shape, class distribution, preprocessing, and leakage controls."
        ),
        "7.2 Handling Class Imbalance": (
            "See §4.2.2 and §5.6 for stratified splitting, per-class metrics, and class-weight settings."
        ),
        "7.3 Threshold Tuning": (
            "See §4.3.1 (Isolation Forest contamination, Autoencoder reconstruction threshold) and "
            "Dashboard → Settings → AI Config for runtime thresholds."
        ),
        "7.4 Explainability Strategy": (
            "See §3.3.1 and stored per-prediction XAI JSON on Detection, Alerts, and AI Assistant pages."
        ),
        "7.5 Report Generation": (
            "See §8.4 and Dashboard → Reports (Export PDF/CSV, Live Evidence tabs)."
        ),
    }
    ch7 = find_exact(doc, "Chapter 7 - Practical Implementation Notes")
    if ch7:
        intro_done = find_contains(doc, "This chapter retains brief operator notes only")
        if not intro_done:
            insert_after(
                ch7,
                "This chapter retains brief operator notes only. Dataset construction, feature justification, "
                "model configuration, evaluation metrics, and XAI methodology are documented in Chapters 4–5 "
                "as requested by the supervisor.",
                "Normal",
            )
    for heading, pointer in pointers.items():
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() != heading:
                continue
            if not (p.style and "Heading" in (p.style.name or "")):
                continue
            j = i + 1
            while j < len(doc.paragraphs):
                nxt = doc.paragraphs[j]
                if nxt.style and nxt.style.name.startswith("Heading"):
                    break
                if (nxt.text or "").strip() and not (nxt.text or "").startswith("See §"):
                    set_text(nxt, pointer)
                    # Remove following "Practical note" duplicate if present
                    if j + 1 < len(doc.paragraphs):
                        nxt2 = doc.paragraphs[j + 1]
                        if (nxt2.text or "").startswith("Practical note"):
                            set_text(nxt2, "")
                j += 1
            break


def fix_dfd_wording(doc: Document) -> None:
    """Doctor §14: do not call pipeline diagrams 'Data Flow Diagram' incorrectly."""
    repl = [
        ("3.4 Data Flow Diagram", "3.4 System Data Pipeline"),
        ("Figure 3. Data Flow Diagram", "Figure 3. System Data Pipeline"),
        ("Figure 4. Data Flow Diagram", "Figure 4. System data pipeline (flow extraction → AI engine → storage → dashboard)."),
        ("Data Flow Diagram (DFD Level 1)", "System Data Pipeline (implemented modules)"),
        ("Figure 3. Data Flow Diagram (DFD Level 1)", "Figure 3. System Data Pipeline (implemented modules)."),
    ]
    for p in doc.paragraphs:
        t = p.text or ""
        new = t
        for old, new_s in repl:
            if old in new:
                new = new.replace(old, new_s)
        if new != t:
            set_text(p, new)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    new = t
                    for old, new_s in repl:
                        if old in new:
                            new = new.replace(old, new_s)
                    if new != t:
                        set_text(p, new)


def rebuild_feature_table_full(doc: Document) -> None:
    """Doctor §2: complete feature list with name, meaning, type, relevance, transform, models."""
    old = None
    for t in doc.tables:
        if len(t.rows) > 5 and "duration" in (t.rows[1].cells[0].text or "").lower():
            old = t
            break
    if old is None:
        old = doc.tables[10]
    headers = ["Feature", "Meaning", "Type", "Relevance to NDR", "Transform / scaling", "Used by"]
    nrows = 1 + len(FEATURE_META)
    tbl = doc.add_table(rows=nrows, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(tbl.rows[0].cells[i], h)
    all_models = "RF, XGB, IF, AE (39-d vector); LSTM (10×39 window)"
    for ri, (name, meaning, typ, rel) in enumerate(FEATURE_META, start=1):
        add_row_to = [
            name,
            meaning,
            typ,
            rel,
            "StandardScaler (train fit); protocol_type numeric",
            all_models,
        ]
        for ci, val in enumerate(add_row_to):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    old._tbl.addprevious(tbl._tbl)
    body.remove(old._tbl)
    # Interpretation paragraph after table
    if not find_contains(doc, "Table 5 lists all 39 features actually passed to training"):
        anchor = find_contains(doc, "4.2.1 Complete Feature Set Used in Experiments")
        if anchor:
            for p in doc.paragraphs:
                if "Table 5 lists every feature" in (p.text or ""):
                    insert_after(
                        p,
                        "Table 5 lists all 39 features actually passed to training/inference. The selection strategy "
                        "is domain knowledge from CICIDS-style NDR literature (flow duration, byte/packet statistics, "
                        "TCP flag counts, host/service connection rates) — not an automatic RFE/PCA step applied only "
                        "for the report. Recursive Feature Elimination and PCA were not used because they would break "
                        "train/serve parity with the saved StandardScaler and deployed models.",
                        "Normal",
                    )
                    break


def add_lstm_sequence_section(doc: Document) -> None:
    """Doctor §5: explain LSTM sequence construction."""
    if find_contains(doc, "4.3.2 LSTM Sequence Input Construction"):
        return
    anchor = find_contains(doc, "4.3.1 Per-Model Training Configuration")
    if not anchor:
        anchor = find_exact(doc, "4.3 Model Development")
    if not anchor:
        return
    h = insert_after(anchor, "4.3.2 LSTM Sequence Input Construction", "Heading 3")
    insert_after(
        h,
        "LSTM requires sequential input, not a single flow row. The implementation uses "
        "training/data_preprocessing.prepare_sequences(): after StandardScaler fitting on the training split, "
        "consecutive scaled flow vectors are grouped with window_size=10 (configurable in train_lstm.py). "
        "Each training sample is a tensor of shape (10, 39) — ten consecutive flows × 39 features — and the "
        "label is the class of the flow immediately after the window (the prediction target). "
        "Random Forest, XGBoost, and Isolation Forest consume one row per flow; Autoencoder reconstructs a "
        "single 39-dimensional vector; only LSTM uses the sliding window. At inference, lstm_prediction.py "
        "buffers the last ten scaled flows before calling the saved PyTorch/sklearn sequence model. "
        "Holdout metrics for LSTM are not reported in Table 11 because the stratified evaluation script "
        "(training/evaluate_models.py) scores row-based models on the same split; LSTM evaluation requires "
        "a separate window-aware holdout and is marked as implemented but not fully evaluated.",
        "Normal",
    )


def add_hybrid_fusion_section(doc: Document) -> None:
    """Doctor §6: Hybrid as fusion, not independent classifier."""
    if find_contains(doc, "4.4.1 Hybrid Decision Fusion Pipeline"):
        return
    fusion = find_exact(doc, "4.4 Decision Fusion")
    if not fusion:
        return
    h = insert_after(fusion, "4.4.1 Hybrid Decision Fusion Pipeline", "Heading 3")
    insert_after(
        h,
        "Hybrid AI is not a sixth trained classifier. The live path is: network flow → feature vector → "
        "individual models (RF, XGB, IF, AE, LSTM each produce label/confidence) → decision_engine.fuse_decisions() "
        "→ final_label, confidence, threat_score, severity → alert_manager. Fusion uses majority vote on attack "
        "vs Normal with confidence weighting (see detection/decision_engine.py and Figure 2). "
        "Benefit: when RF and XGB agree on an attack label, false negatives decrease compared with a single model; "
        "when only IF flags anomaly, severity can be raised for analyst review without forcing a wrong class label. "
        "No standalone holdout accuracy is reported for Hybrid because fusion depends on live model agreement, not a "
        "separate trained weight matrix; demonstrating improvement requires comparing single-model vs fused alerts "
        "on Threat Simulation campaigns (qualitative SOC benefit, not a single percentage in Table 11).",
        "Normal",
    )


def add_model_usage_table(doc: Document, m: dict) -> None:
    """Doctor §4: per-model input/train/test/hyperparameters/output."""
    if find_contains(doc, "Table 5A. Per-Model Dataset Usage Summary"):
        return
    anchor = find_contains(doc, "4.3.2 LSTM Sequence Input Construction") or find_contains(
        doc, "4.3.1 Per-Model Training Configuration"
    )
    if not anchor:
        return
    h = insert_after(anchor, "4.3.3 Per-Model Dataset Usage Summary", "Heading 3")
    p = insert_after(
        h,
        "Table 5A summarises how each model consumes the same underlying CSV but different tensor shapes.",
        "Normal",
    )
    ds = m["dataset"]
    rows_n = ds["rows"]
    tbl = doc.add_table(rows=7, cols=7)
    tbl.style = "Table Grid"
    hdr = [
        "Model",
        "Input",
        "Train data",
        "Test data",
        "Key hyperparameters",
        "Output",
        "Holdout result",
    ]
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    data = [
        hdr,
        [
            "Random Forest",
            f"1×39 scaled vector",
            f"75% stratified ({rows_n:,} rows)",
            "25% holdout",
            "n_estimators=200, max_depth=20",
            "Class + probabilities",
            f"Acc {pct(rf['accuracy'])}",
        ],
        [
            "XGBoost",
            f"1×39 scaled vector",
            f"75% stratified",
            "25% holdout",
            "n_estimators=200, max_depth=8, lr=0.1",
            "Class probabilities",
            f"Acc {pct(xgb['accuracy'])}",
        ],
        [
            "Isolation Forest",
            f"1×39 scaled vector",
            "Normal rows only (train)",
            "25% holdout (binary proxy)",
            "n_estimators=200, contamination=0.1",
            "Inlier/outlier → Normal/Attack",
            f"Acc {pct(iso['accuracy'])} (binary)",
        ],
        [
            "Autoencoder",
            f"1×39 scaled vector",
            "Normal rows only (train)",
            "Not scored on Table 11 holdout",
            "Reconstruction threshold",
            "Anomaly score",
            STATUS_YELLOW,
        ],
        [
            "LSTM",
            f"10×39 sequence window",
            "75% rows → sliding windows",
            "Not scored on Table 11 holdout",
            "window_size=10, LSTM 64→32",
            "Class probabilities",
            STATUS_YELLOW,
        ],
        [
            "Hybrid AI",
            "Fused model votes",
            "N/A (runtime fusion)",
            "No separate holdout score",
            "Majority + confidence weighting",
            "final_label, severity",
            "Fusion — see §4.4.1",
        ],
    ]
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    p._p.addnext(tbl._tbl)
    insert_after(
        p,
        "Interpretation. Supervised RF and XGB share identical row-level inputs and the same stratified holdout, "
        "so their metrics are directly comparable. Isolation Forest is trained only on Normal traffic and evaluated "
        "as a binary Normal-vs-Attack proxy, which explains its different score profile. Autoencoder and LSTM are "
        "loaded at inference but were not evaluated with evaluate_models.py on the same split.",
        "Normal",
    )


def add_section_56_metrics(doc: Document, m: dict) -> None:
    """Doctor §8–9: full metrics including ROC-AUC, FPR, FNR + interpretation."""
    anchor = find_exact(doc, "5.6 Experimental Performance Results")
    if not anchor:
        anchor = find_contains(doc, "5.6 Expected Performance Comparison")
        if anchor:
            set_text(anchor, "5.6 Experimental Performance Results")
    if not anchor:
        return

    if find_contains(doc, "5.6.1 Complete Holdout Metrics"):
        _update_table9_roc(doc, m)
        _update_metrics_table_56(doc, m)
        return

    h = insert_after(anchor, "5.6.1 Complete Holdout Metrics (Reproducible)", "Heading 3")
    intro = insert_after(
        h,
        "All values below are produced by python -m training.evaluate_models on datasets/dataset.csv "
        "(stratified 25% test, random_state=42). Macro-averaged FPR/FNR are computed one-vs-rest from "
        "the multi-class confusion matrix for RF/XGB; Isolation Forest uses binary Normal-vs-Attack FPR/FNR.",
        "Normal",
    )

    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    tbl = doc.add_table(rows=4, cols=8)
    tbl.style = "Table Grid"
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "FPR", "FNR"]
    rows = [
        headers,
        [
            "Random Forest",
            pct(rf["accuracy"]),
            pct(rf["precision"]),
            pct(rf["recall"]),
            pct(rf["f1"]),
            f"{rf.get('roc_auc', 0)*100:.2f}%",
            pct(rf.get("fpr", 0)),
            pct(rf.get("fnr", 0)),
        ],
        [
            "XGBoost",
            pct(xgb["accuracy"]),
            pct(xgb["precision"]),
            pct(xgb["recall"]),
            pct(xgb["f1"]),
            f"{xgb.get('roc_auc', 0)*100:.2f}%",
            pct(xgb.get("fpr", 0)),
            pct(xgb.get("fnr", 0)),
        ],
        [
            "Isolation Forest",
            pct(iso["accuracy"]),
            pct(iso["precision"]),
            pct(iso["recall"]),
            pct(iso["f1"]),
            f"{iso.get('roc_auc', 0)*100:.2f}%" if iso.get("roc_auc") else "Binary proxy",
            pct(iso.get("fpr", 0)),
            pct(iso.get("fnr", 0)),
        ],
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    intro._p.addnext(tbl._tbl)

    insert_after(
        intro,
        "Interpretation. Weighted F1 exceeds 97% for RF/XGB because the evaluation set is synthetic CICIDS-style "
        "data with separable attack signatures and strict leakage controls (train-only scaler, deduplication, no "
        "label in features). XGBoost slightly outperforms RF on minority classes such as SQLInjection. Isolation "
        "Forest shows lower multi-class proxy scores because it was trained only on Normal traffic. From a SOC "
        "perspective, supervised models drive labelled alerts; IF adds an unsupervised signal for unknown patterns. "
        "Confusion between DoS and DDoS appears in the RF matrix (Figure 17) because volumetric features overlap — "
        "analysts should treat severity and MITRE context alongside the raw label.",
        "Normal",
    )

    if not find_contains(doc, "5.6.2 High-Score Justification and Leakage Checklist"):
        h2 = insert_after(
            find_contains(doc, "From a SOC perspective, supervised models drive labelled alerts"),
            "5.6.2 High-Score Justification and Leakage Checklist",
            "Heading 3",
        )
        insert_after(
            h2,
            "The supervisor asked why scores exceed 95%. Verified controls: (1) LabelEncoder and StandardScaler "
            "fit on training split only; (2) duplicate rows removed before split; (3) label column excluded from "
            "features; (4) stratified holdout preserves class ratios; (5) dataset is synthetic with known signatures "
            "— not an official CICIDS2017 leaderboard dump. These controls are documented in §4.2.2 and reproduced "
            "by the evaluation script.",
            "Normal",
        )

    _update_table9_roc(doc, m)


def _update_metrics_table_56(doc: Document, m: dict) -> None:
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    for t in doc.tables:
        if len(t.rows) < 4:
            continue
        hdr = [(c.text or "").strip() for c in t.rows[0].cells]
        if hdr[:8] == ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "FPR", "FNR"]:
            data = {
                "Random Forest": rf,
                "XGBoost": xgb,
                "Isolation Forest": iso,
            }
            for row in t.rows[1:]:
                name = (row.cells[0].text or "").strip()
                if name not in data:
                    continue
                mod = data[name]
                vals = [
                    pct(mod["accuracy"]),
                    pct(mod["precision"]),
                    pct(mod["recall"]),
                    pct(mod["f1"]),
                    f"{mod.get('roc_auc', 0)*100:.2f}%" if mod.get("roc_auc") else "N/A",
                    pct(mod.get("fpr", 0)),
                    pct(mod.get("fnr", 0)),
                ]
                for i, val in enumerate(vals, start=1):
                    if i < len(row.cells):
                        set_cell(row.cells[i], val)
            return


def _update_table9_roc(doc: Document, m: dict) -> None:
    rf, xgb, iso = m["RandomForest"], m["XGBoost"], m["IsolationForest"]
    perf = None
    for t in doc.tables:
        if len(t.rows) >= 2 and (t.rows[0].cells[0].text or "").strip() == "Metric":
            hdr = [(c.text or "").strip() for c in t.rows[0].cells]
            if "Hybrid AI" in hdr:
                perf = t
                break
    if perf is None:
        return
    for row in perf.rows:
        key = (row.cells[0].text or "").strip()
        if key != "ROC-AUC":
            continue
        vals = [
            f"{rf.get('roc_auc', 0)*100:.2f}%" if rf.get("roc_auc") else "See §5.6.1",
            f"{xgb.get('roc_auc', 0)*100:.2f}%" if xgb.get("roc_auc") else "See §5.6.1",
            f"{iso.get('roc_auc', 0)*100:.2f}%" if iso.get("roc_auc") else "Binary proxy",
        ]
        for i, val in enumerate(vals, start=1):
            if i < len(row.cells):
                set_cell(row.cells[i], val)
        break


def add_implementation_status_table(doc: Document) -> None:
    """Doctor §18: green/yellow/blue status for major components."""
    if find_contains(doc, "Table 3A. Component Implementation Status"):
        return
    anchor = find_contains(doc, "3.3.2 Implementation and Evaluation Status")
    if not anchor:
        return
    h = insert_after(anchor, "Table 3A. Component Implementation Status", "Heading 3")
    p = insert_after(
        h,
        "Legend: "
        + STATUS_GREEN
        + " | "
        + STATUS_YELLOW
        + " | "
        + STATUS_BLUE,
        "Normal",
    )
    tbl = doc.add_table(rows=12, cols=3)
    tbl.style = "Table Grid"
    rows = [
        ["Component", "Status", "Evidence"],
        ["Random Forest / XGBoost / IF", STATUS_GREEN, "Holdout metrics §5.6.1; models on disk"],
        ["Autoencoder", STATUS_YELLOW, "Loaded at inference; no Table 11 holdout row"],
        ["LSTM sequences", STATUS_YELLOW, "window_size=10; no Table 11 holdout row"],
        ["Hybrid fusion", STATUS_YELLOW, "Live fuse_decisions(); Threat Simulation alerts"],
        ["Explainable AI (stored JSON)", STATUS_GREEN, "Detection/Alerts/Copilot pages"],
        ["Threat Intelligence / MITRE", STATUS_GREEN, "mitre_map.py; SOC Ops tab tested"],
        ["SOAR playbooks", STATUS_GREEN, "alert_manager playbook hooks"],
        ["Online incremental learning", STATUS_YELLOW, "SOC Ops tab; does not replace core models"],
        ["Threat Simulation", STATUS_GREEN, "Campaign tests on TEST-NET IPs"],
        ["AI Security Assistant", STATUS_YELLOW, "Copilot page; context from DB, not separate ML"],
        ["Email / SMTP alerts", STATUS_BLUE, "Not implemented; Telegram used instead"],
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    p._p.addnext(tbl._tbl)
    insert_after(
        p,
        "Interpretation. Architecture diagrams match the running Streamlit/FastAPI codebase: every green item was "
        "observed during Threat Simulation or holdout evaluation. Yellow items exist in software but lack the same "
        "reproducible holdout script as RF/XGB/IF. Blue items remain future work.",
        "Normal",
    )


def add_table9_interpretation(doc: Document) -> None:
    """Doctor §16: paragraph after Table 9, not table-only."""
    if find_contains(doc, "Table 9 reports measured holdout metrics only for RF"):
        return
    for p in doc.paragraphs:
        if "Table 9. Model performance metrics" in (p.text or "") or (
            (p.text or "").strip().startswith("Table 9") and "performance" in (p.text or "").lower()
        ):
            insert_after(
                p,
                "Table 9 reports measured holdout metrics only for RF, XGBoost, and Isolation Forest. "
                "Autoencoder, LSTM, and Hybrid AI rows intentionally show 'Not on holdout' because those "
                "components were not scored with training/evaluate_models.py. Earlier placeholder percentages "
                "(e.g., 99.2% Hybrid) were removed.",
                "Normal",
            )
            break


def expand_appendix_b_defense(doc: Document) -> None:
    """Doctor §19: complete demo path + Q&A prep."""
    outline = doc.tables[26]
    slides = [
        ("1", "Title — AUCE — Mahmoud Talal Kanaan — Dr. Hasan Noureddine"),
        ("2", "Problem → Methodology → Design → Implementation → Experiments → Results → Discussion"),
        ("3", f"Dataset: {json.loads(METRICS_PATH.read_text())['dataset']['rows']:,} rows, 39 features, stratified holdout"),
        ("4", "Models: RF/XGB/IF measured; AE/LSTM/Hybrid status honest"),
        (
            "5",
            "Live demo: Traffic/CSV → Features → Preprocess → Models → Hybrid fusion → Alert → XAI → "
            "MITRE → SOAR/Telegram → DB → Dashboard",
        ),
        ("6", "Attack demo: Threat Simulation DoS + anomaly/IF noisy flow scenario"),
        (
            "7",
            "Q&A: Why these models/features? Why 12k not 720? Leakage controls? Why high scores? "
            "Hybrid benefit? Limitations?",
        ),
    ]
    for i, (num, content) in enumerate(slides, start=1):
        if i < len(outline.rows):
            set_cell(outline.rows[i].cells[0], num)
            set_cell(outline.rows[i].cells[1], content)


def ensure_attack_matrix_table(doc: Document) -> None:
    """Doctor §10: structured attack table."""
    if find_contains(doc, "Tested (Threat Simulation)"):
        return
    anchor = find_contains(doc, "5.3.1 Structured Attack Testing Matrix")
    if not anchor:
        return
    rows = [
        ["Attack", "Model(s)", "Expected", "Alert/Response", "Status"],
        ["DoS/DDoS", "RF + XGB", "Attack label", "Medium+ alert, playbook, Telegram", "Tested (Threat Simulation)"],
        ["PortScan", "RF", "PortScan", "High alert, TI enrich", "Tested"],
        ["BruteForce", "RF/XGB", "BruteForce", "Alert + optional block", "Tested"],
        ["SQLInjection", "RF/XGB", "SQLInjection", "Alert + MITRE", "Tested"],
        ["Malware/C2", "IF + TI", "Anomaly / suspicious IP", "Analyst review", "Discussed / partial TI"],
        ["Zero-day", "IF + AE", "Anomaly score high", "Medium alert, no forced label", "Simulated noise only"],
    ]
    tbl = doc.add_table(rows=len(rows), cols=5)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    anchor._p.addnext(tbl._tbl)
    insert_after(
        anchor,
        "Interpretation. Rows marked Tested were executed via Threat Simulation on RFC 5737 TEST-NET IPs. "
        "Malware/C2 and zero-day rows remain theoretical or partial because no live malware sample was injected.",
        "Normal",
    )


def page_break_before_chapters(doc: Document) -> None:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if p.style and p.style.name == "Heading 1" and t.startswith("Chapter "):
            prev = p._p.getprevious()
            if prev is not None and prev.tag.endswith("p"):
                # avoid double breaks
                pass
            br = p.insert_paragraph_before("")
            br.add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    paper = Path(sys.argv[1]) if len(sys.argv) > 1 else PAPER
    m = json.loads(METRICS_PATH.read_text(encoding="utf-8"))

    doc = Document(str(paper))

    fix_dfd_wording(doc)
    rebuild_feature_table_full(doc)
    add_lstm_sequence_section(doc)
    add_hybrid_fusion_section(doc)
    add_model_usage_table(doc, m)
    add_section_56_metrics(doc, m)
    add_implementation_status_table(doc)
    add_table9_interpretation(doc)
    slim_chapter7(doc)
    ensure_attack_matrix_table(doc)
    expand_appendix_b_defense(doc)
    page_break_before_chapters(doc)

    # Fix figure captions
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if "Data Flow Diagram" in t and t.startswith("Figure"):
            set_text(p, t.replace("Data Flow Diagram", "System Data Pipeline"))

    doc.save(str(paper))
    print("Doctor complete revision saved:", paper)


if __name__ == "__main__":
    main()
