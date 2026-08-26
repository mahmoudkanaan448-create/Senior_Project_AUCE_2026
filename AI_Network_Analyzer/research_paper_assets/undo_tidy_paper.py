"""Undo the last tidy_paper_layout pass on the Desktop thesis (same file)."""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
ROOT = Path(__file__).resolve().parent
EVAL = ROOT / "evaluation"
SHOTS = ROOT / "user_guide_shots"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Static TOC exactly as before the tidy step (incl. the old 5.6 line).
OLD_TOC = [
    ("toc 1", "Declaration\t3"),
    ("toc 1", "Acknowledgment\t4"),
    ("toc 1", "Abstract\t5"),
    ("toc 1", "List of Figures\t8"),
    ("toc 1", "List of Tables\t9"),
    ("toc 1", "List of Acronyms\tError! Bookmark not defined."),
    ("toc 1", "Chapter 1 - Introduction\t11"),
    ("toc 2", "1.1 Background\t11"),
    ("toc 2", "1.2 Problem Statement\t11"),
    ("toc 2", "1.3 Proposed Solution\t11"),
    ("toc 2", "1.4 Research Objectives\t11"),
    ("toc 2", "1.5 Research Questions\t11"),
    ("toc 2", "1.6 Scope\t11"),
    ("toc 2", "1.7 Methodology\t12"),
    ("toc 2", "1.8 Thesis Organization\t12"),
    ("toc 1", "Chapter 2 - Literature Review and State of the Art\t13"),
    ("toc 2", "2.1 Traditional Intrusion Detection Systems\t13"),
    ("toc 2", "2.2 AI-Based Intrusion Detection\t13"),
    ("toc 2", "2.3 Machine Learning and Deep Learning Techniques\t13"),
    ("toc 2", "2.4 Datasets Used in Network Intrusion Detection\t13"),
    ("toc 2", "2.5 State of the Art: 20 Applications\t13"),
    ("toc 2", "2.6 Literature Gap\t14"),
    ("toc 2", "2.7 Detailed State-of-the-Art Gap Analysis\t15"),
    ("toc 3", "Darktrace\t15"),
    ("toc 3", "Vectra AI\t15"),
    ("toc 3", "ExtraHop RevealX\t15"),
    ("toc 3", "Corelight Open NDR\t15"),
    ("toc 3", "Cisco Secure Network Analytics\t16"),
    ("toc 3", "Palo Alto Cortex XSIAM\t16"),
    ("toc 3", "Microsoft Sentinel\t16"),
    ("toc 3", "Splunk Enterprise Security\t16"),
    ("toc 3", "IBM QRadar\t17"),
    ("toc 3", "Elastic Security\t17"),
    ("toc 3", "CrowdStrike Falcon\t17"),
    ("toc 3", "Trend Micro Vision One\t17"),
    ("toc 3", "Google Security Operations\t18"),
    ("toc 3", "Rapid7 InsightIDR\t18"),
    ("toc 3", "Security Onion 2\t18"),
    ("toc 3", "Zeek\t18"),
    ("toc 3", "Suricata\t19"),
    ("toc 3", "Wazuh\t19"),
    ("toc 3", "Snort 3\t19"),
    ("toc 3", "Cisco XDR\t19"),
    ("toc 1", "Chapter 3 - Proposed System Architecture\t20"),
    ("toc 2", "3.1 Architecture Overview\t20"),
    ("toc 2", "3.2 AI Detection Pipeline\t20"),
    ("toc 2", "3.3 Main Modules\t21"),
    ("toc 2", "3.4 Data Flow Diagram\t23"),
    ("toc 2", "3.5 UML Diagrams\t23"),
    ("toc 2", "3.6 Threat Model\t25"),
    ("toc 1", "Chapter 4 - Implementation Methodology\t28"),
    ("toc 2", "4.1 Technology Stack\t28"),
    ("toc 2", "4.2 Feature Engineering\t28"),
    ("toc 2", "4.3 Model Development\t28"),
    ("toc 2", "4.4 Decision Fusion\t28"),
    ("toc 2", "4.5 Database Design\t29"),
    ("toc 2", "4.6 Dashboard Design\t29"),
    ("toc 2", "4.7 Implementation Workflow\t30"),
    ("toc 2", "4.8 Backend API Design\t31"),
    ("toc 2", "4.9 Algorithmic Workflow\t31"),
    ("toc 2", "4.10 Model Training Pseudocode\t31"),
    ("toc 2", "4.11 Project Folder Structure\t31"),
    ("toc 1", "Chapter 5 - Testing and Evaluation\t32"),
    ("toc 2", "5.1 Testing Objectives\t32"),
    ("toc 2", "5.2 Testing Environment\t32"),
    ("toc 2", "5.3 Attack Simulation Scenarios\t32"),
    ("toc 2", "5.4 Evaluation Metrics\t32"),
    ("toc 2", "5.5 Confusion Matrix and ROC Curve\t33"),
    ("toc 2", "5.6 Expected Performance Comparison\t345.6 Experimental Performance Results\t34"),
    ("toc 2", "5.7 Success Criteria\t39"),
    ("toc 2", "5.8 Detailed Test Cases\t39"),
    ("toc 2", "5.9 Data Validation Checklist\t39"),
    ("toc 2", "5.10 Error Handling and Reliability\t40"),
    ("toc 2", "5.11 Academic Honesty in Results\t40"),
    ("toc 1", "Chapter 6 - Deployment, Limitations and Future Work\t41"),
    ("toc 2", "6.1 Deployment Architecture\t41"),
    ("toc 2", "6.2 Security Considerations\t41"),
    ("toc 2", "6.3 Limitations\t41"),
    ("toc 2", "6.4 Future Work\t41"),
    ("toc 2", "6.5 Ethical and Legal Considerations\t42"),
    ("toc 1", "Chapter 7 - Practical Implementation Notes\t43"),
    ("toc 2", "7.1 Dataset Preparation Notes\t43"),
    ("toc 2", "7.2 Handling Class Imbalance\t43"),
    ("toc 2", "7.3 Threshold Tuning\t43"),
    ("toc 2", "7.4 Explainability Strategy\t43"),
    ("toc 2", "7.5 Report Generation\t43"),
    ("toc 2", "7.6 Final Demonstration Scenario\t43"),
    ("toc 1", "Chapter 8 - User Manual and Operation Guide\t45"),
    ("toc 2", "8.1 Starting the System\t45"),
    ("toc 2", "8.2 Uploading Data or Capturing Live Traffic\t45"),
    ("toc 2", "8.3 Reading the Dashboard\t45"),
    ("toc 2", "8.4 Responding to Alerts\t45"),
    ("toc 2", "8.5 Maintenance Procedures\t45"),
    ("toc 2", "8.6 Minimal Functional Requirements\t45"),
    ("toc 2", "8.7 Non-Functional Requirements\t46"),
    ("toc 1", "Conclusion\t46"),
    ("toc 1", "Appendix A - Defense Preparation Notes\t47"),
    ("toc 1", "Appendix B - Presentation Outline\t47"),
    ("toc 1", "Appendix C - Additional Defense Questions\t47"),
    ("toc 1", "Appendix D - Implementation Timeline\t48"),
    ("toc 1", "Appendix E - Practical Code Skeleton\t49"),
    ("toc 1", "Appendix F - Final Submission Checklist\t49"),
    ("toc 1", "Appendix G - Glossary of Important Terms\t49"),
    ("toc 1", "References\t58"),
]


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_empty_heading(parent_el, style_name: str = "Heading 1") -> None:
    new_p = OxmlElement("w:p")
    parent_el.insert(parent_el.index(parent_el.getchildren()[-1]) if False else 0, new_p)
    # append at end handled by caller


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 6), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            r_id = blips[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            part = paras[j].part.related_parts.get(r_id) if r_id else None
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def restore_toc(doc: Document) -> None:
    for p in list(doc.paragraphs):
        st = p.style.name if p.style else ""
        if st.lower().startswith("toc"):
            delete_paragraph(p)

    anchor_idx = next(
        i for i, p in enumerate(doc.paragraphs) if (p.text or "").startswith("Keywords:")
    )
    insert_after = doc.paragraphs[anchor_idx + 1]
    for style_name, text in OLD_TOC:
        new_p = OxmlElement("w:p")
        insert_after._p.addnext(new_p)
        para = Paragraph(new_p, insert_after._parent)
        try:
            para.style = doc.styles[style_name]
        except Exception:
            para.style = doc.styles["Normal"]
        para.add_run(text)
        insert_after = para


def restore_front_matter_spacing(doc: Document) -> None:
    # Empty Heading 1 placeholders before Acknowledgment (as before tidy)
    ack_idx = next(i for i, p in enumerate(doc.paragraphs) if (p.text or "").strip() == "Acknowledgment")
    for _ in range(2):
        new_p = OxmlElement("w:p")
        doc.paragraphs[ack_idx]._p.addprevious(new_p)
        Paragraph(new_p, doc.paragraphs[ack_idx]._parent).style = doc.styles["Heading 1"]

    abs_idx = next(i for i, p in enumerate(doc.paragraphs) if (p.text or "").strip() == "Abstract")
    for _ in range(2):
        new_p = OxmlElement("w:p")
        doc.paragraphs[abs_idx]._p.addprevious(new_p)
        Paragraph(new_p, doc.paragraphs[abs_idx]._parent).style = doc.styles["Heading 1"]

    # Before List of Figures: empty Heading 1 (as before tidy)
    lof_idx = next(i for i, p in enumerate(doc.paragraphs) if (p.text or "").strip() == "List of Figures")
    new_p = OxmlElement("w:p")
    doc.paragraphs[lof_idx]._p.addprevious(new_p)
    Paragraph(new_p, doc.paragraphs[lof_idx]._parent).style = doc.styles["Heading 1"]

    # Extra page break before Chapter 1
    ch1_idx = next(i for i, p in enumerate(doc.paragraphs) if (p.text or "").startswith("Chapter 1 - Introduction"))
    prev = doc.paragraphs[ch1_idx - 1]
    if not any(br.get(f"{NS_W}type") == "page" for br in prev._p.findall(f".//{NS_W}br")):
        prev.add_run().add_break(WD_BREAK.PAGE)


def remove_figure12_image(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs):
        if not (p.text or "").startswith("Figure 12."):
            continue
        if i == 0:
            break
        prev = doc.paragraphs[i - 1]
        if prev._p.findall(f".//{NS_A}blip") and not (prev.text or "").strip():
            delete_paragraph(prev)
        break


def restore_appendix_spacing(doc: Document) -> None:
    idx = next(
        (i for i, p in enumerate(doc.paragraphs) if (p.text or "").startswith("Appendix A - Defense")),
        None,
    )
    if idx is None:
        return
    anchor = doc.paragraphs[idx]
    for _ in range(18):
        new_p = OxmlElement("w:p")
        anchor._p.addprevious(new_p)
        Paragraph(new_p, anchor._parent)


def main() -> None:
    subprocess.run(
        [sys.executable, str(ROOT / "render_eval_figures.py")],
        check=True,
        cwd=str(ROOT),
    )

    doc = Document(str(PAPER))
    remove_figure12_image(doc)

    swapped = []
    mapping = [
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", SHOTS / "00_home.png"),
        ("Figure 19.", SHOTS / "03_simulation.png"),
        ("Figure 20.", SHOTS / "10_soc.png"),
        ("Figure 21.", SHOTS / "07_models.png"),
        ("Figure 11.", EVAL / "shot_home.png"),
    ]
    for cap, img in mapping:
        if replace_image_before_caption(doc, cap, img):
            swapped.append(cap.split(".")[0])

    restore_toc(doc)
    restore_front_matter_spacing(doc)
    restore_appendix_spacing(doc)

    doc.save(str(PAPER))
    print("Reverted tidy on", PAPER)
    print("figures", swapped)


if __name__ == "__main__":
    main()
