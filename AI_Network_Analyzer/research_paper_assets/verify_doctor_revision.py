"""Verify all 19 doctor email requirements on the Desktop thesis."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

PAPER = Path(
    r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
)
METRICS = Path(__file__).resolve().parent / "evaluation" / "metrics.json"


def text(doc: Document) -> str:
    parts = [(p.text or "") for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts).lower()


def run_checks(paper: Path) -> tuple[int, list]:
    doc = Document(str(paper))
    blob = text(doc)
    m = json.loads(METRICS.read_text(encoding="utf-8"))
    rf = f"{m['RandomForest']['accuracy']*100:.1f}"

    checks = [
        ("1_dataset_12k", "12,000" in blob and "4.2.2" in blob, "dataset construction §4.2.2"),
        ("2_features_6col", "relevance to ndr" in blob and "transform / scaling" in blob, "full feature table"),
        ("3_feature_justification", "domain knowledge" in blob and "rfe" in blob, "feature selection justified"),
        ("4_model_usage_table", "per-model dataset usage" in blob, "model usage summary"),
        ("5_lstm_sequences", "window_size=10" in blob and "prepare_sequences" in blob, "LSTM sequence explained"),
        ("6_hybrid_fusion", "4.4.1" in blob and "fuse_decisions" in blob, "Hybrid fusion pipeline"),
        (
            "7_no_fake_hybrid_992",
            "hybrid ai" not in blob or not re.search(r"hybrid\s*ai.*99\.2|99\.2.*hybrid", blob),
            "no fake Hybrid 99.2% accuracy",
        ),
        ("8_metrics_fpr", "fpr" in blob or "false positive rate" in blob, "FPR reported"),
        ("9_leakage", "leakage" in blob and "5.6.2" in blob, "leakage checklist"),
        ("10_attack_table", "tested (threat simulation)" in blob, "attack matrix"),
        ("11_fault_tc", "tc11" in blob, "fault test cases"),
        (
            "12_ch7_slim",
            "chapter 7 - practical implementation notes" not in blob
            and ("chapter 7 - user manual" in blob or "see §4.2.2" in blob),
            "Ch7 redundant chapter removed; user manual renumbered",
        ),
        (
            "13_chapter_breaks",
            "chapter 1" in blob and "chapter 7" in blob and "chapter 8" not in blob,
            "chapters 1–7 present; old Ch.8 removed",
        ),
        ("14_no_mislabel_dfd", "data flow diagram (dfd level 1)" not in blob, "DFD mislabel fixed"),
        ("15_diagrams_academic", "system data pipeline" in blob, "pipeline not generic DFD"),
        ("16_table_interpretation", "interpretation." in blob, "table interpretation paragraphs"),
        ("17_implementation_status", "component implementation status" in blob, "status table §18"),
        ("18_ae_lstm_honest", "not on holdout" in blob, "AE/LSTM honest"),
        ("19_defense_demo", "live demo:" in blob or "q&a:" in blob, "defense outline"),
        ("rf_metric", rf in blob, f"RF {rf}%"),
    ]
    passed = sum(1 for _, ok, _ in checks if ok)
    lines = [f"{'PASS' if ok else 'FAIL'} {name}: {detail}" for name, ok, detail in checks]
    return passed, lines


def main() -> None:
    paper = Path(sys.argv[1]) if len(sys.argv) > 1 else PAPER
    passed, lines = run_checks(paper)
    print(f"Doctor verification {passed}/{len(lines)}")
    for line in lines:
        print(line)
    if passed < len(lines):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
