"""
Final master build: doctor feedback + all figures + image swap + content integrity + single Desktop docx.
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
ASSETS = Path(__file__).resolve().parent
DESKTOP = Path(r"C:\Users\mohamad\OneDrive\Desktop")
FINAL = DESKTOP / "AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx"
BASE = DESKTOP / "AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED_pre_doctor_20260819_2336.docx"
if not BASE.exists():
    BASE = ASSETS / "WORKING_LOF.docx"
STAGING = ASSETS / "_FINAL_STAGING.docx"
METRICS = ASSETS / "evaluation" / "metrics.json"
EVAL = ASSETS / "evaluation"
FIGS = ASSETS / "new_figures"
SHOTS = ASSETS / "user_guide_shots"
SNIPS = ASSETS / "code_snippets"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def find_exact(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def insert_after(p: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    if style:
        try:
            para.style = style
        except Exception:
            pass
    if text:
        set_text(para, text)
    return para


def set_cell(cell, text: str) -> None:
    set_text(cell.paragraphs[0], text)


def replace_image_before_caption(doc: Document, caption_contains: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    paras = list(doc.paragraphs)
    blob = image_path.read_bytes()
    for i, p in enumerate(paras):
        if caption_contains not in (p.text or ""):
            continue
        for j in range(i, max(-1, i - 8), -1):
            if j < 0:
                break
            blips = paras[j]._p.findall(f".//{NS_A}blip")
            if not blips:
                continue
            r_id = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = paras[j].part.related_parts.get(r_id) if r_id else None
            if part is None:
                continue
            part._blob = blob
            return True
    return False


def insert_image_before_caption(doc: Document, caption_contains: str, image_path: Path, width_in: float = 6.2) -> bool:
    if not image_path.exists():
        return False
    for i, p in enumerate(doc.paragraphs):
        if caption_contains not in (p.text or ""):
            continue
        prev = doc.paragraphs[i - 1] if i else None
        if prev and prev._p.findall(f".//{NS_A}blip"):
            return replace_image_before_caption(doc, caption_contains, image_path)
        new_p = OxmlElement("w:p")
        p._p.addprevious(new_p)
        para = Paragraph(new_p, p._parent)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(image_path), width=Inches(width_in))
        return True
    return False


def full_text(doc: Document) -> str:
    parts = [(p.text or "") for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def restore_ch7_from_backup(doc: Document, backup: Document) -> None:
    headings = [
        "7.1 Dataset Preparation Notes",
        "7.2 Handling Class Imbalance",
        "7.3 Threshold Tuning",
        "7.4 Explainability Strategy",
        "7.5 Report Generation",
    ]

    def body_after(d: Document, heading: str) -> str | None:
        for i, p in enumerate(d.paragraphs):
            if (p.text or "").strip() == heading and p.style and "Heading" in (p.style.name or ""):
                for j in range(i + 1, len(d.paragraphs)):
                    nxt = d.paragraphs[j]
                    if nxt.style and nxt.style.name.startswith("Heading"):
                        break
                    t = (nxt.text or "").strip()
                    if t and not t.startswith("See §"):
                        return t
        return None

    for h in headings:
        original = body_after(backup, h)
        if not original:
            continue
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() == h and p.style and "Heading" in (p.style.name or ""):
                note = (
                    f"Practical note (methodology detail also in Chapters 4–5): {original[:300]}..."
                    if len(original) > 300
                    else f"Practical note (methodology detail also in Chapters 4–5): {original}"
                )
                if i + 1 < len(doc.paragraphs):
                    nxt = doc.paragraphs[i + 1]
                    if not (nxt.style and nxt.style.name.startswith("Heading")):
                        # prepend restoration if body was shortened
                        cur = (nxt.text or "").strip()
                        if cur.startswith("See §") or len(cur) < 120:
                            set_text(nxt, original)
                            insert_after(nxt, note, "Normal")
                break


def add_attack_table(doc: Document) -> None:
    anchor = find_contains(doc, "5.3.1 Structured Attack Testing Matrix")
    if not anchor or find_contains(doc, "Tested (Threat Simulation)"):
        return
    rows = [
        ["Attack", "Model(s)", "Expected", "Alert/Response", "Status"],
        ["DoS/DDoS", "RF + XGB", "Attack label", "Medium+ alert, playbook, Telegram", "Tested (Threat Simulation)"],
        ["PortScan", "RF", "PortScan", "High alert, TI enrich", "Tested"],
        ["BruteForce", "RF/XGB", "BruteForce", "Alert + optional block", "Tested"],
        ["SQLInjection", "RF/XGB", "SQLInjection", "Alert + MITRE", "Tested"],
        ["Malware/C2", "IF + TI", "Anomaly / suspicious IP", "Analyst review", "Discussed / partial TI"],
        ["Zero-day", "IF + AE", "Anomaly score high", "Medium alert, no forced label", "Simulated noise only"],
    ]
    tbl = doc.add_table(rows=len(rows), cols=5)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val)
    body = doc.element.body
    body.remove(tbl._tbl)
    anchor._p.addnext(tbl._tbl)


def update_lof_captions(doc: Document, ds_rows: int) -> None:
    mapping = {
        "Figure 4": "System data pipeline (implemented capture → AI → database → dashboard).",
        "Figure 14": f"Measured holdout performance on {ds_rows:,}-row evaluation set.",
        "Figure 15": f"Class distribution ({ds_rows:,} rows).",
    }
    for row in doc.tables[0].rows:
        no = (row.cells[0].text or "").strip()
        for k, v in mapping.items():
            if no.startswith(k):
                set_cell(row.cells[1], v)


def swap_all_images(doc: Document) -> list[str]:
    from PIL import Image

    def crop(src: Path, dst: Path) -> None:
        if not src.exists():
            return
        im = Image.open(src).convert("RGB")
        w, h = im.size
        im.crop((0, 0, w, min(h, int(h * 0.92)))).save(dst, quality=95)

    for s, d in [
        (SHOTS / "00_home.png", EVAL / "shot_home_print.png"),
        (SHOTS / "03_simulation.png", EVAL / "shot_sim_print.png"),
        (SHOTS / "10_soc.png", EVAL / "shot_soc_print.png"),
        (SHOTS / "07_models.png", EVAL / "shot_models_print.png"),
    ]:
        crop(s, d)

    mapping = [
        ("Figure 1.", FIGS / "fig01_proposed_architecture.png"),
        ("Figure 2.", FIGS / "fig02_ai_pipeline.png"),
        ("Figure 3.", FIGS / "fig_expanded_architecture.png"),
        ("Figure 4.", FIGS / "fig04_data_pipeline.png"),
        ("Figure 5.", FIGS / "fig05_use_case.png"),
        ("Figure 6.", FIGS / "fig06_sequence.png"),
        ("Figure 7.", FIGS / "fig07_activity.png"),
        ("Figure 8.", FIGS / "fig08_stride.png"),
        ("Figure 9.", FIGS / "fig_mitre_soar.png"),
        ("Figure 10.", FIGS / "fig_detection_pipeline.png"),
        ("Figure 11.", EVAL / "shot_home_print.png"),
        ("Figure 12.", EVAL / "fig_cm_example.png"),
        ("Figure 13.", EVAL / "fig_roc_example.png"),
        ("Figure 14.", EVAL / "fig_model_performance.png"),
        ("Figure 15.", EVAL / "fig_dataset_distribution.png"),
        ("Figure 16.", EVAL / "fig_model_performance.png"),
        ("Figure 17.", EVAL / "fig_confusion_rf.png"),
        ("Figure 18.", EVAL / "shot_home_print.png"),
        ("Figure 19.", EVAL / "shot_sim_print.png"),
        ("Figure 20.", EVAL / "shot_soc_print.png"),
        ("Figure 21.", EVAL / "shot_models_print.png"),
        ("Figure 22.", FIGS / "fig_server_supervisor.png"),
        ("Figure H1", SNIPS / "01_hybrid_ai_decision_fusion.png"),
        ("Figure H2", SNIPS / "09_mitre_map.png"),
        ("Figure H3", SNIPS / "10_soar_playbooks.png"),
        ("Figure H4", SNIPS / "12_online_learning.png"),
        ("Figure H5", SNIPS / "14_threat_simulation.png"),
        ("Figure H6", SNIPS / "13_alert_manager.png"),
        ("Figure H7", SNIPS / "16_telegram_alert.png"),
        ("Figure H8", SNIPS / "15_supervisor.png"),
        ("Figure H9", SNIPS / "17_ai_assistant.png"),
    ]
    done = []
    for cap, img in mapping:
        if replace_image_before_caption(doc, cap, img) or insert_image_before_caption(doc, cap, img):
            done.append(cap.split(".")[0].split()[0])
    return done


def integrity_check(doc: Document, backup: Document) -> tuple[int, int]:
    b = full_text(backup).lower()
    d = full_text(doc).lower()
    must = [
        "declaration", "harvard", "chapter 1", "chapter 8", "appendix h",
        "darktrace", "random forest", "references", "conclusion",
        "mitre", "telegram", "threat simulation", "copilot",
    ]
    ok = sum(1 for m in must if m in d)
    bw = len(re.findall(r"\w+", b))
    dw = len(re.findall(r"\w+", d))
    total = len(must)
    if dw < int(bw * 0.92):
        ok = min(ok, total - 1)
    return ok, total


def word_toc_update(path: Path) -> None:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        wdoc = word.Documents.Open(str(path), ReadOnly=False)
        wdoc.Fields.Update()
        for i in range(1, wdoc.TablesOfContents.Count + 1):
            wdoc.TablesOfContents(i).Update()
        wdoc.Save()
        wdoc.Close()
        word.Quit()
    except Exception as exc:
        print("Word TOC skip:", exc)


def clean_desktop(keep: Path) -> None:
    for p in DESKTOP.glob("*AUCE*"):
        if p.resolve() != keep.resolve():
            try:
                p.unlink()
                print("Removed", p.name)
            except Exception as exc:
                print("Could not remove", p.name, exc)


def main() -> None:
    print("=== 1. Regenerate metrics & figures ===")
    subprocess.run([sys.executable, "-m", "training.evaluate_models"], check=True, cwd=str(ROOT))
    subprocess.run([sys.executable, str(ASSETS / "generate_new_figures.py")], check=True, cwd=str(ASSETS))
    subprocess.run([sys.executable, str(ASSETS / "generate_academic_diagrams.py")], check=True, cwd=str(ASSETS))
    subprocess.run([sys.executable, str(ASSETS / "render_eval_figures.py")], check=True, cwd=str(ASSETS))

    print("=== 2. Base document from backup (preserve content) ===")
    shutil.copy2(BASE, STAGING)
    backup_doc = Document(str(BASE))

    print("=== 3. Doctor revision ===")
    import doctor_email_revision as dr

    dr.PAPER = STAGING
    dr.BACKUP = ASSETS / "_dr_revision_backup.docx"
    dr.main()

    print("=== 4. Complete doctor revision (19 points) ===")
    import doctor_complete_revision as dcr

    dcr.PAPER = STAGING
    dcr.main()

    doc = Document(str(STAGING))
    m = json.loads(METRICS.read_text(encoding="utf-8"))
    update_lof_captions(doc, m["dataset"]["rows"])

    # Hybrid fusion paragraph + diagram reference
    fusion = find_exact(doc, "4.4 Decision Fusion")
    if fusion and not find_contains(doc, "fig_hybrid_fusion"):
        insert_after(
            fusion,
            "Figure 2 and the expanded architecture show how decision_engine.fuse_decisions() aggregates "
            "model votes into final_label, confidence, threat_score, and severity. Hybrid AI is not trained "
            "as a separate classifier; its benefit is reduced false negatives when multiple models agree on an attack.",
            "Normal",
        )

    print("=== 5. Swap all images ===")
    swapped = swap_all_images(doc)
    print("Swapped", len(swapped), "figures")

    # Page breaks before chapters
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if p.style and p.style.name == "Heading 1" and t.startswith("Chapter "):
            br = p.insert_paragraph_before("")
            br.add_run().add_break(WD_BREAK.PAGE)

    doc.save(str(STAGING))

    print("=== 6. Integrity check ===")
    doc = Document(str(STAGING))
    ok, total = integrity_check(doc, backup_doc)
    print(f"Integrity {ok}/{total}")
    if ok < total:
        raise SystemExit("Content integrity check failed")

    print("=== 7. Doctor verify x10 ===")
    import verify_doctor_revision as vd

    vd.PAPER = STAGING
    for i in range(10):
        passed, lines = vd.run_checks(STAGING)
        if passed < len(lines):
            raise SystemExit(f"Verify failed run {i+1}")
    print("10/10 verification passes")

    print("=== 8. Word TOC + save Desktop ===")
    shutil.copy2(STAGING, FINAL)
    word_toc_update(FINAL)

    print("=== 9. Clean Desktop (Word only) ===")
    clean_desktop(FINAL)
    print("DONE:", FINAL, "size", FINAL.stat().st_size)


if __name__ == "__main__":
    main()
