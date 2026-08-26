"""Second-pass citation fixer: two-phase apply + cite remaining bibliography."""
from __future__ import annotations

import re
import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt

ROOT = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
SRC = ROOT / "WORKING_UPDATED_v2.docx"
OUT = ROOT / "WORKING_UPDATED_v3.docx"
DESK = Path(r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx")
REPORT = ROOT / "_refs_figures_review_pass2.txt"


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return new_para


def set_para_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def ensure_cite(text: str, cite: str) -> str:
    # cite may be like "[9]" or "[3][24]" — only append missing pieces
    parts = re.findall(r"\[\d+\]", cite)
    if not parts:
        return text
    missing = [p for p in parts if p not in text]
    if not missing:
        return text
    return text.rstrip() + " " + "".join(missing)


def find_refs_heading(doc: Document):
    for p in doc.paragraphs:
        if (p.text or "").strip() == "References":
            return p
    return None


def is_bib(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^\[\d+\]\s", t))


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    log = []

    vendor = [
        ("Darktrace was reviewed", "[9]"),
        ("Vectra AI was reviewed", "[30]"),
        ("ExtraHop RevealX was reviewed", "[11]"),
        ("Corelight Open NDR was reviewed", "[8]"),
        ("Cisco Secure Network Analytics was reviewed", "[6]"),
        ("Palo Alto Cortex XSIAM was reviewed", "[19]"),
        ("Microsoft Sentinel was reviewed", "[16]"),
        ("Splunk Enterprise Security was reviewed", "[27]"),
        ("IBM QRadar was reviewed", "[13]"),
        ("Elastic Security was reviewed", "[10]"),
        ("CrowdStrike Falcon was reviewed", "[41]"),
        ("Trend Micro Vision One was reviewed", "[29]"),
        ("Google Security Operations was reviewed", "[12]"),
        ("Rapid7 InsightIDR was reviewed", "[21]"),
        ("Security Onion 2 was reviewed", "[23]"),
        ("Zeek was reviewed", "[32]"),
        ("Suricata was reviewed", "[28]"),
        ("Wazuh was reviewed", "[31]"),
        ("Snort 3 was reviewed", "[25]"),
        ("Cisco XDR was reviewed", "[7]"),
    ]

    # Additional claim-level citations (needle in paragraph -> cite)
    claims = [
        ("An Intrusion Detection System monitors", "[14]"),
        ("AI-based IDS uses Machine Learning", "[1][26]"),
        ("signature-based detection compares traffic", "[14][25]"),
        ("Datasets Used in Network Intrusion Detection", None),  # heading skip
        ("CICIDS2017", "[3][24]"),
        ("NSL-KDD", "[4]"),
        ("UNSW-NB15", "[17]"),
        ("NIST Cybersecurity Framework", "[18]"),
        ("Isolation Forest", "[15][22]"),
        ("Random Forest", "[2]"),
        ("XGBoost", "[5]"),
        ("scikit-learn", "[20]"),
        ("Scikit-learn", "[20]"),
        ("FastAPI", "[33]"),
        ("Streamlit", "[34]"),
        ("Telegram", "[35]"),
        ("MITRE ATT&CK", "[36]"),
        ("SQLite", "[37]"),
        ("uvicorn", "[38]"),
        ("Uvicorn", "[38]"),
        ("AbuseIPDB", "[39]"),
        ("partial_fit", "[40]"),
        ("SGDClassifier", "[40]"),
        ("run_server.bat", "[33][38]"),
    ]

    refs_h = find_refs_heading(doc)
    updates: list[tuple[Paragraph, str]] = []

    # Phase 1 collect vendor
    for p in list(doc.paragraphs):
        if refs_h and p._p is refs_h._p:
            break
        t = (p.text or "").strip()
        if not t or is_bib(t) or t.startswith("Figure ") or t.startswith("Table "):
            continue
        if p.style and str(p.style.name).startswith("Heading"):
            # still allow claim cites on some long headings? skip
            continue
        new = t
        for needle, cite in vendor:
            if needle in new:
                new = ensure_cite(new, cite)
                break
        for needle, cite in claims:
            if cite and needle in new:
                new = ensure_cite(new, cite)
        if new != t:
            updates.append((p, new))

    # Phase 2 apply
    for p, new in updates:
        set_para_text(p, new)
        log.append(f"Updated: {new[:90]}...")

    # Ensure literature gap / methodology mentions cite surveys
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("The reviewed systems show that modern cybersecurity"):
            new = ensure_cite(t, "[1][14]")
            if new != t:
                set_para_text(p, new)
                log.append("Cited literature gap with [1][14]")
        if "Design Science Research" in t:
            # no specific ref required
            pass
        if "Threat Intelligence enrichment" in t and "AbuseIPDB" not in t and "geo-location" in t.lower():
            new = ensure_cite(t, "[39]")
            if new != t:
                set_para_text(p, new)

    # If AbuseIPDB never appears, add a TI sentence with cite near Threat Intelligence module text
    body = "\n".join(p.text or "" for p in doc.paragraphs)
    if "AbuseIPDB" not in body.split("References")[0]:
        anchor = None
        for p in doc.paragraphs:
            if "Threat Intelligence enriches source IPs" in (p.text or ""):
                anchor = p
                break
        if anchor:
            set_para_text(
                anchor,
                (anchor.text or "").rstrip()
                + " Reputation checks can use providers such as AbuseIPDB [39].",
            )
            log.append("Added AbuseIPDB [39] mention")

    if "[38]" not in "\n".join(p.text or "" for p in doc.paragraphs if True):
        pass
    # Ensure uvicorn cited in deployment paragraph
    for p in doc.paragraphs:
        t = p.text or ""
        if "run_server.bat" in t and "[38]" not in t:
            set_para_text(p, ensure_cite(t, "[38]"))
            log.append("Cited uvicorn/server with [38]")
            break

    doc.save(str(OUT))

    # ===== DOUBLE VERIFY =====
    doc2 = Document(str(OUT))
    refs_h = find_refs_heading(doc2)
    ref_map = {}
    started = False
    for p in doc2.paragraphs:
        if refs_h and p._p is refs_h._p:
            started = True
            continue
        if not started:
            continue
        t = (p.text or "").strip()
        m = re.match(r"^\[(\d+)\]\s+(.+)$", t)
        if m:
            ref_map[int(m.group(1))] = m.group(2)
        elif t:
            log.append(f"ERROR unnumbered ref: {t[:80]}")

    body_parts = []
    for p in doc2.paragraphs:
        if refs_h and p._p is refs_h._p:
            break
        body_parts.append(p.text or "")
    body = "\n".join(body_parts)
    cites = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", body)))
    missing = [c for c in cites if c not in ref_map]
    unused = [n for n in sorted(ref_map) if n not in cites]

    # LoF check
    body_figs = []
    after = False
    for p in doc2.paragraphs:
        if (p.text or "").strip() == "Chapter 1 - Introduction":
            after = True
        if not after:
            continue
        m = re.match(r"^Figure\s+(\d+)\.\s*(.+)$", (p.text or "").strip())
        if m:
            body_figs.append((int(m.group(1)), m.group(2).strip()))
    table_figs = []
    for row in doc2.tables[0].rows[1:]:
        m = re.match(r"Figure\s+(\d+)", row.cells[0].text.strip())
        if m:
            table_figs.append((int(m.group(1)), row.cells[1].text.strip()))

    acr_n = len(doc2.tables[2].rows) - 1

    log.append("\n===== VERIFY PASS 1 =====")
    log.append(f"Figures body={len(body_figs)} table={len(table_figs)} match={body_figs==table_figs}")
    log.append(f"Figure nums sequential={ [n for n,_ in body_figs]==list(range(1,len(body_figs)+1)) }")
    log.append(f"Acronyms={acr_n}")
    log.append(f"Refs={len(ref_map)} contiguous={list(ref_map)==list(range(1,max(ref_map)+1))}")
    log.append(f"Inline cites={cites}")
    log.append(f"Missing defs={missing or 'NONE'}")
    log.append(f"Unused refs={unused or 'NONE'}")

    log.append("\n===== VERIFY PASS 2 (critical) =====")
    checks = vendor + [
        ("Random Forest", "[2]"),
        ("XGBoost", "[5]"),
        ("Isolation Forest", "[15]"),
        ("Streamlit", "[34]"),
        ("FastAPI", "[33]"),
        ("SQLite", "[37]"),
        ("Telegram", "[35]"),
        ("MITRE", "[36]"),
        ("AbuseIPDB", "[39]"),
        ("An Intrusion Detection System monitors", "[14]"),
        ("AI-based IDS uses Machine Learning", "[1]"),
    ]
    fails = 0
    for needle, cite in checks:
        ok = any(needle in (p.text or "") and cite in (p.text or "") for p in doc2.paragraphs)
        # for cite like [1][26], require first at least
        if not ok and cite.startswith("["):
            first = re.match(r"\[\d+\]", cite)
            if first:
                ok = any(needle in (p.text or "") and first.group(0) in (p.text or "") for p in doc2.paragraphs)
        status = "OK" if ok else "FAIL"
        if not ok:
            fails += 1
        log.append(f"{status}: {needle} :: {cite}")

    log.append(f"\nPASS2 failures={fails}")
    log.append("Leftover letter figures=" + str([
        (p.text or "").strip() for p in doc2.paragraphs
        if re.match(r"^Figure\s+[A-Z]\.", (p.text or "").strip() or "")
    ]))

    REPORT.write_text("\n".join(log), encoding="utf-8")
    shutil.copy2(OUT, DESK)
    shutil.copy2(OUT, ROOT / "WORKING_UPDATED.docx")  # keep working sync
    print("failures", fails)
    print("unused", unused)
    print("figures", len(body_figs), "acronyms", acr_n, "refs", len(ref_map))
    print("desktop", DESK)


if __name__ == "__main__":
    main()
