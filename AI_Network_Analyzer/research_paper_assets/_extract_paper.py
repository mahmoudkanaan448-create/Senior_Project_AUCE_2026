from docx import Document
from pathlib import Path
import json
import re

path = Path(r"c:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL.docx")
out_dir = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
out_dir.mkdir(parents=True, exist_ok=True)

doc = Document(str(path))
print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))

heads = []
full_lines = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    text = (p.text or "").strip()
    if text:
        full_lines.append(f"[{i}|{style}] {text}")
    if not text:
        continue
    if style.startswith("Heading") or style.startswith("Title"):
        heads.append((i, style, text[:160]))

print("=== HEADINGS ===")
for h in heads:
    print(f"{h[0]:4d} | {h[1]:20s} | {h[2]}")

words = sum(len((p.text or "").split()) for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            words += len((cell.text or "").split())
print("approx_words", words)

nimg = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        nimg += 1
print("images", nimg)

# Save full text for analysis
(out_dir / "_paper_extract.txt").write_text("\n".join(full_lines), encoding="utf-8")

# Extract images
img_dir = out_dir / "paper_existing_images"
img_dir.mkdir(exist_ok=True)
idx = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        idx += 1
        blob = rel.target_part.blob
        ext = Path(rel.target_ref).suffix or ".png"
        (img_dir / f"img_{idx:02d}{ext}").write_bytes(blob)
print("extracted_images", idx)

# Keywords audit
text_all = "\n".join(full_lines).lower()
keys = [
    "mitre", "att&ck", "soar", "playbook", "online learning", "incremental",
    "telegram", "whatsapp", "threat simulation", "isolation forest", "xgboost",
    "random forest", "lstm", "autoencoder", "streamlit", "fastapi", "explainable",
    "xai", "firewall", "supervisor", "watchdog", "self-heal", "abuseipdb",
    "virustotal", "sqlite", "hybrid", "confidence", "reference",
]
print("=== KEYWORD PRESENCE ===")
for k in keys:
    print(f"{k}: {'YES' if k in text_all else 'NO'}")

# References section sample
print("=== LAST 40 NON-EMPTY PARAS ===")
nonempty = [l for l in full_lines if l.split("] ",1)[-1].strip()]
for l in nonempty[-40:]:
    print(l[:200])
