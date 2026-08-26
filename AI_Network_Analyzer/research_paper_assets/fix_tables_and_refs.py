"""
Fix List of Figures + List of Acronyms as proper tables,
renumber figures sequentially, strengthen citations, and double-verify.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\mohamad\.vscode\Senior_Project_AUCE_2026\AI_Network_Analyzer\research_paper_assets")
SRC = ROOT / "WORKING_UPDATED.docx"
OUT_LOCAL = ROOT / "WORKING_UPDATED_v2.docx"
OUT_DESKTOP = Path(r"C:\Users\mohamad\OneDrive\Desktop\AI_Powered_Network_Traffic_Analyser_Anomaly_Detector_AUCE_FINAL_UPDATED.docx")
FIGS = ROOT / "new_figures"
REPORT = ROOT / "_refs_figures_review.txt"

# Body caption text (current) -> new sequential number and clean title
# Filled dynamically after scan.

VENDOR_CITES = [
    ("Darktrace was reviewed", "[9]"),
    ("Vectra AI was reviewed", "[30]"),
    ("ExtraHop RevealX was reviewed", "[11]"),
    ("Corelight Open NDR was reviewed", "[8]"),
    ("Cisco Secure Network Analytics was reviewed", "[6]"),
    ("Palo Alto Cortex XSIAM was reviewed", "[19]"),
    ("Microsoft Sentinel was reviewed", "[16]"),
    ("Splunk Enterprise Security was reviewed", "[27]"),
    ("IBM QRadar was reviewed", "[13]"),
    ("Elastic Security was reviewed", "[10]"),
    ("Trend Micro Vision One was reviewed", "[29]"),
    ("Google Security Operations was reviewed", "[12]"),
    ("Rapid7 InsightIDR was reviewed", "[21]"),
    ("Security Onion 2 was reviewed", "[23]"),
    ("Zeek was reviewed", "[32]"),
    ("Suricata was reviewed", "[28]"),
    ("Wazuh was reviewed", "[31]"),
    ("Snort 3 was reviewed", "[25]"),
    ("Cisco XDR was reviewed", "[7]"),
]

TECH_CITES = [
    # (substring to find in paragraph, citation to ensure present, only if paragraph also contains hint)
    ("Random Forest", "[2]", None),
    ("XGBoost", "[5]", None),
    ("Isolation Forest", "[15]", None),
    ("CICIDS2017", "[3][24]", None),
    ("NSL-KDD", "[4]", None),
    ("UNSW-NB15", "[17]", None),
    ("scikit-learn", "[20]", None),
    ("Scikit-learn", "[20]", None),
    ("FastAPI", "[33]", None),
    ("Streamlit", "[34]", None),
    ("Telegram", "[35]", None),
    ("MITRE ATT&CK", "[36]", None),
    ("SQLite", "[37]", None),
    ("AbuseIPDB", "[39]", None),
    ("partial_fit", "[40]", None),
    ("SGDClassifier", "[40]", None),
    ("NIST Cybersecurity Framework", "[18]", None),
]


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_in: float, caption: str) -> Paragraph:
    p = insert_after(paragraph, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = insert_after(p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)
    return cap


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if (p.text or "").strip() == exact:
            return p
    return None


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    return None


def set_para_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clear_table_keep_header(table) -> None:
    # Remove all rows except header (index 0)
    tbl = table._tbl
    trs = list(tbl.tr_lst)
    for tr in trs[1:]:
        tbl.remove(tr)


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = val


def ensure_cite(text: str, cite: str) -> str:
    """Append citation if none of its numbers already present nearby."""
    nums = re.findall(r"\d+", cite)
    if any(f"[{n}]" in text for n in nums):
        return text
    # Avoid putting cites on pure headings
    if len(text) < 25:
        return text
    if text.rstrip().endswith(cite):
        return text
    return text.rstrip() + f" {cite}"


def is_bibliography_line(text: str) -> bool:
    t = text.strip()
    if re.match(r"^\[\d+\]\s", t):
        return True
    if t.startswith(("Ahmad,", "Breiman,", "Canadian", "Chen,", "Cisco.", "Corelight",
                     "Darktrace", "Elastic", "ExtraHop", "Google", "IBM.", "Khraisat",
                     "Liu,", "Microsoft", "Moustafa", "National", "Palo Alto", "Pedregosa",
                     "Rapid7", "Scikit-learn", "Security Onion", "Sharafaldin", "Snort",
                     "Sommer", "Splunk", "Suricata", "Trend Micro", "Vectra", "Wazuh",
                     "Zeek", "FastAPI", "Streamlit", "Telegram", "MITRE Corporation",
                     "The SQLite", "Uvicorn", "AbuseIPDB")):
        return True
    return False


def collect_body_figures(doc: Document) -> list[tuple[Paragraph, str, str]]:
    """Return list of (paragraph, old_label, title) for figure captions in body order."""
    out = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        m = re.match(r"^Figure\s+([0-9A-Z]+)\.\s*(.+)$", t)
        if not m:
            continue
        # Skip LoF list duplicates that sit under List of Figures before Chapter 1
        # We still collect all; filtering by section later
        out.append((p, m.group(1), m.group(2).strip()))
    return out


def rebuild_list_of_figures(doc: Document, report: list[str]) -> dict[str, int]:
    """
    Renumber body figures in appearance order (excluding the temporary LoF bullet list),
    rebuild Table 0, remove loose LoF paragraphs.
    Returns mapping old_label_unique_key -> new_number (by paragraph id order).
    """
    lof = find_para(doc, "List of Figures")
    lot = find_para(doc, "List of Tables")
    ch1 = find_para(doc, "Chapter 1 - Introduction")

    # Delete loose Normal lines between LoF and LoT that look like figure entries
    if lof and lot:
        to_delete = []
        started = False
        for p in doc.paragraphs:
            if p._p is lof._p:
                started = True
                continue
            if not started:
                continue
            if p._p is lot._p:
                break
            t = (p.text or "").strip()
            if t.startswith("Figure "):
                to_delete.append(p)
        for p in to_delete:
            delete_paragraph(p)
        report.append(f"Removed {len(to_delete)} loose LoF paragraph entries")

    # Ensure Figure C exists in body (pipeline diagram)
    if not find_contains(doc, "End-to-end detection and response pipeline implemented"):
        fuse = find_contains(doc, "The proposed system does not depend on one model only")
        img = FIGS / "fig_detection_pipeline.png"
        if fuse and img.exists():
            insert_picture_after(
                fuse, img, 6.3,
                "Figure C. End-to-end detection and response pipeline implemented in software.",
            )
            report.append("Inserted missing Figure C (pipeline) into body")

    # Collect body figures after Chapter 1 only (true captions)
    body_figs = []
    after_ch1 = False
    for p in doc.paragraphs:
        if ch1 and p._p is ch1._p:
            after_ch1 = True
        if not after_ch1:
            # also allow figures that appear before ch1? Architecture starts in ch3
            continue
        t = (p.text or "").strip()
        m = re.match(r"^Figure\s+([0-9A-Z]+)\.\s*(.+)$", t)
        if m:
            body_figs.append((p, m.group(1), m.group(2).strip()))

    # Renumber sequentially
    mapping_titles = []
    for idx, (p, old, title) in enumerate(body_figs, start=1):
        new_cap = f"Figure {idx}. {title}"
        set_para_text(p, new_cap)
        mapping_titles.append((idx, title, old))
        report.append(f"Caption renumber: Figure {old} -> Figure {idx} | {title}")

    # Rebuild LoF table (tables[0])
    table = doc.tables[0]
    clear_table_keep_header(table)
    # Ensure header
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Figure Title"
    for num, title, _old in mapping_titles:
        add_table_row(table, [f"Figure {num}", title])
    report.append(f"LoF table rebuilt with {len(mapping_titles)} figures")
    return {f"{old}:{title}": num for num, title, old in mapping_titles}


def rebuild_acronyms(doc: Document, report: list[str]) -> None:
    loa = find_para(doc, "List of Acronyms")
    ch1 = find_para(doc, "Chapter 1 - Introduction")

    # Remove loose acronym paragraphs between LoA heading and Chapter 1
    if loa and ch1:
        to_delete = []
        started = False
        for p in doc.paragraphs:
            if p._p is loa._p:
                started = True
                continue
            if not started:
                continue
            if p._p is ch1._p:
                break
            t = (p.text or "").strip()
            if " - " in t or t.startswith(("ATT&CK", "MITRE", "SOAR", "SGD", "SOC", "SQLite")):
                to_delete.append(p)
        for p in to_delete:
            delete_paragraph(p)
        report.append(f"Removed {len(to_delete)} loose LoA paragraph entries")

    # Desired full acronym list (alphabetical by acronym)
    acronyms = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("ATT&CK", "Adversarial Tactics, Techniques, and Common Knowledge"),
        ("AUC", "Area Under Curve"),
        ("CSV", "Comma-Separated Values"),
        ("DDoS", "Distributed Denial of Service"),
        ("DFD", "Data Flow Diagram"),
        ("DL", "Deep Learning"),
        ("DoS", "Denial of Service"),
        ("IDS", "Intrusion Detection System"),
        ("IF", "Isolation Forest"),
        ("IPS", "Intrusion Prevention System"),
        ("JWT", "JSON Web Token"),
        ("LSTM", "Long Short-Term Memory"),
        ("MITRE", "MITRE Corporation (ATT&CK knowledge base)"),
        ("ML", "Machine Learning"),
        ("NDR", "Network Detection and Response"),
        ("NIDS", "Network Intrusion Detection System"),
        ("REST", "Representational State Transfer"),
        ("RF", "Random Forest"),
        ("ROC", "Receiver Operating Characteristic"),
        ("SGD", "Stochastic Gradient Descent"),
        ("SIEM", "Security Information and Event Management"),
        ("SOAR", "Security Orchestration, Automation and Response"),
        ("SOC", "Security Operations Center"),
        ("SQL", "Structured Query Language"),
        ("SQLite", "Embedded Relational Database Engine"),
        ("STRIDE", "Spoofing, Tampering, Repudiation, Information Disclosure, DoS, Elevation of Privilege"),
        ("TCP", "Transmission Control Protocol"),
        ("TI", "Threat Intelligence"),
        ("UDP", "User Datagram Protocol"),
        ("UML", "Unified Modeling Language"),
        ("XAI", "Explainable Artificial Intelligence"),
        ("XDR", "Extended Detection and Response"),
        ("XGB", "Extreme Gradient Boosting (XGBoost)"),
    ]

    table = doc.tables[2]
    clear_table_keep_header(table)
    table.rows[0].cells[0].text = "Acronym"
    table.rows[0].cells[1].text = "Meaning"
    for acr, meaning in acronyms:
        add_table_row(table, [acr, meaning])
    report.append(f"Acronym table rebuilt with {len(acronyms)} entries")

    # Update List of Tables entry if needed (Table 1 remains List of Acronyms)
    # tables[1] is LoT - leave as is


def strengthen_citations(doc: Document, report: list[str]) -> None:
    changed = 0

    # Vendor first-mention paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t or is_bibliography_line(t):
            continue
        for needle, cite in VENDOR_CITES:
            if needle in t:
                new = ensure_cite(t, cite)
                if new != t:
                    set_para_text(p, new)
                    changed += 1
                    report.append(f"Vendor cite: {needle} -> {cite}")
                break

    # Tech citations (skip captions/headings/short)
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t or is_bibliography_line(t):
            continue
        if t.startswith("Figure ") or (p.style and str(p.style.name).startswith("Heading")):
            continue
        if t.startswith("Table "):
            continue
        new = t
        for needle, cite, _ in TECH_CITES:
            if needle in new:
                new = ensure_cite(new, cite)
        if new != t:
            set_para_text(p, new)
            changed += 1
    report.append(f"Citation updates applied (paragraph touches≈{changed})")

    # Fix abstract stack to match implemented app + cites
    for p in doc.paragraphs:
        t = p.text or ""
        if "PostgreSQL" in t and "FastAPI" in t:
            set_para_text(
                p,
                "The proposed system integrates multiple Artificial Intelligence models, including Random Forest [2], "
                "XGBoost [5], Isolation Forest [15], Autoencoder, and LSTM. The design combines supervised detection, "
                "unsupervised anomaly detection, sequential attack prediction, Explainable AI (XAI), Threat Intelligence, "
                "geo-location, a real-time Streamlit dashboard [34], SQLite storage [37], Telegram notifications [35], "
                "MITRE ATT&CK enrichment [36], SOAR-style playbooks, online incremental learning [40], and automatic IP "
                "blocking. The system is implemented using open-source technologies such as Python, FastAPI [33], "
                "SQLite [37], Scikit-learn [20], XGBoost [5], Streamlit [34], and related networking libraries.",
            )
            report.append("Abstract updated to match implemented stack + citations")
            break

    # CrowdStrike has no bibliography entry — add [41] and cite
    refs_head = find_para(doc, "References")
    if refs_head:
        # check if CrowdStrike ref exists
        has_cs = any("CrowdStrike" in (p.text or "") for p in doc.paragraphs)
        # add bibliography if missing numbered CrowdStrike
        has_cs_ref = any(re.match(r"^\[41\].*CrowdStrike", (p.text or "").strip()) for p in doc.paragraphs)
        if not has_cs_ref:
            # append after last ref
            last = None
            started = False
            for p in doc.paragraphs:
                if p._p is refs_head._p:
                    started = True
                    continue
                if started and (p.text or "").strip():
                    last = p
            if last:
                insert_after(
                    last,
                    "[41] CrowdStrike. (n.d.). CrowdStrike Falcon platform. https://www.crowdstrike.com/platform/",
                )
                report.append("Added reference [41] CrowdStrike")
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t.startswith("CrowdStrike Falcon was reviewed") and "[41]" not in t:
                set_para_text(p, ensure_cite(t, "[41]"))
                report.append("Cited CrowdStrike with [41]")
                break


def verify_twice(doc: Document, report: list[str]) -> None:
    report.append("\n========== REVIEW PASS 1 ==========")
    # Figures: table vs body
    body_caps = []
    ch1 = find_para(doc, "Chapter 1 - Introduction")
    after = False
    for p in doc.paragraphs:
        if ch1 and p._p is ch1._p:
            after = True
        if not after:
            continue
        t = (p.text or "").strip()
        m = re.match(r"^Figure\s+(\d+)\.\s*(.+)$", t)
        if m:
            body_caps.append((int(m.group(1)), m.group(2).strip()))

    table_figs = []
    for row in doc.tables[0].rows[1:]:
        no = row.cells[0].text.strip()
        title = row.cells[1].text.strip()
        m = re.match(r"Figure\s+(\d+)", no)
        if m:
            table_figs.append((int(m.group(1)), title))

    report.append(f"Body figure captions: {len(body_caps)}")
    report.append(f"LoF table rows: {len(table_figs)}")
    if [x[0] for x in body_caps] != list(range(1, len(body_caps) + 1)):
        report.append("ERROR: body figure numbers not sequential")
    else:
        report.append("OK: body figure numbers sequential 1..N")
    if len(body_caps) != len(table_figs):
        report.append(f"ERROR: LoF count mismatch body={len(body_caps)} table={len(table_figs)}")
    else:
        mismatches = [(a, b) for a, b in zip(body_caps, table_figs) if a != b]
        if mismatches:
            report.append(f"ERROR: LoF/body mismatches: {mismatches[:5]}")
        else:
            report.append("OK: LoF table matches body captions")

    # Acronyms
    acr_rows = [(r.cells[0].text.strip(), r.cells[1].text.strip()) for r in doc.tables[2].rows[1:]]
    report.append(f"Acronym rows: {len(acr_rows)}")
    # loose leftovers?
    loa = find_para(doc, "List of Acronyms")
    ch1 = find_para(doc, "Chapter 1 - Introduction")
    loose = 0
    if loa and ch1:
        started = False
        for p in doc.paragraphs:
            if p._p is loa._p:
                started = True
                continue
            if not started:
                continue
            if p._p is ch1._p:
                break
            if (p.text or "").strip():
                # ignore empty; table is not a paragraph
                loose += 1
    report.append(f"Loose LoA paragraphs (should be 0): {loose}")

    # References map
    ref_map = {}
    started = False
    refs_head = find_para(doc, "References")
    for p in doc.paragraphs:
        if refs_head and p._p is refs_head._p:
            started = True
            continue
        if not started:
            continue
        t = (p.text or "").strip()
        m = re.match(r"^\[(\d+)\]\s+(.+)$", t)
        if m:
            ref_map[int(m.group(1))] = m.group(2)
        elif t:
            report.append(f"ERROR: unnumbered bibliography line: {t[:80]}")

    expected = list(range(1, max(ref_map) + 1)) if ref_map else []
    missing_nums = [n for n in expected if n not in ref_map]
    report.append(f"References numbered: {len(ref_map)} (max={max(ref_map) if ref_map else 0})")
    if missing_nums:
        report.append(f"ERROR: missing reference numbers: {missing_nums}")
    else:
        report.append("OK: reference numbers contiguous")

    # Inline cites vs refs (exclude bibliography)
    body_text = []
    started = True
    for p in doc.paragraphs:
        if refs_head and p._p is refs_head._p:
            break
        body_text.append(p.text or "")
    body = "\n".join(body_text)
    cites = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", body)))
    missing_defs = [c for c in cites if c not in ref_map]
    unused = [n for n in ref_map if n not in cites]
    report.append(f"Inline unique citations: {cites}")
    report.append(f"Cited but missing in References: {missing_defs or 'NONE'}")
    report.append(f"In References but never cited in body: {unused or 'NONE'}")

    report.append("\n========== REVIEW PASS 2 ==========")
    # Re-check critical tech mentions have cites
    checks = [
        ("Random Forest", "[2]"),
        ("XGBoost", "[5]"),
        ("Isolation Forest", "[15]"),
        ("Streamlit", "[34]"),
        ("FastAPI", "[33]"),
        ("SQLite", "[37]"),
        ("Telegram", "[35]"),
        ("MITRE", "[36]"),
        ("Darktrace was reviewed", "[9]"),
        ("Zeek was reviewed", "[32]"),
        ("Suricata was reviewed", "[28]"),
        ("Snort 3 was reviewed", "[25]"),
    ]
    for needle, cite in checks:
        ok = False
        for p in doc.paragraphs:
            t = p.text or ""
            if refs_head and p._p is refs_head._p:
                break
            if needle in t and cite in t:
                ok = True
                break
            # for model names, any paragraph containing both is enough once
        report.append(f"PASS2 {'OK' if ok else 'FAIL'}: {needle} with {cite}")

    # Figure letters leftover?
    leftover_letters = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if re.match(r"^Figure\s+[A-Z]\.", t):
            leftover_letters.append(t)
    report.append(f"Leftover lettered figures: {leftover_letters or 'NONE'}")

    report.append(f"LoF table final count={len(doc.tables[0].rows)-1}")
    report.append(f"Acronym table final count={len(doc.tables[2].rows)-1}")
    for i, (n, title) in enumerate(table_figs, 1):
        report.append(f"  LoF[{i}] Figure {n}: {title}")


def main() -> None:
    shutil.copy2(SRC, OUT_LOCAL)
    doc = Document(str(OUT_LOCAL))
    report: list[str] = []

    rebuild_list_of_figures(doc, report)
    rebuild_acronyms(doc, report)
    strengthen_citations(doc, report)

    doc.save(str(OUT_LOCAL))

    # Reload for clean verification (pass 1+2)
    doc2 = Document(str(OUT_LOCAL))
    verify_twice(doc2, report)

    REPORT.write_text("\n".join(report), encoding="utf-8")
    # Copy to Desktop
    shutil.copy2(OUT_LOCAL, OUT_DESKTOP)
    print("Saved local:", OUT_LOCAL)
    print("Saved desktop:", OUT_DESKTOP)
    print("Report:", REPORT)
    # print summary tail
    for line in report[-40:]:
        print(line)


if __name__ == "__main__":
    main()
